import os
import re
import io
import time
from datetime import datetime, timedelta
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
import threading
import json
import calendar
import gc
import queue
import customtkinter as ctk
from customtkinter import CTkScrollableFrame, CTkButton, CTkLabel, CTkEntry, CTkTextbox, CTkFrame, CTkComboBox

multiprocessing = None
ProcessPoolExecutor = None
ThreadPoolExecutor = None
as_completed = None

def _ensure_concurrent_imports():
    """Гарантирует что multiprocessing и concurrent.futures загружены"""
    global multiprocessing, ProcessPoolExecutor, ThreadPoolExecutor, as_completed
    if multiprocessing is None:
        import multiprocessing as _mp
        from concurrent.futures import ProcessPoolExecutor as _PPE
        from concurrent.futures import ThreadPoolExecutor as _TPE
        from concurrent.futures import as_completed as _ac
        multiprocessing = _mp
        ProcessPoolExecutor = _PPE
        ThreadPoolExecutor = _TPE
        as_completed = _ac

try:
    from pypdf import PdfMerger
except ImportError:
    try:
        from PyPDF2 import PdfMerger
    except ImportError:
        PdfMerger = None

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    TKDND_AVAILABLE = True
except ImportError:
    TKDND_AVAILABLE = False

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.utils import ImageReader
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Windows OCR (встроен в Windows 10+)
try:
    from winsdk.windows.media.ocr import OcrEngine
    from winsdk.windows.storage.streams import InMemoryRandomAccessStream, DataWriter
    from winsdk.windows.graphics.imaging import BitmapDecoder, BitmapPixelFormat, BitmapAlphaMode
    from winsdk.windows.globalization import Language
    import asyncio
    WINDOWS_OCR_AVAILABLE = True
except ImportError:
    WINDOWS_OCR_AVAILABLE = False
except Exception:
    WINDOWS_OCR_AVAILABLE = False

# ─────────────────────────────────────────────────────────────────────────────
# ПРОВЕРКА OCR КОМПОНЕНТОВ
# ─────────────────────────────────────────────────────────────────────────────

def get_ocr_status():
    """Возвращает статус OCR компонентов"""
    return {
        'pymupdf': PYMUPDF_AVAILABLE,          # Для конвертации PDF в изображения
        'windows_ocr': WINDOWS_OCR_AVAILABLE,  # Windows OCR (встроен в Windows 10+)
        'reportlab': REPORTLAB_AVAILABLE,      # Для создания PDF с текстом
        'pillow': PIL_AVAILABLE                # Для работы с изображениями
    }

def is_ocr_available():
    """Проверяет доступность OCR"""
    status = get_ocr_status()
    return all([
        status['pymupdf'],
        status['windows_ocr'],
        status['reportlab'],
        status['pillow']
    ])

# ─────────────────────────────────────────────────────────────────────────────
# КОНФИГУРАЦИЯ
# ─────────────────────────────────────────────────────────────────────────────

CONFIG_FILE = "generation_doc_config.json"

COLORS = {
    "primary": "#1E3A8A",
    "primary_light": "#3B82F6",
    "primary_dark": "#1E293B",
    "primary_hover": "#2563EB",
    "accent": "#6366F1",
    "accent_light": "#818CF8",
    "success": "#10B981",
    "success_hover": "#059669",
    "success_light": "#D1FAE5",
    "warning": "#F59E0B",
    "warning_hover": "#D97706",
    "warning_light": "#FEF3C7",
    "danger": "#EF4444",
    "danger_hover": "#DC2626",
    "danger_light": "#FEE2E2",
    "bg_primary": "#FFFFFF",
    "bg_secondary": "#F8FAFC",
    "bg_tertiary": "#F1F5F9",
    "bg_hover": "#E2E8F0",
    "card_bg": "#FFFFFF",
    "card_border": "#E2E8F0",
    "card_shadow": "#64748B20",
    "text_primary": "#0F172A",
    "text_secondary": "#475569",
    "text_tertiary": "#94A3B8",
    "text_disabled": "#CBD5E1",
    "text_on_primary": "#FFFFFF",
    "border": "#E2E8F0",
    "border_focus": "#3B82F6",
    "border_hover": "#CBD5E1",
    "dark_bg_primary": "#0F172A",
    "dark_bg_secondary": "#1E293B",
    "dark_text_primary": "#F1F5F9",
    "dark_text_secondary": "#94A3B8",
}

FONTS = {
    "title": ("Segoe UI", 18, "bold"),
    "subtitle": ("Segoe UI", 14, "bold"),
    "heading": ("Segoe UI", 12, "bold"),
    "body": ("Segoe UI", 10),
    "body_medium": ("Segoe UI", 10, "normal"),
    "small": ("Segoe UI", 9),
    "tiny": ("Segoe UI", 8),
    "button": ("Segoe UI", 10, "bold"),
    "button_small": ("Segoe UI", 9),
    "mono": ("Consolas", 9),
}

SPACING = {
    "xs": 2,
    "sm": 4,
    "md": 8,
    "lg": 12,
    "xl": 16,
    "xxl": 24,
}

BORDER_RADIUS = {
    "sm": 4,
    "md": 8,
    "lg": 12,
    "xl": 16,
    "full": 999,
}

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

# ─────────────────────────────────────────────────────────────────────────────
# СТИЛИ TREEVIEW
# ─────────────────────────────────────────────────────────────────────────────

_TREEVIEW_STYLE_CONFIGURED = False

def setup_modern_treeview_style():
    """Глобальная настройка стилей для Treeview (вызывается один раз)"""
    global _TREEVIEW_STYLE_CONFIGURED
    
    if _TREEVIEW_STYLE_CONFIGURED:
        return
    
    style = ttk.Style()
    
    try:
        style.theme_use("clam")
    except:
        pass
    
    style.configure(
        "Modern.Treeview",
        background=COLORS["card_bg"],
        fieldbackground=COLORS["card_bg"],
        foreground=COLORS["text_primary"],
        rowheight=36,
        font=FONTS["body"],
        borderwidth=0,
        relief=tk.FLAT
    )
    
    style.configure(
        "Modern.Treeview.Heading",
        background=COLORS["primary"],
        foreground="white",
        font=FONTS["button"],
        borderwidth=0,
        relief=tk.FLAT,
        padding=(12, 8)
    )
    
    style.map(
        "Modern.Treeview.Heading",
        background=[("active", COLORS["primary_hover"])],
        foreground=[("active", "white")]
    )
    
    style.map(
        "Modern.Treeview",
        background=[
            ("selected", COLORS["primary_light"]),
            ("!selected", COLORS["card_bg"])
        ],
        foreground=[
            ("selected", "white"),
            ("!selected", COLORS["text_primary"])
        ]
    )
    
    style.configure(
        "Modern.Vertical.TScrollbar",
        background=COLORS["bg_tertiary"],
        troughcolor=COLORS["bg_secondary"],
        borderwidth=0,
        arrowsize=0
    )
    
    style.map(
        "Modern.Vertical.TScrollbar",
        background=[
            ("active", COLORS["primary"]),
            ("!active", COLORS["border"])
        ]
    )
    
    style.configure(
        "Modern.Horizontal.TScrollbar",
        background=COLORS["bg_tertiary"],
        troughcolor=COLORS["bg_secondary"],
        borderwidth=0,
        arrowsize=0
    )
    
    style.map(
        "Modern.Horizontal.TScrollbar",
        background=[
            ("active", COLORS["primary"]),
            ("!active", COLORS["border"])
        ]
    )
    
    _TREEVIEW_STYLE_CONFIGURED = True

def create_modern_treeview(parent, columns, show="headings", height=15, selectable=True, **kwargs):
    """
    Создание современной таблицы Treeview с готовой стилизацией.
    
    Args:
        parent: Родительский виджет
        columns: Список колонок
        show: Что показывать ("headings", "tree headings", "tree", "")
        height: Высота в строках
        selectable: Можно ли выделять строки
        **kwargs: Дополнительные параметры для Treeview
    
    Returns:
        tuple: (tree_frame, tree) - фрейм с таблицей и сама таблица
    """
    setup_modern_treeview_style()
    
    tree_frame = tk.Frame(
        parent,
        bg=COLORS["card_bg"],
        highlightthickness=1,
        highlightbackground=COLORS["border"],
        highlightcolor=COLORS["border_focus"]
    )
    
    v_scrollbar = ttk.Scrollbar(
        tree_frame,
        orient=tk.VERTICAL,
        style="Modern.Vertical.TScrollbar"
    )
    v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    h_scrollbar = ttk.Scrollbar(
        tree_frame,
        orient=tk.HORIZONTAL,
        style="Modern.Horizontal.TScrollbar"
    )
    h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
    
    tree = ttk.Treeview(
        tree_frame,
        columns=columns,
        show=show,
        height=height,
        style="Modern.Treeview",
        yscrollcommand=v_scrollbar.set,
        xscrollcommand=h_scrollbar.set,
        **kwargs
    )
    tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    v_scrollbar.config(command=tree.yview)
    h_scrollbar.config(command=tree.xview)
    
    tree.tag_configure("oddrow", background=COLORS["bg_secondary"])
    tree.tag_configure("evenrow", background=COLORS["card_bg"])
    tree.tag_configure("selected", background=COLORS["primary_light"], foreground="white")
    
    return tree_frame, tree

def configure_treeview_columns(tree, columns_config):
    """
    Настройка колонок Treeview.
    
    Args:
        tree: Виджет Treeview
        columns_config: Словарь {column_name: {"text": "Заголовок", "width": 100, "anchor": tk.W}}
    """
    for col, config in columns_config.items():
        tree.heading(col, text=config.get("text", col))
        tree.column(
            col,
            width=config.get("width", 100),
            minwidth=config.get("minwidth", 50),
            anchor=config.get("anchor", tk.W),
            stretch=config.get("stretch", False)
        )

def insert_treeview_row(tree, values, tags=None, auto_tag=True):
    """
    Вставка строки в Treeview с автоматическим чередованием цветов.
    
    Args:
        tree: Виджет Treeview
        values: Значения строки
        tags: Дополнительные теги
        auto_tag: Автоматически добавлять теги четности
    
    Returns:
        str: ID вставленной строки
    """
    if auto_tag:
        row_count = len(tree.get_children())
        parity_tag = "oddrow" if row_count % 2 == 0 else "evenrow"
        
        if tags:
            if isinstance(tags, str):
                tags = (tags, parity_tag)
            else:
                tags = tuple(tags) + (parity_tag,)
        else:
            tags = (parity_tag,)
    
    return tree.insert("", tk.END, values=values, tags=tags)

def format_cell_value(value, date_format="%d.%m.%Y"):
    """
    Форматирование значения ячейки для отображения в таблице.
    Даты преобразуются в формат "дд.мм.гггг", NaN/None становятся пустой строкой.
    
    Args:
        value: Значение ячейки (любой тип)
        date_format: Формат даты (по умолчанию "%d.%m.%Y")
    
    Returns:
        str: Отформатированная строка для отображения
    """
    import pandas as pd
    from datetime import datetime, date
    
    if pd.isna(value) or value is None:
        return ""
    
    if hasattr(value, 'strftime'):
        try:
            return value.strftime(date_format)
        except:
            return str(value)
    
    if isinstance(value, date):
        return value.strftime(date_format)
    
    try:
        import numpy as np
        if isinstance(value, np.datetime64):
            return pd.Timestamp(value).strftime(date_format)
    except:
        pass
    
    value_str = str(value).strip()
    
    if not value_str:
        return ""
    
    if '-' in value_str and len(value_str) >= 10:
        iso_patterns = [
            '%Y-%m-%d %H:%M:%S',
            '%Y-%m-%d %H:%M:%S.%f',
            '%Y-%m-%d %H:%M',
            '%Y-%m-%d',
        ]
        for pattern in iso_patterns:
            try:
                dt = datetime.strptime(value_str, pattern)
                return dt.strftime(date_format)
            except ValueError:
                continue
    
    if '.' in value_str and len(value_str) == 10:
        try:
            datetime.strptime(value_str, '%d.%m.%Y')
            return value_str
        except ValueError:
            pass
    
    return value_str

def format_dataframe_row(row, date_format="%d.%m.%Y"):
    """Форматирование всей строки DataFrame для отображения в таблице."""
    return [format_cell_value(val, date_format) for val in row]

# ─────────────────────────────────────────────────────────────────────────────
# DRAG & DROP
# ─────────────────────────────────────────────────────────────────────────────

def parse_drop_files(data):
    """
    Парсинг данных о перетащенных файлах из DND события.
    
    Args:
        data: Строка с данными о файлах из tkinterdnd2
    
    Returns:
        list: Список путей к файлам
    """
    files = []
    
    if data.startswith('{'):
        import re
        files = re.findall(r'\{([^}]+)\}', data)
    else:
        files = data.split()
    
    files = [f.strip('"').strip("'") for f in files]
    return files

def setup_file_drop(entry_widget, string_var, file_types=None, on_drop_callback=None, parent_window=None):
    """
    Настройка drag & drop для поля ввода файла с визуальной индикацией.
    
    Args:
        entry_widget: Виджет Entry/CTkEntry для отображения пути
        string_var: StringVar привязанный к полю
        file_types: Список допустимых расширений (например ['.xlsx', '.xls'])
        on_drop_callback: Функция обратного вызова при успешном drop(file_path)
        parent_window: Родительское окно для messagebox
    """
    if not TKDND_AVAILABLE:
        return
    
    try:
        # Определяем виджет для регистрации drop
        actual_widget = entry_widget
        
        # Для CTkEntry пробуем разные варианты
        if hasattr(entry_widget, '_entry'):
            # Пробуем получить внутренний Entry
            actual_widget = entry_widget._entry
        elif hasattr(entry_widget, 'winfo_children'):
            # Иногда нужен первый дочерний виджет
            children = entry_widget.winfo_children()
            if children:
                actual_widget = children[0]
        
        # Сохраняем исходные цвета для восстановления
        original_bg = None
        original_state = None
        
        # Сохраняем текущее состояние
        if hasattr(entry_widget, 'cget'):
            try:
                original_bg = entry_widget.cget('fg_color')
                original_state = entry_widget.cget('state')
            except:
                original_bg = COLORS["bg_tertiary"]
                original_state = 'normal'
        
        # Регистрируем виджет для приёма файлов
        actual_widget.drop_target_register(DND_FILES)
        
        def on_drag_enter(event):
            """Визуальная индикация при наведении файла"""
            try:
                if hasattr(entry_widget, 'configure'):
                    entry_widget.configure(fg_color=COLORS["primary_hover"])
            except:
                pass
        
        def on_drag_leave(event):
            """Восстановление цвета при уходе курсора"""
            try:
                if hasattr(entry_widget, 'configure'):
                    entry_widget.configure(fg_color=original_bg)
            except:
                pass
        
        def on_drop(event):
            """Обработка drop файла"""
            # Восстанавливаем цвет
            try:
                if hasattr(entry_widget, 'configure'):
                    entry_widget.configure(fg_color=original_bg)
            except:
                pass
            
            files = parse_drop_files(event.data)
            
            if not files:
                return
            
            file_path = files[0]
            
            # Проверка типа файла
            if file_types:
                valid = any(file_path.lower().endswith(ext) for ext in file_types)
                if not valid:
                    types_str = ", ".join(file_types)
                    parent = parent_window if parent_window else entry_widget.winfo_toplevel()
                    messagebox.showwarning(
                        "Неверный тип файла",
                        f"Принимаются только файлы: {types_str}\n\nПолучен: {os.path.basename(file_path)}",
                        parent=parent
                    )
                    return
            
            # Временно делаем поле редактируемым для установки значения
            was_readonly = False
            if original_state == 'readonly':
                try:
                    entry_widget.configure(state='normal')
                    was_readonly = True
                except:
                    pass
            
            # Устанавливаем путь
            string_var.set(file_path)
            
            # Возвращаем readonly если было
            if was_readonly:
                try:
                    entry_widget.configure(state='readonly')
                except:
                    pass
            
            # Вызываем callback
            if on_drop_callback:
                try:
                    on_drop_callback(file_path)
                except Exception as e:
                    print(f"Ошибка в callback: {e}")
        
        # Привязываем события drag-and-drop
        actual_widget.dnd_bind('<<DragEnter>>', on_drag_enter)
        actual_widget.dnd_bind('<<DragLeave>>', on_drag_leave)
        actual_widget.dnd_bind('<<Drop>>', on_drop)
        
    except Exception as e:
        pass

def setup_folder_drop(entry_widget, string_var, on_drop_callback=None, parent_window=None):
    """
    Настройка drag & drop для поля ввода папки с визуальной индикацией.
    
    Args:
        entry_widget: Виджет Entry/CTkEntry для отображения пути
        string_var: StringVar привязанный к полю
        on_drop_callback: Функция обратного вызова при успешном drop(folder_path)
        parent_window: Родительское окно для messagebox (опционально)
    """
    if not TKDND_AVAILABLE:
        return
    
    try:
        actual_widget = entry_widget
        if hasattr(entry_widget, '_entry'):
            actual_widget = entry_widget._entry
        
        # Сохраняем исходные цвета
        original_bg = None
        if hasattr(entry_widget, 'cget'):
            try:
                original_bg = entry_widget.cget('fg_color')
            except:
                original_bg = COLORS["bg_tertiary"]
        
        actual_widget.drop_target_register(DND_FILES)
        
        def on_drag_enter(event):
            """Визуальная индикация при наведении"""
            try:
                if hasattr(entry_widget, 'configure'):
                    entry_widget.configure(fg_color=COLORS["success_hover"])
            except:
                pass
        
        def on_drag_leave(event):
            """Восстановление цвета"""
            try:
                if hasattr(entry_widget, 'configure'):
                    entry_widget.configure(fg_color=original_bg)
            except:
                pass
        
        def on_drop(event):
            """Обработка drop папки/файла"""
            # Восстанавливаем цвет
            try:
                if hasattr(entry_widget, 'configure'):
                    entry_widget.configure(fg_color=original_bg)
            except:
                pass
            
            files = parse_drop_files(event.data)
            
            if not files:
                return
            
            # Берём первый путь
            path = files[0]
            
            # Если это файл, берём его папку
            if os.path.isfile(path):
                path = os.path.dirname(path)
            
            string_var.set(path)
            
            # Вызываем callback если есть
            if on_drop_callback:
                on_drop_callback(path)
        
        # Привязываем события
        actual_widget.dnd_bind('<<DragEnter>>', on_drag_enter)
        actual_widget.dnd_bind('<<DragLeave>>', on_drag_leave)
        actual_widget.dnd_bind('<<Drop>>', on_drop)
        
    except Exception as e:
        pass

# ── ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ СОВРЕМЕННОГО UI ────────────────────────────────────────────────────────

def create_modern_button(parent, text="", command=None, style="primary", width=None, height=36, icon=None, tooltip=""):
    """Создание современной кнопки CustomTkinter с встроенными эффектами"""
    style_config = {
        "primary": {
            "fg_color": COLORS["primary"],
            "hover_color": COLORS["primary_hover"],
            "text_color": "white",
        },
        "success": {
            "fg_color": COLORS["success"],
            "hover_color": COLORS["success_hover"],
            "text_color": "white",
        },
        "warning": {
            "fg_color": COLORS["warning"],
            "hover_color": COLORS["warning_hover"],
            "text_color": "white",
        },
        "danger": {
            "fg_color": COLORS["danger"],
            "hover_color": COLORS["danger_hover"],
            "text_color": "white",
        },
        "secondary": {
            "fg_color": COLORS["bg_tertiary"],
            "hover_color": COLORS["bg_hover"],
            "text_color": COLORS["text_primary"],
        },
    }
    
    config = style_config.get(style, style_config["primary"])
    button_text = f"{icon} {text}" if icon else text
    
    btn = CTkButton(
        parent,
        text=button_text,
        command=command,
        fg_color=config["fg_color"],
        hover_color=config["hover_color"],
        text_color=config["text_color"],
        font=FONTS["button"],
        corner_radius=BORDER_RADIUS["md"],
        border_width=0,
        height=height,
        width=width if width else 120,
        cursor="hand2"
    )
    
    if tooltip:
        ToolTip(btn, tooltip)
    
    return btn

def create_icon_button(parent, icon="", command=None, tooltip="", style="primary", width=32, height=28):
    """Создание иконочной кнопки CustomTkinter для toolbar"""
    style_config = {
        "primary": {
            "fg_color": COLORS["primary_dark"],
            "hover_color": COLORS["primary_hover"],
        },
        "accent": {
            "fg_color": COLORS["accent"],
            "hover_color": COLORS["accent_light"],
        },
        "success": {
            "fg_color": COLORS["success"],
            "hover_color": COLORS["success_hover"],
        },
    }
    
    config = style_config.get(style, style_config["primary"])
    
    btn = CTkButton(
        parent,
        text=icon,
        command=command,
        fg_color=config["fg_color"],
        hover_color=config["hover_color"],
        text_color="white",
        font=FONTS["body"],
        width=width,
        height=height,
        corner_radius=BORDER_RADIUS["sm"],
        border_width=0,
        cursor="hand2"
    )
    
    if tooltip:
        ToolTip(btn, tooltip)
    
    return btn

def create_card_frame(parent, **kwargs):
    """Создание современной карточки с тенью"""
    shadow_frame = tk.Frame(
        parent,
        bg=COLORS["bg_secondary"],
        **kwargs
    )
    
    card = tk.Frame(
        shadow_frame,
        bg=COLORS["card_bg"],
        highlightthickness=1,
        highlightbackground=COLORS["card_border"],
        highlightcolor=COLORS["border_focus"]
    )
    card.pack(padx=1, pady=1, fill=tk.BOTH, expand=True)
    
    return shadow_frame, card

def create_section(parent, title, icon=""):
    """Создание современной секции с заголовком и разделителем"""
    # Контейнер секции
    section_frame = tk.Frame(parent, bg=COLORS["bg_secondary"])
    
    header_frame = tk.Frame(section_frame, bg=COLORS["bg_secondary"], height=40)
    header_frame.pack(fill=tk.X, pady=(0, SPACING["sm"]))
    header_frame.pack_propagate(False)
    
    title_text = f"{icon}  {title}" if icon else title
    title_label = tk.Label(
        header_frame,
        text=title_text,
        font=FONTS["heading"],
        bg=COLORS["bg_secondary"],
        fg=COLORS["primary"],
        anchor="w"
    )
    title_label.pack(side=tk.LEFT, fill=tk.Y, padx=(SPACING["md"], 0))
    
    # Разделительная линия
    separator = tk.Frame(
        header_frame,
        height=2,
        bg=COLORS["primary"],
        relief=tk.FLAT
    )
    separator.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(SPACING["md"], 0), pady=15)
    
    # Контент секции (карточка)
    shadow_frame, content_frame = create_card_frame(section_frame)
    shadow_frame.pack(fill=tk.BOTH, expand=True)
    
    # Внутренний padding для контента
    content = tk.Frame(content_frame, bg=COLORS["card_bg"], padx=SPACING["lg"], pady=SPACING["lg"])
    content.pack(fill=tk.BOTH, expand=True)
    
    return section_frame, content

# ── ФУНКЦИЯ ДЛЯ ПОДДЕРЖКИ БУФЕРА ОБМЕНА ────────────────────────────
def enable_clipboard_shortcuts(widget):
    """Включает комбинации Ctrl+C (копировать) и Ctrl+A (выделить всё)
    Используется только для полей только для чтения (логи)"""
    def copy(event=None):
        if isinstance(widget, (tk.Text, ScrolledText)):
            was_disabled = str(widget.cget('state')) == 'disabled'
            
            if not was_disabled:
                return None
            
            widget.config(state=tk.NORMAL)
            
            try:
                selected_text = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
                if selected_text:
                    root = widget.winfo_toplevel()
                    root.clipboard_clear()
                    root.clipboard_append(selected_text)
                    root.update()
            except tk.TclError:
                pass
            finally:
                widget.config(state=tk.DISABLED)
            
            return "break"
        else:
            return None
    
    def select_all(event=None):
        try:
            if isinstance(widget, (tk.Text, ScrolledText)):
                was_disabled = str(widget.cget('state')) == 'disabled'
                
                if not was_disabled:
                    return None
                
                widget.config(state=tk.NORMAL)
                
                widget.tag_add(tk.SEL, "1.0", tk.END)
                widget.mark_set(tk.INSERT, "1.0")
                widget.see(tk.INSERT)
                
                widget.config(state=tk.DISABLED)
                return "break"
            else:
                return None
        except:
            pass
        return None
    
    widget.bind("<Control-c>", copy)
    widget.bind("<Control-C>", copy)
    widget.bind("<Control-a>", select_all)
    widget.bind("<Control-A>", select_all)

def enable_field_shortcuts(widget, readonly=False):
    """Включает полный набор горячих клавиш для полей ввода
    
    Примечание: С версии 2026 используются глобальные горячие клавиши.
    Эта функция оставлена для обратной совместимости, но не выполняет действий.
    """
    pass

def add_context_menu(widget, readonly=False):
    """Добавляет контекстное меню с операциями буфера обмена
    
    Args:
        widget: tk.Entry, CTkEntry, tk.Text или ScrolledText
        readonly: если True, только копирование и выделение
    """
    
    def get_actual_widget(w):
        if isinstance(w, ctk.CTkEntry) and hasattr(w, '_entry'):
            return w._entry
        return w
    
    def show_context_menu(event):
        menu = ModernContextMenu(widget)
        
        has_selection = False
        try:
            w = get_actual_widget(widget)
            if isinstance(w, (tk.Text, ScrolledText)):
                w.get(tk.SEL_FIRST, tk.SEL_LAST)
                has_selection = True
            elif isinstance(w, tk.Entry):
                if w.selection_present():
                    has_selection = True
        except:
            pass
        
        if not readonly:
            menu.add_command(
                label="✂ Вырезать",
                command=lambda: cut_text(widget)
            )
        
        menu.add_command(
            label="📋 Копировать",
            command=lambda: copy_text(widget)
        )
        
        if not readonly:
            menu.add_command(
                label="📄 Вставить",
                command=lambda: paste_text(widget)
            )
        
        menu.add_separator()
        
        menu.add_command(
            label="🔘 Выделить всё",
            command=lambda: select_all_text(widget)
        )
        
        menu.post(event.x_root, event.y_root)
    
    def cut_text(w):
        """Вырезать выделенный текст"""
        try:
            actual = get_actual_widget(w)
            if isinstance(actual, (tk.Text, ScrolledText)):
                if actual.tag_ranges(tk.SEL):
                    text = actual.get(tk.SEL_FIRST, tk.SEL_LAST)
                    actual.clipboard_clear()
                    actual.clipboard_append(text)
                    actual.delete(tk.SEL_FIRST, tk.SEL_LAST)
            elif isinstance(actual, tk.Entry):
                if actual.selection_present():
                    text = actual.selection_get()
                    actual.clipboard_clear()
                    actual.clipboard_append(text)
                    actual.delete(tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass
    
    def copy_text(w):
        """Копировать выделенный текст"""
        try:
            actual = get_actual_widget(w)
            if isinstance(actual, (tk.Text, ScrolledText)):
                if actual.tag_ranges(tk.SEL):
                    text = actual.get(tk.SEL_FIRST, tk.SEL_LAST)
                    actual.clipboard_clear()
                    actual.clipboard_append(text)
            elif isinstance(actual, tk.Entry):
                if actual.selection_present():
                    text = actual.selection_get()
                    actual.clipboard_clear()
                    actual.clipboard_append(text)
        except:
            pass
    
    def paste_text(w):
        """Вставить текст из буфера обмена"""
        try:
            actual = get_actual_widget(w)
            clipboard_text = actual.clipboard_get()
            if isinstance(actual, (tk.Text, ScrolledText)):
                if actual.tag_ranges(tk.SEL):
                    actual.delete(tk.SEL_FIRST, tk.SEL_LAST)
                actual.insert(tk.INSERT, clipboard_text)
            elif isinstance(actual, tk.Entry):
                if actual.selection_present():
                    actual.delete(tk.SEL_FIRST, tk.SEL_LAST)
                actual.insert(actual.index(tk.INSERT), clipboard_text)
        except:
            pass
    
    def select_all_text(w):
        """Выделить весь текст в виджете"""
        try:
            actual = get_actual_widget(w)
            if isinstance(actual, (tk.Text, ScrolledText)):
                actual.tag_remove(tk.SEL, "1.0", tk.END)
                actual.tag_add(tk.SEL, "1.0", tk.END)
                actual.mark_set(tk.INSERT, "1.0")
                actual.see(tk.INSERT)
            elif isinstance(actual, tk.Entry):
                actual.select_range(0, tk.END)
                actual.icursor(tk.END)
        except:
            pass
    
    widget.bind("<Button-3>", show_context_menu)

def enable_text_selection_in_disabled(widget):
    """Включает возможность выделения текста мышью даже в состоянии DISABLED"""
    if not isinstance(widget, tk.Text):
        return
    
    def on_button_press(event):
        # Временно включаем виджет для выделения
        was_disabled = str(widget.cget('state')) == 'disabled'
        if was_disabled:
            widget.config(state=tk.NORMAL)
        
        widget.mark_set("sel_start", f"@{event.x},{event.y}")
        widget.tag_remove(tk.SEL, "1.0", tk.END)
        
        if was_disabled:
            widget.config(state=tk.DISABLED)
        
        return "break"
    
    def on_button_motion(event):
        # Выделяем текст при движении мыши с зажатой кнопкой
        was_disabled = str(widget.cget('state')) == 'disabled'
        if was_disabled:
            widget.config(state=tk.NORMAL)
        
        try:
            widget.tag_remove(tk.SEL, "1.0", tk.END)
            widget.tag_add(tk.SEL, "sel_start", f"@{event.x},{event.y}")
        except:
            pass
        
        if was_disabled:
            widget.config(state=tk.DISABLED)
        
        return "break"
    
    widget.bind("<Button-1>", on_button_press)
    widget.bind("<B1-Motion>", on_button_motion)

# ── СОВРЕМЕННЫЙ КЛАСС ДЛЯ ТУЛТИПОВ ─────────────────────────────────────────────────────────────────
class ToolTip:
    """Современный класс для создания всплывающих подсказок"""
    def __init__(self, widget, text, delay=500):
        self.widget = widget
        self.text = text
        self.delay = delay
        self.tooltip_window = None
        self.show_timer = None
        self.widget.bind("<Enter>", self.schedule_show)
        self.widget.bind("<Leave>", self.hide_tooltip)
        self.widget.bind("<Button>", self.hide_tooltip)
    
    def schedule_show(self, event=None):
        """ Отложенный показ tooltip"""
        if self.show_timer:
            self.widget.after_cancel(self.show_timer)
        self.show_timer = self.widget.after(self.delay, self.show_tooltip)
    
    def show_tooltip(self, event=None):
        if self.tooltip_window or not self.text:
            return
        
        # Позиционирование
        x = self.widget.winfo_rootx() + self.widget.winfo_width() // 2
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 8
        
        self.tooltip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_withdraw()
        
        # Внешний фрейм для эффекта тени
        shadow_frame = tk.Frame(
            tw,
            bg=COLORS["text_primary"],
            highlightthickness=0
        )
        shadow_frame.pack(padx=2, pady=2)
        
        # Внутренний фрейм
        content_frame = tk.Frame(
            shadow_frame,
            bg=COLORS["text_primary"],
            highlightthickness=0
        )
        content_frame.pack()
        
        label = tk.Label(
            content_frame,
            text=self.text,
            justify=tk.LEFT,
            background=COLORS["text_primary"],
            foreground="white",
            font=FONTS["small"],
            padx=12,
            pady=6,
            relief=tk.FLAT
        )
        label.pack()
        
        tw.update_idletasks()
        tw_width = tw.winfo_width()
        x = x - tw_width // 2
        tw.wm_geometry(f"+{x}+{y}")
        
        tw.wm_deiconify()
        tw.attributes('-alpha', 0.0)
        self._fade_in(tw, 0.0)
    
    def _fade_in(self, window, alpha):
        """ Плавное появление """
        if alpha < 0.95:
            alpha += 0.15
            try:
                window.attributes('-alpha', alpha)
                window.after(20, lambda: self._fade_in(window, alpha))
            except:
                pass
        else:
            try:
                window.attributes('-alpha', 0.95)
            except:
                pass
    
    def hide_tooltip(self, event=None):
        if self.show_timer:
            self.widget.after_cancel(self.show_timer)
            self.show_timer = None
        if self.tooltip_window:
            try:
                self.tooltip_window.destroy()
            except:
                pass
            self.tooltip_window = None

# ── ФУНКЦИЯ ДЛЯ УСТАНОВКИ КУРСОРА НА СТРЕЛОЧКУ COMBOBOX ────────────────────────────────────────────
def set_combobox_cursor(combobox, cursor="hand2"):
    """Устанавливает курсор только для кнопки-стрелочки CTkComboBox"""
    try:
        # В CTkComboBox стрелочка находится в области справа на основном canvas
        # Нам нужно отслеживать позицию мыши на canvas и менять курсор только для правой части
        def find_main_canvas(widget):
            for child in widget.winfo_children():
                if child.__class__.__name__ == 'CTkCanvas':
                    return child
            return None
        
        main_canvas = find_main_canvas(combobox)
        
        if main_canvas:
            # Определяем ширину кнопки (обычно ~28-30 пикселей справа)
            button_width = 30
            
            def on_motion(event):
                canvas_width = main_canvas.winfo_width()
                # Если курсор в правой части (где стрелочка)
                if event.x > canvas_width - button_width:
                    main_canvas.configure(cursor=cursor)
                else:
                    main_canvas.configure(cursor="")
            
            def on_leave(event):
                main_canvas.configure(cursor="")
            
            main_canvas.bind("<Motion>", on_motion)
            main_canvas.bind("<Leave>", on_leave)
                
    except Exception as e:
        pass

# ── СОВРЕМЕННЫЙ КЛАСС ДЛЯ КОНТЕКСТНЫХ МЕНЮ ─────────────────────────────────────────────────────────
class ModernContextMenu:
    """Современный класс для создания контекстных меню"""
    def __init__(self, parent):
        self.parent = parent
        self.menu_window = None
        self.items = []
    
    def add_command(self, label, command):
        """Добавить команду в меню"""
        self.items.append({"type": "command", "label": label, "command": command})
    
    def add_separator(self):
        """Добавить разделитель"""
        self.items.append({"type": "separator"})
    
    def post(self, x, y):
        """Показать меню в указанных координатах"""
        if self.menu_window:
            self.hide()
        
        self.menu_window = tw = tk.Toplevel(self.parent)
        tw.wm_overrideredirect(True)
        tw.wm_withdraw()
        tw.attributes('-topmost', True)
        
        # Внешний фрейм для тени
        shadow_frame = tk.Frame(
            tw,
            bg="#94A3B8",
            highlightthickness=0
        )
        shadow_frame.pack(padx=2, pady=2)
        
        # Внутренний фрейм с меню
        menu_frame = tk.Frame(
            shadow_frame,
            bg=COLORS["card_bg"],
            highlightbackground=COLORS["border"],
            highlightthickness=1,
            relief=tk.FLAT
        )
        menu_frame.pack()
        
        for item in self.items:
            if item["type"] == "separator":
                separator = tk.Frame(
                    menu_frame,
                    bg=COLORS["border"],
                    height=1
                )
                separator.pack(fill=tk.X, padx=4, pady=4)
            else:
                btn = tk.Label(
                    menu_frame,
                    text=item["label"],
                    font=FONTS["body"],
                    bg=COLORS["card_bg"],
                    fg=COLORS["text_primary"],
                    anchor=tk.W,
                    padx=16,
                    pady=8,
                    cursor="hand2"
                )
                btn.pack(fill=tk.X)
                
                # Привязываем команду
                def on_click(event, cmd=item["command"]):
                    self.hide()
                    if cmd:
                        cmd()
                
                btn.bind("<Button-1>", on_click)
                
                # Hover эффект
                def on_enter(event, b=btn):
                    b.configure(bg=COLORS["primary_light"], fg="white")
                
                def on_leave(event, b=btn):
                    b.configure(bg=COLORS["card_bg"], fg=COLORS["text_primary"])
                
                btn.bind("<Enter>", on_enter)
                btn.bind("<Leave>", on_leave)
        
        # Позиционирование
        tw.update_idletasks()
        tw_width = tw.winfo_width()
        tw_height = tw.winfo_height()
        
        screen_width = tw.winfo_screenwidth()
        screen_height = tw.winfo_screenheight()
        
        if x + tw_width > screen_width:
            x = screen_width - tw_width - 10
        if y + tw_height > screen_height:
            y = screen_height - tw_height - 10
        
        tw.wm_geometry(f"+{x}+{y}")
        
        tw.wm_deiconify()
        tw.attributes('-alpha', 0.0)
        self._fade_in(tw, 0.0)
        
        # Закрытие при клике вне меню
        def close_on_click_outside(event):
            if self.menu_window and event.widget not in [self.menu_window] + list(self.menu_window.winfo_children()):
                self.hide()
        
        # Задержка биндинга для предотвращения немедленного закрытия
        tw.after(100, lambda: tw.bind_all("<Button-1>", close_on_click_outside, add="+"))
        tw.bind("<FocusOut>", lambda e: self.hide())
    
    def _fade_in(self, window, alpha):
        """Плавное появление"""
        if alpha < 0.95:
            alpha += 0.15
            try:
                window.attributes('-alpha', alpha)
                window.after(20, lambda: self._fade_in(window, alpha))
            except:
                pass
        else:
            try:
                window.attributes('-alpha', 0.95)
            except:
                pass
    
    def hide(self):
        """Скрыть меню"""
        if self.menu_window:
            try:
                self.menu_window.unbind_all("<Button-1>")
                self.menu_window.destroy()
            except:
                pass
            self.menu_window = None
        self.items = []

# ── КЛАСС ДЛЯ ВСПЛЫВАЮЩЕЙ ПОДСКАЗКИ СО СТАТУСОМ ЗАДАЧ ────────────────
class TabStatusTooltip:
    """Всплывающая подсказка с информацией о статусе выполнения задач"""
    def __init__(self, widget, app):
        self.widget = widget
        self.app = app
        self.tooltip_window = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)
    
    def show_tooltip(self, event=None):
        if self.tooltip_window:
            return
        
        total_tabs = len(self.app.tabs)
        processing_tabs = [tab for tab in self.app.tabs if tab.is_processing]
        idle_tabs = [tab for tab in self.app.tabs if not tab.is_processing]
        
        lines = []
        lines.append(f"📊 Всего вкладок: {total_tabs}/{self.app.max_tabs}")
        lines.append("")
        
        if processing_tabs:
            lines.append("⏳ Выполняются задачи:")
            for tab in processing_tabs:
                lines.append(f"  • {tab.tab_name}")
        else:
            lines.append("✓ Нет активных задач")
        
        if idle_tabs:
            lines.append("")
            lines.append("💤 Ожидают запуска:")
            for tab in idle_tabs:
                lines.append(f"  • {tab.tab_name}")
        
        tooltip_text = "\n".join(lines)
        
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        
        self.tooltip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_withdraw()
        
        label = tk.Label(
            tw,
            text=tooltip_text,
            justify=tk.LEFT,
            background="#2C3E50",
            foreground="white",
            relief=tk.SOLID,
            borderwidth=1,
            font=FONTS["small"],
            padx=12,
            pady=8
        )
        label.pack()
        
        tw.wm_geometry(f"+{x}+{y}")
        tw.wm_deiconify()
    
    def hide_tooltip(self, event=None):
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

# ══════════════════════════════════════════════════════════════════════
# СИСТЕМА КЭШИРОВАНИЯ ДОКУМЕНТОВ
# ══════════════════════════════════════════════════════════════════════

import hashlib
import pickle

class DocumentCache:
    """Система кэширования для избежания повторного создания одинаковых документов"""
    def __init__(self, cache_dir=None):
        if cache_dir is None:
            cache_dir = os.path.join(os.path.dirname(__file__), ".cache")
        
        self.cache_dir = cache_dir
        self.cache_file = os.path.join(cache_dir, "document_cache.pkl")
        self.cache_data = {}
        
        # Создаем директорию если её нет
        os.makedirs(cache_dir, exist_ok=True)
        
        # Загружаем кэш
        self.load_cache()
    
    def load_cache(self):
        """Загрузка кэша из файла"""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'rb') as f:
                    self.cache_data = pickle.load(f)
            except:
                self.cache_data = {}
    
    def save_cache(self):
        """Сохранение кэша в файл"""
        try:
            with open(self.cache_file, 'wb') as f:
                pickle.dump(self.cache_data, f)
        except:
            pass
    
    def get_file_hash(self, file_path):
        """Вычисление хэша файла"""
        if not os.path.exists(file_path):
            return None
        
        hasher = hashlib.md5()
        try:
            with open(file_path, 'rb') as f:
                # Читаем файл порциями для экономии памяти
                for chunk in iter(lambda: f.read(8192), b''):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except:
            return None
    
    def get_data_hash(self, data_dict):
        """Вычисление хэша данных (словаря)"""
        # Сериализуем словарь в строку и вычисляем хэш
        data_str = json.dumps(data_dict, sort_keys=True, ensure_ascii=False)
        return hashlib.md5(data_str.encode('utf-8')).hexdigest()
    
    def get_document_key(self, template_path, data, output_name):
        """
        Создание ключа для документа
        
        Args:
            template_path: путь к шаблону
            data: данные для заполнения (словарь)
            output_name: имя выходного файла
        
        Returns:
            str: уникальный ключ документа
        """
        template_hash = self.get_file_hash(template_path)
        data_hash = self.get_data_hash(data)
        
        if template_hash is None:
            return None
        
        # Комбинируем хэши
        combined = f"{template_hash}_{data_hash}_{output_name}"
        return hashlib.md5(combined.encode('utf-8')).hexdigest()
    
    def should_create_document(self, template_path, data, output_path):
        """
        Проверка нужно ли создавать документ
        
        Returns:
            bool: True если нужно создать, False если уже существует и актуален
        """
        # Если выходной файл не существует - нужно создать
        if not os.path.exists(output_path):
            return True
        
        output_name = os.path.basename(output_path)
        doc_key = self.get_document_key(template_path, data, output_name)
        
        if doc_key is None:
            return True
        
        # Проверяем есть ли в кэше
        if doc_key in self.cache_data:
            cache_entry = self.cache_data[doc_key]
            cached_output = cache_entry.get('output_path')
            
            # Если выходной файл совпадает и существует
            if cached_output == output_path and os.path.exists(output_path):
                # Проверяем что файл не был изменен
                current_hash = self.get_file_hash(output_path)
                cached_hash = cache_entry.get('output_hash')
                
                if current_hash == cached_hash:
                    return False  # Документ актуален, не нужно создавать
        
        return True
    
    def register_document(self, template_path, data, output_path):
        """
        Регистрация созданного документа в кэше
        
        Args:
            template_path: путь к шаблону
            data: данные для заполнения
            output_path: путь к созданному документу
        """
        output_name = os.path.basename(output_path)
        doc_key = self.get_document_key(template_path, data, output_name)
        
        if doc_key is None:
            return
        
        output_hash = self.get_file_hash(output_path)
        
        self.cache_data[doc_key] = {
            'template_path': template_path,
            'output_path': output_path,
            'output_hash': output_hash,
            'created_at': time.time()
        }
        
        # Сохраняем кэш
        self.save_cache()
    
    def clear_old_entries(self, max_age_days=30):
        """Удаление старых записей из кэша"""
        current_time = time.time()
        max_age_seconds = max_age_days * 24 * 3600
        
        keys_to_remove = []
        for key, entry in self.cache_data.items():
            age = current_time - entry.get('created_at', 0)
            if age > max_age_seconds:
                keys_to_remove.append(key)
        
        for key in keys_to_remove:
            del self.cache_data[key]
        
        if keys_to_remove:
            self.save_cache()
        
        return len(keys_to_remove)
    
    def clear_cache(self):
        """Полная очистка кэша"""
        self.cache_data = {}
        self.save_cache()

# Глобальный экземпляр кэша
document_cache = DocumentCache()

# ── КЛАСС ДЛЯ ФОНОВОЙ ПРЕДЗАГРУЗКИ WORD ДОКУМЕНТОВ ───────────────────
class WordPreloadManager:
    """Менеджер для фоновой конвертации Word документов в PDF"""
    def __init__(self):
        self.cache = {}  # {file_path: {'temp_pdf_path': str, 'status': str, 'error': str, 'timestamp': float}}
        self.queue = queue.Queue()
        self.worker_thread = None
        self.running = False
        self.max_cache_age = 3600  # Максимальный возраст кэша в секундах (1 час)
        self.max_cache_size = 10  # Максимальное количество закэшированных файлов
        
    def start(self):
        """Запускает фоновый поток обработки"""
        if not self.running:
            self.running = True
            self.worker_thread = threading.Thread(target=self._worker, daemon=True)
            self.worker_thread.start()
    
    def stop(self):
        """Останавливает фоновый поток"""
        self.running = False
        self.clear_cache()
    
    def _worker(self):
        """Фоновый поток для конвертации Word документов"""
        while self.running:
            try:
                file_path = self.queue.get(timeout=0.5)
                
                if file_path in self.cache:
                    status = self.cache[file_path]['status']
                    if status in ('processing', 'ready'):
                        self.queue.task_done()
                        continue
                
                self.cache[file_path] = {
                    'temp_pdf_path': None,
                    'status': 'processing',
                    'error': None,
                    'timestamp': datetime.now().timestamp()
                }
                
                temp_pdf_path = self._convert_word_to_pdf(file_path)
                
                if temp_pdf_path:
                    self.cache[file_path]['temp_pdf_path'] = temp_pdf_path
                    self.cache[file_path]['status'] = 'ready'
                else:
                    self.cache[file_path]['status'] = 'error'
                    self.cache[file_path]['error'] = 'Не удалось конвертировать файл'
                
                self.queue.task_done()
                
                self._cleanup_old_cache()
                
            except queue.Empty:
                continue
            except Exception as e:
                if file_path in self.cache:
                    self.cache[file_path]['status'] = 'error'
                    self.cache[file_path]['error'] = str(e)
                continue
    
    def _convert_word_to_pdf(self, file_path):
        """Конвертирует Word документ в PDF и возвращает путь к временному файлу"""
        try:
            import tempfile
            
            if not os.path.exists(file_path):
                return None
            
            temp_pdf_fd, temp_pdf_path = tempfile.mkstemp(suffix='.pdf', prefix='word_preview_')
            os.close(temp_pdf_fd)
            
            # Пробуем конвертировать через win32com (Windows)
            if WIN32COM_AVAILABLE:
                try:
                    import win32com.client
                    import pythoncom
                    
                    pythoncom.CoInitialize()
                    
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    
                    # Открываем Word документ
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    
                    doc.SaveAs(os.path.abspath(temp_pdf_path), FileFormat=17)
                    doc.Close()
                    word.Quit()
                    
                    pythoncom.CoUninitialize()
                    
                    return temp_pdf_path
                    
                except Exception as e:
                    # Если win32com не сработал, пробуем docx2pdf
                    if os.path.exists(temp_pdf_path):
                        try:
                            os.unlink(temp_pdf_path)
                        except:
                            pass
            
            # Используем docx2pdf как запасной вариант
            if DOCX2PDF_AVAILABLE:
                try:
                    from docx2pdf import convert
                    convert(file_path, temp_pdf_path)
                    return temp_pdf_path
                except Exception as e:
                    if os.path.exists(temp_pdf_path):
                        try:
                            os.unlink(temp_pdf_path)
                        except:
                            pass
            
            # Если ничего не сработало
            if os.path.exists(temp_pdf_path):
                try:
                    os.unlink(temp_pdf_path)
                except:
                    pass
            
            return None
            
        except Exception as e:
            return None
    
    def preload(self, file_path):
        """Добавляет файл в очередь на предзагрузку"""
        if not file_path or not os.path.exists(file_path):
            return
        
        if not file_path.lower().endswith(('.docx', '.doc')):
            return
        
        if file_path in self.cache:
            status = self.cache[file_path]['status']
            if status in ('processing', 'ready'):
                return  # Уже обрабатывается или готов
        
        if not self.running:
            self.start()
        
        self.queue.put(file_path)
    
    def get_cached_pdf(self, file_path):
        """Возвращает путь к кэшированному PDF или None"""
        if file_path not in self.cache:
            return None
        
        cache_entry = self.cache[file_path]
        if cache_entry['status'] == 'ready':
            cache_entry['timestamp'] = datetime.now().timestamp()
            return cache_entry['temp_pdf_path']
        
        return None
    
    def get_status(self, file_path):
        """Возвращает статус конвертации: 'processing', 'ready', 'error', или None"""
        if file_path not in self.cache:
            return None
        return self.cache[file_path]['status']
    
    def _cleanup_old_cache(self):
        """Очищает старые записи из кэша"""
        current_time = datetime.now().timestamp()
        
        # Удаляем записи старше max_cache_age
        expired = []
        for file_path, entry in self.cache.items():
            if current_time - entry['timestamp'] > self.max_cache_age:
                expired.append(file_path)
        
        for file_path in expired:
            self._remove_from_cache(file_path)
        
        # Если кэш слишком большой, удаляем самые старые записи
        if len(self.cache) > self.max_cache_size:
            # Сортируем по timestamp
            sorted_entries = sorted(self.cache.items(), key=lambda x: x[1]['timestamp'])
            to_remove = len(self.cache) - self.max_cache_size
            
            for file_path, _ in sorted_entries[:to_remove]:
                self._remove_from_cache(file_path)
    
    def _remove_from_cache(self, file_path):
        """Удаляет запись из кэша и очищает временный файл"""
        if file_path in self.cache:
            entry = self.cache[file_path]
            temp_pdf = entry.get('temp_pdf_path')
            
            if temp_pdf and os.path.exists(temp_pdf):
                try:
                    os.unlink(temp_pdf)
                except:
                    pass
            
            del self.cache[file_path]
    
    def clear_cache(self):
        """Полностью очищает кэш"""
        for file_path in list(self.cache.keys()):
            self._remove_from_cache(file_path)

# Глобальный экземпляр менеджера предзагрузки
word_preload_manager = WordPreloadManager()

# ── КЛАСС ДЛЯ ВКЛАДКИ ЗАДАЧИ ──────────────────────────────────────────
class TabTask:
    """Класс для одной вкладки с задачей генерации документов"""
    def __init__(self, parent_frame, app, tab_id):
        self.parent_frame = parent_frame
        self.app = app  # Ссылка на главное приложение
        self.tab_id = tab_id
        
        self.tab_name = f"Задача {tab_id}"
        
        self.excel_path = tk.StringVar()
        self.word_template_path = tk.StringVar()
        self.excel_template_path = tk.StringVar()  # Путь к Excel шаблону
        self.output_folder = tk.StringVar(value="документы")
        self.filename_base = tk.StringVar(value="наименование")
        self.filename_pattern = tk.StringVar(value="наименование {i:04d}{suffix}.docx")
        self.filename_column = tk.StringVar(value="")
        self.filename_mode = tk.StringVar(value="standard")
        
        self.custom_list_vars = {}
        self.custom_list_combos = {}
        
        self.excel_columns = []
        
        self.last_excel_dir = self.app.last_excel_dir
        self.last_word_dir = self.app.last_word_dir
        self.last_output_dir = self.app.last_output_dir
        
        self.is_processing = False
        self.should_stop = False  # Флаг для остановки обработки
        
        self.create_widgets()
    
    def create_widgets(self):
        """Создание современного интерфейса вкладки"""
        # Контейнер для canvas и scrollbar
        container = tk.Frame(self.parent_frame, bg=COLORS["bg_secondary"])
        container.pack(fill=tk.BOTH, expand=True)
        
        # Canvas для прокрутки
        canvas = tk.Canvas(
            container,
            bg=COLORS["bg_secondary"],
            highlightthickness=0,
            borderwidth=0
        )
        
        # Scrollbar
        scrollbar = tk.Scrollbar(
            container,
            orient="vertical",
            command=canvas.yview
        )
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Фрейм внутри canvas
        scrollable_frame = tk.Frame(canvas, bg=COLORS["bg_secondary"])
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        # Обновление области прокрутки при изменении размера
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        def on_canvas_configure(event):
            canvas.itemconfig(canvas_window, width=event.width)
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas.bind("<Configure>", on_canvas_configure)
        
        # Поддержка прокрутки колесом мыши (везде в пределах canvas)
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_to_mousewheel(widget):
            """Рекурсивно привязываем прокрутку ко всем виджетам, КРОМЕ Text и Listbox"""
            try:
                # НЕ привязываем к Text и Listbox - у них свой скролл
                if not isinstance(widget, (tk.Text, tk.Listbox)):
                    widget.bind("<MouseWheel>", on_mousewheel)
                for child in widget.winfo_children():
                    bind_to_mousewheel(child)
            except:
                pass
        
        canvas.bind("<MouseWheel>", on_mousewheel)
        scrollable_frame.bind("<MouseWheel>", on_mousewheel)
        
        self._bind_to_mousewheel = bind_to_mousewheel
        
        main_frame = tk.Frame(scrollable_frame, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH, expand=True, padx=SPACING["lg"], pady=SPACING["lg"])
        
        # ══════════════════════════════════════════════════════════════
        # СЕКЦИЯ 1: БАЗОВЫЕ ДАННЫЕ
        # ══════════════════════════════════════════════════════════════
        db_section, db_content = create_section(main_frame, "Базовые данные", "📅")
        db_section.pack(fill=tk.X, pady=(0, SPACING["md"]))
        
        db_content.grid_columnconfigure(1, weight=1)
        
        tk.Label(
            db_content,
            text="Дата:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=0, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        date_frame = tk.Frame(db_content, bg=COLORS["card_bg"])
        date_frame.grid(row=0, column=1, sticky="w", pady=SPACING["sm"])
        
        self.selected_date = SimpleDatePicker(date_frame, width=12)
        self.selected_date.pack()
        
        self.db_frame = db_content
        self.custom_lists_row = 1
        self.refresh_custom_list_widgets()
        
        # ══════════════════════════════════════════════════════════════
        # СЕКЦИЯ 2: ФАЙЛЫ
        # ══════════════════════════════════════════════════════════════
        files_section, files_content = create_section(main_frame, "Файлы", "📁")
        files_section.pack(fill=tk.X, pady=(0, SPACING["md"]))
        
        files_content.grid_columnconfigure(1, weight=1)
        
        tk.Label(
            files_content,
            text="Excel файл:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=0, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        excel_entry = ctk.CTkEntry(
            files_content,
            textvariable=self.excel_path,
            font=FONTS["body"],
            state="readonly",
            fg_color=COLORS["bg_tertiary"],
            border_color=COLORS["border"],
            height=28
        )
        excel_entry.grid(row=0, column=1, sticky="ew", pady=SPACING["sm"], padx=(0, SPACING["sm"]))
        enable_field_shortcuts(excel_entry, readonly=True)
        add_context_menu(excel_entry, readonly=True)
        ToolTip(excel_entry, "Путь к Excel файлу с данными")
        
        def on_excel_drop(file_path):
            self.last_excel_dir = os.path.dirname(file_path)
            self.log(f"Excel файл выбран: {file_path}")
            self.load_excel_columns(file_path)
        
        setup_file_drop(
            excel_entry, 
            self.excel_path, 
            file_types=['.xlsx', '.xls'],
            on_drop_callback=on_excel_drop
        )
        
        excel_btn_frame = tk.Frame(files_content, bg=COLORS["card_bg"])
        excel_btn_frame.grid(row=0, column=2, pady=SPACING["sm"])
        
        excel_btn = create_modern_button(
            excel_btn_frame,
            text="Обзор",
            command=self.browse_excel,
            style="primary",
            width=70,
            height=28
        )
        excel_btn.pack(side=tk.LEFT, padx=(0, SPACING["xs"]))
        ToolTip(excel_btn, "Открыть диалог выбора Excel файла")
        
        excel_preview_btn = create_icon_button(
            excel_btn_frame,
            icon="👁",
            command=self.preview_excel,
            tooltip="Предпросмотр Excel файла",
            width=28,
            height=28
        )
        excel_preview_btn.pack(side=tk.LEFT)
        
        # Word шаблон
        tk.Label(
            files_content,
            text="Word шаблон:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=1, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        self.word_entry = ctk.CTkEntry(
            files_content,
            textvariable=self.word_template_path,
            font=FONTS["body"],
            state="readonly",
            fg_color=COLORS["bg_tertiary"],
            border_color=COLORS["border"],
            height=28
        )
        self.word_entry.grid(row=1, column=1, sticky="ew", pady=SPACING["sm"], padx=(0, SPACING["sm"]))
        enable_field_shortcuts(self.word_entry, readonly=True)
        add_context_menu(self.word_entry, readonly=True)
        ToolTip(self.word_entry, "Путь к шаблону Word документа")
        
        def on_word_drop(file_path):
            self.last_word_dir = os.path.dirname(file_path)
            self.log(f"Word шаблон выбран: {file_path}")
            word_preload_manager.preload(file_path)
        
        setup_file_drop(
            self.word_entry,
            self.word_template_path,
            file_types=['.docx'],
            on_drop_callback=on_word_drop
        )
        
        self.word_btn_frame = tk.Frame(files_content, bg=COLORS["card_bg"])
        self.word_btn_frame.grid(row=1, column=2, pady=SPACING["sm"])
        
        word_btn = create_modern_button(
            self.word_btn_frame,
            text="Обзор",
            command=self.browse_word_template,
            style="primary",
            width=70,
            height=28
        )
        word_btn.pack(side=tk.LEFT, padx=(0, SPACING["xs"]))
        ToolTip(word_btn, "Открыть диалог выбора Word шаблона")
        
        word_preview_btn = create_icon_button(
            self.word_btn_frame,
            icon="👁",
            command=self.preview_word_template,
            tooltip="Предпросмотр Word шаблона",
            width=28,
            height=28
        )
        word_preview_btn.pack(side=tk.LEFT)
        
        # Excel шаблон
        tk.Label(
            files_content,
            text="Excel шаблон:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=2, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        self.excel_template_entry = ctk.CTkEntry(
            files_content,
            textvariable=self.excel_template_path,
            font=FONTS["body"],
            state="readonly",
            fg_color=COLORS["bg_tertiary"],
            border_color=COLORS["border"],
            height=28
        )
        self.excel_template_entry.grid(row=2, column=1, sticky="ew", pady=SPACING["sm"], padx=(0, SPACING["sm"]))
        enable_field_shortcuts(self.excel_template_entry, readonly=True)
        add_context_menu(self.excel_template_entry, readonly=True)
        ToolTip(self.excel_template_entry, "Путь к шаблону Excel документа")
        
        def on_excel_template_drop(file_path):
            self.last_word_dir = os.path.dirname(file_path)
            self.log(f"Excel шаблон выбран: {file_path}")
        
        setup_file_drop(
            self.excel_template_entry,
            self.excel_template_path,
            file_types=['.xlsx', '.xls'],
            on_drop_callback=on_excel_template_drop
        )
        
        self.excel_template_btn_frame = tk.Frame(files_content, bg=COLORS["card_bg"])
        self.excel_template_btn_frame.grid(row=2, column=2, pady=SPACING["sm"])
        
        excel_template_btn = create_modern_button(
            self.excel_template_btn_frame,
            text="Обзор",
            command=self.browse_excel_template,
            style="primary",
            width=70,
            height=28
        )
        excel_template_btn.pack(side=tk.LEFT, padx=(0, SPACING["xs"]))
        ToolTip(excel_template_btn, "Открыть диалог выбора Excel шаблона")
        
        excel_template_preview_btn = create_icon_button(
            self.excel_template_btn_frame,
            icon="👁",
            command=self.preview_excel_template,
            tooltip="Предпросмотр Excel шаблона",
            width=28,
            height=28
        )
        excel_template_preview_btn.pack(side=tk.LEFT)
        
        # Подсказка о выборе шаблонов
        hint_label = tk.Label(
            files_content,
            text="💡 Можно выбрать один или оба шаблона одновременно",
            font=("Segoe UI", 8, "italic"),
            bg=COLORS["card_bg"],
            fg=COLORS["text_secondary"]
        )
        hint_label.grid(row=3, column=0, columnspan=3, sticky="w", pady=(0, SPACING["sm"]))
        
        tk.Label(
            files_content,
            text="Папка сохранения:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=4, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        output_entry = ctk.CTkEntry(
            files_content,
            textvariable=self.output_folder,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28
        )
        output_entry.grid(row=4, column=1, sticky="ew", pady=SPACING["sm"], padx=(0, SPACING["sm"]))
        enable_field_shortcuts(output_entry)
        add_context_menu(output_entry)
        ToolTip(output_entry, "Путь к папке для сохранения результатов")
        
        def on_output_drop(folder_path):
            self.last_output_dir = folder_path
            self.log(f"Папка сохранения выбрана: {folder_path}")
        
        setup_folder_drop(
            output_entry,
            self.output_folder,
            on_drop_callback=on_output_drop
        )
        
        output_btn = create_modern_button(
            files_content,
            text="Обзор",
            command=self.browse_output_folder,
            style="primary",
            width=70,
            height=28
        )
        output_btn.grid(row=4, column=2, pady=SPACING["sm"])
        ToolTip(output_btn, "Выбрать папку для сохранения")
        
        # ══════════════════════════════════════════════════════════════
        # СЕКЦИЯ 3: НАСТРОЙКИ НАИМЕНОВАНИЯ ФАЙЛОВ
        # ══════════════════════════════════════════════════════════════
        naming_section, naming_content = create_section(main_frame, "Настройки наименования файлов", "✏️")
        naming_section.pack(fill=tk.X, pady=(0, SPACING["md"]))
        
        naming_content.grid_columnconfigure(1, weight=1)
        
        tk.Label(
            naming_content,
            text="Базовое имя:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=0, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        filename_base_entry = ctk.CTkEntry(
            naming_content,
            textvariable=self.filename_base,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28
        )
        filename_base_entry.grid(row=0, column=1, sticky="ew", pady=SPACING["sm"])
        enable_field_shortcuts(filename_base_entry)
        add_context_menu(filename_base_entry)
        filename_base_entry.bind("<KeyRelease>", lambda e: self.update_filename_pattern())
        ToolTip(filename_base_entry, "Базовое наименование для генерируемых файлов")
        
        tk.Label(
            naming_content,
            text="Режим:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=1, column=0, sticky="nw", pady=SPACING["md"], padx=(0, SPACING["md"]))
        
        naming_mode_frame = tk.Frame(naming_content, bg=COLORS["card_bg"])
        naming_mode_frame.grid(row=1, column=1, sticky="w", pady=SPACING["sm"])
        
        tk.Radiobutton(
            naming_mode_frame,
            text="Стандартный (имя_0001.docx)",
            variable=self.filename_mode,
            value="standard",
            command=self.update_filename_pattern,
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            activebackground=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).pack(anchor="w", pady=2)
        
        tk.Radiobutton(
            naming_mode_frame,
            text="С данными (имя_значение.docx)",
            variable=self.filename_mode,
            value="column",
            command=self.update_filename_pattern,
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            activebackground=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).pack(anchor="w", pady=2)
        
        tk.Radiobutton(
            naming_mode_frame,
            text="Смешанное (имя_значение_0001.docx)",
            variable=self.filename_mode,
            value="mixed",
            command=self.update_filename_pattern,
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            activebackground=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).pack(anchor="w", pady=2)
        
        tk.Label(
            naming_content,
            text="Столбец данных:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=2, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        self.filename_column_combo = ctk.CTkComboBox(
            naming_content,
            variable=self.filename_column,
            values=[""],
            state="readonly",
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            button_color=COLORS["primary"],
            button_hover_color=COLORS["primary_hover"],
            dropdown_fg_color=COLORS["bg_primary"],
            height=28,
            command=lambda e: self.update_filename_pattern()
        )
        self.filename_column_combo.grid(row=2, column=1, sticky="ew", pady=SPACING["sm"])
        self.filename_column_combo.set("")
        set_combobox_cursor(self.filename_column_combo)
        ToolTip(self.filename_column_combo, "Выберите столбец Excel для имени файла")
        
        tk.Label(
            naming_content,
            text="Пример имени:",
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"]
        ).grid(row=3, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
        
        self.filename_entry = ctk.CTkEntry(
            naming_content,
            textvariable=self.filename_pattern,
            font=FONTS["body"],
            state="readonly",
            fg_color=COLORS["bg_tertiary"],
            border_color=COLORS["border"],
            height=28
        )
        self.filename_entry.grid(row=3, column=1, sticky="ew", pady=SPACING["sm"])
        enable_field_shortcuts(self.filename_entry, readonly=True)
        add_context_menu(self.filename_entry, readonly=True)
        ToolTip(self.filename_entry, "Пример итогового имени файла")
        
        self.update_filename_pattern()
        
        # ══════════════════════════════════════════════════════════════
        # КНОПКА ЗАПУСКА
        # ══════════════════════════════════════════════════════════════
        btn_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        btn_frame.pack(pady=SPACING["xl"])
        
        self.start_btn = create_modern_button(
            btn_frame,
            text="Начать обработку",
            icon="▶",
            command=self.start_processing,
            style="success",
            width=220,
            height=48
        )
        self.start_btn.pack()
        ToolTip(self.start_btn, "Запустить процесс генерации документов\nВо время обработки можно остановить нажатием на эту кнопку")
        
        # ══════════════════════════════════════════════════════════════
        # СЕКЦИЯ 4: ЛОГ ВЫПОЛНЕНИЯ
        # ══════════════════════════════════════════════════════════════
        log_section, log_content = create_section(main_frame, "Лог выполнения", "📋")
        log_section.pack(fill=tk.BOTH, expand=True)
        
        self.log_text = ScrolledText(
            log_content,
            height=10,
            wrap=tk.WORD,
            bg=COLORS["card_bg"],
            font=FONTS["mono"],
            relief=tk.FLAT,
            borderwidth=0
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)
        self.log_text.config(state=tk.DISABLED)
        
        # Контекстное меню для лога
        def show_context_menu(event):
            menu = ModernContextMenu(self.log_text)
            menu.add_command(label="Копировать", command=self.copy_log_text)
            menu.add_command(label="Выделить всё", command=self.select_all_log)
            menu.post(event.x_root, event.y_root)
        
        self.log_text.bind("<Button-3>", show_context_menu)
        
        # Отложенная привязка скролла после создания всех виджетов
        self.parent_frame.after(100, lambda: bind_to_mousewheel(scrollable_frame))
    
    def refresh_custom_list_widgets(self):
        """Обновление виджетов выпадающих списков"""
        # Удаляем старые виджеты
        for key in list(self.custom_list_vars.keys()):
            if key in self.custom_list_combos:
                combo = self.custom_list_combos[key]
                combo.grid_forget()
                combo.destroy()
            del self.custom_list_vars[key]
            del self.custom_list_combos[key]
        
        # Удаляем метки и кнопки
        for widget in self.db_frame.grid_slaves():
            info = widget.grid_info()
            if info and int(info['row']) >= self.custom_lists_row:
                widget.grid_forget()
                widget.destroy()
        
        row = self.custom_lists_row
        for key, list_data in self.app.CUSTOM_LISTS.items():
            if isinstance(list_data, dict):
                display_name = list_data.get("display_name", key)
                values = list_data.get("values", [])
                is_active = list_data.get("active", True)
            else:
                display_name = key
                values = list_data
                is_active = True
            
            if not is_active:
                continue
            
            tk.Label(
                self.db_frame, 
                text=f"{display_name}:", 
                anchor="w",
                font=FONTS["body"],
                bg=COLORS["card_bg"],
                fg=COLORS["text_primary"]
            ).grid(row=row, column=0, sticky="w", pady=SPACING["sm"], padx=(0, SPACING["md"]))
            
            var = tk.StringVar()
            combo = ctk.CTkComboBox(
                self.db_frame, 
                variable=var, 
                values=values, 
                state="readonly",
                font=FONTS["body"],
                fg_color=COLORS["bg_primary"],
                border_color=COLORS["border"],
                button_color=COLORS["primary"],
                button_hover_color=COLORS["primary_hover"],
                dropdown_fg_color=COLORS["bg_primary"],
                height=28,
                width=200
            )
            combo.grid(row=row, column=1, pady=SPACING["sm"], sticky="w")
            if values:
                combo.set(values[0])
            set_combobox_cursor(combo)
            ToolTip(combo, f"Выберите значение для {display_name}")
            
            self.custom_list_vars[key] = var
            self.custom_list_combos[key] = combo
            row += 1
    
    def update_filename_pattern(self):
        """Обновление шаблона имени файла"""
        mode = self.filename_mode.get()
        base_name = self.filename_base.get() or "документ"
        
        if mode == "standard":
            self.filename_pattern.set(f"{base_name} {{i:04d}}{{suffix}}.docx")
        elif mode == "column":
            self.filename_pattern.set(f"{base_name} {{column}}{{suffix}}.docx")
        elif mode == "mixed":
            self.filename_pattern.set(f"{base_name} {{column}} {{i:04d}}{{suffix}}.docx")
        
        self.filename_entry.configure(state="normal")
        self.filename_entry.configure(state="readonly")
    
    def browse_excel(self):
        """Выбор Excel файла"""
        filename = filedialog.askopenfilename(
            title="Выберите Excel файл",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("Все файлы", "*.*")],
            initialdir=self.last_excel_dir
        )
        if filename:
            self.excel_path.set(filename)
            self.last_excel_dir = os.path.dirname(filename)
            self.log(f"Excel файл выбран: {filename}")
            self.load_excel_columns(filename)
    
    def load_excel_columns(self, filename):
        """Загрузка списка столбцов из Excel файла"""
        if self.app._pandas_loaded and self.app._pandas:
            pd = self.app._pandas
        else:
            import pandas as pd
        
        try:
            df = pd.read_excel(filename, engine='openpyxl', nrows=0)
            self.excel_columns = list(df.columns)
            
            column_values = [""] + self.excel_columns
            self.filename_column_combo.configure(values=column_values)
            self.filename_column_combo.set("")
            
            self.log(f"Найдено столбцов: {len(self.excel_columns)}")
        except Exception as e:
            self.log(f"⚠ Не удалось прочитать заголовки: {e}")
            self.excel_columns = []
            self.filename_column_combo.configure(values=[""])
            self.filename_column_combo.set("")
    
    def browse_word_template(self):
        """Выбор Word шаблона"""
        filename = filedialog.askopenfilename(
            title="Выберите Word шаблон",
            filetypes=[("Word files", "*.docx"), ("Все файлы", "*.*")],
            initialdir=self.last_word_dir
        )
        if filename:
            self.word_template_path.set(filename)
            self.last_word_dir = os.path.dirname(filename)
            self.log(f"Word шаблон выбран: {filename}")
            
            # Запускаем фоновую предзагрузку для быстрого просмотра
            word_preload_manager.preload(filename)
    
    def preview_excel(self):
        """Предварительный просмотр Excel файла"""
        excel_file = self.excel_path.get()
        if not excel_file or not os.path.exists(excel_file):
            messagebox.showwarning("Предупреждение", "Сначала выберите Excel файл!")
            return
        
        try:
            PreviewWindow(self.parent_frame, excel_file, f"Просмотр: {os.path.basename(excel_file)}", data_manager=self.app)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{str(e)}")
    
    def preview_word_template(self):
        """Предварительный просмотр Word шаблона"""
        word_file = self.word_template_path.get()
        if not word_file or not os.path.exists(word_file):
            messagebox.showwarning("Предупреждение", "Сначала выберите Word шаблон!")
            return
        
        try:
            PreviewWindow(self.parent_frame, word_file, f"Просмотр: {os.path.basename(word_file)}", data_manager=self.app)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{str(e)}")
    
    def browse_excel_template(self):
        """Выбор Excel шаблона"""
        filename = filedialog.askopenfilename(
            title="Выберите Excel шаблон",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("Все файлы", "*.*")],
            initialdir=self.last_word_dir
        )
        if filename:
            self.excel_template_path.set(filename)
            self.last_word_dir = os.path.dirname(filename)
            self.log(f"Excel шаблон выбран: {filename}")
    
    def preview_excel_template(self):
        """Предварительный просмотр Excel шаблона"""
        excel_file = self.excel_template_path.get()
        if not excel_file or not os.path.exists(excel_file):
            messagebox.showwarning("Предупреждение", "Сначала выберите Excel шаблон!")
            return
        
        try:
            PreviewWindow(self.parent_frame, excel_file, f"Просмотр: {os.path.basename(excel_file)}", data_manager=self.app)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{str(e)}")
    
    def browse_output_folder(self):
        """Выбор папки для сохранения"""
        folder = filedialog.askdirectory(
            title="Выберите папку для сохранения документов",
            initialdir=self.last_output_dir
        )
        if folder:
            self.output_folder.set(folder)
            self.last_output_dir = folder
            self.log(f"Папка сохранения выбрана: {folder}")
    
    def log(self, message):
        """Добавление сообщения в лог"""
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.config(state=tk.DISABLED)
    
    def copy_log_text(self):
        """Копирование выделенного текста"""
        try:
            selected_text = self.log_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.app.root.clipboard_clear()
            self.app.root.clipboard_append(selected_text)
        except tk.TclError:
            pass
    
    def select_all_log(self):
        """Выделить весь текст в логе"""
        self.log_text.tag_add(tk.SEL, "1.0", tk.END)
        self.log_text.mark_set(tk.INSERT, "1.0")
        self.log_text.see(tk.INSERT)
    
    def start_processing(self):
        """Запуск или остановка обработки документов"""
        # Если обработка идет - останавливаем
        if self.is_processing:
            self.should_stop = True
            self.start_btn.configure(text="⏹ Остановка...")
            self.log("\n⚠️ Запрошена остановка обработки...")
            return
        
        if not self.excel_path.get():
            messagebox.showerror("Ошибка", "Выберите Excel файл!")
            return
        
        # Проверяем наличие хотя бы одного шаблона
        has_word = bool(self.word_template_path.get())
        has_excel = bool(self.excel_template_path.get())
        
        if not has_word and not has_excel:
            messagebox.showerror("Ошибка", "Выберите хотя бы один шаблон (Word или Excel)!")
            return
        
        if not self.output_folder.get():
            messagebox.showerror("Ошибка", "Укажите папку для сохранения!")
            return
        
        mode = self.filename_mode.get()
        if mode in ("column", "mixed") and not self.filename_column.get():
            result = messagebox.askyesno(
                "Предупреждение", 
                f"Для режима '{('С данными' if mode == 'column' else 'Смешанное')}' не выбран столбец.\n\n"
                f"Вместо значений из столбца будет использоваться 'строкаN'.\n\n"
                f"Продолжить?"
            )
            if not result:
                return
        
        self.is_processing = True
        self.should_stop = False
        self.start_btn.configure(text="⏹ Остановить")
        thread = threading.Thread(target=self.process_documents)
        thread.daemon = True
        thread.start()
    
    def process_documents(self):
        """Обработка документов (вызывается в отдельном потоке)"""
        # Передаём управление главному классу
        self.app.process_documents_for_tab(self)

# ══════════════════════════════════════════════════════════════════════
# КЛАСС ДЛЯ ВКЛАДКИ ОБЪЕДИНЕНИЯ И КОНВЕРТАЦИИ ДОКУМЕНТОВ
# ══════════════════════════════════════════════════════════════════════

class MergeTabTask:
    """Класс для одной вкладки с задачей объединения/конвертации документов"""
    def __init__(self, parent_frame, window, tab_id):
        self.parent_frame = parent_frame
        self.window = window  # Ссылка на окно MergeDocumentsWindow
        self.tab_id = tab_id
        
        self.tab_name = f"Задача {tab_id}"
        
        self.file_list = []
        self.doc_type = tk.StringVar(value="word")
        self.use_ocr = tk.BooleanVar(value=True)  # Применять OCR по умолчанию
        self.is_processing = False
        self.should_stop = False  # Флаг для остановки обработки
        
        # Новые параметры для качества и размещения
        self.ocr_resolution = tk.StringVar(value="1.5")  # Разрешение OCR (1.0-3.0)
        self.max_image_size = tk.StringVar(value="4000")  # Макс. размер изображения
        self.fit_mode = tk.StringVar(value="центр")  # Режим размещения
        
        # Параметры нумерации/штампа
        self.use_numbering = tk.BooleanVar(value=False)  # Включить нумерацию (по умолчанию выключена)
        self.numbering_line1 = tk.StringVar(value="")  # Строка 1
        self.numbering_line2 = tk.StringVar(value="")  # Строка 2 (с автоинкрементом)
        self.numbering_line3 = tk.StringVar(value="")  # Строка 3
        self.numbering_position = tk.StringVar(value="правый-нижний")  # Позиция штампа
        self.numbering_border = tk.BooleanVar(value=True)  # Обводить в рамку
        self.numbering_increment_mode = tk.StringVar(value="per_document")  # Режим нумерации: per_document или per_page
        
        # Переменные для прогресса и ETA
        self.start_time = None
        self.total_items = 0
        self.processed_items = 0
        
        self.create_widgets()
    
    def create_widgets(self):
        """Создание интерфейса вкладки с прокруткой"""
        # Создаём Canvas с прокруткой
        container = tk.Frame(self.parent_frame, bg=COLORS["bg_secondary"])
        container.pack(fill=tk.BOTH, expand=True)
        
        # Canvas для прокрутки
        self.merge_canvas = tk.Canvas(container, bg=COLORS["bg_secondary"], highlightthickness=0)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=self.merge_canvas.yview)
        
        self.scrollable_frame = tk.Frame(self.merge_canvas, padx=18, pady=18, bg=COLORS["bg_secondary"])
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.merge_canvas.configure(scrollregion=self.merge_canvas.bbox("all"))
        )
        
        self.canvas_window = self.merge_canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        
        # Привязка изменения ширины canvas к ширине содержимого
        self.merge_canvas.bind('<Configure>', self._on_canvas_configure)
        
        self.merge_canvas.configure(yscrollcommand=scrollbar.set)
        
        scrollbar.pack(side="right", fill="y")
        self.merge_canvas.pack(side="left", fill="both", expand=True)
        
        # Привязка колеса мыши к прокрутке
        def _on_mousewheel(event):
            self.merge_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        self._mousewheel_bound = True
        self._mousewheel_func = _on_mousewheel
        
        # Привязка ко всем виджетам в окне, КРОМЕ Text и Listbox
        def _bind_mousewheel_to_widget(widget):
            try:
                # НЕ привязываем к Text и Listbox - у них свой скролл
                if not isinstance(widget, (tk.Text, tk.Listbox)):
                    widget.bind("<MouseWheel>", _on_mousewheel)
                for child in widget.winfo_children():
                    _bind_mousewheel_to_widget(child)
            except:
                pass
        
        # Отложенная привязка после создания всех виджетов
        self.parent_frame.after(100, lambda: _bind_mousewheel_to_widget(self.parent_frame))
        
        main_frame = self.scrollable_frame
        
        # Выбор функции
        type_frame = tk.LabelFrame(
            main_frame, 
            text=" Функция ", 
            font=FONTS["heading"], 
            padx=12, 
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        type_frame.pack(fill=tk.X, pady=(0, 12))
        
        # Варианты функций с отображаемым текстом и внутренним значением
        function_options = [
            ("Объединить Word документы", "word"),
            ("Объединить PDF документы", "pdf"),
            ("Конвертировать Word → PDF (раздельно)", "convert"),
            ("Конвертировать Word → единый PDF", "convert_merge"),
            ("Конвертировать изображения → PDF (раздельно)", "image"),
            ("Конвертировать изображения → единый PDF", "image_merge"),
            ("Пронумеровать PDF (раздельно)", "number_separate"),
            ("Пронумеровать PDF (с объединением)", "number_merge"),
            ("Разделить PDF файл", "split_pdf"),
            ("Повернуть страницы PDF", "rotate_pdf"),
            ("Извлечь страницы из PDF", "extract_pdf"),
            ("Извлечь данные в Excel", "extract_to_excel"),
        ]
        
        function_label = tk.Label(
            type_frame,
            text="Выберите функцию:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            anchor="w"
        )
        function_label.pack(anchor="w", pady=(0, 5))
        
        self.function_combo = ctk.CTkComboBox(
            type_frame,
            values=[opt[0] for opt in function_options],
            state="readonly",
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            button_color=COLORS["primary"],
            button_hover_color=COLORS["primary_hover"],
            dropdown_fg_color=COLORS["bg_primary"],
            width=400,
            height=32
        )
        self.function_combo.pack(anchor="w", pady=(0, 5))
        
        # Метод для управления видимостью настроек нумерации
        def toggle_numbering_visibility():
            """Показывает/скрывает настройки нумерации на основе чекбокса"""
            doc_type_value = self.doc_type.get()
            
            # Для функций прямой нумерации чекбокс не влияет (настройки всегда видимы)
            if doc_type_value in ['number_separate', 'number_merge']:
                return
            
            # Для остальных функций управляем видимостью через чекбокс
            if self.use_numbering.get():
                self.numbering_subframe.pack(fill=tk.X, pady=(10, 0))
            else:
                self.numbering_subframe.pack_forget()
        
        # Сохраняем метод в self для доступа из других мест
        self.toggle_numbering_visibility = toggle_numbering_visibility
        
        def toggle_split_mode():
            """Показывает/скрывает поле диапазонов в зависимости от режима разделения"""
            if self.split_mode.get() == "ranges":
                self.split_ranges_frame.pack(fill=tk.X, pady=(0, 10))
                self.split_help_label.pack(pady=(5, 0))
            else:
                self.split_ranges_frame.pack_forget()
                self.split_help_label.pack_forget()
        
        # Сохраняем метод в self для доступа из других мест
        self.toggle_split_mode = toggle_split_mode
        
        # Устанавливаем соответствие между label и value
        def on_function_change(choice):
            for label, value in function_options:
                if label == choice:
                    self.doc_type.set(value)
                    
                    # Сначала скрываем все дополнительные фреймы
                    self.ocr_frame.pack_forget()
                    self.numbering_checkbox_frame.pack_forget()
                    self.numbering_subframe.pack_forget()
                    self.split_pdf_frame.pack_forget()
                    self.rotate_pdf_frame.pack_forget()
                    self.extract_pdf_frame.pack_forget()
                    self.extract_to_excel_frame.pack_forget()
                    
                    # Показываем нужные фреймы в зависимости от функции
                    if value == 'word':
                        # Для Word не нужны дополнительные настройки
                        pass
                    
                    elif value == 'split_pdf':
                        # Разделение PDF - показываем настройки разделения
                        self.split_pdf_frame.pack(fill=tk.X, pady=(0, 12))
                    
                    elif value == 'rotate_pdf':
                        # Поворот страниц - показываем настройки поворота
                        self.rotate_pdf_frame.pack(fill=tk.X, pady=(0, 12))
                    
                    elif value == 'extract_pdf':
                        # Извлечение страниц - показываем настройки извлечения
                        self.extract_pdf_frame.pack(fill=tk.X, pady=(0, 12))
                    
                    elif value == 'extract_to_excel':
                        # Извлечение данных в Excel - показываем настройки
                        self.extract_to_excel_frame.pack(fill=tk.X, pady=(0, 12))
                    
                    elif value in ['number_separate', 'number_merge']:
                        # Для функций прямой нумерации:
                        # - показываем ocr_frame
                        # - скрываем чекбокс (он не нужен)
                        # - показываем настройки нумерации (всегда)
                        self.ocr_frame.pack(fill=tk.X, pady=(0, 12))
                        self.numbering_subframe.pack(fill=tk.X, pady=(10, 0))
                        # Включаем нумерацию принудительно
                        self.use_numbering.set(True)
                    
                    else:
                        # Для всех остальных функций (конвертация, объединение):
                        # - показываем ocr_frame
                        # - показываем чекбокс нумерации
                        # - управление видимостью через чекбокс
                        self.ocr_frame.pack(fill=tk.X, pady=(0, 12))
                        self.numbering_checkbox_frame.pack(fill=tk.X, pady=(10, 0))
                        
                        # Применяем текущее состояние чекбокса
                        if self.use_numbering.get():
                            self.numbering_subframe.pack(fill=tk.X, pady=(10, 0))
                    break
        
        self.function_combo.configure(command=on_function_change)
        
        # Устанавливаем начальное отображаемое значение
        current_value = self.doc_type.get()
        for label, value in function_options:
            if value == current_value:
                self.function_combo.set(label)
                break
        
        set_combobox_cursor(self.function_combo)
        ToolTip(
            self.function_combo,
            "Выберите функцию для работы с документами:\n"
            "• Объединение - склеивает несколько файлов в один\n"
            "• Конвертация - переводит формат документов\n"
            "• Нумерация - добавляет штампы нумерации к PDF"
        )
        
        # Список файлов
        files_frame = tk.LabelFrame(
            main_frame, 
            text=" Файлы ", 
            font=FONTS["heading"], 
            padx=12, 
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        files_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 12))
        
        # Listbox с прокруткой
        list_container = tk.Frame(files_frame, bg=COLORS["bg_secondary"])
        list_container.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = tk.Scrollbar(list_container)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.files_listbox = tk.Listbox(
            list_container, 
            yscrollcommand=scrollbar.set, 
            font=FONTS["body"], 
            relief=tk.SOLID, 
            borderwidth=1
        )
        self.files_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.files_listbox.yview)
        
        btn_files_frame = tk.Frame(files_frame, bg=COLORS["bg_secondary"])
        btn_files_frame.pack(fill=tk.X, pady=(10, 0))
        
        add_btn = create_modern_button(
            btn_files_frame, 
            text="+ Добавить", 
            command=self.add_files, 
            style="primary",
            width=12, 
            tooltip="Добавить файлы в конец списка"
        )
        add_btn.pack(side=tk.LEFT, padx=2)
        
        insert_btn = create_modern_button(
            btn_files_frame, 
            text="⊕ Вставить после", 
            command=self.insert_files_after, 
            style="primary",
            width=14, 
            tooltip="Добавить файлы после выбранного элемента"
        )
        insert_btn.pack(side=tk.LEFT, padx=2)
        
        up_btn = create_modern_button(
            btn_files_frame, 
            text="⬆ Вверх", 
            command=self.move_up, 
            style="secondary",
            width=10, 
            tooltip="Переместить выбранный файл вверх"
        )
        up_btn.pack(side=tk.LEFT, padx=2)
        
        down_btn = create_modern_button(
            btn_files_frame, 
            text="⬇ Вниз", 
            command=self.move_down, 
            style="secondary",
            width=10, 
            tooltip="Переместить выбранный файл вниз"
        )
        down_btn.pack(side=tk.LEFT, padx=2)
        
        preview_btn = tk.Button(
            btn_files_frame,
            text="👁",
            command=self.preview_selected_file,
            width=3,
            font=FONTS["button"],
            bg=COLORS["success"],
            fg="white",
            relief=tk.FLAT,
            cursor="hand2",
            activebackground=COLORS["success_hover"]
        )
        preview_btn.pack(side=tk.LEFT, padx=2)
        ToolTip(preview_btn, "Предварительный просмотр выбранного файла")
        
        # Счетчик файлов
        self.file_counter_label = tk.Label(
            btn_files_frame,
            text="Файлов: 0",
            font=FONTS["body"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"]
        )
        self.file_counter_label.pack(side=tk.RIGHT, padx=5)
        
        # Вторая строка кнопок
        btn_files_frame2 = tk.Frame(files_frame, bg=COLORS["bg_secondary"])
        btn_files_frame2.pack(fill=tk.X, pady=(5, 10))
        
        del_btn = create_modern_button(
            btn_files_frame2, 
            text="Удалить", 
            command=self.remove_file, 
            style="danger",
            width=10, 
            tooltip="Удалить выбранный файл из списка"
        )
        del_btn.pack(side=tk.LEFT, padx=2)
        
        clear_all_btn = create_modern_button(
            btn_files_frame2, 
            text="🗑 Очистить все", 
            command=self.clear_all_files, 
            style="danger",
            width=14, 
            tooltip="Удалить все файлы из списка"
        )
        clear_all_btn.pack(side=tk.LEFT, padx=2)
        
        cache_btn = create_modern_button(
            btn_files_frame2, 
            text="💾 Очистить кэш", 
            command=self.clear_document_cache, 
            style="warning",
            width=14, 
            tooltip="Очистить кэш созданных документов"
        )
        cache_btn.pack(side=tk.LEFT, padx=2)
        
        # Подсказка о Drag and Drop
        if TKDND_AVAILABLE:
            hint_text = "💡 Вы можете перетаскивать файлы из системы в нужное место списка"
            hint_color = COLORS["text_secondary"]
        else:
            hint_text = "💡 Установите tkinterdnd2 для добавления файлов перетаскиванием: pip install tkinterdnd2"
            hint_color = COLORS["warning"]
        
        hint_label = tk.Label(
            files_frame,
            text=hint_text,
            font=FONTS["small"],
            fg=hint_color,
            bg=COLORS["bg_secondary"]
        )
        hint_label.pack(pady=(5, 0))
        
        # Настройка Drag and Drop
        self.setup_drag_and_drop()
        
        # OCR настройки
        self.ocr_frame = tk.LabelFrame(
            main_frame, 
            text=" Настройки обработки PDF ", 
            font=FONTS["heading"], 
            padx=12, 
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        self.ocr_frame.pack(fill=tk.X, pady=(0, 12))
        
        ocr_checkbox = tk.Checkbutton(
            self.ocr_frame,
            text="Применять текстовый слой (OCR) при объединении/конвертации PDF",
            variable=self.use_ocr,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            activebackground=COLORS["bg_secondary"],
            selectcolor=COLORS["bg_primary"]
        )
        ocr_checkbox.pack(anchor="w", pady=3)
        ToolTip(
            ocr_checkbox, 
            "Включите для распознавания текста в сканах и изображениях.\n"
            "Отключите для быстрого объединения/конвертации без распознавания текста.\n\n"
            "Примечание: OCR увеличивает время обработки, но позволяет копировать текст из PDF."
        )
        
        # ═══ КАЧЕСТВО И РАЗМЕЩЕНИЕ ═══
        quality_subframe = tk.Frame(self.ocr_frame, bg=COLORS["bg_secondary"])
        quality_subframe.pack(fill=tk.X, pady=(8, 0))
        
        # Разрешение OCR
        ocr_res_frame = tk.Frame(quality_subframe, bg=COLORS["bg_secondary"])
        ocr_res_frame.pack(fill=tk.X, pady=3)
        
        tk.Label(
            ocr_res_frame, 
            text="Качество OCR (разрешение):", 
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        ocr_res_entry = ctk.CTkEntry(
            ocr_res_frame,
            textvariable=self.ocr_resolution,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            width=80,
            height=28
        )
        ocr_res_entry.pack(side=tk.LEFT, padx=(0, 5))
        enable_field_shortcuts(ocr_res_entry)
        add_context_menu(ocr_res_entry)
        ToolTip(
            ocr_res_entry,
            "Множитель разрешения для OCR (1.0-3.0):\n"
            "• 1.0 = быстро, низкое качество\n"
            "• 1.5 = оптимально (по умолчанию)\n"
            "• 2.0-3.0 = медленно, высокое качество"
        )
        
        tk.Label(
            ocr_res_frame,
            text="(1.0-3.0, по умолч. 1.5)",
            font=FONTS["small"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"]
        ).pack(side=tk.LEFT)
        
        # Максимальный размер изображения
        img_size_frame = tk.Frame(quality_subframe, bg=COLORS["bg_secondary"])
        img_size_frame.pack(fill=tk.X, pady=3)
        
        tk.Label(
            img_size_frame,
            text="Макс. размер изображения (px):",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        img_size_entry = ctk.CTkEntry(
            img_size_frame,
            textvariable=self.max_image_size,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            width=80,
            height=28
        )
        img_size_entry.pack(side=tk.LEFT, padx=(0, 5))
        enable_field_shortcuts(img_size_entry)
        add_context_menu(img_size_entry)
        ToolTip(
            img_size_entry,
            "Максимальный размер изображения в пикселях:\n"
            "• 4000 = оптимально (по умолчанию)\n"
            "• 0 или пусто = без ограничений (полное качество)\n"
            "• Меньше = быстрее, но ниже качество"
        )
        
        tk.Label(
            img_size_frame,
            text="(0=без огр., 4000 по умолч.)",
            font=FONTS["small"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"]
        ).pack(side=tk.LEFT)
        
        # Режим размещения на листе
        fit_mode_frame = tk.Frame(quality_subframe, bg=COLORS["bg_secondary"])
        fit_mode_frame.pack(fill=tk.X, pady=3)
        
        tk.Label(
            fit_mode_frame,
            text="Размещение на листе:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            width=20,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        fit_mode_options = [
            ("Центрировать", "центр"),
            ("Заполнить лист", "заполнить"),
            ("Вписать в лист", "вписать")
        ]
        
        fit_mode_combo = ctk.CTkComboBox(
            fit_mode_frame,
            values=[opt[0] for opt in fit_mode_options],
            state="readonly",
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            button_color=COLORS["primary"],
            button_hover_color=COLORS["primary_hover"],
            dropdown_fg_color=COLORS["bg_primary"],
            width=150,
            height=28
        )
        fit_mode_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        # Устанавливаем соответствие между label и value
        def on_fit_mode_change(choice):
            for label, value in fit_mode_options:
                if label == choice:
                    self.fit_mode.set(value)
                    break
        
        fit_mode_combo.configure(command=on_fit_mode_change)
        
        # Устанавливаем начальное отображаемое значение
        current_value = self.fit_mode.get()
        for label, value in fit_mode_options:
            if value == current_value:
                fit_mode_combo.set(label)
                break
        
        set_combobox_cursor(fit_mode_combo)
        ToolTip(
            fit_mode_combo,
            "Режим размещения изображения на листе PDF:\n"
            "• Центрировать = по центру с сохранением пропорций\n"
            "• Заполнить лист = растянуть на весь лист (может исказить)\n"
            "• Вписать в лист = максимально вписать с сохранением пропорций"
        )
        
        # ═══ ЧЕКБОКС НУМЕРАЦИИ ═══
        # Создаем фрейм-контейнер для чекбокса нумерации
        self.numbering_checkbox_frame = tk.Frame(self.ocr_frame, bg=COLORS["bg_secondary"])
        self.numbering_checkbox_frame.pack(fill=tk.X, pady=(10, 0))
        
        self.numbering_checkbox = tk.Checkbutton(
            self.numbering_checkbox_frame,
            text="Добавить нумерацию",
            variable=self.use_numbering,
            command=self.toggle_numbering_visibility,
            font=("Segoe UI", 10),
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            selectcolor=COLORS["bg_primary"],
            activebackground=COLORS["bg_secondary"],
            activeforeground=COLORS["text_primary"]
        )
        self.numbering_checkbox.pack(anchor=tk.W, padx=10, pady=5)
        ToolTip(
            self.numbering_checkbox,
            "При включении появятся настройки для добавления\n"
            "штампа с нумерацией на каждую страницу документа"
        )
        
        # ═══ НУМЕРАЦИЯ / ШТАМП ═══
        # Создаем фрейм нумерации, который будет показываться/скрываться
        self.numbering_subframe = tk.LabelFrame(
            self.ocr_frame,
            text=" Нумерация ",
            font=FONTS["body"],
            padx=8,
            pady=6,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        )
        # По умолчанию скрыт - будет показан при выборе соответствующей функции
        # self.numbering_subframe.pack(fill=tk.X, pady=(10, 0))
        
        # Строка 1
        line1_frame = tk.Frame(self.numbering_subframe, bg=COLORS["bg_secondary"])
        line1_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(
            line1_frame,
            text="Строка 1:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            width=15,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        line1_entry = ctk.CTkEntry(
            line1_frame,
            textvariable=self.numbering_line1,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28
        )
        line1_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 0))
        enable_field_shortcuts(line1_entry)
        add_context_menu(line1_entry)
        ToolTip(line1_entry, "Первая строка (например: ООО 'Компания')")
        
        # Строка 2 (с автоинкрементом)
        line2_frame = tk.Frame(self.numbering_subframe, bg=COLORS["bg_secondary"])
        line2_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(
            line2_frame,
            text="Строка 2:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            width=15,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        line2_entry = ctk.CTkEntry(
            line2_frame,
            textvariable=self.numbering_line2,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28
        )
        line2_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 0))
        enable_field_shortcuts(line2_entry)
        add_context_menu(line2_entry)
        ToolTip(line2_entry, "Вторая строка с автоинкрементом (123, АБВ/1319, № 1819-А)")
        
        # Строка 3
        line3_frame = tk.Frame(self.numbering_subframe, bg=COLORS["bg_secondary"])
        line3_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(
            line3_frame,
            text="Строка 3:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            width=15,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        line3_entry = ctk.CTkEntry(
            line3_frame,
            textvariable=self.numbering_line3,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28
        )
        line3_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 0))
        enable_field_shortcuts(line3_entry)
        add_context_menu(line3_entry)
        ToolTip(line3_entry, "Третья строка (например: от 25.12.2024)")
        
        # Расположение штампа
        position_frame = tk.Frame(self.numbering_subframe, bg=COLORS["bg_secondary"])
        position_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(
            position_frame,
            text="Расположение:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            width=15,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        position_options = [
            ("Правый нижний", "правый-нижний"),
            ("Центр снизу", "центр-нижний"),
            ("Левый нижний", "левый-нижний"),
            ("Левый верхний", "левый-верхний"),
            ("Центр сверху", "центр-верхний"),
            ("Правый верхний", "правый-верхний")
        ]
        
        position_combo = ctk.CTkComboBox(
            position_frame,
            values=[opt[0] for opt in position_options],
            state="readonly",
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            button_color=COLORS["primary"],
            button_hover_color=COLORS["primary_hover"],
            dropdown_fg_color=COLORS["bg_primary"],
            width=150,
            height=28
        )
        position_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        # Устанавливаем соответствие между label и value
        def on_position_change(choice):
            for label, value in position_options:
                if label == choice:
                    self.numbering_position.set(value)
                    break
        
        position_combo.configure(command=on_position_change)
        
        # Устанавливаем начальное отображаемое значение
        current_value = self.numbering_position.get()
        for label, value in position_options:
            if value == current_value:
                position_combo.set(label)
                break
        
        set_combobox_cursor(position_combo)
        ToolTip(position_combo, "Выберите расположение штампа на странице")
        
        # Чекбокс "Поставить штамп" (рамка)
        border_frame = tk.Frame(self.numbering_subframe, bg=COLORS["bg_secondary"])
        border_frame.pack(fill=tk.X, pady=2)
        
        border_check = tk.Checkbutton(
            border_frame,
            text="Поставить штамп (обводить рамкой)",
            variable=self.numbering_border,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            selectcolor=COLORS["bg_primary"]
        )
        border_check.pack(side=tk.LEFT)
        ToolTip(border_check, "Обводить текст белым прямоугольником с черной рамкой")
        
        # Режим инкремента
        increment_frame = tk.Frame(self.numbering_subframe, bg=COLORS["bg_secondary"])
        increment_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(
            increment_frame,
            text="Инкремент строки 2:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            width=15,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        increment_doc = tk.Radiobutton(
            increment_frame,
            text="На каждом документе",
            variable=self.numbering_increment_mode,
            value="per_document",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            selectcolor=COLORS["bg_primary"]
        )
        increment_doc.pack(side=tk.LEFT, padx=(5, 10))
        ToolTip(increment_doc, "Увеличивать номер на каждом новом документе")
        
        increment_doc_first = tk.Radiobutton(
            increment_frame,
            text="На первом листе документа",
            variable=self.numbering_increment_mode,
            value="per_document_first_page",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            selectcolor=COLORS["bg_primary"]
        )
        increment_doc_first.pack(side=tk.LEFT, padx=(0, 10))
        ToolTip(increment_doc_first, "Увеличивать номер документа, но штамповать только первый лист")
        
        increment_page = tk.Radiobutton(
            increment_frame,
            text="На каждом листе",
            variable=self.numbering_increment_mode,
            value="per_page",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            selectcolor=COLORS["bg_primary"]
        )
        increment_page.pack(side=tk.LEFT)
        ToolTip(increment_page, "Увеличивать номер на каждом новом листе")
        
        # Кнопки управления пресетами
        preset_buttons_frame = tk.Frame(self.numbering_subframe, bg=COLORS["bg_secondary"])
        preset_buttons_frame.pack(fill=tk.X, pady=(8, 5))
        
        save_preset_btn = create_icon_button(
            preset_buttons_frame,
            icon="💾",
            command=self.save_stamp_preset,
            tooltip="Сохранить пресет настроек штампа",
            style="success",
            width=32,
            height=28
        )
        save_preset_btn.pack(side=tk.LEFT, padx=2)
        
        load_preset_btn = create_icon_button(
            preset_buttons_frame,
            icon="📂",
            command=self.load_stamp_preset,
            tooltip="Загрузить пресет",
            style="primary",
            width=32,
            height=28
        )
        load_preset_btn.pack(side=tk.LEFT, padx=2)
        
        tk.Label(
            self.numbering_subframe,
            text="💡 Строка 2 автоматически увеличивается (1→2, АБВ/1319→АБВ/1320, № 1819-А→№ 1820-А)",
            font=FONTS["small"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"],
            wraplength=600,
            justify="left"
        ).pack(pady=(5, 0))
        
        # ═══════════════════════════════════════════════════════════
        # НАСТРОЙКИ РАЗДЕЛЕНИЯ PDF
        # ═══════════════════════════════════════════════════════════
        self.split_pdf_frame = tk.LabelFrame(
            main_frame,
            text=" Настройки разделения PDF ",
            font=FONTS["heading"],
            padx=12,
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        
        # Режим разделения
        mode_frame = tk.Frame(self.split_pdf_frame, bg=COLORS["bg_secondary"])
        mode_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            mode_frame,
            text="Режим разделения:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 15))
        
        self.split_mode = tk.StringVar(value="pages")
        
        tk.Radiobutton(
            mode_frame,
            text="По отдельным страницам",
            variable=self.split_mode,
            value="pages",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            selectcolor=COLORS["bg_secondary"],
            command=self.toggle_split_mode
        ).pack(side=tk.LEFT, padx=(0, 15))
        
        tk.Radiobutton(
            mode_frame,
            text="По диапазонам",
            variable=self.split_mode,
            value="ranges",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            selectcolor=COLORS["bg_secondary"],
            command=self.toggle_split_mode
        ).pack(side=tk.LEFT)
        
        # Поле для ввода диапазонов
        self.split_ranges_frame = tk.Frame(self.split_pdf_frame, bg=COLORS["bg_secondary"])
        self.split_ranges_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            self.split_ranges_frame,
            text="Диапазоны страниц:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        self.split_ranges = tk.StringVar(value="1-2; 3-5")
        ranges_entry = CTkEntry(
            self.split_ranges_frame,
            textvariable=self.split_ranges,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28
        )
        ranges_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        enable_field_shortcuts(ranges_entry)
        add_context_menu(ranges_entry)
        
        self.split_help_label = tk.Label(
            self.split_pdf_frame,
            text="💡 Введите диапазоны через точку с запятой или запятую.\nПримеры: 1-2; 3-5; 6; 7-13  или  1-3, 4-6, 7",
            font=FONTS["small"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"],
            wraplength=600,
            justify="left"
        )
        self.split_help_label.pack(pady=(5, 0))
        
        # Скрыть поле для диапазонов по умолчанию
        self.split_ranges_frame.pack_forget()
        self.split_help_label.pack_forget()
        
        # ═══════════════════════════════════════════════════════════
        # НАСТРОЙКИ ПОВОРОТА СТРАНИЦ PDF
        # ═══════════════════════════════════════════════════════════
        self.rotate_pdf_frame = tk.LabelFrame(
            main_frame,
            text=" Настройки поворота страниц ",
            font=FONTS["heading"],
            padx=12,
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        
        angle_frame = tk.Frame(self.rotate_pdf_frame, bg=COLORS["bg_secondary"])
        angle_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            angle_frame,
            text="Угол поворота:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        self.rotate_angle = tk.StringVar(value="90")
        for angle in ["90", "180", "270"]:
            tk.Radiobutton(
                angle_frame,
                text=f"{angle}°",
                variable=self.rotate_angle,
                value=angle,
                font=FONTS["body"],
                bg=COLORS["bg_secondary"],
                fg=COLORS["text_primary"],
                selectcolor=COLORS["bg_primary"]
            ).pack(side=tk.LEFT, padx=(0, 10))
        
        pages_range_frame = tk.Frame(self.rotate_pdf_frame, bg=COLORS["bg_secondary"])
        pages_range_frame.pack(fill=tk.X, pady=(5, 0))
        
        tk.Label(
            pages_range_frame,
            text="Страницы (пусто=все):",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        self.rotate_pages = tk.StringVar(value="")
        rotate_entry = CTkEntry(
            pages_range_frame,
            textvariable=self.rotate_pages,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28,
            width=300
        )
        rotate_entry.pack(side=tk.LEFT)
        enable_field_shortcuts(rotate_entry)
        add_context_menu(rotate_entry)
        
        tk.Label(
            self.rotate_pdf_frame,
            text="💡 Примеры: 1,3,5 или 1-10 или 1-5,10,15-20",
            font=FONTS["small"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"]
        ).pack(pady=(5, 0))
        
        # ═══════════════════════════════════════════════════════════
        # НАСТРОЙКИ ИЗВЛЕЧЕНИЯ СТРАНИЦ PDF
        # ═══════════════════════════════════════════════════════════
        self.extract_pdf_frame = tk.LabelFrame(
            main_frame,
            text=" Настройки извлечения страниц ",
            font=FONTS["heading"],
            padx=12,
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        
        extract_info_frame = tk.Frame(self.extract_pdf_frame, bg=COLORS["bg_secondary"])
        extract_info_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            extract_info_frame,
            text="Укажите страницы для извлечения:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        self.extract_pages = tk.StringVar(value="")
        extract_entry = CTkEntry(
            extract_info_frame,
            textvariable=self.extract_pages,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=28,
            width=400
        )
        extract_entry.pack(side=tk.LEFT)
        enable_field_shortcuts(extract_entry)
        add_context_menu(extract_entry)
        
        tk.Label(
            self.extract_pdf_frame,
            text="💡 Примеры: 1-5 (с 1 по 5), 1,3,5 (страницы 1, 3 и 5), 1-5,10,15-20",
            font=FONTS["small"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"],
            wraplength=600,
            justify="left"
        ).pack(pady=(5, 0))
        
        # ═══════════════════════════════════════════════════════════
        # НАСТРОЙКИ ИЗВЛЕЧЕНИЯ ДАННЫХ В EXCEL
        # ═══════════════════════════════════════════════════════════
        self.extract_to_excel_frame = tk.LabelFrame(
            main_frame,
            text=" Настройки извлечения данных ",
            font=FONTS["heading"],
            padx=12,
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        
        extract_mode_frame = tk.Frame(self.extract_to_excel_frame, bg=COLORS["bg_secondary"])
        extract_mode_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            extract_mode_frame,
            text="Метод извлечения:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        self.extract_method = tk.StringVar(value="full_text")
        extraction_methods = [
            ("Полный текст", "full_text"),
            ("Только таблицы", "tables_only"),
            ("По регулярному выражению", "regex")
        ]
        
        for label, value in extraction_methods:
            tk.Radiobutton(
                self.extract_to_excel_frame,
                text=label,
                variable=self.extract_method,
                value=value,
                font=FONTS["body"],
                bg=COLORS["bg_secondary"],
                fg=COLORS["text_primary"],
                selectcolor=COLORS["bg_primary"]
            ).pack(anchor="w", pady=2)
        
        # Поле для регулярного выражения (опционально)
        regex_frame = tk.Frame(self.extract_to_excel_frame, bg=COLORS["bg_secondary"])
        regex_frame.pack(fill=tk.X, pady=(10, 0))
        
        tk.Label(
            regex_frame,
            text="Регулярное выражение (опционально):",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(anchor="w")
        
        self.extract_regex = tk.StringVar(value="")
        regex_entry = tk.Entry(
            regex_frame,
            textvariable=self.extract_regex,
            font=FONTS["body"],
            width=50
        )
        regex_entry.pack(fill=tk.X, pady=(5, 0))
        
        tk.Label(
            self.extract_to_excel_frame,
            text="💡 Извлекает текст из Word/PDF документов и сохраняет в Excel.\n"
                 "Каждый документ = строка в Excel. Полезно для анализа множества файлов.",
            font=FONTS["small"],
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"],
            wraplength=600,
            justify="left"
        ).pack(pady=(10, 0))
        
        # Кнопка запуска (перед логами)
        btn_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        btn_frame.pack(fill=tk.X, pady=(12, 12))
        
        self.merge_btn = tk.Button(
            btn_frame,
            text="▶ Начать",
            command=self.merge_documents,
            font=FONTS["button"],
            bg=COLORS["success"],
            fg="white",
            activebackground=COLORS["success_hover"],
            activeforeground="white",
            relief=tk.FLAT,
            cursor="hand2",
            padx=30,
            pady=12,
            width=30
        )
        self.merge_btn.pack(pady=5)
        ToolTip(self.merge_btn, "Запустить процесс объединения или конвертации документов")
        
        # Прогресс-бар и ETA
        progress_container = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        progress_container.pack(fill=tk.X, pady=(0, 12))
        
        # Лейбл с инфо о прогрессе
        self.progress_label = tk.Label(
            progress_container,
            text="",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            anchor="w"
        )
        self.progress_label.pack(fill=tk.X, pady=(0, 5))
        
        # Прогресс-бар
        self.progress_bar = ttk.Progressbar(
            progress_container,
            orient="horizontal",
            length=100,
            mode="determinate"
        )
        self.progress_bar.pack(fill=tk.X)
        
        # ETA лейбл (оставшееся время)
        self.eta_label = tk.Label(
            progress_container,
            text="",
            font=FONTS["small"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_secondary"],
            anchor="e"
        )
        self.eta_label.pack(fill=tk.X, pady=(5, 0))
        
        # Скрываем прогресс по умолчанию
        progress_container.pack_forget()
        self.progress_container = progress_container
        
        # Лог выполнения
        log_frame = tk.LabelFrame(
            main_frame, 
            text=" Лог выполнения ", 
            font=FONTS["heading"], 
            padx=12, 
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        log_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 0))
        
        self.log_text = ScrolledText(
            log_frame, 
            height=6, 
            wrap=tk.WORD, 
            bg=COLORS["bg_secondary"],
            font=FONTS["small"],
            relief=tk.FLAT,
            borderwidth=0
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)
        self.log_text.config(state=tk.DISABLED)
        
        # Контекстное меню для лога
        def show_context_menu(event):
            menu = ModernContextMenu(self.log_text)
            menu.add_command(label="Копировать", command=self.copy_log_text)
            menu.add_command(label="Выделить всё", command=self.select_all_log)
            menu.post(event.x_root, event.y_root)
        
        self.log_text.bind("<Button-3>", show_context_menu)
        
        # Применяем начальную логику показа/скрытия фреймов (после создания всех элементов)
        current_value = self.doc_type.get()
        if current_value == 'word':
            # По умолчанию word - скрываем настройки PDF
            self.ocr_frame.pack_forget()
        elif current_value in ['number_separate', 'number_merge']:
            # Для функций нумерации - скрываем чекбокс, показываем настройки, включаем нумерацию
            self.numbering_checkbox_frame.pack_forget()
            self.numbering_subframe.pack(fill=tk.X, pady=(10, 0))
            self.use_numbering.set(True)
        else:
            # Для остальных функций - чекбокс выключен по умолчанию, настройки нумерации скрыты
            self.numbering_subframe.pack_forget()
            self.use_numbering.set(False)
    
    def _on_canvas_configure(self, event):
        """Изменение ширины canvas при изменении размера окна"""
        self.merge_canvas.itemconfig(self.canvas_window, width=event.width)
    
    def setup_drag_and_drop(self):
        """Настройка поддержки перетаскивания файлов"""
        # Внешний drag-and-drop (файлов из системы)
        if TKDND_AVAILABLE:
            try:
                self.files_listbox.drop_target_register(DND_FILES)
                self.files_listbox.dnd_bind('<<Drop>>', self.on_drop)
            except Exception as e:
                pass
        
        # Внутренний drag-and-drop (перестановка элементов в списке) - работает всегда
        self.drag_data = {"index": None, "y": 0}
        
        def on_drag_start(event):
            # Запоминаем индекс элемента под курсором
            index = self.files_listbox.nearest(event.y)
            if index >= 0 and index < len(self.file_list):
                self.drag_data["index"] = index
                self.drag_data["y"] = event.y
        
        def on_drag_motion(event):
            # Проверяем, что перетаскивание началось
            if self.drag_data["index"] is None:
                return
            
            # Определяем новую позицию
            new_index = self.files_listbox.nearest(event.y)
            old_index = self.drag_data["index"]
            
            if new_index != old_index and 0 <= new_index < len(self.file_list):
                # Перемещаем элемент
                item = self.file_list.pop(old_index)
                self.file_list.insert(new_index, item)
                
                # Обновляем listbox
                self.files_listbox.delete(old_index)
                self.files_listbox.insert(new_index, os.path.basename(item))
                self.files_listbox.selection_clear(0, tk.END)
                self.files_listbox.selection_set(new_index)
                
                # Обновляем сохраненный индекс
                self.drag_data["index"] = new_index
        
        def on_drag_end(event):
            # Сбрасываем данные перетаскивания
            self.drag_data["index"] = None
        
        # Привязываем события
        self.files_listbox.bind('<Button-1>', on_drag_start)
        self.files_listbox.bind('<B1-Motion>', on_drag_motion)
        self.files_listbox.bind('<ButtonRelease-1>', on_drag_end)
    
    def on_drop(self, event):
        """Обработка перетаскивания файлов из системы"""
        files = parse_drop_files(event.data)  # Используем глобальную функцию
        doc_type = self.doc_type.get()
        
        # Определяем позицию вставки по координатам курсора
        try:
            # Получаем координаты мыши относительно listbox
            widget_x = self.files_listbox.winfo_pointerx() - self.files_listbox.winfo_rootx()
            widget_y = self.files_listbox.winfo_pointery() - self.files_listbox.winfo_rooty()
            
            # Если список пустой, добавляем в начало
            if not self.file_list:
                drop_index = 0
            else:
                # Получаем индекс элемента под курсором
                drop_index = self.files_listbox.nearest(widget_y)
                
                # Проверяем, находится ли курсор над текстом файла или справа в пустой области
                try:
                    bbox = self.files_listbox.bbox(drop_index)
                    if bbox:
                        item_x = bbox[0]
                        item_y = bbox[1]
                        item_width = bbox[2]
                        item_height = bbox[3]
                        
                        # Получаем текст элемента для определения ширины текста
                        item_text = self.files_listbox.get(drop_index)
                        
                        # Примерная ширина одного символа (можно использовать font.measure для точности)
                        char_width = 8  # Примерная ширина символа
                        text_width = len(item_text) * char_width + 10  # +10 для отступов
                        
                        # Если курсор справа от текста (в пустой области), добавляем в конец
                        if widget_x > item_x + text_width:
                            drop_index = len(self.file_list)
                        # Если курсор над текстом, определяем где вставить (до или после)
                        else:
                            # Если курсор в нижней половине элемента, вставляем после
                            if widget_y > item_y + item_height / 2:
                                drop_index += 1
                            # Иначе вставляем перед элементом (drop_index остается как есть)
                    else:
                        # Если не удалось получить bbox, добавляем в конец
                        drop_index = len(self.file_list)
                except:
                    # Если ошибка с bbox, добавляем в конец
                    drop_index = len(self.file_list)
                
                # Ограничиваем индекс размером списка
                if drop_index > len(self.file_list):
                    drop_index = len(self.file_list)
        except Exception as e:
            # В случае ошибки добавляем в конец
            drop_index = len(self.file_list)
        
        added_count = 0
        invalid_count = 0
        
        for file_path in files:
            valid = False
            if doc_type in ['word', 'convert', 'convert_merge']:
                valid = file_path.lower().endswith('.docx')
            elif doc_type in ['pdf', 'number_separate', 'number_merge', 'split_pdf', 'rotate_pdf', 'extract_pdf']:
                valid = file_path.lower().endswith('.pdf')
            elif doc_type in ['image', 'image_merge']:
                valid_exts = ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.gif')
                valid = file_path.lower().endswith(valid_exts)
            elif doc_type == 'extract_to_excel':
                valid = file_path.lower().endswith(('.docx', '.pdf'))
            
            if not valid:
                invalid_count += 1
                continue
            
            if file_path not in self.file_list:
                # Вставляем в определенную позицию
                self.file_list.insert(drop_index, file_path)
                self.files_listbox.insert(drop_index, os.path.basename(file_path))
                drop_index += 1  # Увеличиваем индекс для следующего файла
                added_count += 1
                
                # Предзагрузка Word файлов
                if file_path.lower().endswith(('.docx', '.doc')):
                    word_preload_manager.preload(file_path)
        
        self.update_file_counter()
        
        if invalid_count > 0:
            if doc_type in ['word', 'convert', 'convert_merge']:
                file_type = "Word (.docx)"
            elif doc_type in ['pdf', 'number_separate', 'number_merge', 'split_pdf', 'rotate_pdf', 'extract_pdf']:
                file_type = "PDF (.pdf)"
            elif doc_type in ['image', 'image_merge']:
                file_type = "изображения (.jpg, .png, .bmp, .tiff, .gif)"
            elif doc_type == 'extract_to_excel':
                file_type = "Word (.docx) или PDF (.pdf)"
            else:
                file_type = "поддерживаемые"
            
            messagebox.showwarning(
                "Неподходящие файлы",
                f"Пропущено файлов: {invalid_count}\n\n"
                f"Принимаются только {file_type} файлы.",
                parent=self.window.window
            )
        
        if added_count > 0:
            self.log(f"Добавлено файлов перетаскиванием: {added_count}")
    
    def add_files(self):
        """Добавить файлы в список"""
        doc_type = self.doc_type.get()
        if doc_type in ["word", "convert", "convert_merge"]:
            filetypes = [("Word файлы", "*.docx"), ("Все файлы", "*.*")]
        elif doc_type in ["image", "image_merge"]:
            filetypes = [
                ("Изображения", "*.jpg *.jpeg *.png *.bmp *.tiff *.tif *.gif"),
                ("JPEG", "*.jpg *.jpeg"),
                ("PNG", "*.png"),
                ("BMP", "*.bmp"),
                ("TIFF", "*.tiff *.tif"),
                ("GIF", "*.gif"),
                ("Все файлы", "*.*")
            ]
        elif doc_type == "extract_to_excel":
            filetypes = [
                ("Документы", "*.docx *.pdf"),
                ("Word файлы", "*.docx"),
                ("PDF файлы", "*.pdf"),
                ("Все файлы", "*.*")
            ]
        else:
            filetypes = [("PDF файлы", "*.pdf"), ("Все файлы", "*.*")]
        
        files = filedialog.askopenfilenames(
            title="Выберите файлы",
            filetypes=filetypes
        )
        
        if not files:
            return
        
        # Временно отключаем listbox для быстрой вставки
        added_count = 0
        try:
            if len(files) > 10:  # Оптимизация только для большого количества файлов
                self.files_listbox.config(state=tk.DISABLED)
            
            for file in files:
                if file not in self.file_list:
                    self.file_list.append(file)
                    self.files_listbox.insert(tk.END, os.path.basename(file))
                    added_count += 1
                    
                    if file.lower().endswith(('.docx', '.doc')):
                        word_preload_manager.preload(file)
            
            if len(files) > 10:
                self.files_listbox.config(state=tk.NORMAL)
        except Exception as e:
            self.files_listbox.config(state=tk.NORMAL)
            raise e
        
        self.update_file_counter()
        if added_count > 0:
            self.log(f"Добавлено файлов: {added_count}")
    
    def insert_files_after(self):
        """Вставить файлы после выбранного элемента"""
        selection = self.files_listbox.curselection()
        if not selection:
            messagebox.showinfo(
                "Информация", 
                "Выберите элемент, после которого нужно вставить файлы.\n\nЕсли элемент не выбран, файлы будут добавлены в конец списка.",
                parent=self.window.window
            )
            self.add_files()
            return
        
        insert_index = selection[0] + 1
        
        doc_type = self.doc_type.get()
        if doc_type in ["word", "convert", "convert_merge"]:
            filetypes = [("Word файлы", "*.docx"), ("Все файлы", "*.*")]
        elif doc_type in ["image", "image_merge"]:
            filetypes = [
                ("Изображения", "*.jpg *.jpeg *.png *.bmp *.tiff *.tif *.gif"),
                ("JPEG", "*.jpg *.jpeg"),
                ("PNG", "*.png"),
                ("BMP", "*.bmp"),
                ("TIFF", "*.tiff *.tif"),
                ("GIF", "*.gif"),
                ("Все файлы", "*.*")
            ]
        elif doc_type == "extract_to_excel":
            filetypes = [
                ("Документы", "*.docx *.pdf"),
                ("Word файлы", "*.docx"),
                ("PDF файлы", "*.pdf"),
                ("Все файлы", "*.*")
            ]
        else:
            filetypes = [("PDF файлы", "*.pdf"), ("Все файлы", "*.*")]
        
        files = filedialog.askopenfilenames(
            title="Выберите файлы для вставки",
            filetypes=filetypes
        )
        
        if not files:
            return
        
        added_count = 0
        for file in files:
            if file not in self.file_list:
                self.file_list.insert(insert_index, file)
                self.files_listbox.insert(insert_index, os.path.basename(file))
                insert_index += 1
                added_count += 1
                
                if file.lower().endswith(('.docx', '.doc')):
                    word_preload_manager.preload(file)
        
        self.update_file_counter()
        if added_count > 0:
            self.log(f"Вставлено файлов: {added_count}")
    
    def remove_file(self):
        """Удалить выбранный файл"""
        selection = self.files_listbox.curselection()
        if selection:
            index = selection[0]
            self.files_listbox.delete(index)
            del self.file_list[index]
            self.update_file_counter()
    
    def clear_all_files(self):
        """Очистить все файлы из списка"""
        if not self.file_list:
            return
        
        count = len(self.file_list)
        result = messagebox.askyesno(
            "Подтверждение",
            f"Вы уверены, что хотите удалить все файлы из списка?\n\nВсего файлов: {count}",
            parent=self.window.window
        )
        
        if result:
            try:
                # Сохраняем размер перед очисткой
                listbox_size = self.files_listbox.size()
                
                # Очищаем данные  
                self.file_list.clear()
                
                # Явно удаляем все элементы из listbox
                if listbox_size > 0:
                    self.files_listbox.delete(0, listbox_size - 1)
                
                # Обновляем счетчик
                self.update_file_counter()
                
                self.log(f"Удалено {count} файлов из списка")
            except Exception as e:
                self.log(f"Ошибка при очистке: {str(e)}")
                messagebox.showerror(
                    "Ошибка", 
                    f"Не удалось очистить список файлов:\n{str(e)}", 
                    parent=self.window.window
                )
    
    def clear_document_cache(self):
        """Очистить кэш созданных документов"""
        try:
            cache = DocumentCache()
            entries_count = len(cache.cache)
            
            if entries_count == 0:
                messagebox.showinfo(
                    "Информация",
                    "Кэш уже пуст",
                    parent=self.window.window
                )
                return
            
            result = messagebox.askyesno(
                "Подтверждение",
                f"Очистить кэш созданных документов?\n\nВсего записей: {entries_count}\n\n"
                f"После очистки все документы будут создаваться заново при следующей генерации.",
                parent=self.window.window
            )
            
            if result:
                cache.clear_all()
                self.log(f"✓ Кэш очищен ({entries_count} записей)")
                messagebox.showinfo(
                    "Успех",
                    f"Кэш успешно очищен!\n\nУдалено записей: {entries_count}",
                    parent=self.window.window
                )
        except Exception as e:
            self.log(f"❌ Ошибка при очистке кэша: {str(e)}")
            messagebox.showerror(
                "Ошибка",
                f"Не удалось очистить кэш:\n{str(e)}", 
                    parent=self.window.window
                )
    
    def move_up(self):
        """Переместить файл вверх"""
        selection = self.files_listbox.curselection()
        if not selection or selection[0] == 0:
            return
            
        index = selection[0]
        
        # Меняем местами в списке
        self.file_list[index], self.file_list[index-1] = self.file_list[index-1], self.file_list[index]
        
        # Перерисовываем только затронутые элементы
        self.files_listbox.delete(index-1, index)
        self.files_listbox.insert(index-1, os.path.basename(self.file_list[index-1]))
        self.files_listbox.insert(index, os.path.basename(self.file_list[index]))
        
        # Выделяем перемещённый элемент
        self.files_listbox.selection_clear(0, tk.END)
        self.files_listbox.selection_set(index-1)
        self.files_listbox.see(index-1)
    
    def move_down(self):
        """Переместить файл вниз"""
        selection = self.files_listbox.curselection()
        if not selection or selection[0] >= len(self.file_list) - 1:
            return
            
        index = selection[0]
        
        # Меняем местами в списке
        self.file_list[index], self.file_list[index+1] = self.file_list[index+1], self.file_list[index]
        
        # Перерисовываем только затронутые элементы
        self.files_listbox.delete(index, index+1)
        self.files_listbox.insert(index, os.path.basename(self.file_list[index]))
        self.files_listbox.insert(index+1, os.path.basename(self.file_list[index+1]))
        
        # Выделяем перемещённый элемент
        self.files_listbox.selection_clear(0, tk.END)
        self.files_listbox.selection_set(index+1)
        self.files_listbox.see(index+1)
    
    def refresh_listbox(self):
        """Обновить отображение списка файлов"""
        # Сохраняем размер до очистки
        listbox_size = self.files_listbox.size()
        
        # Оптимизация для больших списков
        if len(self.file_list) > 50:
            self.files_listbox.config(state=tk.DISABLED)
        
        # Очищаем listbox
        if listbox_size > 0:
            self.files_listbox.delete(0, listbox_size - 1)
        
        # Заполняем заново
        for i, file in enumerate(self.file_list):
            self.files_listbox.insert(i, os.path.basename(file))
        
        if len(self.file_list) > 50:
            self.files_listbox.config(state=tk.NORMAL)
        
        self.update_file_counter()
    
    def update_file_counter(self):
        """Обновить счетчик файлов"""
        count = len(self.file_list)
        self.file_counter_label.config(text=f"Файлов: {count}")
    
    def preview_selected_file(self):
        """Предварительный просмотр выбранного файла"""
        selection = self.files_listbox.curselection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Сначала выберите файл из списка!", parent=self.window.window)
            return
        
        index = selection[0]
        file_path = self.file_list[index]
        
        if not os.path.exists(file_path):
            messagebox.showerror("Ошибка", f"Файл не найден:\n{file_path}", parent=self.window.window)
            return
        
        try:
            PreviewWindow(self.window.window, file_path, f"Просмотр: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{str(e)}", parent=self.window.window)
    
    def copy_log_text(self):
        """Копирование выделенного текста в буфер обмена"""
        try:
            selected_text = self.log_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.window.window.clipboard_clear()
            self.window.window.clipboard_append(selected_text)
        except tk.TclError:
            pass
    
    def select_all_log(self):
        """Выделить весь текст в логе"""
        self.log_text.config(state=tk.NORMAL)
        self.log_text.tag_add(tk.SEL, "1.0", tk.END)
        self.log_text.mark_set(tk.INSERT, "1.0")
        self.log_text.see(tk.INSERT)
        self.log_text.config(state=tk.DISABLED)
    
    def log(self, message):
        """Добавить сообщение в лог"""
        def _log():
            self.log_text.config(state=tk.NORMAL)
            self.log_text.insert(tk.END, message + "\n")
            self.log_text.see(tk.END)
            self.log_text.config(state=tk.DISABLED)
        
        try:
            self.window.window.after(0, _log)
        except:
            _log()
    
    def merge_documents(self):
        """Запуск или остановка объединения/конвертации документов"""
        # Если обработка идет - останавливаем
        if self.is_processing:
            self.should_stop = True
            self.merge_btn.configure(text="⏹ Остановка...")
            self.log("\n⚠️ Запрошена остановка обработки...")
            return
        
        if not self.file_list:
            messagebox.showwarning("Предупреждение", "Добавьте файлы для обработки!", parent=self.window.window)
            return
        
        doc_type = self.doc_type.get()
        
        # Проверка количества файлов
        if doc_type in ["convert", "image", "number_separate", "extract_to_excel"]:
            pass  # Можно обрабатывать любое количество файлов
        elif doc_type in ["split_pdf", "rotate_pdf", "extract_pdf"]:
            if len(self.file_list) != 1:
                messagebox.showwarning("Предупреждение", "Выберите ровно один PDF файл для обработки!", parent=self.window.window)
                return
        elif doc_type in ["convert_merge", "image_merge", "number_merge"] and len(self.file_list) < 2:
            messagebox.showwarning("Предупреждение", "Для объединения в единый PDF добавьте минимум 2 файла!", parent=self.window.window)
            return
        elif doc_type in ["word", "pdf"] and len(self.file_list) < 2:
            messagebox.showwarning("Предупреждение", "Добавьте минимум 2 файла для объединения!", parent=self.window.window)
            return
        
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state=tk.DISABLED)
        
        # Определяем выходной путь
        if doc_type in ["convert", "image", "number_separate"]:
            output_folder = filedialog.askdirectory(
                title="Выберите папку для сохранения PDF файлов"
            )
            if not output_folder:
                return
            output_path = output_folder
        elif doc_type == "split_pdf":
            output_folder = filedialog.askdirectory(
                title="Выберите папку для сохранения разделенных PDF файлов"
            )
            if not output_folder:
                return
            output_path = output_folder
        elif doc_type == "extract_to_excel":
            output_file = filedialog.asksaveasfilename(
                title="Сохранить данные в Excel",
                defaultextension=".xlsx",
                filetypes=[("Excel файлы", "*.xlsx"), ("Все файлы", "*.*")]
            )
            if not output_file:
                return
            output_path = output_file
        elif doc_type in ["rotate_pdf", "extract_pdf"]:
            output_file = filedialog.asksaveasfilename(
                title="Сохранить обработанный PDF файл",
                defaultextension=".pdf",
                filetypes=[("PDF файлы", "*.pdf"), ("Все файлы", "*.*")]
            )
            if not output_file:
                return
            output_path = output_file
        elif doc_type in ["convert_merge", "image_merge", "number_merge"]:
            output_file = filedialog.asksaveasfilename(
                title="Сохранить объединенный PDF файл",
                defaultextension=".pdf",
                filetypes=[("PDF файлы", "*.pdf"), ("Все файлы", "*.*")]
            )
            if not output_file:
                return
            output_path = output_file
        else:
            default_ext = ".docx" if doc_type == "word" else ".pdf"
            filetypes = [("Word файлы", "*.docx")] if doc_type == "word" else [("PDF файлы", "*.pdf")]
            output_file = filedialog.asksaveasfilename(
                title="Сохранить объединенный документ",
                defaultextension=default_ext,
                filetypes=filetypes + [("Все файлы", "*.*")]
            )
            if not output_file:
                return
            output_path = output_file
        
        # Блокируем кнопку
        self.is_processing = True
        self.should_stop = False
        self.merge_btn.configure(text="⏹ Остановить")
        
        # Получаем настройку OCR
        use_ocr = self.use_ocr.get()
        
        # Запускаем обработку в отдельном потоке
        thread = threading.Thread(target=self.process_in_thread, args=(doc_type, output_path, use_ocr))
        thread.daemon = True
        thread.start()
    
    def show_message_safe(self, msg_type, title, message):
        """Безопасный вызов messagebox из фонового потока через главный поток UI"""
        def show_in_main_thread():
            try:
                if msg_type == "info":
                    messagebox.showinfo(title, message, parent=self.window.window)
                elif msg_type == "warning":
                    messagebox.showwarning(title, message, parent=self.window.window)
                elif msg_type == "error":
                    messagebox.showerror(title, message, parent=self.window.window)
            except:
                pass  # Если окно уже закрыто
        
        # Планируем выполнение в главном потоке UI
        try:
            self.window.window.after(0, show_in_main_thread)
        except:
            pass  # Если окно уже закрыто
    
    def process_in_thread(self, doc_type, output_path, use_ocr=True):
        """Обработка документов в отдельном потоке"""
        try:
            # Получаем параметры качества и нумерации из UI
            try:
                ocr_resolution = float(self.ocr_resolution.get())
                ocr_resolution = max(1.0, min(3.0, ocr_resolution))  # Ограничиваем 1.0-3.0
            except:
                ocr_resolution = 1.5  # По умолчанию
            
            try:
                max_img_size_str = self.max_image_size.get().strip()
                if max_img_size_str in ["", "0", "None"]:
                    max_img_size = None  # Без ограничений
                else:
                    max_img_size = int(max_img_size_str)
            except:
                max_img_size = 4000  # По умолчанию
            
            fit_mode = self.fit_mode.get() if self.fit_mode.get() in ['центр', 'заполнить', 'вписать'] else 'центр'
            
            # Параметры нумерации (только если включена нумерация)
            # Для функций прямой нумерации (number_separate, number_merge) - всегда используем
            # Для остальных - только если включен чекбокс use_numbering
            if doc_type in ['number_separate', 'number_merge'] or self.use_numbering.get():
                numbering_line1 = self.numbering_line1.get().strip() or None
                numbering_line2 = self.numbering_line2.get().strip() or None
                numbering_line3 = self.numbering_line3.get().strip() or None
                numbering_position = self.numbering_position.get()
                numbering_border = self.numbering_border.get()
                numbering_increment_mode = self.numbering_increment_mode.get()
            else:
                # Нумерация отключена - все параметры None
                numbering_line1 = None
                numbering_line2 = None
                numbering_line3 = None
                numbering_position = None
                numbering_border = False
                numbering_increment_mode = None

            
            self.log("═" * 60)
            self.log("Начало обработки...")
            self.log(f"Режим: {self.get_mode_name(doc_type)}")
            self.log(f"Файлов в очереди: {len(self.file_list)}")
            self.log(f"Применение OCR: {'Да' if use_ocr else 'Нет (быстрый режим)'}")
            if doc_type in ['image', 'image_merge', 'pdf']:
                self.log(f"Разрешение OCR: {ocr_resolution}x")
                self.log(f"Макс. размер изображения: {max_img_size if max_img_size else 'без ограничений'}")
            if doc_type in ['image', 'image_merge']:
                self.log(f"Размещение на листе: {fit_mode}")
            
            # Логирование нумерации для всех режимов кроме объединения Word
            if doc_type != 'word':
                if numbering_line1 or numbering_line2 or numbering_line3:
                    self.log("Нумерация:")
                    if numbering_line1:
                        self.log(f"  Строка 1: {numbering_line1}")
                    if numbering_line2:
                        self.log(f"  Строка 2: {numbering_line2} (с автоинкрементом)")
                    if numbering_line3:
                        self.log(f"  Строка 3: {numbering_line3}")
                    self.log(f"  Позиция: {numbering_position}")
                    self.log(f"  Рамка: {'Да' if numbering_border else 'Нет'}")
                    self.log(f"  Инкремент: {numbering_increment_mode}")
                else:
                    self.log("Нумерация: не задана (все строки пустые)")
            self.log("═" * 60)
            
            # Проверка остановки перед началом
            if self.should_stop:
                self.log("\n⚠️ Обработка отменена до начала")
                return
            
            # Инициализация прогресса
            self.start_time = time.time()
            self.total_items = len(self.file_list)
            self.processed_items = 0
            self.update_progress(0, self.total_items, "Подготовка...")
            
            if doc_type == "convert":
                self.log(f"Папка для сохранения: {output_path}")
                converted_files = GenerationDocApp.convert_word_to_pdf(
                    self.file_list, output_path, self.log,
                    numbering_line1, numbering_line2, numbering_line3,
                    numbering_position, numbering_border, numbering_increment_mode,
                    progress_callback=self.update_progress
                )
                
                self.log("═" * 60)
                self.log(f"✅ Успешно конвертировано файлов: {len(converted_files)}")
                for f in converted_files:
                    self.log(f"  ✓ {os.path.basename(f)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех", 
                    f"Успешно конвертировано файлов: {len(converted_files)}\n\n"
                    f"Файлы сохранены в:\n{output_path}"
                )
            
            elif doc_type == "image":
                self.log(f"Папка для сохранения: {output_path}")
                converted_files = GenerationDocApp.convert_images_to_pdf(
                    self.file_list, output_path, self.log, 
                    use_ocr=use_ocr,
                    max_image_size=max_img_size,
                    fit_mode=fit_mode,
                    numbering_line1=numbering_line1,
                    numbering_line2=numbering_line2,
                    numbering_line3=numbering_line3,
                    numbering_position=numbering_position,
                    numbering_border=numbering_border,
                    numbering_increment_mode=numbering_increment_mode,
                    progress_callback=self.update_progress
                )
                
                self.log("═" * 60)
                self.log(f"✅ Успешно конвертировано файлов: {len(converted_files)}")
                for f in converted_files:
                    self.log(f"  ✓ {os.path.basename(f)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех", 
                    f"Успешно конвертировано файлов: {len(converted_files)}\n\n"
                    f"Файлы сохранены в:\n{output_path}"
                )
            
            elif doc_type == "convert_merge":
                self.log(f"Файл для сохранения: {output_path}")
                GenerationDocApp.convert_and_merge_word_to_pdf(
                    self.file_list, output_path, self.log,
                    numbering_line1, numbering_line2, numbering_line3,
                    numbering_position, numbering_border, numbering_increment_mode
                )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Файл сохранен: {os.path.basename(output_path)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех", 
                    f"Word документы успешно конвертированы и объединены!\n\n"
                    f"Обработано файлов: {len(self.file_list)}\n\n"
                    f"Файл сохранен:\n{output_path}"
                )
            
            elif doc_type == "image_merge":
                self.log(f"Файл для сохранения: {output_path}")
                GenerationDocApp.convert_and_merge_images_to_pdf(
                    self.file_list, output_path, self.log, 
                    use_ocr=use_ocr,
                    max_image_size=max_img_size,
                    fit_mode=fit_mode,
                    numbering_line1=numbering_line1,
                    numbering_line2=numbering_line2,
                    numbering_line3=numbering_line3,
                    numbering_position=numbering_position,
                    numbering_border=numbering_border,
                    numbering_increment_mode=numbering_increment_mode
                )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Файл сохранен: {os.path.basename(output_path)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех", 
                    f"Изображения успешно конвертированы и объединены!\n\n"
                    f"Обработано файлов: {len(self.file_list)}\n\n"
                    f"Файл сохранен:\n{output_path}"
                )
            
            elif doc_type == "split_pdf":
                # Разделение PDF файла
                self.log(f"Папка для сохранения: {output_path}")
                input_file = self.file_list[0]
                
                split_mode = self.split_mode.get()
                ranges_text = self.split_ranges.get() if split_mode == "ranges" else "all_pages"
                
                self.log("Разделение PDF файла...")
                created_files = GenerationDocApp.split_pdf_file(
                    input_file, output_path, 
                    ranges_text=ranges_text,
                    log_callback=self.log
                )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Создано файлов: {len(created_files)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех",
                    f"PDF файл успешно разделен!\n\n"
                    f"Создано файлов: {len(created_files)}\n\n"
                    f"Файлы сохранены в:\n{output_path}"
                )
            
            elif doc_type == "rotate_pdf":
                # Поворот страниц PDF
                self.log(f"Файл для сохранения: {output_path}")
                input_file = self.file_list[0]
                
                angle = self.rotate_angle.get()
                pages = self.rotate_pages.get()
                
                self.log("Поворот страниц PDF...")
                GenerationDocApp.rotate_pdf_pages(
                    input_file, output_path,
                    angle=angle,
                    pages_range=pages,
                    log_callback=self.log
                )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Файл сохранен: {os.path.basename(output_path)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех",
                    f"Страницы PDF успешно повернуты!\n\n"
                    f"Угол поворота: {angle}°\n\n"
                    f"Файл сохранен:\n{output_path}"
                )
            
            elif doc_type == "extract_pdf":
                # Извлечение страниц PDF
                self.log(f"Файл для сохранения: {output_path}")
                input_file = self.file_list[0]
                
                pages = self.extract_pages.get()
                
                if not pages or pages.strip() == "":
                    self.show_message_safe(
                        "warning",
                        "Предупреждение",
                        "Укажите страницы для извлечения!\n\nПример: 1-5 или 1,3,5 или 1-5,10,15-20"
                    )
                    return
                
                self.log("Извлечение страниц PDF...")
                GenerationDocApp.extract_pdf_pages(
                    input_file, output_path,
                    pages_range=pages,
                    log_callback=self.log
                )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Файл сохранен: {os.path.basename(output_path)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех",
                    f"Страницы PDF успешно извлечены!\n\n"
                    f"Диапазон: {pages}\n\n"
                    f"Файл сохранен:\n{output_path}"
                )
            
            elif doc_type == "extract_to_excel":
                # Извлечение данных в Excel
                self.log(f"Файл Excel для сохранения: {output_path}")
                
                extract_method = self.extract_method.get()
                extract_regex = self.extract_regex.get() if extract_method == "regex" else None
                
                self.log(f"Метод извлечения: {extract_method}")
                if extract_regex:
                    self.log(f"Регулярное выражение: {extract_regex}")
                
                GenerationDocApp.extract_data_to_excel(
                    self.file_list, output_path,
                    method=extract_method,
                    regex_pattern=extract_regex,
                    log_callback=self.log,
                    progress_callback=self.update_progress
                )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Данные извлечены")
                self.log(f"  Обработано файлов: {len(self.file_list)}")
                self.log(f"  Сохранено в: {os.path.basename(output_path)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех",
                    f"Данные успешно извлечены!\n\n"
                    f"Обработано документов: {len(self.file_list)}\n"
                    f"Метод: {extract_method}\n\n"
                    f"Файл сохранен:\n{output_path}"
                )
            
            elif doc_type == "number_separate":
                # Нумерация PDF раздельно
                self.log(f"Папка для сохранения: {output_path}")
                self.log("Нумерация PDF файлов (раздельно)...")
                
                numbered_files = []
                skipped_files = []
                current_line2 = numbering_line2
                
                for idx, pdf_file in enumerate(self.file_list, 1):
                    if self.should_stop:
                        break
                    
                    self.log(f"\n[{idx}/{len(self.file_list)}] {os.path.basename(pdf_file)}")
                    
                    # Проверяем существование исходного файла
                    if not os.path.exists(pdf_file):
                        self.log(f"  ⚠ ПРОПУЩЕН: файл не найден по пути {pdf_file}")
                        skipped_files.append(pdf_file)
                        continue
                    
                    # Формируем имя выходного файла
                    base_name = os.path.splitext(os.path.basename(pdf_file))[0]
                    output_file = os.path.join(output_path, f"{base_name}_numbered.pdf")
                    
                    try:
                        # Добавляем нумерацию к PDF
                        GenerationDocApp.add_numbering_to_existing_pdf(
                            pdf_file, output_file,
                            numbering_line1, current_line2, numbering_line3,
                            numbering_position, numbering_border,
                            numbering_increment_mode, self.log
                        )
                        
                        numbered_files.append(output_file)
                        self.log(f"  ✓ Сохранен: {os.path.basename(output_file)}")
                        
                        # Инкрементируем для следующего документа если режим per_document
                        if numbering_increment_mode in ['per_document', 'per_document_first_page'] and current_line2:
                            current_line2 = GenerationDocApp.increment_line2(current_line2)
                    except Exception as e:
                        self.log(f"  ✗ ОШИБКА при обработке: {str(e)}")
                        skipped_files.append(pdf_file)
                        continue
                
                self.log("═" * 60)
                self.log(f"✅ Успешно пронумеровано файлов: {len(numbered_files)}")
                for f in numbered_files:
                    self.log(f"  ✓ {os.path.basename(f)}")
                
                if skipped_files:
                    self.log(f"\n⚠ Пропущено файлов: {len(skipped_files)}")
                    for f in skipped_files:
                        self.log(f"  ⚠ {os.path.basename(f)}")
                
                self.log("═" * 60)
                
                result_message = f"Успешно пронумеровано файлов: {len(numbered_files)}\n\n"
                if skipped_files:
                    result_message += f"Пропущено файлов: {len(skipped_files)}\n(файлы не найдены или возникли ошибки)\n\n"
                result_message += f"Файлы сохранены в:\n{output_path}"
                
                self.show_message_safe(
                    "info" if not skipped_files else "warning",
                    "Результат нумерации", 
                    result_message
                )
            
            elif doc_type == "number_merge":
                # Нумерация PDF с объединением
                self.log(f"Файл для сохранения: {output_path}")
                self.log("Нумерация и объединение PDF документов...")
                
                GenerationDocApp.merge_pdf_documents(
                    self.file_list, output_path, self.log, use_ocr=use_ocr,
                    numbering_line1=numbering_line1, numbering_line2=numbering_line2, numbering_line3=numbering_line3,
                    numbering_position=numbering_position, numbering_border=numbering_border,
                    numbering_increment_mode=numbering_increment_mode
                )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Файл сохранен: {os.path.basename(output_path)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех", 
                    f"PDF документы успешно пронумерованы и объединены!\n\n"
                    f"Обработано файлов: {len(self.file_list)}\n\n"
                    f"Файл сохранен:\n{output_path}"
                )
            
            else:
                # word или pdf - объединение
                self.log(f"Файл для сохранения: {output_path}")
                if doc_type == "word":
                    self.log("Объединение Word документов...")
                    GenerationDocApp.merge_word_documents(self.file_list, output_path, self.log)
                elif doc_type == "pdf":
                    self.log("Объединение PDF документов...")
                    GenerationDocApp.merge_pdf_documents(
                        self.file_list, output_path, self.log, use_ocr=use_ocr,
                        numbering_line1=numbering_line1, numbering_line2=numbering_line2, numbering_line3=numbering_line3,
                        numbering_position=numbering_position, numbering_border=numbering_border,
                        numbering_increment_mode=numbering_increment_mode
                    )
                
                self.log("═" * 60)
                self.log(f"✅ ГОТОВО! Файл сохранен: {os.path.basename(output_path)}")
                self.log("═" * 60)
                
                self.show_message_safe(
                    "info",
                    "Успех", 
                    f"Документы успешно объединены!\n\nФайл сохранен:\n{output_path}"
                )
        
        except Warning as w:
            self.log(f"⚠️ Частичный успех: {str(w)}")
            self.show_message_safe("warning", "Частичный успех", str(w))
        except Exception as e:
            if not self.should_stop:
                self.log(f"❌ ОШИБКА: {str(e)}")
                self.show_message_safe("error", "Ошибка", f"Ошибка при обработке документов:\n{str(e)}")
        finally:
            if self.should_stop:
                self.log("\n⏹ Обработка остановлена пользователем")
            
            # Сброс прогресса
            self.reset_progress()
            
            self.is_processing = False
            self.should_stop = False
            try:
                if self.window.window.winfo_exists():
                    self.merge_btn.configure(text="▶ Начать")
            except:
                pass
    
    def update_progress(self, current, total, message=""):
        """
        Обновление прогресс-бара и ETA
        
        Args:
            current: текущее количество обработанных элементов
            total: общее количество элементов
            message: дополнительное сообщение для отображения
        """
        def update_ui():
            try:
                if not self.window.window.winfo_exists():
                    return
                
                # Показываем прогресс-контейнер если он скрыт
                if not self.progress_container.winfo_viewable():
                    self.progress_container.pack(fill=tk.X, pady=(0, 12), before=self.log_text.master)
                
                # Обновляем прогресс-бар
                if total > 0:
                    progress_percent = (current / total) * 100
                    self.progress_bar["value"] = progress_percent
                    self.progress_bar["maximum"] = 100
                
                # Обновляем текст прогресса
                progress_text = f"Обработано: {current} из {total}"
                if message:
                    progress_text += f" - {message}"
                self.progress_label.config(text=progress_text)
                
                # Рассчитываем и показываем ETA
                if self.start_time and current > 0 and current < total:
                    elapsed_time = time.time() - self.start_time
                    avg_time_per_item = elapsed_time / current
                    remaining_items = total - current
                    eta_seconds = avg_time_per_item * remaining_items
                    
                    # Форматируем ETA
                    if eta_seconds < 60:
                        eta_str = f"Осталось: ~{int(eta_seconds)} сек."
                    elif eta_seconds < 3600:
                        minutes = int(eta_seconds / 60)
                        seconds = int(eta_seconds % 60)
                        eta_str = f"Осталось: ~{minutes} мин. {seconds} сек."
                    else:
                        hours = int(eta_seconds / 3600)
                        minutes = int((eta_seconds % 3600) / 60)
                        eta_str = f"Осталось: ~{hours} ч. {minutes} мин."
                    
                    self.eta_label.config(text=eta_str)
                elif current >= total:
                    self.eta_label.config(text="Завершено!")
                else:
                    self.eta_label.config(text="Расчет времени...")
                
            except Exception as e:
                pass  # Игнорируем ошибки обновления UI
        
        # Планируем обновление в главном потоке
        try:
            self.window.window.after(0, update_ui)
        except:
            pass
    
    def reset_progress(self):
        """Сброс прогресса"""
        def reset_ui():
            try:
                if self.window.window.winfo_exists():
                    self.progress_container.pack_forget()
                    self.progress_bar["value"] = 0
                    self.progress_label.config(text="")
                    self.eta_label.config(text="")
            except:
                pass
        
        try:
            self.window.window.after(0, reset_ui)
        except:
            pass
        
        self.start_time = None
        self.total_items = 0
        self.processed_items = 0
    
    def save_stamp_preset(self):
        """Сохранить текущие настройки штампа как пресет"""
        # Проверяем, есть ли что сохранять
        line1 = self.numbering_line1.get().strip()
        line2 = self.numbering_line2.get().strip()
        line3 = self.numbering_line3.get().strip()
        
        if not line1 and not line2 and not line3:
            messagebox.showwarning(
                "Пустой пресет",
                "Заполните хотя бы одну строку штампа перед сохранением пресета.",
                parent=self.window.window
            )
            return
        
        # Запрашиваем имя пресета через SimpleInputDialog
        dialog = SimpleInputDialog(
            self.window.window,
            "Сохранить пресет штампа",
            "Введите название пресета:"
        )
        self.window.window.wait_window(dialog.top)
        preset_name = dialog.result
        
        if not preset_name:
            return
        
        # Создаем папку для пресетов штампов если её нет
        presets_dir = os.path.join(os.path.dirname(__file__), "presets", "stamps")
        os.makedirs(presets_dir, exist_ok=True)
        
        # Формируем данные пресета
        preset_data = {
            "line1": line1,
            "line2": line2,
            "line3": line3,
            "position": self.numbering_position.get(),
            "border": self.numbering_border.get(),
            "increment_mode": self.numbering_increment_mode.get()
        }
        
        # Сохраняем в JSON
        preset_file = os.path.join(presets_dir, f"{preset_name}.json")
        try:
            with open(preset_file, 'w', encoding='utf-8') as f:
                json.dump(preset_data, f, ensure_ascii=False, indent=2)
            
            messagebox.showinfo(
                "Успех",
                f"Пресет '{preset_name}' успешно сохранен!",
                parent=self.window.window
            )
            self.log(f"✓ Пресет штампа '{preset_name}' сохранен")
        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Не удалось сохранить пресет:\n{str(e)}",
                parent=self.window.window
            )
    
    def load_stamp_preset(self):
        """Загрузить сохраненный пресет штампа"""
        presets_dir = os.path.join(os.path.dirname(__file__), "presets", "stamps")
        os.makedirs(presets_dir, exist_ok=True)
        
        # Используем PresetSelectionDialog для выбора пресета
        dialog = PresetSelectionDialog(self.window.window, presets_dir)
        self.window.window.wait_window(dialog.top)
        
        preset_path = dialog.result
        if not preset_path:
            return
        
        # Загружаем выбранный пресет
        try:
            with open(preset_path, 'r', encoding='utf-8') as f:
                preset_data = json.load(f)
            
            # Применяем настройки
            self.numbering_line1.set(preset_data.get("line1", ""))
            self.numbering_line2.set(preset_data.get("line2", ""))
            self.numbering_line3.set(preset_data.get("line3", ""))
            self.numbering_position.set(preset_data.get("position", "правый-нижний"))
            self.numbering_border.set(preset_data.get("border", True))
            self.numbering_increment_mode.set(preset_data.get("increment_mode", "per_document"))
            
            preset_name = os.path.basename(preset_path).replace('.json', '')
            messagebox.showinfo(
                "Успех",
                f"Пресет '{preset_name}' успешно загружен!",
                parent=self.window.window
            )
            self.log(f"✓ Пресет штампа '{preset_name}' загружен")
        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Не удалось загрузить пресет:\n{str(e)}",
                parent=self.window.window
            )
    

    
    def get_mode_name(self, doc_type):
        """Получить название режима"""
        modes = {
            "word": "Объединение Word документов",
            "pdf": "Объединение PDF документов",
            "convert": "Конвертация Word → PDF (раздельно)",
            "convert_merge": "Конвертация Word → единый PDF",
            "image": "Конвертация изображений → PDF (раздельно)",
            "image_merge": "Конвертация изображений → единый PDF",
            "number_separate": "Нумерация PDF (раздельно)",
            "number_merge": "Нумерация PDF (с объединением)",
            "split_pdf": "Разделение PDF файла",
            "rotate_pdf": "Поворот страниц PDF",
            "extract_pdf": "Извлечение страниц из PDF",
            "extract_to_excel": "Извлечение данных в Excel"
        }
        return modes.get(doc_type, "Неизвестный режим")

# ── РУССКИЕ НАЗВАНИЯ МЕСЯЦЕВ ─────────────────────────────────────────
RUSSIAN_MONTHS = [
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
]

RUSSIAN_DAYS = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]

# ── ПАДЕЖИ ──────────────────────────────────────────────────────────
RUSSIAN_CASES = {
    "nomn": "Именительный (кто? что?)",
    "gent": "Родительный (кого? чего?)",
    "datv": "Дательный (кому? чему?)",
    "accs": "Винительный (кого? что?)",
    "ablt": "Творительный (кем? чем?)",
    "loct": "Предложный (о ком? о чём?)"
}

# ── ПОЛЬЗОВАТЕЛЬСКИЕ СПИСКИ (ПУСТЫЕ ПО УМОЛЧАНИЮ) ───────────────────
DEFAULT_CUSTOM_LISTS = {}

# ── ПЛЕЙСХОЛДЕРЫ (ПУСТЫЕ ПО УМОЛЧАНИЮ) ──────────────────────────────
DEFAULT_PLACEHOLDERS = []

# ── ФУНКЦИИ ДЛЯ ПАРАЛЛЕЛЬНОЙ ОБРАБОТКИ ──────────────────────────────

def create_placeholder_pattern(placeholder):
    """
    Создаёт regex паттерн для поиска плейсхолдера.
    Использует word boundary (\\b) только если плейсхолдер начинается/заканчивается 
    буквенно-цифровым символом, иначе ищет точное совпадение.
    
    Примеры:
    - "дата" -> r'\\bдата\\b' (обычное слово)
    - "{дата}" -> r'{дата}' (спецсимволы в начале/конце)
    - "[значение]" -> r'\\[значение\\]' (экранированные скобки)
    """
    escaped = re.escape(placeholder)
    # Проверяем первый и последний символ
    starts_with_word = placeholder and placeholder[0].isalnum()
    ends_with_word = placeholder and placeholder[-1].isalnum()
    
    # Добавляем \b только там, где это имеет смысл
    prefix = r'\b' if starts_with_word else ''
    suffix = r'\b' if ends_with_word else ''
    
    return prefix + escaped + suffix

def _normalize_paragraph_runs(paragraph):
    """Объединяет смежные runs с одинаковым форматированием.
    
    Word часто разбивает текст на множество runs даже при одинаковом форматировании,
    что мешает поиску плейсхолдеров. Эта функция склеивает такие runs.
    """
    if len(paragraph.runs) <= 1:
        return
    
    i = 0
    while i < len(paragraph.runs) - 1:
        current = paragraph.runs[i]
        next_run = paragraph.runs[i + 1]
        
        # Проверяем одинаковое ли форматирование
        if (current.bold == next_run.bold and
            current.italic == next_run.italic and
            current.underline == next_run.underline and
            current.font.size == next_run.font.size and
            current.font.name == next_run.font.name and
            current.font.color.rgb == next_run.font.color.rgb):
            
            # Объединяем runs
            current.text = current.text + next_run.text
            
            # Удаляем следующий run
            p = next_run._element.getparent()
            p.remove(next_run._element)
        else:
            i += 1

def _replace_placeholders_in_paragraph(paragraph, replacements):
    """Вспомогательная функция замены плейсхолдеров (для использования в процессах)
    
    Сначала нормализует параграф (объединяет runs с одинаковым форматированием),
    затем делает замену. Это решает проблему разбитых плейсхолдеров И сохраняет форматирование.
    """
    from docx.oxml.ns import qn
    import re
    
    # Проверяем есть ли хоть один плейсхолдер в полном тексте
    full_text = paragraph.text
    if not any(ph in full_text for ph in replacements):
        return
    
    # Нормализуем runs (объединяем смежные с одинаковым форматированием)
    _normalize_paragraph_runs(paragraph)
    
    # Теперь делаем замену в каждом run
    for run in paragraph.runs:
        # Проверяем есть ли в run встроенные объекты (картинки, фигуры)
        has_objects = False
        if hasattr(run._element, 'xpath'):
            drawings = run._element.xpath('.//w:drawing')
            pictures = run._element.xpath('.//w:pict')
            has_objects = len(drawings) > 0 or len(pictures) > 0
        
        if has_objects:
            # Если есть объекты, работаем на уровне XML элементов текста
            for text_elem in run._element.findall(qn('w:t')):
                if text_elem.text:
                    modified_text = text_elem.text
                    for placeholder, replacement in replacements.items():
                        pattern = re.escape(placeholder)
                        modified_text = re.sub(pattern, str(replacement), modified_text)
                    text_elem.text = modified_text
        else:
            # Обычная замена для run без объектов
            text = run.text
            for placeholder, replacement in replacements.items():
                pattern = re.escape(placeholder)
                text = re.sub(pattern, str(replacement), text)
            run.text = text

def _convert_single_image(args):
    """
    Конвертация одного изображения в PDF (функция для параллельного выполнения).
    
    Args:
        args: кортеж (image_file, output_folder, use_ocr, max_image_size, fit_mode, 
                     numbering_line1, numbering_line2, numbering_line3, numbering_position, numbering_border)
    
    Returns:
        dict: {
            'success': bool,
            'pdf_file': str or None,
            'image_file': str,
            'error': str or None
        }
    """
    image_file, output_folder, use_ocr, max_image_size, fit_mode, numbering_line1, numbering_line2, numbering_line3, numbering_position, numbering_border = args
    
    try:
        if not os.path.exists(image_file):
            raise FileNotFoundError(f"Файл не найден: {image_file}")
        
        image_file = os.path.abspath(image_file)
        
        base_name = os.path.splitext(os.path.basename(image_file))[0]
        
        if output_folder:
            os.makedirs(output_folder, exist_ok=True)
            pdf_file = os.path.join(output_folder, base_name + ".pdf")
        else:
            pdf_file = os.path.join(os.path.dirname(image_file), base_name + ".pdf")
        
        # Конвертация
        if use_ocr:
            GenerationDocApp.image_to_pdf_with_ocr(image_file, pdf_file, None, max_image_size, fit_mode,
                                                   numbering_line1, numbering_line2, numbering_line3,
                                                   numbering_position, numbering_border)  # log_callback=None для параллельных задач
        else:
            GenerationDocApp.image_to_pdf_simple(image_file, pdf_file, None, max_image_size, fit_mode,
                                                 numbering_line1, numbering_line2, numbering_line3,
                                                 numbering_position, numbering_border)
        
        if os.path.exists(pdf_file):
            return {
                'success': True,
                'pdf_file': pdf_file,
                'image_file': image_file,
                'error': None
            }
        else:
            return {
                'success': False,
                'pdf_file': None,
                'image_file': image_file,
                'error': 'PDF файл не был создан'
            }
    
    except Exception as e:
        return {
            'success': False,
            'pdf_file': None,
            'image_file': image_file,
            'error': str(e)
        }


def _convert_single_pdf(args):
    """
    Конвертация одного DOCX файла в PDF (функция для параллельного выполнения).
    
    Args:
        args: кортеж (docx_file, output_folder)
    
    Returns:
        dict: {
            'success': bool,
            'pdf_file': str or None,
            'docx_file': str,
            'error': str or None
        }
    """
    docx_file, output_folder = args
    
    try:
        if not os.path.exists(docx_file):
            raise FileNotFoundError(f"Файл не найден: {docx_file}")
        
        if not docx_file.lower().endswith('.docx'):
            raise ValueError("Файл должен иметь расширение .docx")
        
        docx_file = os.path.abspath(docx_file)
        
        if output_folder:
            os.makedirs(output_folder, exist_ok=True)
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            pdf_file = os.path.abspath(os.path.join(output_folder, base_name + ".pdf"))
        else:
            pdf_file = os.path.abspath(os.path.splitext(docx_file)[0] + ".pdf")
        
        if os.path.exists(pdf_file):
            try:
                os.remove(pdf_file)
            except:
                pass
        
        success = False
        last_error = None
        word = None
        doc = None
        
        if WIN32COM_AVAILABLE:
            try:
                import win32com.client
                import pythoncom
                
                pythoncom.CoInitialize()
                try:
                    # Используем DispatchEx для нового экземпляра Word
                    word = win32com.client.DispatchEx("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = 0  # Выключаем все диалоги
                    
                    # Открываем документ с таймаутом
                    doc = word.Documents.Open(docx_file, ReadOnly=True, AddToRecentFiles=False)
                    
                    # Сохраняем в PDF
                    doc.SaveAs(pdf_file, FileFormat=17)  # 17 = wdFormatPDF
                    
                    success = True
                    
                except Exception as e:
                    last_error = f"win32com: {str(e)}"
                    
                finally:
                    # Гарантированное закрытие документа и Word
                    try:
                        if doc:
                            doc.Close(SaveChanges=False)
                    except:
                        pass
                    
                    try:
                        if word:
                            word.Quit()
                            word = None
                    except:
                        pass
                    
                    try:
                        pythoncom.CoUninitialize()
                    except:
                        pass
                    
                    # Очищаем ссылки
                    doc = None
                    word = None
                    gc.collect()
                    
            except Exception as e:
                last_error = f"win32com: {str(e)}"
        
        if not success and DOCX2PDF_AVAILABLE:
            try:
                from docx2pdf import convert
                convert(docx_file, pdf_file)
                success = True
            except Exception as e:
                if last_error:
                    last_error += f"; docx2pdf: {str(e)}"
                else:
                    last_error = f"docx2pdf: {str(e)}"
        
        if success and os.path.exists(pdf_file):
            return {
                'success': True,
                'pdf_file': pdf_file,
                'docx_file': docx_file,
                'error': None
            }
        else:
            if not last_error:
                last_error = "PDF файл не был создан"
            raise Exception(last_error)
            
    except Exception as e:
        error_text = str(e)
        if "NoneType" in error_text or "COM" in error_text:
            error_text += " (Попробуйте закрыть все окна Word)"
        
        return {
            'success': False,
            'pdf_file': None,
            'docx_file': docx_file,
            'error': error_text
        }

def _process_single_document(args):
    """
    Обработка одного документа (функция для параллельного выполнения).
    
    Args:
        args: кортеж (row_index, row_data, word_template, output_folder, 
                     filename_pattern, required_columns, placeholders, 
                     filename_column)
    
    Returns:
        dict: результат обработки {
            'success': bool,
            'index': int,
            'filename': str,
            'is_incomplete': bool,
            'error': str or None,
            'logs': list of str  # Логи для вывода
        }
    """
    import pandas as pd
    from docx import Document
    import os
    
    logs = []
    row_index = None
    
    # Создаем инстанс кэша для проверки дублирования
    cache = DocumentCache()
    
    try:
        (row_index, row_data, word_template, output_folder, filename_pattern,
         required_columns, placeholders, filename_column) = args
        
        doc = Document(word_template)
        
        is_incomplete = any(
            pd.isna(row_data.get(col)) or str(row_data.get(col, "")).strip() == ""
            for col in required_columns
        )
        suffix = "_пусто" if is_incomplete else ""
        
        if is_incomplete:
            logs.append(f"   ⚠ Обнаружены пустые обязательные поля")
        
        column_value = ""
        if filename_column and filename_column in row_data:
            column_value = row_data.get(filename_column, "")
            if pd.isna(column_value):
                column_value = ""
            else:
                column_value = str(column_value).strip()
                invalid_chars = '<>:"/\\|?*'
                for char in invalid_chars:
                    column_value = column_value.replace(char, '')
                column_value = column_value.rstrip('.')
                if not column_value:
                    column_value = f"строка{row_index + 1}"
        
        if not column_value and '{column}' in filename_pattern:
            column_value = f"строка{row_index + 1}"
        
        replacements = {}
        for ph in placeholders:
            if not ph.get("active", True):
                continue
            
            value = row_data.get(ph["name"], "")
            
            # Гарантируем что ключ содержит фигурные скобки
            placeholder_key = ph["name"]
            if not placeholder_key.startswith('{'):
                placeholder_key = f"{{{placeholder_key}}}"
            
            replacements[placeholder_key] = value
        
        for paragraph in doc.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, replacements)
        
        for table in doc.tables:
            for table_row in table.rows:
                for cell in table_row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_placeholders_in_paragraph(paragraph, replacements)
        
        filename = filename_pattern.format(i=row_index + 1, suffix=suffix, column=column_value)
        name_part, ext = os.path.splitext(filename)
        if len(name_part) > 200:
            name_part = name_part[:200]
            filename = name_part + ext
        
        output_folder = output_folder.strip()  # Удаляем пробелы в конце
        filepath = os.path.join(output_folder, filename)
        
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Проверяем кэш перед созданием документа
        if not cache.should_create_document(word_template, dict(row_data), filepath):
            logs.append(f"💾 Пропущен (уже существует): {filename}")
            return {
                'success': True,
                'index': row_index,
                'filename': filename,
                'is_incomplete': is_incomplete,
                'error': None,
                'logs': logs
            }
        
        doc.save(filepath)
        logs.append(f"💾 Сохранен: {filename}")
        
        # Регистрируем документ в кэше после успешного сохранения
        cache.register_document(word_template, dict(row_data), filepath)
        
        del doc
        gc.collect()
        
        return {
            'success': True,
            'index': row_index,
            'filename': filename,
            'is_incomplete': is_incomplete,
            'error': None,
            'logs': logs
        }
        
    except Exception as e:
        error_msg = str(e)
        logs.append(f"   ❌ ОШИБКА: {error_msg}")
        
        return {
            'success': False,
            'index': row_index if row_index is not None else -1,
            'filename': None,
            'is_incomplete': False,
            'error': error_msg,
            'logs': logs
        }


def _process_single_excel_document(args):
    """
    Обработка одного Excel документа (функция для параллельного выполнения).
    
    Args:
        args: кортеж (row_index, row_data, excel_template, output_folder, 
                     filename_pattern, required_columns, placeholders, 
                     filename_column)
    
    Returns:
        dict: результат обработки {
            'success': bool,
            'index': int,
            'filename': str,
            'is_incomplete': bool,
            'error': str or None,
            'logs': list of str
        }
    """
    import pandas as pd
    from openpyxl import load_workbook
    import os
    import gc  
    import re
    
    logs = []
    row_index = None
    
    # Создаем инстанс кэша для проверки дублирования
    cache = DocumentCache()
    
    try:
        (row_index, row_data, excel_template, output_folder, filename_pattern,
         required_columns, placeholders, filename_column) = args
        
        # Загружаем Excel шаблон
        wb = load_workbook(excel_template)
        
        # Проверяем обязательные поля
        is_incomplete = any(
            pd.isna(row_data.get(col)) or str(row_data.get(col, "")).strip() == ""
            for col in required_columns
        )
        suffix = "_пусто" if is_incomplete else ""
        
        if is_incomplete:
            logs.append(f"   ⚠ Обнаружены пустые обязательные поля")
        
        # Формируем имя файла
        column_value = ""
        if filename_column and filename_column in row_data:
            column_value = row_data.get(filename_column, "")
            if pd.isna(column_value):
                column_value = ""
            else:
                column_value = str(column_value).strip()
                # Убираем недопустимые символы
                invalid_chars = '<>:"/\\|?*'
                for char in invalid_chars:
                    column_value = column_value.replace(char, '')
                column_value = column_value.rstrip('.')
                if not column_value:
                    column_value = f"строка{row_index + 1}"
        
        if not column_value and '{column}' in filename_pattern:
            column_value = f"строка{row_index + 1}"
        
        # Создаём замены для плейсхолдеров
        replacements = {}
        for ph in placeholders:
            if not ph.get("active", True):
                continue
            
            value = row_data.get(ph["name"], "")
            
            # Гарантируем что ключ содержит фигурные скобки
            placeholder_key = ph["name"]
            if not placeholder_key.startswith('{'):
                placeholder_key = f"{{{placeholder_key}}}"
            
            replacements[placeholder_key] = value
        
        # Заменяем плейсхолдеры во всех ячейках всех листов
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        cell_text = cell.value
                        # Заменяем все плейсхолдеры
                        for placeholder, replacement in replacements.items():
                            if placeholder in cell_text:
                                cell_text = cell_text.replace(placeholder, str(replacement))
                                logs.append(f"   ✓ Замена в ячейке {cell.coordinate}: {placeholder}")
                        cell.value = cell_text
        
        # Формируем имя файла
        filename = filename_pattern.format(i=row_index + 1, suffix=suffix, column=column_value)
        name_part, ext = os.path.splitext(filename)
        if len(name_part) > 200:
            name_part = name_part[:200]
            filename = name_part + ext
        
        output_folder = output_folder.strip()
        filepath = os.path.join(output_folder, filename)
        
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Проверяем кэш перед созданием документа
        if not cache.should_create_document(excel_template, dict(row_data), filepath):
            logs.append(f"💾 Пропущен (уже существует): {filename}")
            wb.close()
            return {
                'success': True,
                'index': row_index,
                'filename': filename,
                'is_incomplete': is_incomplete,
                'error': None,
                'logs': logs
            }
        
        # Сохраняем файл
        wb.save(filepath)
        logs.append(f"💾 Сохранен: {filename}")
        
        # Регистрируем документ в кэше после успешного сохранения
        cache.register_document(excel_template, dict(row_data), filepath)
        
        wb.close()
        del wb
        gc.collect()
        
        return {
            'success': True,
            'index': row_index,
            'filename': filename,
            'is_incomplete': is_incomplete,
            'error': None,
            'logs': logs
        }
        
    except Exception as e:
        error_msg = str(e)
        logs.append(f"   ❌ ОШИБКА: {error_msg}")
        
        return {
            'success': False,
            'index': row_index if row_index is not None else -1,
            'filename': None,
            'is_incomplete': False,
            'error': error_msg,
            'logs': logs
        }


class SimpleDatePicker(tk.Frame):
    """Простой выбор даты с календарём на русском языке"""
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=COLORS["card_bg"])
        self.selected_date = datetime.now()
        
        # Поле отображения даты (современное)
        self.date_var = tk.StringVar(value=self.selected_date.strftime('%d.%m.%Y'))
        self.entry = tk.Entry(
            self, 
            textvariable=self.date_var, 
            width=kwargs.get('width', 12), 
            state='readonly', 
            justify='center',
            font=FONTS["body"],
            relief=tk.FLAT,
            borderwidth=0,
            bg=COLORS["bg_tertiary"],
            fg=COLORS["text_primary"],
            readonlybackground=COLORS["bg_tertiary"],
            highlightthickness=1,
            highlightbackground=COLORS["border"],
            highlightcolor=COLORS["border_focus"]
        )
        self.entry.pack(side=tk.LEFT, padx=(0, SPACING["xs"]), ipady=2)
        enable_field_shortcuts(self.entry, readonly=True)
        add_context_menu(self.entry, readonly=True)
        
        self.btn = tk.Button(
            self, 
            text="📅", 
            command=self.open_calendar, 
            width=2,
            height=1,
            font=FONTS["body"],
            bg=COLORS["primary"],
            fg="white",
            relief=tk.FLAT,
            borderwidth=0,
            cursor="hand2",
            activebackground=COLORS["primary_hover"],
            activeforeground="white"
        )
        self.btn.pack(side=tk.LEFT)
        ToolTip(self.btn, "Выбрать дату из календаря")
        
        self.calendar_window = None
    
    def open_calendar(self):
        if self.calendar_window:
            return
        
        self.calendar_window = tk.Toplevel(self)
        self.calendar_window.withdraw()
        self.calendar_window.title("Выбор даты")
        self.calendar_window.resizable(False, False)
        self.calendar_window.configure(bg=COLORS["card_bg"])
        
        self.calendar_window.update_idletasks()
        x = self.winfo_toplevel().winfo_x() + (self.winfo_toplevel().winfo_width() // 2) - 150
        y = self.winfo_toplevel().winfo_y() + (self.winfo_toplevel().winfo_height() // 2) - 200
        self.calendar_window.geometry(f"+{x}+{y}")
        
        self.calendar_window.deiconify()
        self.calendar_window.grab_set()
        
        self.current_year = self.selected_date.year
        self.current_month = self.selected_date.month
        
        header_frame = tk.Frame(self.calendar_window, bg=COLORS["primary"], height=50)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        prev_btn = tk.Button(
            header_frame, 
            text="◀", 
            command=self.prev_month, 
            width=3,
            bg=COLORS["primary"],
            fg="white",
            font=FONTS["button"],
            relief=tk.FLAT,
            borderwidth=0,
            cursor="hand2",
            activebackground=COLORS["primary_hover"],
            activeforeground="white"
        )
        prev_btn.pack(side=tk.LEFT, padx=SPACING["md"], pady=SPACING["md"])
        ToolTip(prev_btn, "Предыдущий месяц")
        
        self.month_year_label = tk.Label(
            header_frame, 
            text="", 
            font=FONTS["heading"], 
            bg=COLORS["primary"],
            fg="white"
        )
        self.month_year_label.pack(side=tk.LEFT, expand=True, fill=tk.BOTH, pady=SPACING["md"])
        
        next_btn = tk.Button(
            header_frame, 
            text="▶", 
            command=self.next_month, 
            width=3,
            bg=COLORS["primary"],
            fg="white",
            font=FONTS["button"],
            relief=tk.FLAT,
            borderwidth=0,
            cursor="hand2",
            activebackground=COLORS["primary_hover"],
            activeforeground="white"
        )
        next_btn.pack(side=tk.RIGHT, padx=SPACING["md"], pady=SPACING["md"])
        ToolTip(next_btn, "Следующий месяц")
        
        # Календарная сетка (с современным фоном)
        self.calendar_frame = tk.Frame(self.calendar_window, bg=COLORS["card_bg"])
        self.calendar_frame.pack(padx=SPACING["lg"], pady=SPACING["lg"])
        
        self.draw_calendar()
        
        self.calendar_window.protocol("WM_DELETE_WINDOW", self.close_calendar)
    
    def draw_calendar(self):
        # Очистка
        for widget in self.calendar_frame.winfo_children():
            widget.destroy()
        
        # Обновление заголовка
        month_name = RUSSIAN_MONTHS[self.current_month - 1]
        self.month_year_label.config(text=f"{month_name} {self.current_year}")
        
        # Дни недели (современный стиль)
        for i, day in enumerate(RUSSIAN_DAYS):
            lbl = tk.Label(
                self.calendar_frame, 
                text=day, 
                font=FONTS["button"], 
                width=5,
                height=2,
                bg=COLORS["bg_secondary"],
                fg=COLORS["text_secondary"],
                relief=tk.FLAT,
                borderwidth=0
            )
            lbl.grid(row=0, column=i, padx=1, pady=1, sticky="nsew")
        
        cal = calendar.monthcalendar(self.current_year, self.current_month)
        
        # Рисуем дни (современный стиль)
        for week_num, week in enumerate(cal):
            for day_num, day in enumerate(week):
                if day == 0:
                    lbl = tk.Label(
                        self.calendar_frame, 
                        text="", 
                        width=5,
                        height=2,
                        bg=COLORS["card_bg"],
                        relief=tk.FLAT
                    )
                    lbl.grid(row=week_num + 1, column=day_num, padx=1, pady=1)
                else:
                    is_selected = (day == self.selected_date.day and 
                                 self.current_month == self.selected_date.month and 
                                 self.current_year == self.selected_date.year)
                    
                    bg_color = COLORS["primary"] if is_selected else COLORS["card_bg"]
                    fg_color = "white" if is_selected else COLORS["text_primary"]
                    
                    btn = tk.Button(
                        self.calendar_frame, 
                        text=str(day), 
                        width=5,
                        height=2,
                        bg=bg_color, 
                        fg=fg_color,
                        font=FONTS["body"],
                        relief=tk.FLAT,
                        borderwidth=0,
                        cursor="hand2",
                        activebackground=COLORS["primary_hover"],
                        activeforeground="white",
                        command=lambda d=day: self.select_date(d)
                    )
                    btn.grid(row=week_num + 1, column=day_num, padx=1, pady=1, sticky="nsew")
                    
                    # Эффект наведения для невыбранных дней
                    if not is_selected:
                        btn.bind("<Enter>", lambda e, b=btn: b.config(bg=COLORS["primary_light"], fg="white"))
                        btn.bind("<Leave>", lambda e, b=btn: b.config(bg=COLORS["card_bg"], fg=COLORS["text_primary"]))
    
    def prev_month(self):
        if self.current_month == 1:
            self.current_month = 12
            self.current_year -= 1
        else:
            self.current_month -= 1
        self.draw_calendar()
    
    def next_month(self):
        if self.current_month == 12:
            self.current_month = 1
            self.current_year += 1
        else:
            self.current_month += 1
        self.draw_calendar()
    
    def select_date(self, day):
        self.selected_date = datetime(self.current_year, self.current_month, day)
        self.date_var.set(self.selected_date.strftime('%d.%m.%Y'))
        self.close_calendar()
    
    def close_calendar(self):
        if self.calendar_window:
            self.calendar_window.destroy()
            self.calendar_window = None
    
    def get_date(self):
        """Возвращает выбранную дату как объект datetime"""
        return self.selected_date
    
    def set_date(self, date_obj):
        """Устанавливает дату программно"""
        if isinstance(date_obj, datetime):
            self.selected_date = date_obj
            self.date_var.set(self.selected_date.strftime('%d.%m.%Y'))

class PlaceholderEditorDialog:
    """Диалоговое окно для управления настройками (плейсхолдеры и списки)"""
    def __init__(self, parent, placeholders, custom_lists):
        self.result = None
        self.lists_result = None
        self.top = tk.Toplevel(parent)
        self.top.withdraw()
        self.top.title("Настройки")
        self.top.geometry("1000x800")
        self.top.resizable(True, True)
        self.top.transient(parent)
        
        self.top.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (1000 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (800 // 2)
        self.top.geometry(f"+{x}+{y}")
        
        self.top.deiconify()
        self.top.grab_set()
        
        self.placeholders = [p.copy() for p in placeholders]
        self.custom_lists = {k: v.copy() for k, v in custom_lists.items()}
        
        header = tk.Frame(self.top, bg=COLORS["primary"], height=55)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(
            header,
            text="⚙️ Настройки плейсх олдеров и списков",
            font=FONTS["title"],
            bg=COLORS["primary"],
            fg="white"
        ).pack(pady=12)
        
        tabs_container = tk.Frame(self.top, bg=COLORS["bg_secondary"])
        tabs_container.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)
        
        self.tabs_frame = tk.Frame(tabs_container, bg=COLORS["bg_tertiary"], height=45)
        self.tabs_frame.pack(fill=tk.X)
        self.tabs_frame.pack_propagate(False)
        
        self.content_frame = tk.Frame(tabs_container, bg=COLORS["bg_secondary"])
        self.content_frame.pack(fill=tk.BOTH, expand=True, pady=(SPACING["md"], 0))
        
        self.placeholders_frame = None
        self.lists_frame = None
        
        self.tab_buttons = []
        self.current_tab = 0
        
        self.create_tab_button("📋 Плейсхолдеры", 0)
        self.create_tab_button("📝 Выпадающие списки", 1)
        
        self.create_placeholders_tab()
        
        self.create_custom_lists_tab()
        
        self.switch_tab(0)
        
        self.switch_tab(0)
        
        bottom_frame = tk.Frame(self.top, bg=COLORS["bg_tertiary"])
        bottom_frame.pack(fill=tk.X, padx=12, pady=12)
        
        cancel_btn = create_modern_button(
            bottom_frame, 
            text="✗ Отмена", 
            command=self.cancel, 
            style="secondary",
            width=100,
            height=48,
            tooltip="Отменить изменения и закрыть"
        )
        cancel_btn.pack(side=tk.RIGHT, padx=5)
        
        save_btn = create_modern_button(
            bottom_frame, 
            text="✓ Сохранить", 
            command=self.save, 
            style="success",
            width=120,
            height=48,
            tooltip="Сохранить все изменения"
        )
        save_btn.pack(side=tk.RIGHT, padx=5)
    
    def create_tab_button(self, text, tab_id):
        """Создание кнопки вкладки"""
        btn = tk.Button(
            self.tabs_frame,
            text=text,
            font=FONTS["button"],
            bg=COLORS["bg_hover"],
            activebackground=COLORS["primary_light"],
            fg=COLORS["text_primary"],
            activeforeground="white",
            relief=tk.FLAT,
            cursor="hand2",
            padx=20,
            pady=10,
            command=lambda: self.switch_tab(tab_id)
        )
        btn.pack(side=tk.LEFT, padx=2)
        self.tab_buttons.append(btn)
        return btn
    
    def switch_tab(self, tab_id):
        """Переключение между вкладками"""
        self.current_tab = tab_id
        
        for i, btn in enumerate(self.tab_buttons):
            if i == tab_id:
                btn.configure(
                    bg=COLORS["primary"],
                    fg="white",
                    font=("Segoe UI", 10, "bold")
                )
            else:
                btn.configure(
                    bg=COLORS["bg_hover"],
                    fg=COLORS["text_primary"],
                    font=FONTS["button"]
                )
        
        if tab_id == 0:
            if self.placeholders_frame:
                self.placeholders_frame.pack(fill=tk.BOTH, expand=True)
            if self.lists_frame:
                self.lists_frame.pack_forget()
        else:
            if self.placeholders_frame:
                self.placeholders_frame.pack_forget()
            if self.lists_frame:
                self.lists_frame.pack(fill=tk.BOTH, expand=True)
    
    def create_placeholders_tab(self):
        """Создание вкладки с плейсхолдерами"""
        tab_frame = tk.Frame(self.content_frame, bg=COLORS["bg_secondary"])
        self.placeholders_frame = tab_frame
        
        # Список плейсхолдеров
        list_frame = tk.Frame(tab_frame, padx=SPACING["lg"], pady=SPACING["lg"], bg=COLORS["bg_secondary"])
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(
            list_frame, 
            text="📋 Список плейсхолдеров:", 
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["primary"]
        ).pack(anchor="w", pady=(0, SPACING["xs"]))
        
        tk.Label(
            list_frame, 
            text="Создавайте и редактируйте плейсхолдеры для автозаполнения документов.", 
            font=FONTS["small"], 
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor="w", pady=(0, SPACING["md"]))
        
        # Таблица плейсхолдеров с современным стилем
        columns = ("Плейсхолдер", "Источник", "Значение", "Падеж", "Обязат.", "Активный")
        tree_container, self.tree = create_modern_treeview(
            list_frame,
            columns=columns,
            height=12
        )
        tree_container.pack(fill=tk.BOTH, expand=True)
        
        # Настройка колонок
        columns_config = {
            "Плейсхолдер": {"text": "Плейсхолдер", "width": 160, "anchor": tk.W},
            "Источник": {"text": "Источник", "width": 100, "anchor": tk.W},
            "Значение": {"text": "Значение", "width": 180, "anchor": tk.W},
            "Падеж": {"text": "Падеж", "width": 100, "anchor": tk.W},
            "Обязат.": {"text": "Обязат.", "width": 70, "anchor": tk.CENTER},
            "Активный": {"text": "Активный", "width": 80, "anchor": tk.CENTER}
        }
        configure_treeview_columns(self.tree, columns_config)
        
        self.refresh_placeholders_list()
        
        btn_frame = tk.Frame(tab_frame, padx=SPACING["lg"], pady=SPACING["md"], bg=COLORS["bg_secondary"])
        btn_frame.pack(fill=tk.X)
        
        add_btn = create_modern_button(
            btn_frame, 
            text="+ Добавить", 
            command=self.add_placeholder, 
            style="success",
            width=110,
            height=36,
            tooltip="Добавить новый плейсхолдер"
        )
        add_btn.pack(side=tk.LEFT, padx=SPACING["sm"])
        
        edit_btn = create_modern_button(
            btn_frame, 
            text="✏ Изменить", 
            command=self.edit_placeholder, 
            style="warning",
            width=110,
            height=36,
            tooltip="Редактировать выбранный плейсхолдер"
        )
        edit_btn.pack(side=tk.LEFT, padx=SPACING["sm"])
        
        delete_btn = create_modern_button(
            btn_frame, 
            text="🗑 Удалить", 
            command=self.delete_placeholder, 
            style="danger",
            width=110,
            height=36,
            tooltip="Удалить выбранный плейсхолдер"
        )
        delete_btn.pack(side=tk.LEFT, padx=SPACING["sm"])
    
    def create_custom_lists_tab(self):
        """Создание вкладки с выпадающими списками"""
        tab_frame = tk.Frame(self.content_frame, bg=COLORS["bg_secondary"])
        self.lists_frame = tab_frame
        
        list_frame = tk.Frame(tab_frame, padx=SPACING["lg"], pady=SPACING["lg"], bg=COLORS["bg_secondary"])
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(
            list_frame, 
            text="📝 Управление выпадающими списками:", 
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["primary"]
        ).pack(anchor="w", pady=(0, SPACING["xs"]))
        
        tk.Label(
            list_frame, 
            text="Создавайте и редактируйте выпадающие списки для использования в интерфейсе.", 
            font=FONTS["small"], 
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor="w", pady=(0, SPACING["md"]))
        
        columns = ("Название", "Ключ", "Кол-во элементов", "Элементы", "Активный")
        tree_container, self.lists_tree = create_modern_treeview(
            list_frame,
            columns=columns,
            height=10
        )
        tree_container.pack(fill=tk.BOTH, expand=True)
        
        columns_config = {
            "Название": {"text": "Название", "width": 180, "anchor": tk.W},
            "Ключ": {"text": "Ключ", "width": 130, "anchor": tk.W},
            "Кол-во элементов": {"text": "Кол-во", "width": 80, "anchor": tk.CENTER},
            "Элементы": {"text": "Предпросмотр", "width": 280, "anchor": tk.W},
            "Активный": {"text": "Активный", "width": 80, "anchor": tk.CENTER}
        }
        configure_treeview_columns(self.lists_tree, columns_config)
        
        self.refresh_custom_lists()
        
        btn_frame = tk.Frame(tab_frame, padx=SPACING["lg"], pady=SPACING["md"], bg=COLORS["bg_secondary"])
        btn_frame.pack(fill=tk.X)
        
        add_list_btn = create_modern_button(
            btn_frame, 
            text="+ Добавить", 
            command=self.add_custom_list, 
            style="success",
            width=110,
            height=36,
            tooltip="Добавить новый выпадающий список"
        )
        add_list_btn.pack(side=tk.LEFT, padx=SPACING["sm"])
        
        edit_list_btn = create_modern_button(
            btn_frame, 
            text="✏ Изменить", 
            command=self.edit_custom_list, 
            style="warning",
            width=110,
            height=36,
            tooltip="Редактировать выбранный список"
        )
        edit_list_btn.pack(side=tk.LEFT, padx=SPACING["sm"])
        
        delete_list_btn = create_modern_button(
            btn_frame, 
            text="🗑 Удалить", 
            command=self.delete_custom_list, 
            style="danger",
            width=110,
            height=36,
            tooltip="Удалить выбранный список"
        )
        delete_list_btn.pack(side=tk.LEFT, padx=SPACING["sm"])
    
    def refresh_placeholders_list(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        source_type_names = {
            "excel": "Excel",
            "dropdown": "Список",
            "date": "Дата",
            "static": "Статика"
        }
        
        for ph in self.placeholders:
            case_name = RUSSIAN_CASES.get(ph.get("case", "nomn"), "Именительный").split(" ")[0]
            is_active = ph.get("active", True)
            insert_treeview_row(self.tree, (
                ph["name"],
                source_type_names.get(ph["source_type"], ph["source_type"]),
                ph["source_value"],
                case_name,
                "✓" if ph["required"] else "",
                "✓" if is_active else "✗"
            ))
    
    def refresh_custom_lists(self):
        for item in self.lists_tree.get_children():
            self.lists_tree.delete(item)
        
        for key, list_data in self.custom_lists.items():
            if isinstance(list_data, dict):
                display_name = list_data.get("display_name", key)
                values = list_data.get("values", [])
                is_active = list_data.get("active", True)
            else:
                display_name = key
                values = list_data
                is_active = True
            
            preview = ", ".join(values[:3])
            if len(values) > 3:
                preview += f", ... (+{len(values)-3})"
            insert_treeview_row(self.lists_tree, (
                display_name,
                key,
                len(values),
                preview,
                "✓" if is_active else "✗"
            ))
    
    def add_placeholder(self):
        dialog = PlaceholderItemDialog(self.top, "Добавить плейсхолдер", None)
        self.top.wait_window(dialog.top)
        
        if dialog.result:
            self.placeholders.append(dialog.result)
            self.refresh_placeholders_list()
    
    def edit_placeholder(self):
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите плейсхолдер для изменения", parent=self.top)
            return
        
        index = self.tree.index(selection[0])
        old_data = self.placeholders[index]
        
        dialog = PlaceholderItemDialog(self.top, "Изменить плейсхолдер", old_data)
        self.top.wait_window(dialog.top)
        
        if dialog.result:
            self.placeholders[index] = dialog.result
            self.refresh_placeholders_list()
    
    def delete_placeholder(self):
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите плейсхолдер для удаления", parent=self.top)
            return
        
        index = self.tree.index(selection[0])
        ph_name = self.placeholders[index]["name"]
        
        if messagebox.askyesno("Подтверждение", f"Удалить плейсхолдер '{ph_name}'?", parent=self.top):
            del self.placeholders[index]
            self.refresh_placeholders_list()
    
    def save(self):
        self.result = self.placeholders
        self.lists_result = self.custom_lists
        self.top.destroy()
    
    def cancel(self):
        self.result = None
        self.lists_result = None
        self.top.destroy()
    
    def add_custom_list(self):
        """Создать новый выпадающий список"""
        dialog = CustomListDialog(self.top, "Создать список", None, None, None, True)
        self.top.wait_window(dialog.top)
        
        if dialog.result:
            key, display_name, values, is_active = dialog.result
            if key in self.custom_lists:
                messagebox.showerror("Ошибка", f"Список с ключом '{key}' уже существует", parent=self.top)
                return
            self.custom_lists[key] = {
                "display_name": display_name,
                "values": values,
                "active": is_active
            }
            self.refresh_custom_lists()
    
    def edit_custom_list(self):
        """Редактировать выбранный список"""
        selection = self.lists_tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите список для редактирования", parent=self.top)
            return
        
        item_values = self.lists_tree.item(selection[0], 'values')
        display_name = item_values[0]
        key = item_values[1]  # Ключ во второй колонке
        
        list_data = self.custom_lists.get(key, {})
        if isinstance(list_data, dict):
            current_display_name = list_data.get("display_name", key)
            values = list_data.get("values", [])
            current_is_active = list_data.get("active", True)
        else:
            current_display_name = key
            values = list_data
            current_is_active = True
        
        dialog = CustomListDialog(self.top, "Редактировать список", key, values, current_display_name, current_is_active)
        self.top.wait_window(dialog.top)
        
        if dialog.result:
            new_key, new_display_name, new_values, is_active = dialog.result
            
            if new_key != key and new_key in self.custom_lists:
                messagebox.showerror("Ошибка", f"Список с ключом '{new_key}' уже существует", parent=self.top)
                return
            
            if new_key != key:
                del self.custom_lists[key]
            
            self.custom_lists[new_key] = {
                "display_name": new_display_name,
                "values": new_values,
                "active": is_active
            }
            self.refresh_custom_lists()
    
    def delete_custom_list(self):
        """Удалить выбранный список"""
        selection = self.lists_tree.selection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите список для удаления", parent=self.top)
            return
        
        item_values = self.lists_tree.item(selection[0], 'values')
        display_name = item_values[0]
        key = item_values[1]  # Ключ во второй колонке
        
        if messagebox.askyesno("Подтверждение", f"Удалить список '{display_name}'?", parent=self.top):
            del self.custom_lists[key]
            self.refresh_custom_lists()

class CustomListDialog:
    """Диалог создания/редактирования выпадающего списка"""
    def __init__(self, parent, title, key=None, values=None, display_name=None, is_active=True):
        self.result = None
        
        self.top = tk.Toplevel(parent)
        self.top.withdraw()
        self.top.title(title)
        self.top.geometry("600x600")
        self.top.transient(parent)
        
        self.top.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (600 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (600 // 2)
        self.top.geometry(f"+{x}+{y}")
        
        self.top.deiconify()
        self.top.grab_set()
        
        header = tk.Frame(self.top, bg=COLORS["primary"], height=40)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        tk.Label(header, text=title, font=FONTS["title"], bg=COLORS["primary"], fg="white").pack(pady=8)
        
        main_frame = tk.Frame(self.top, padx=15, pady=15, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH)
        
        tk.Label(main_frame, text="Ключ списка (для плейсхолдеров):", font=FONTS["body"], bg=COLORS["bg_secondary"], fg=COLORS["text_primary"]).pack(pady=(0,5), anchor="w")
        self.key_entry = ctk.CTkEntry(
            main_frame,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=32
        )
        self.key_entry.pack(fill=tk.X, pady=(0,12))
        if key:
            self.key_entry.insert(0, key)
        enable_field_shortcuts(self.key_entry)
        add_context_menu(self.key_entry)
        ToolTip(self.key_entry, "Уникальное имя для использования в плейсхолдерах")
        
        tk.Label(main_frame, text="Название (как видит пользователь):", font=FONTS["body"], bg=COLORS["bg_secondary"], fg=COLORS["text_primary"]).pack(pady=(0,5), anchor="w")
        self.display_name_entry = ctk.CTkEntry(
            main_frame,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=32
        )
        self.display_name_entry.pack(fill=tk.X, pady=(0,12))
        if display_name:
            self.display_name_entry.insert(0, display_name)
        enable_field_shortcuts(self.display_name_entry)
        add_context_menu(self.display_name_entry)
        ToolTip(self.display_name_entry, "Название, которое отображается в интерфейсе")
        
        tk.Label(main_frame, text="Элементы списка (по одному на строку):", font=FONTS["body"], bg=COLORS["bg_secondary"]).pack(pady=(0,5), anchor="w")
        
        text_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        text_frame.pack(fill=tk.BOTH, pady=(0,12))
        
        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.text_widget = tk.Text(text_frame, wrap=tk.WORD, yscrollcommand=scrollbar.set, font=FONTS["body"], relief=tk.SOLID, borderwidth=1, height=12)
        self.text_widget.pack(side=tk.LEFT, fill=tk.BOTH)
        scrollbar.config(command=self.text_widget.yview)
        enable_field_shortcuts(self.text_widget)
        add_context_menu(self.text_widget)
        
        if values:
            self.text_widget.insert("1.0", "\n".join(values))
        
        self.active_var = tk.BooleanVar(value=is_active)
        active_check = tk.Checkbutton(
            main_frame, 
            text="Активный (отображать этот список в интерфейсе)", 
            variable=self.active_var, 
            font=FONTS["body"], 
            bg=COLORS["bg_secondary"]
        )
        active_check.pack(pady=(0, 12), anchor="w")
        ToolTip(active_check, "Если включено, список будет отображаться на главном экране. Отключите для создания пресетов")
        
        btn_frame = tk.Frame(self.top, bg=COLORS["bg_tertiary"], padx=12, pady=12)
        btn_frame.pack(fill=tk.X)
        
        cancel_btn = create_modern_button(
            btn_frame, 
            text="✗ Отмена", 
            command=self.cancel, 
            style="secondary",
            width=90,
            height=40,
            tooltip="Отменить изменения"
        )
        cancel_btn.pack(side=tk.RIGHT, padx=3)
        
        save_btn = create_modern_button(
            btn_frame, 
            text="✓ Сохранить", 
            command=self.ok, 
            style="success",
            width=100,
            height=40,
            tooltip="Сохранить список"
        )
        save_btn.pack(side=tk.RIGHT, padx=3)
    
    def ok(self):
        key = self.key_entry.get().strip()
        if not key:
            messagebox.showerror("Ошибка", "Укажите ключ списка", parent=self.top)
            return
        
        display_name = self.display_name_entry.get().strip()
        if not display_name:
            messagebox.showerror("Ошибка", "Укажите название списка", parent=self.top)
            return
        
        text = self.text_widget.get("1.0", tk.END).strip()
        if not text:
            messagebox.showerror("Ошибка", "Добавьте элементы в список", parent=self.top)
            return
        
        values = [line.strip() for line in text.split("\n") if line.strip()]
        if not values:
            messagebox.showerror("Ошибка", "Список не может быть пустым", parent=self.top)
            return
        
        self.result = (key, display_name, values, self.active_var.get())
        self.top.destroy()
    
    def cancel(self):
        self.result = None
        self.top.destroy()

class PlaceholderItemDialog:
    """Диалог для редактирования одного плейсхолдера"""
    def __init__(self, parent, title, data):
        self.result = None
        self.top = tk.Toplevel(parent)
        self.top.withdraw()
        self.top.title(title)
        self.top.geometry("650x700")
        self.top.resizable(False, False)
        self.top.transient(parent)
        
        self.top.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (650 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (700 // 2)
        self.top.geometry(f"+{x}+{y}")
        
        self.top.deiconify()
        self.top.grab_set()
        
        header = tk.Frame(self.top, bg=COLORS["primary"], height=40)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        tk.Label(header, text=title, font=FONTS["title"], bg=COLORS["primary"], fg="white").pack(pady=8)
        
        main_frame = tk.Frame(self.top, padx=15, pady=15, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH)
        
        # Имя плейсхолдера
        tk.Label(main_frame, text="Имя плейсхолдера:", font=FONTS["heading"], bg=COLORS["bg_secondary"], fg=COLORS["primary"]).grid(row=0, column=0, sticky="w", pady=(0, 5))
        self.name_var = tk.StringVar(value=data["name"] if data else "")
        name_entry = ctk.CTkEntry(
            main_frame,
            textvariable=self.name_var,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=32
        )
        name_entry.grid(row=1, column=0, columnspan=2, pady=(0, 15), sticky="ew")
        enable_field_shortcuts(name_entry)
        add_context_menu(name_entry)
        ToolTip(name_entry, "Имя плейсхолдера в формате: {имя}")
        
        # Источник данных
        tk.Label(main_frame, text="Источник данных:", font=FONTS["heading"], bg=COLORS["bg_secondary"]).grid(row=2, column=0, sticky="w", pady=(0, 5))
        self.source_type_var = tk.StringVar(value=data["source_type"] if data else "excel")
        source_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        source_frame.grid(row=3, column=0, columnspan=2, sticky="w", pady=(0, 15))
        
        tk.Radiobutton(source_frame, text="Excel столбец", variable=self.source_type_var, value="excel", command=self.update_source_state, font=FONTS["body"], bg=COLORS["bg_secondary"]).pack(anchor="w", pady=2)
        tk.Radiobutton(source_frame, text="Выпадающий список", variable=self.source_type_var, value="dropdown", command=self.update_source_state, font=FONTS["body"], bg=COLORS["bg_secondary"]).pack(anchor="w", pady=2)
        tk.Radiobutton(source_frame, text="Дата", variable=self.source_type_var, value="date", command=self.update_source_state, font=FONTS["body"], bg=COLORS["bg_secondary"]).pack(anchor="w", pady=2)
        tk.Radiobutton(source_frame, text="Статическое значение", variable=self.source_type_var, value="static", command=self.update_source_state, font=FONTS["body"], bg=COLORS["bg_secondary"]).pack(anchor="w", pady=2)
        
        # Значение источника
        tk.Label(main_frame, text="Значение источника:", font=FONTS["heading"], bg=COLORS["bg_secondary"], fg=COLORS["primary"]).grid(row=4, column=0, sticky="w", pady=(0, 5))
        tk.Label(main_frame, text="(для Excel - имя столбца, для списка - ключ)", font=FONTS["small"], fg=COLORS["text_secondary"], bg=COLORS["bg_secondary"]).grid(row=4, column=1, sticky="w", padx=(10, 0), pady=(0, 5))
        self.source_value_var = tk.StringVar(value=data["source_value"] if data else "")
        self.source_value_entry = ctk.CTkEntry(
            main_frame,
            textvariable=self.source_value_var,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=32
        )
        self.source_value_entry.grid(row=5, column=0, columnspan=2, pady=(0, 15), sticky="ew")
        enable_field_shortcuts(self.source_value_entry)
        add_context_menu(self.source_value_entry)
        ToolTip(self.source_value_entry, "Имя столбца Excel, ключ списка или статическое значение")
        
        # Падеж
        tk.Label(main_frame, text="Падеж для склонения:", font=FONTS["heading"], bg=COLORS["bg_secondary"], fg=COLORS["primary"]).grid(row=6, column=0, sticky="w", pady=(0, 5))
        self.case_var = tk.StringVar(value=data.get("case", "nomn") if data else "nomn")
        case_values = [(key, val) for key, val in RUSSIAN_CASES.items()]
        case_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        case_frame.grid(row=7, column=0, columnspan=2, sticky="w", pady=(0, 15))
        
        tk.Label(case_frame, text="Выберите падеж для автоматического склонения текста:", font=FONTS["small"], bg=COLORS["bg_secondary"], fg=COLORS["text_secondary"]).pack(anchor="w", pady=(0, 5))
        self.case_combo = ctk.CTkComboBox(
            case_frame,
            variable=self.case_var,
            values=[f"{key} - {val}" for key, val in case_values],
            state="readonly",
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            button_color=COLORS["primary"],
            button_hover_color=COLORS["primary_hover"],
            dropdown_fg_color=COLORS["bg_primary"],
            height=32,
            width=450
        )
        # Установка текущего значения
        current_case = data.get("case", "nomn") if data else "nomn"
        for i, (key, val) in enumerate(case_values):
            if key == current_case:
                self.case_combo.set(f"{key} - {val}")
                break
        self.case_combo.pack(anchor="w", fill=tk.X)
        set_combobox_cursor(self.case_combo)
        ToolTip(self.case_combo, "Падеж для автоматического склонения текста")
        
        # Опции
        tk.Label(main_frame, text="Опции:", font=FONTS["heading"], bg=COLORS["bg_secondary"]).grid(row=8, column=0, sticky="w", pady=(10, 5))
        
        self.required_var = tk.BooleanVar(value=data["required"] if data else False)
        req_check = tk.Checkbutton(main_frame, text="Обязательное поле (проверка на пустоту в Excel)", variable=self.required_var, font=FONTS["body"], bg=COLORS["bg_secondary"])
        req_check.grid(row=9, column=0, columnspan=2, sticky="w", pady=(5, 5))
        ToolTip(req_check, "Если включено, будет выдано предупреждение при пустом значении в Excel")
        
        self.active_var = tk.BooleanVar(value=data.get("active", True) if data else True)
        active_check = tk.Checkbutton(main_frame, text="Активный (использовать этот плейсхолдер)", variable=self.active_var, font=FONTS["body"], bg=COLORS["bg_secondary"])
        active_check.grid(row=10, column=0, columnspan=2, sticky="w", pady=(5, 15))
        ToolTip(active_check, "Если включено, плейсхолдер будет использоваться при генерации. Отключите для создания пресетов")
        
        btn_frame = tk.Frame(self.top, bg=COLORS["bg_tertiary"], padx=12, pady=12)
        btn_frame.pack(fill=tk.X)
        
        cancel_btn = create_modern_button(
            btn_frame, 
            text="✗ Отмена", 
            command=self.cancel, 
            style="secondary",
            width=90,
            height=40,
            tooltip="Отменить изменения"
        )
        cancel_btn.pack(side=tk.RIGHT, padx=3)
        
        ok_btn = create_modern_button(
            btn_frame, 
            text="✓ Сохранить", 
            command=self.ok, 
            style="success",
            width=100,
            height=40,
            tooltip="Сохранить плейсхолдер"
        )
        ok_btn.pack(side=tk.RIGHT, padx=3)
        
        self.update_source_state()
    
    def update_source_state(self):
        source_type = self.source_type_var.get()
        # Для типа "дата" отключаем поле значения
        if source_type == "date":
            self.source_value_entry.configure(state="disabled")
            self.source_value_var.set("")
        else:
            self.source_value_entry.configure(state="normal")
    
    def ok(self):
        name = self.name_var.get().strip()
        if not name:
            messagebox.showerror("Ошибка", "Введите имя плейсхолдера!", parent=self.top)
            return
        
        source_type = self.source_type_var.get()
        source_value = self.source_value_var.get().strip()
        
        if source_type != "date" and not source_value:
            messagebox.showerror("Ошибка", "Введите значение источника!", parent=self.top)
            return
        
        # Извлекаем ключ падежа из выбранного значения (формат: "gent - Родительный (кого? чего?)")
        case_str = self.case_var.get()
        case_key = case_str.split(" - ")[0] if " - " in case_str else case_str
        
        self.result = {
            "name": name,
            "source_type": source_type,
            "source_value": source_value,
            "case": case_key,
            "required": self.required_var.get(),
            "active": self.active_var.get()
        }
        self.top.destroy()
    
    def cancel(self):
        self.result = None
        self.top.destroy()

class EditListDialog:
    """Диалоговое окно для редактирования списка"""
    def __init__(self, parent, title, items):
        self.result = None
        self.top = tk.Toplevel(parent)
        self.top.withdraw()
        self.top.title(title)
        self.top.geometry("450x500")
        self.top.resizable(False, False)
        self.top.transient(parent)
        
        self.top.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (450 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (500 // 2)
        self.top.geometry(f"+{x}+{y}")
        
        self.top.deiconify()
        self.top.grab_set()
        
        self.items = items.copy()
        
        header = tk.Frame(self.top, bg=COLORS["primary"], height=40)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        tk.Label(header, text=title, font=FONTS["title"], bg=COLORS["primary"], fg="white").pack(pady=8)
        
        list_frame = tk.Frame(self.top, padx=12, pady=12, bg=COLORS["bg_secondary"])
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(list_frame, text="Список элементов:", font=FONTS["heading"], bg=COLORS["bg_secondary"], fg=COLORS["primary"]).pack(anchor="w", pady=(0, 8))
        
        listbox_container = tk.Frame(list_frame, bg=COLORS["card_bg"], relief=tk.FLAT)
        listbox_container.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = tk.Scrollbar(
            listbox_container, 
            orient=tk.VERTICAL
        )
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.listbox = tk.Listbox(
            listbox_container, 
            yscrollcommand=scrollbar.set, 
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"],
            selectbackground=COLORS["primary_light"],
            selectforeground="white",
            relief=tk.FLAT,
            borderwidth=0,
            highlightthickness=1,
            highlightbackground=COLORS["border"],
            highlightcolor=COLORS["border_focus"]
        )
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.configure(command=self.listbox.yview)
        
        self.refresh_list()
        
        btn_frame = tk.Frame(self.top, padx=12, pady=12, bg=COLORS["bg_secondary"])
        btn_frame.pack(fill=tk.X)
        
        add_btn = create_modern_button(
            btn_frame, 
            text="+ Добавить", 
            command=self.add_item, 
            style="success",
            width=12, 
            tooltip="Добавить новый элемент"
        )
        add_btn.pack(side=tk.LEFT, padx=2)
        
        edit_btn = create_modern_button(
            btn_frame, 
            text="Изменить", 
            command=self.edit_item, 
            style="warning",
            width=12, 
            tooltip="Редактировать выбранный элемент"
        )
        edit_btn.pack(side=tk.LEFT, padx=2)
        
        del_btn = create_modern_button(
            btn_frame, 
            text="Удалить", 
            command=self.delete_item, 
            style="danger",
            width=12, 
            tooltip="Удалить выбранный элемент"
        )
        del_btn.pack(side=tk.LEFT, padx=2)
        
        bottom_frame = tk.Frame(self.top, padx=12, pady=12, bg=COLORS["bg_tertiary"])
        bottom_frame.pack(fill=tk.X)
        
        cancel_btn = create_modern_button(
            bottom_frame, 
            text="✗ Отмена", 
            command=self.cancel, 
            style="secondary",
            width=100,
            height=40,
            tooltip="Отменить изменения"
        )
        cancel_btn.pack(side=tk.RIGHT, padx=3)
        
        save_btn = create_modern_button(
            bottom_frame, 
            text="✓ Сохранить", 
            command=self.save, 
            style="success",
            width=110,
            height=40,
            tooltip="Сохранить все изменения"
        )
        save_btn.pack(side=tk.RIGHT, padx=3)
    
    def refresh_list(self):
        self.listbox.delete(0, tk.END)
        for item in self.items:
            self.listbox.insert(tk.END, item)
    
    def add_item(self):
        dialog = SimpleInputDialog(self.top, "Добавить", "Введите новое значение:")
        self.top.wait_window(dialog.top)
        new_item = dialog.result
        if new_item and new_item.strip():
            self.items.append(new_item.strip())
            self.refresh_list()
    
    def edit_item(self):
        selection = self.listbox.curselection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите элемент для изменения", parent=self.top)
            return
        
        index = selection[0]
        old_value = self.items[index]
        dialog = SimpleInputDialog(self.top, "Изменить", "Новое значение:", old_value)
        self.top.wait_window(dialog.top)
        new_value = dialog.result
        if new_value and new_value.strip():
            self.items[index] = new_value.strip()
            self.refresh_list()
            self.listbox.selection_set(index)
    
    def delete_item(self):
        selection = self.listbox.curselection()
        if not selection:
            messagebox.showwarning("Предупреждение", "Выберите элемент для удаления", parent=self.top)
            return
        
        if len(self.items) <= 1:
            messagebox.showwarning("Предупреждение", "Нельзя удалить последний элемент", parent=self.top)
            return
        
        index = selection[0]
        if messagebox.askyesno("Подтверждение", f"Удалить '{self.items[index]}'?", parent=self.top):
            del self.items[index]
            self.refresh_list()
    
    def save(self):
        self.result = self.items
        self.top.destroy()
    
    def cancel(self):
        self.result = None
        self.top.destroy()

class PresetSelectionDialog:
    """Диалоговое окно для выбора пресета"""
    def __init__(self, parent, presets_dir="presets"):
        self.result = None
        self.presets_dir = presets_dir
        self.top = tk.Toplevel(parent)
        self.top.withdraw()
        self.top.title("Выбор пресета")
        self.top.geometry("500x450")
        self.top.resizable(False, False)
        self.top.transient(parent)
        
        self.top.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (500 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (450 // 2)
        self.top.geometry(f"+{x}+{y}")
        
        self.top.deiconify()
        self.top.grab_set()
        
        header = tk.Frame(self.top, bg=COLORS["primary"], height=50)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        tk.Label(
            header, 
            text="Загрузка пресета", 
            font=FONTS["title"], 
            bg=COLORS["primary"], 
            fg="white"
        ).pack(pady=12)
        
        main_frame = tk.Frame(self.top, padx=15, pady=15, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        info_label = tk.Label(
            main_frame, 
            text="Выберите пресет для загрузки:", 
            font=FONTS["heading"], 
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        )
        info_label.pack(anchor="w", pady=(0, 10))
        
        # Список пресетов с современным стилем
        list_frame = tk.Frame(
            main_frame, 
            bg=COLORS["card_bg"],
            highlightthickness=1,
            highlightbackground=COLORS["border"],
            highlightcolor=COLORS["border_focus"]
        )
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.listbox = tk.Listbox(
            list_frame, 
            yscrollcommand=scrollbar.set, 
            font=FONTS["body"],
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"],
            selectbackground=COLORS["primary_light"],
            selectforeground="white",
            relief=tk.FLAT,
            borderwidth=0,
            highlightthickness=0,
            selectmode=tk.SINGLE,
            height=12,
            activestyle="none"
        )
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=1, pady=1)
        scrollbar.config(command=self.listbox.yview)
        
        # Двойной клик для загрузки
        self.listbox.bind("<Double-Button-1>", lambda e: self.load_preset())
        
        self.refresh_presets()
        
        btn_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        btn_frame.pack(fill=tk.X, pady=(5, 0))
        
        delete_btn = create_modern_button(
            btn_frame, 
            text="🗑 Удалить", 
            command=self.delete_preset, 
            style="danger",
            width=14, 
            tooltip="Удалить выбранный пресет"
        )
        delete_btn.pack(side=tk.LEFT, padx=2)
        
        bottom_frame = tk.Frame(self.top, padx=15, pady=15, bg=COLORS["bg_tertiary"])
        bottom_frame.pack(fill=tk.X)
        
        cancel_btn = create_modern_button(
            bottom_frame, 
            text="✗ Отмена", 
            command=self.cancel, 
            style="secondary",
            width=110,
            height=40,
            tooltip="Закрыть без загрузки"
        )
        cancel_btn.pack(side=tk.RIGHT, padx=3)
        
        load_btn = create_modern_button(
            bottom_frame, 
            text="✓ Загрузить", 
            command=self.load_preset, 
            style="success",
            width=110,
            height=40,
            tooltip="Загрузить выбранный пресет"
        )
        load_btn.pack(side=tk.RIGHT, padx=3)
    
    def refresh_presets(self):
        """Обновление списка пресетов"""
        self.listbox.delete(0, tk.END)
        
        os.makedirs(self.presets_dir, exist_ok=True)
        
        preset_files = [f for f in os.listdir(self.presets_dir) if f.endswith('.json')]
        
        if not preset_files:
            self.listbox.insert(tk.END, "  (нет сохраненных пресетов)")
            self.listbox.itemconfig(0, fg=COLORS["text_secondary"])
        else:
            for preset_file in sorted(preset_files):
                preset_name = preset_file[:-5]
                self.listbox.insert(tk.END, preset_name)
    
    def get_selected_preset(self):
        """Получение выбранного пресета"""
        selection = self.listbox.curselection()
        if not selection:
            return None
        
        preset_name = self.listbox.get(selection[0])
        if preset_name.startswith("  ("):
            return None
        
        return preset_name
    
    def load_preset(self):
        """Загрузка выбранного пресета"""
        preset_name = self.get_selected_preset()
        if not preset_name:
            messagebox.showwarning(
                "Предупреждение", 
                "Выберите пресет для загрузки", 
                parent=self.top
            )
            return
        
        preset_path = os.path.join(self.presets_dir, f"{preset_name}.json")
        self.result = preset_path
        self.top.destroy()
    
    def delete_preset(self):
        """Удаление выбранного пресета"""
        preset_name = self.get_selected_preset()
        if not preset_name:
            messagebox.showwarning(
                "Предупреждение", 
                "Выберите пресет для удаления", 
                parent=self.top
            )
            return
        
        # Подтверждение удаления
        if not messagebox.askyesno(
            "Подтверждение удаления", 
            f"Вы уверены, что хотите удалить пресет '{preset_name}'?\n\nЭто действие нельзя отменить.",
            parent=self.top
        ):
            return
        
        # Удаляем файл
        preset_path = os.path.join(self.presets_dir, f"{preset_name}.json")
        try:
            os.remove(preset_path)
            self.refresh_presets()
            messagebox.showinfo(
                "Успех", 
                f"Пресет '{preset_name}' успешно удален", 
                parent=self.top
            )
        except Exception as e:
            messagebox.showerror(
                "Ошибка", 
                f"Не удалось удалить пресет:\n{e}", 
                parent=self.top
            )
    
    def cancel(self):
        """Отмена выбора"""
        self.result = None
        self.top.destroy()

class SimpleInputDialog:
    """Простое диалоговое окно для ввода текста"""
    def __init__(self, parent, title, prompt, initial_value="", default_value=""):
        # Поддержка обоих параметров для обратной совместимости
        value = default_value if default_value else initial_value
        
        self.result = None
        self.top = tk.Toplevel(parent)
        self.top.withdraw()
        self.top.title(title)
        self.top.geometry("520x240")
        self.top.resizable(False, False)
        self.top.transient(parent)
        
        self.top.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (520 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (240 // 2)
        self.top.geometry(f"+{x}+{y}")
        
        self.top.deiconify()
        self.top.grab_set()
        
        # Фон окна
        self.top.configure(bg=COLORS["bg_secondary"])
        
        header = tk.Frame(self.top, bg=COLORS["primary"], height=50)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        tk.Label(
            header,
            text=title,
            font=FONTS["title"],
            bg=COLORS["primary"],
            fg="white"
        ).pack(pady=12)
        
        main_frame = tk.Frame(self.top, padx=25, pady=20, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Подсказка
        tk.Label(
            main_frame,
            text=prompt,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            wraplength=460,
            justify=tk.LEFT
        ).pack(anchor="w", pady=(0, 12))
        
        # Обёртка для поля ввода
        entry_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        entry_frame.pack(fill=tk.X, pady=(0, 15))
        
        entry_wrapper = tk.Frame(entry_frame, bg=COLORS["border"], padx=1, pady=1)
        entry_wrapper.pack(fill=tk.X)
        
        self.entry_var = tk.StringVar(value=value)
        self.entry = tk.Entry(
            entry_wrapper,
            textvariable=self.entry_var,
            font=FONTS["body"],
            relief=tk.FLAT,
            borderwidth=0,
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"],
            insertbackground=COLORS["text_primary"]
        )
        self.entry.pack(fill=tk.X, ipady=6, padx=2, pady=2)
        enable_field_shortcuts(self.entry)
        add_context_menu(self.entry)
        self.entry.focus_set()
        self.entry.select_range(0, tk.END)
        
        # Привязываем Enter к OK
        self.entry.bind("<Return>", lambda e: self.ok())
        self.entry.bind("<Escape>", lambda e: self.cancel())
        
        btn_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        btn_frame.pack(fill=tk.X, pady=(8, 0))
        
        cancel_btn = create_modern_button(
            btn_frame,
            text="✗ Отмена",
            command=self.cancel,
            style="secondary",
            width=110,
            height=36,
            tooltip="Отменить изменения"
        )
        cancel_btn.pack(side=tk.RIGHT, padx=5)
        
        ok_btn = create_modern_button(
            btn_frame,
            text="✓ Сохранить",
            command=self.ok,
            style="success",
            width=120,
            height=36,
            tooltip="Подтвердить изменения"
        )
        ok_btn.pack(side=tk.RIGHT, padx=5)
    
    def ok(self):
        self.result = self.entry_var.get().strip()
        self.top.destroy()
    
    def cancel(self):
        self.result = None
        self.top.destroy()

class PerformanceSettingsDialog:
    """Диалоговое окно настроек производительности"""
    def __init__(self, parent, app):
        import platform
        
        self.app = app
        self.top = tk.Toplevel(parent)
        self.top.title("Настройки производительности")
        
        self.top.withdraw()
        
        self.top.geometry("600x550")
        self.top.resizable(False, False)
        self.top.transient(parent)
        
        self.top.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - self.top.winfo_width()) // 2
        y = parent.winfo_y() + (parent.winfo_height() - self.top.winfo_height()) // 2
        self.top.geometry(f"+{x}+{y}")
        
        self.top.deiconify()
        self.top.grab_set()
        
        main_frame = tk.Frame(self.top, padx=20, pady=20, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        title_label = tk.Label(
            main_frame,
            text="⚡ Настройки производительности",
            font=FONTS["title"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        )
        title_label.pack(anchor="w", pady=(0, 15))
        
        cpu_cores = self.app.cpu_cores
        cpu_name = self.get_cpu_name()
        
        info_frame = tk.LabelFrame(
            main_frame,
            text=" 💻 Информация о системе ",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            padx=15,
            pady=12
        )
        info_frame.pack(fill=tk.X, pady=(0, 15))
        
        if not self.app._cpu_info_loaded:
            tk.Label(
                info_frame,
                text="⏳ Определение характеристик процессора...",
                font=FONTS["small"],
                bg=COLORS["bg_secondary"],
                fg="#FFC107"
            ).pack(anchor="w", pady=2)
        
        tk.Label(
            info_frame,
            text=f"Процессор: {cpu_name}",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            wraplength=520,
            justify=tk.LEFT
        ).pack(anchor="w", pady=2)
        
        tk.Label(
            info_frame,
            text=f"Доступно ядер: {cpu_cores}",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        ).pack(anchor="w", pady=2)
        
        tk.Label(
            info_frame,
            text=f"Рекомендуется: {max(1, cpu_cores - 1)} рабочих процессов (оставляет 1 ядро для системы)",
            font=FONTS["small"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_secondary"],
            wraplength=520,
            justify=tk.LEFT
        ).pack(anchor="w", pady=2)
        
        perf_frame = tk.LabelFrame(
            main_frame,
            text=" ⚙️ Количество рабочих процессов ",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            padx=15,
            pady=12
        )
        perf_frame.pack(fill=tk.X, pady=(0, 15))
        
        scale_frame = tk.Frame(perf_frame, bg=COLORS["bg_secondary"])
        scale_frame.pack(fill=tk.X, pady=(5, 10))
        
        self.worker_scale = tk.Scale(
            scale_frame,
            from_=1,
            to=cpu_cores,
            orient=tk.HORIZONTAL,
            variable=self.app.worker_processes,
            length=400,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            highlightthickness=0,
            troughcolor=COLORS["bg_tertiary"],
            command=self.update_info_label
        )
        self.worker_scale.pack(fill=tk.X)
        
        self.info_label = tk.Label(
            perf_frame,
            text="",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["primary"],
            justify=tk.LEFT
        )
        self.info_label.pack(anchor="w", pady=(5, 0))
        
        self.update_info_label()
        
        explain_frame = tk.LabelFrame(
            main_frame,
            text=" 💡 Рекомендации и пояснения ",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            padx=15,
            pady=10
        )
        explain_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        recommendations = [
            "• 1 процесс - последовательная обработка (медленно, минимальная нагрузка)",
            "• Средние значения - баланс скорости и стабильности системы",
            "• Максимум - максимальная скорость (может нагружать систему)",
            "• Рекомендуется оставлять 1-2 ядра для операционной системы"
        ]
        
        for rec in recommendations:
            if rec == "":
                tk.Label(
                    explain_frame,
                    text="",
                    bg=COLORS["bg_secondary"],
                    font=FONTS["small"]
                ).pack(anchor="w")
            else:
                tk.Label(
                    explain_frame,
                    text=rec,
                    font=FONTS["small"],
                    bg=COLORS["bg_secondary"],
                    fg=COLORS["text_secondary"],
                    justify=tk.LEFT,
                    wraplength=520
                ).pack(anchor="w", pady=1)
        
        btn_frame = tk.Frame(self.top, bg=COLORS["bg_tertiary"], padx=20, pady=15)
        btn_frame.pack(fill=tk.X, side=tk.BOTTOM)
        
        ok_btn = create_modern_button(
            btn_frame,
            text="✓ Применить",
            command=self.ok,
            style="success",
            width=110,
            height=40,
            tooltip="Сохранить настройки"
        )
        ok_btn.pack(side=tk.RIGHT, padx=3)
        
        cancel_btn = create_modern_button(
            btn_frame,
            text="✗ Отмена",
            command=self.cancel,
            style="secondary",
            width=100,
            height=40,
            tooltip="Отменить изменения"
        )
        cancel_btn.pack(side=tk.RIGHT, padx=3)
    
    def get_cpu_name(self):
        """
        Получение точного названия процессора из Windows.
        Использует WMI через subprocess или реестр Windows для получения полного названия (например, "Intel Core i7-10700K"),
        если не удается - возвращает fallback значения.
        """
        import platform
        
        cpu_name = ""
        
        if not cpu_name:
            try:
                import subprocess
                result = subprocess.run(
                    ['wmic', 'cpu', 'get', 'name'],
                    capture_output=True,
                    text=True,
                    timeout=3,
                    creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
                )
                if result.returncode == 0:
                    lines = result.stdout.strip().split('\n')
                    if len(lines) > 1:
                        cpu_name = lines[1].strip()
            except:
                pass
        
        if not cpu_name:
            try:
                import winreg
                key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, 
                                    r"HARDWARE\DESCRIPTION\System\CentralProcessor\0")
                cpu_name = winreg.QueryValueEx(key, "ProcessorNameString")[0].strip()
                winreg.CloseKey(key)
            except:
                pass
        
        # Fallback: platform.processor() с дополнением
        if not cpu_name or cpu_name == "":
            cpu_name = platform.processor()
            if not cpu_name or cpu_name.strip() == "":
                cpu_name = "Неизвестный процессор"
        
        return cpu_name
    
    def update_info_label(self, *args):
        """Обновление информационной метки"""
        workers = self.app.worker_processes.get()
        cpu_cores = self.app.cpu_cores  # Используем кэшированное значение
        
        if workers == 1:
            info = "Последовательная обработка (медленно, минимальная нагрузка)"
        elif workers == cpu_cores:
            info = f"Используются все {cpu_cores} ядер (максимальная производительность)"
        else:
            percentage = int((workers / cpu_cores) * 100)
            info = f"Используется {workers} из {cpu_cores} ядер (~{percentage}% мощности процессора)"
        
        self.info_label.config(text=info)
    
    def ok(self):
        """Применение настроек"""
        self.app.save_config()
        self.top.destroy()
    
    def cancel(self):
        """Отмена изменений"""
        self.top.destroy()

# ── КЛАСС ИНДИКАТОРА ЗАГРУЗКИ МОДУЛЕЙ ───────────────────────────────
class LoadingProgressTooltip:
    """Всплывающее окно с прогресс-барами загрузки модулей"""
    def __init__(self, widget, app):
        self.widget = widget
        self.app = app
        self.tooltip_window = None
        
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)
        
        self.modules = {
            'pandas': {'progress': 0, 'label': 'pandas (обработка данных)', 'color': '#4CAF50'},
            'docx': {'progress': 0, 'label': 'python-docx (создание документов)', 'color': '#2196F3'},
            'pymorphy3': {'progress': 0, 'label': 'pymorphy3 (склонение слов)', 'color': '#FF9800'}
        }
        
        self.progress_bars = {}
        self.progress_labels = {}
    
    def show_tooltip(self, event=None):
        """Показать всплывающее окно с прогрессом"""
        if self.tooltip_window:
            return
        
        # Если загрузка завершена, показываем простую подсказку
        if self.app._all_modules_loaded:
            x = self.widget.winfo_rootx() + 10
            y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
            
            self.tooltip_window = tw = tk.Toplevel(self.widget)
            tw.wm_overrideredirect(True)
            tw.wm_withdraw()
            
            label = tk.Label(
                tw,
                text="Все модули загружены и готовы к работе",
                justify=tk.LEFT,
                background="#2C3E50",
                foreground="white",
                relief=tk.SOLID,
                borderwidth=1,
                font=FONTS["small"],
                padx=8,
                pady=4
            )
            label.pack()
            
            tw.wm_geometry(f"+{x}+{y}")
            tw.wm_deiconify()
            return
        
        x = self.widget.winfo_rootx() + 10
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        
        self.tooltip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_withdraw()
        
        # Основной контейнер
        main_frame = tk.Frame(
            tw,
            bg="#2C3E50",
            relief=tk.SOLID,
            borderwidth=1,
            padx=15,
            pady=12
        )
        main_frame.pack()
        
        title = tk.Label(
            main_frame,
            text="📦 Загрузка модулей",
            bg="#2C3E50",
            fg="white",
            font=("Segoe UI", 10, "bold")
        )
        title.pack(anchor=tk.W, pady=(0, 8))
        
        for module_key, module_data in self.modules.items():
            # Контейнер для каждого модуля
            module_frame = tk.Frame(main_frame, bg="#2C3E50")
            module_frame.pack(fill=tk.X, pady=3)
            
            # Название модуля
            name_label = tk.Label(
                module_frame,
                text=module_data['label'],
                bg="#2C3E50",
                fg="#B0BEC5",
                font=FONTS["small"],
                anchor=tk.W,
                width=35
            )
            name_label.pack(side=tk.LEFT)
            
            # Прогресс-бар (Canvas)
            progress_canvas = tk.Canvas(
                module_frame,
                width=150,
                height=16,
                bg="#34495E",
                highlightthickness=0
            )
            progress_canvas.pack(side=tk.LEFT, padx=(5, 5))
            
            percent_label = tk.Label(
                module_frame,
                text="0%",
                bg="#2C3E50",
                fg="white",
                font=("Segoe UI", 8),
                width=5
            )
            percent_label.pack(side=tk.LEFT)
            
            self.progress_bars[module_key] = progress_canvas
            self.progress_labels[module_key] = percent_label
            
            # Рисуем начальное состояние
            self.update_progress_bar(module_key, module_data['progress'])
        
        tw.wm_geometry(f"+{x}+{y}")
        tw.wm_deiconify()
    
    def update_progress_bar(self, module_key, progress):
        """Обновить прогресс-бар конкретного модуля"""
        if module_key not in self.progress_bars:
            return
        
        self.modules[module_key]['progress'] = progress
        canvas = self.progress_bars[module_key]
        label = self.progress_labels[module_key]
        color = self.modules[module_key]['color']
        
        canvas.delete("all")
        
        canvas.create_rectangle(0, 0, 150, 16, fill="#34495E", outline="")
        
        width = int(150 * (progress / 100))
        if width > 0:
            canvas.create_rectangle(0, 0, width, 16, fill=color, outline="")
        
            lighter = self.lighten_color(color, 1.2)
            canvas.create_rectangle(0, 0, width, 4, fill=lighter, outline="")
        
        label.config(text=f"{progress}%")
    
    def lighten_color(self, color, factor=1.2):
        """Осветлить цвет для эффекта градиента"""
        # Простая функция для осветления hex цвета
        color = color.lstrip('#')
        rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
        rgb = tuple(min(255, int(c * factor)) for c in rgb)
        return '#{:02x}{:02x}{:02x}'.format(*rgb)
    
    def hide_tooltip(self, event=None):
        """Скрыть всплывающее окно"""
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None
            self.progress_bars.clear()
            self.progress_labels.clear()
    
    def update_all(self):
        """Обновить все прогресс-бары (вызывается из фонового потока)"""
        for module_key, module_data in self.modules.items():
            self.update_progress_bar(module_key, module_data['progress'])

class GenerationDocApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор документов • 2026")
        
        # Современные размеры окна и минимальные размеры
        window_width = 800
        window_height = 900
        
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        self.root.minsize(800, 900)
        self.root.resizable(True, True)
        
        try:
            if os.path.exists("doc.ico"):
                self.root.iconbitmap("doc.ico")
        except:
            pass
        
        # Современный фон окна
        self.root.configure(bg=COLORS["bg_secondary"])
        
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        self.load_config()
        
        # Отложенная инициализация морфологического анализатора (lazy loading)
        self._morph = None
        
        # Флаги загрузки модулей
        self._pandas_loaded = False
        self._docx_loaded = False
        self._pymorphy_loaded = False
        self._modules_loading = False
        self._all_modules_loaded = False
        
        # Кэш для загруженных модулей
        self._pandas = None
        self._Document = None
        self._Pt = None
        self._WD_BREAK = None
        
        # Настройки производительности
        self.cpu_cores = 2
        default_workers = 1
        self.worker_processes = tk.IntVar(value=default_workers)
        self.max_workers = 4
        self._cpu_info_loaded = False
        
        # Запоминание последних директорий
        self.last_excel_dir = os.getcwd()
        self.last_word_dir = os.getcwd()
        self.last_output_dir = os.getcwd()
        
        # Система вкладок
        self.tabs = []  # Список объектов TabTask
        self.tab_counter = 0  # Счетчик для уникальных ID вкладок
        self.max_tabs = 5  # Максимальное количество вкладок
        
        # Текущий загруженный пресет
        self.current_preset_name = None
        
        self.create_widgets()
        
        self.add_tab()
        
        # Запускаем фоновую загрузку модулей
        self.start_background_loading()
        
    def load_config(self):
        """Загрузка конфигурации из файла"""
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.PLACEHOLDERS = config.get("placeholders", [])
                    self.CUSTOM_LISTS = config.get("custom_lists", {})
                    
                    self.last_excel_dir = config.get("last_excel_dir", os.getcwd())
                    self.last_word_dir = config.get("last_word_dir", os.getcwd())
                    self.last_output_dir = config.get("last_output_dir", os.getcwd())
                    
                    saved_workers = config.get("worker_processes", None)
                    if saved_workers is not None and hasattr(self, 'worker_processes'):
                        if 1 <= saved_workers <= 32:
                            self.worker_processes.set(saved_workers)
                    
                    for ph in self.PLACEHOLDERS:
                        if "apply_genitive" in ph and "case" not in ph:
                            ph["case"] = "gent" if ph["apply_genitive"] else "nomn"
                            del ph["apply_genitive"]
            except:
                self.load_defaults()
        else:
            self.load_defaults()
    
    @property
    def morph(self):
        """Lazy loading для pymorphy3 - инициализация только при первом обращении"""
        if self._morph is None:
            import pymorphy3
            self._morph = pymorphy3.MorphAnalyzer()
        return self._morph
    
    def load_defaults(self):
        """Загрузка значений по умолчанию (пустые списки)"""
        self.PLACEHOLDERS = []
        self.CUSTOM_LISTS = {}
    
    def start_background_loading(self):
        """Запуск фоновой загрузки модулей"""
        self._modules_loading = True
        loading_thread = threading.Thread(target=self.load_modules_background, daemon=True)
        loading_thread.start()
    
    def load_modules_background(self):
        """Фоновая загрузка тяжелых модулей с обновлением прогресса"""
        try:
            import time
            _ensure_concurrent_imports()
            
            total_progress = 0
            
            try:
                self.cpu_cores = multiprocessing.cpu_count()
                self.max_workers = self.cpu_cores
                
                def update_workers():
                    saved_workers = self.worker_processes.get()
                    if saved_workers > self.max_workers:
                        self.worker_processes.set(self.max_workers)
                    elif saved_workers == 1 and self.cpu_cores > 1:
                        optimal_workers = max(1, self.cpu_cores - 1)
                        self.worker_processes.set(optimal_workers)
                
                self.root.after(0, update_workers)
                self._cpu_info_loaded = True
            except:
                self.cpu_cores = 2
                self.max_workers = 4
            
            self.update_module_progress('pandas', 0)
            time.sleep(0.1)  # Небольшая задержка для визуализации
            self.update_module_progress('pandas', 30)
            import pandas as pd
            self.update_module_progress('pandas', 70)
            self._pandas = pd
            self._pandas_loaded = True
            self.update_module_progress('pandas', 100)
            total_progress += 40
            self.update_loading_status(total_progress)
            time.sleep(0.05)
            
            self.update_module_progress('docx', 0)
            time.sleep(0.1)
            self.update_module_progress('docx', 40)
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_BREAK
            self._Document = Document
            self._Pt = Pt
            self._WD_BREAK = WD_BREAK
            self._docx_loaded = True
            self.update_module_progress('docx', 100)
            total_progress += 30
            self.update_loading_status(total_progress)
            time.sleep(0.05)
            
            self.update_module_progress('pymorphy3', 0)
            time.sleep(0.1)
            self.update_module_progress('pymorphy3', 20)
            import pymorphy3
            self.update_module_progress('pymorphy3', 50)
            time.sleep(0.05)
            self._morph = pymorphy3.MorphAnalyzer()
            self._pymorphy_loaded = True
            self.update_module_progress('pymorphy3', 100)
            total_progress = 100
            self.update_loading_status(total_progress)
            
            self._all_modules_loaded = True
            self._modules_loading = False
            
            self.root.after(0, self.on_loading_complete)
            
        except Exception as e:
            self._modules_loading = False
    
    def update_module_progress(self, module_key, progress):
        """Обновить прогресс конкретного модуля"""
        def update():
            if hasattr(self, 'loading_tooltip'):
                self.loading_tooltip.modules[module_key]['progress'] = progress
                self.loading_tooltip.update_progress_bar(module_key, progress)
        
        self.root.after(0, update)
    
    def update_loading_status(self, total_progress):
        """Обновить общий статус загрузки"""
        def update():
            if total_progress < 100:
                self.loading_label.config(
                    text=f"⏳ Загрузка {total_progress}%",
                    fg=COLORS["warning"]
                )
            else:
                self.loading_label.config(
                    text="✓ Готово",
                    fg=COLORS["success"]
                )
        
        self.root.after(0, update)
    
    def on_loading_complete(self):
        """Вызывается после завершения загрузки всех модулей"""
        # Изменяем текст и цвет индикатора
        self.loading_label.config(
            text="✓ Готово",
            fg=COLORS["success"],
            cursor="hand2"
        )
        # События наведения остаются активными для показа информации о загруженных модулях
    
    def save_config(self):
        """Сохранение конфигурации в файл"""
        # Читаем существующий конфиг чтобы сохранить excel_presets
        config = {}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            except:
                pass
        
        # Обновляем только нужные поля, сохраняя остальные (включая excel_presets)
        config.update({
            "placeholders": self.PLACEHOLDERS,
            "custom_lists": self.CUSTOM_LISTS,
            "last_excel_dir": self.last_excel_dir,
            "last_word_dir": self.last_word_dir,
            "last_output_dir": self.last_output_dir,
            "worker_processes": self.worker_processes.get()
        })
        
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
    
    def save_preset(self):
        """Сохранение текущих настроек в пресет"""
        current_tab = self.get_current_tab()
        if not current_tab:
            messagebox.showwarning("Предупреждение", "Нет активной вкладки")
            return
        
        # Запрашиваем имя пресета
        dialog = SimpleInputDialog(
            self.root,
            "Сохранить пресет",
            "Введите название пресета:"
        )
        self.root.wait_window(dialog.top)
        preset_name = dialog.result
        
        if not preset_name:
            return
        
        # Собираем текущие значения выпадающих списков
        dropdown_values = {}
        for key, var in current_tab.custom_list_vars.items():
            dropdown_values[key] = var.get()
        
        preset = {
            "placeholders": self.PLACEHOLDERS,
            "custom_lists": self.CUSTOM_LISTS,
            "output_folder": current_tab.output_folder.get(),
            "filename_base": current_tab.filename_base.get(),
            "filename_pattern": current_tab.filename_pattern.get(),
            "filename_mode": current_tab.filename_mode.get(),
            "filename_column": current_tab.filename_column.get(),
            # Дату не сохраняем - всегда используется актуальная
            "dropdown_values": dropdown_values
        }
        
        presets_dir = "presets"
        os.makedirs(presets_dir, exist_ok=True)
        
        preset_filename = f"{preset_name}.json"
        preset_path = os.path.join(presets_dir, preset_filename)
        
        try:
            with open(preset_path, 'w', encoding='utf-8') as f:
                json.dump(preset, f, ensure_ascii=False, indent=2)
            
            # Обновляем имя текущего пресета и индикатор
            self.current_preset_name = preset_name
            self.update_preset_indicator()
            
            current_tab.log(f"✓ Пресет '{preset_name}' успешно сохранен")
            messagebox.showinfo("Успех", f"Пресет '{preset_name}' сохранен!")
        except Exception as e:
            current_tab.log(f"❌ Ошибка сохранения пресета: {e}")
            messagebox.showerror("Ошибка", f"Не удалось сохранить пресет:\n{e}")
    
    def load_preset(self):
        """Загрузка пресета из файла"""
        current_tab = self.get_current_tab()
        if not current_tab:
            messagebox.showwarning("Предупреждение", "Нет активной вкладки")
            return
        
        presets_dir = "presets"
        os.makedirs(presets_dir, exist_ok=True)
        
        dialog = PresetSelectionDialog(self.root, presets_dir)
        self.root.wait_window(dialog.top)
        
        preset_path = dialog.result
        if not preset_path:
            return
        
        try:
            with open(preset_path, 'r', encoding='utf-8') as f:
                preset = json.load(f)
            
            self.PLACEHOLDERS = preset.get("placeholders", [])
            self.CUSTOM_LISTS = preset.get("custom_lists", {})
            
            current_tab.output_folder.set(preset.get("output_folder", "документы"))
            current_tab.filename_base.set(preset.get("filename_base", "документ"))
            current_tab.filename_pattern.set(preset.get("filename_pattern", "наименование {i:04d}{suffix}.docx"))
            current_tab.filename_mode.set(preset.get("filename_mode", "standard"))
            current_tab.filename_column.set(preset.get("filename_column", ""))
            
            current_tab.update_filename_pattern()
            
            self.refresh_all_tabs_custom_lists()
            
            dropdown_values = preset.get("dropdown_values", {})
            for key, value in dropdown_values.items():
                if key in current_tab.custom_list_vars:
                    combo = current_tab.custom_list_combos[key]
                    if value in combo['values']:
                        current_tab.custom_list_vars[key].set(value)
            
            self.save_config()
            
            preset_name = os.path.basename(preset_path).replace('.json', '')
            
            # Сохраняем имя текущего пресета и обновляем индикатор
            self.current_preset_name = preset_name
            self.update_preset_indicator()
            
            current_tab.log(f"✓ Пресет '{preset_name}' успешно загружен")
            messagebox.showinfo("Успех", f"Пресет '{preset_name}' загружен!")
            
        except Exception as e:
            current_tab.log(f"❌ Ошибка загрузки пресета: {e}")
            messagebox.showerror("Ошибка", f"Не удалось загрузить пресет:\n{e}")
    
    def update_preset_indicator(self):
        """Обновление индикатора текущего пресета"""
        if self.current_preset_name:
            # Показываем индикатор с именем пресета
            self.preset_indicator_label.config(
                text=f"📋 {self.current_preset_name}",
                cursor="hand2"
            )
            self.preset_indicator_frame.pack(side=tk.LEFT, padx=SPACING["sm"])
            
            # Добавляем tooltip с дополнительной информацией
            tooltip_text = f"Текущий пресет: {self.current_preset_name}\n\n💡 Нажмите '💾' чтобы сохранить изменения\n💡 Нажмите '📂' чтобы загрузить другой пресет\n💡 Клик для сброса индикатора"
            try:
                if hasattr(self, '_preset_tooltip'):
                    self._preset_tooltip.text = tooltip_text
                else:
                    self._preset_tooltip = ToolTip(self.preset_indicator_label, tooltip_text)
            except:
                pass
            
            # Добавляем обработчик клика для сброса индикатора
            def on_click(event):
                if messagebox.askyesno(
                    "Сброс индикатора пресета",
                    f"Вы работаете с пресетом '{self.current_preset_name}'.\n\nСбросить индикатор?\n(Это не изменит текущие настройки)",
                    parent=self.root
                ):
                    self.clear_preset_indicator()
            
            # Удаляем старые привязки и добавляем новую
            self.preset_indicator_label.unbind("<Button-1>")
            self.preset_indicator_label.bind("<Button-1>", on_click)
        else:
            # Скрываем индикатор если пресет не загружен
            self.preset_indicator_frame.pack_forget()
    
    def clear_preset_indicator(self):
        """Очистка индикатора пресета"""
        self.current_preset_name = None
        self.update_preset_indicator()
        
    def create_widgets(self):
        """Создание современного интерфейса с системой вкладок"""
        
        # ═══════════════════════════════════════════════════════════
        # СОВРЕМЕННАЯ ВЕРХНЯЯ ПАНЕЛЬ С ГРАДИЕНТОМ
        # ═══════════════════════════════════════════════════════════
        
        # Основной заголовок с градиентом
        header_frame = tk.Frame(self.root, bg=COLORS["primary"], height=70)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Левая часть - заголовок и статус
        left_section = tk.Frame(header_frame, bg=COLORS["primary"])
        left_section.pack(side=tk.LEFT, fill=tk.Y, padx=SPACING["xl"])
        
        title_label = tk.Label(
            left_section, 
            text="Генератор документов", 
            font=FONTS["title"],
            bg=COLORS["primary"],
            fg="white"
        )
        title_label.pack(anchor=tk.W, pady=(SPACING["md"], 0))
        
        # Подзаголовок с версией и статусом
        subtitle_frame = tk.Frame(left_section, bg=COLORS["primary"])
        subtitle_frame.pack(anchor=tk.W)
        
        version_label = tk.Label(
            subtitle_frame,
            text="5.0 • 2026",
            font=FONTS["small"],
            bg=COLORS["primary"],
            fg=COLORS["accent_light"]
        )
        version_label.pack(side=tk.LEFT, padx=(0, SPACING["md"]))
        
        # Индикатор загрузки модулей
        self.loading_label = tk.Label(
            subtitle_frame,
            text="⏳ Загрузка 0%",
            font=FONTS["small"],
            bg=COLORS["primary"],
            fg=COLORS["warning"],
            cursor="hand2"
        )
        self.loading_label.pack(side=tk.LEFT)
        
        self.loading_tooltip = LoadingProgressTooltip(self.loading_label, self)
        
        # Индикатор текущего пресета (в левой части заголовка)
        self.preset_indicator_frame = tk.Frame(subtitle_frame, bg=COLORS["primary"])
        self.preset_indicator_frame.pack(side=tk.LEFT, padx=SPACING["md"])
        
        self.preset_indicator_label = tk.Label(
            self.preset_indicator_frame,
            text="",
            font=FONTS["small"],
            bg=COLORS["primary"],
            fg=COLORS["accent_light"],
            padx=SPACING["xs"],
            pady=2
        )
        self.preset_indicator_label.pack()
        
        # Изначально скрываем индикатор
        self.preset_indicator_frame.pack_forget()
        
        # Правая часть - кнопки и автор
        right_section = tk.Frame(header_frame, bg=COLORS["primary"])
        right_section.pack(side=tk.RIGHT, fill=tk.Y, padx=SPACING["xl"])
        
        author_label = tk.Label(
            right_section,
            text="by Канатов М.Э.",
            font=FONTS["tiny"],
            bg=COLORS["primary"],
            fg=COLORS["accent_light"]
        )
        author_label.pack(anchor=tk.E, pady=(SPACING["sm"], 0))
        
        buttons_frame = tk.Frame(right_section, bg=COLORS["primary"])
        buttons_frame.pack(anchor=tk.E, pady=(SPACING["xs"], 0))
        
        save_preset_btn = create_icon_button(
            buttons_frame, 
            icon="💾", 
            command=self.save_preset, 
            tooltip="Сохранить пресет настроек"
        )
        save_preset_btn.pack(side=tk.LEFT, padx=2)
        
        load_preset_btn = create_icon_button(
            buttons_frame, 
            icon="📂", 
            command=self.load_preset, 
            tooltip="Загрузить пресет"
        )
        load_preset_btn.pack(side=tk.LEFT, padx=2)
        
        constructor_btn = create_icon_button(
            buttons_frame, 
            icon="🔧", 
            command=self.open_excel_constructor, 
            tooltip="Конструктор Excel файлов"
        )
        constructor_btn.pack(side=tk.LEFT, padx=2)
        
        merge_btn = create_icon_button(
            buttons_frame, 
            icon="📄", 
            command=self.open_merge_window, 
            tooltip="Объединение документов"
        )
        merge_btn.pack(side=tk.LEFT, padx=2)
        
        settings_btn = create_icon_button(
            buttons_frame, 
            icon="⚙", 
            command=self.edit_placeholders, 
            tooltip="Настройки плейсхолдеров"
        )
        settings_btn.pack(side=tk.LEFT, padx=2)
        
        perf_btn = create_icon_button(
            buttons_frame, 
            icon="⚡", 
            command=self.open_performance_settings, 
            tooltip="Производительность"
        )
        perf_btn.pack(side=tk.LEFT, padx=2)
        
        # ═══════════════════════════════════════════════════════════
        # ОСНОВНАЯ ОБЛАСТЬ КОНТЕНТА
        # ═══════════════════════════════════════════════════════════
        
        # Контейнер для вкладок
        tabs_container = tk.Frame(self.root, bg=COLORS["bg_secondary"])
        tabs_container.pack(fill=tk.BOTH, expand=True, padx=SPACING["lg"], pady=SPACING["md"])
        
        # Панель управления вкладками
        tabs_control = tk.Frame(tabs_container, bg=COLORS["bg_secondary"], height=48)
        tabs_control.pack(fill=tk.X, pady=(0, SPACING["md"]))
        tabs_control.pack_propagate(False)
        
        # Левая часть - кнопка добавления и счетчик
        control_left = tk.Frame(tabs_control, bg=COLORS["bg_secondary"])
        control_left.pack(side=tk.LEFT, fill=tk.Y, pady=SPACING["sm"])
        
        # Современная кнопка добавления вкладки
        self.add_tab_btn = create_icon_button(
            control_left,
            icon="➕",
            command=self.add_tab,
            tooltip="Добавить новую задачу",
            style="success",
            width=32,
            height=32
        )
        self.add_tab_btn.pack(side=tk.LEFT, padx=(0, SPACING["sm"]))
        
        # Счетчик вкладок
        self.tab_count_label = tk.Label(
            control_left,
            text=f"0 / {self.max_tabs}",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_secondary"],
            cursor="hand2",
            padx=SPACING["lg"]
        )
        self.tab_count_label.pack(side=tk.LEFT)
        
        self.tab_status_tooltip = TabStatusTooltip(self.tab_count_label, self)
        
        # ═══════════════════════════════════════════════════════════
        # CUSTOM BUTTON TABS (КАК В PLACEHOLDEREDITOR)
        # ═══════════════════════════════════════════════════════════
        
        # Фрейм для кнопок-вкладок (горизонтальный список)
        self.tabs_buttons_frame = tk.Frame(
            tabs_container,
            bg=COLORS["bg_secondary"],
            height=38
        )
        self.tabs_buttons_frame.pack(fill=tk.X, pady=(0, SPACING["sm"]))
        self.tabs_buttons_frame.pack_propagate(False)
        
        # Фрейм для контента вкладок
        self.tabs_content_frame = tk.Frame(
            tabs_container,
            bg=COLORS["bg_secondary"]
        )
        self.tabs_content_frame.pack(fill=tk.BOTH, expand=True)
        
        # Списки для хранения кнопок и фреймов
        self.tab_buttons = []
        self.tab_frames = []
        self.current_tab_index = 0  # Индекс активной вкладки
    
    def add_tab(self):
        """Добавление новой вкладки"""
        if len(self.tabs) >= self.max_tabs:
            messagebox.showinfo(
                "Ограничение", 
                f"Достигнуто максимальное количество вкладок ({self.max_tabs})"
            )
            return
        
        tab_frame = tk.Frame(self.tabs_content_frame, bg=COLORS["bg_secondary"])
        self.tab_frames.append(tab_frame)
        
        self.tab_counter += 1
        tab = TabTask(tab_frame, self, self.tab_counter)
        self.tabs.append(tab)
        
        tab.tab_name = f"Задача {self.tab_counter}"
        
        self.create_tab_button(tab.tab_name, len(self.tabs) - 1)
        
        self.switch_tab(len(self.tabs) - 1)
        
        self.update_tab_controls()
    
    def close_tab(self, tab_index=None):
        """Закрытие вкладки"""
        if tab_index is None:
            tab_index = self.current_tab_index
        
        if len(self.tabs) <= 1:
            messagebox.showinfo("Информация", "Нельзя закрыть последнюю вкладку")
            return
        
        tab = self.tabs[tab_index]
        if tab.is_processing:
            result = messagebox.askyesno(
                "Подтверждение",
                "На этой вкладке выполняется задача. Закрыть вкладку?"
            )
            if not result:
                return
        
        # Удаляем кнопку вкладки
        self.tab_buttons[tab_index][0].destroy()
        del self.tab_buttons[tab_index]
        
        # Удаляем фрейм контента
        self.tab_frames[tab_index].destroy()
        del self.tab_frames[tab_index]
        
        # Удаляем объект вкладки
        del self.tabs[tab_index]
        
        # Переключаемся на соседнюю вкладку
        if tab_index >= len(self.tabs):
            tab_index = len(self.tabs) - 1
        self.switch_tab(tab_index)
        
        self.update_tab_controls()
    
    def create_tab_button(self, tab_name, tab_index):
        """Создание кнопки вкладки с кнопкой закрытия"""
        btn_container = tk.Frame(
            self.tabs_buttons_frame,
            bg=COLORS["bg_tertiary"],
            highlightthickness=1,
            highlightbackground=COLORS["border"]
        )
        btn_container.pack(side=tk.LEFT, padx=2)
        
        tab_btn = tk.Button(
            btn_container,
            text=tab_name,
            font=FONTS["button"],
            bg=COLORS["bg_tertiary"],
            activebackground=COLORS["primary_light"],
            fg=COLORS["text_primary"],
            activeforeground="white",
            relief=tk.FLAT,
            cursor="hand2",
            padx=10,
            pady=5,
            command=lambda: self.switch_tab(tab_index)
        )
        tab_btn.pack(side=tk.LEFT)
        
        close_btn = tk.Button(
            btn_container,
            text="✕",
            font=("Segoe UI", 8, "bold"),
            bg=COLORS["bg_tertiary"],
            activebackground=COLORS["danger"],
            fg=COLORS["text_secondary"],
            activeforeground="white",
            relief=tk.FLAT,
            cursor="hand2",
            padx=5,
            pady=5,
            command=lambda: self.close_tab(tab_index)
        )
        close_btn.pack(side=tk.LEFT)
        ToolTip(close_btn, "Закрыть вкладку")
        
        def show_context_menu(event):
            current_index = None
            for i, (cont, btn, cls) in enumerate(self.tab_buttons):
                if btn == tab_btn:
                    current_index = i
                    break
            
            if current_index is None:
                return
            
            menu = ModernContextMenu(self.root)
            menu.add_command(
                label="Переименовать вкладку",
                command=lambda: self.rename_tab(current_index)
            )
            menu.add_separator()
            menu.add_command(
                label="Закрыть вкладку",
                command=lambda: self.close_tab(current_index)
            )
            menu.post(event.x_root, event.y_root)
        
        tab_btn.bind("<Button-3>", show_context_menu)
        close_btn.bind("<Button-2>", lambda e: self.close_tab(tab_index))  # Средняя кнопка мыши
        
        self.tab_buttons.append((btn_container, tab_btn, close_btn))
    
    def switch_tab(self, tab_index):
        """Переключение между вкладками"""
        if tab_index < 0 or tab_index >= len(self.tabs):
            return
        
        self.current_tab_index = tab_index
        
        for i, (container, tab_btn, close_btn) in enumerate(self.tab_buttons):
            if i == tab_index:
                # Активная вкладка - синяя
                container.configure(
                    bg=COLORS["primary"],
                    highlightbackground=COLORS["primary"]
                )
                tab_btn.configure(
                    bg=COLORS["primary"],
                    fg="white",
                    font=("Segoe UI", 10, "bold")
                )
                close_btn.configure(
                    bg=COLORS["primary"],
                    fg="white"
                )
            else:
                # Неактивная вкладка - серая
                container.configure(
                    bg=COLORS["bg_tertiary"],
                    highlightbackground=COLORS["border"]
                )
                tab_btn.configure(
                    bg=COLORS["bg_tertiary"],
                    fg=COLORS["text_primary"],
                    font=FONTS["button"]
                )
                close_btn.configure(
                    bg=COLORS["bg_tertiary"],
                    fg=COLORS["text_secondary"]
                )
        
        for i, frame in enumerate(self.tab_frames):
            if i == tab_index:
                frame.pack(fill=tk.BOTH, expand=True)
            else:
                frame.pack_forget()
    
    def rename_tab(self, tab_index):
        """Переименование вкладки"""
        if tab_index < 0 or tab_index >= len(self.tabs):
            return
        
        tab = self.tabs[tab_index]
        current_name = tab.tab_name
        
        dialog = SimpleInputDialog(
            self.root,
            "Переименовать вкладку",
            "Введите новое название вкладки:",
            default_value=current_name
        )
        self.root.wait_window(dialog.top)
        
        if dialog.result and dialog.result.strip():
            new_name = dialog.result.strip()
            tab.tab_name = new_name
            container, tab_btn, close_btn = self.tab_buttons[tab_index]
            tab_btn.configure(text=new_name)
    
    def update_tab_controls(self):
        """Обновление счетчика и состояния кнопки добавления"""
        count = len(self.tabs)
        self.tab_count_label.config(text=f"{count} / {self.max_tabs}")
        
        # Отключаем кнопку, если достигнут лимит
        if count >= self.max_tabs:
            self.add_tab_btn.configure(state="disabled", fg_color=COLORS["text_disabled"])
        else:
            self.add_tab_btn.configure(state="normal", fg_color=COLORS["success"])
    
    def get_current_tab(self):
        """Получение текущей активной вкладки"""
        try:
            return self.tabs[self.current_tab_index]
        except:
            return None
    
    def refresh_all_tabs_custom_lists(self):
        """Обновление выпадающих списков во всех вкладках"""
        for tab in self.tabs:
            tab.refresh_custom_list_widgets()
    
    def edit_list(self, list_key):
        """Открыть диалог редактирования списка"""
        if list_key not in self.CUSTOM_LISTS:
            return
        
        list_data = self.CUSTOM_LISTS[list_key]
        if isinstance(list_data, dict):
            title = f"Редактирование: {list_data.get('display_name', list_key)}"
            items = list_data.get('values', [])
        else:
            title = f"Редактирование: {list_key}"
            items = list_data
        
        dialog = EditListDialog(self.root, title, items)
        self.root.wait_window(dialog.top)
        
        if dialog.result is not None:
            if isinstance(self.CUSTOM_LISTS[list_key], dict):
                self.CUSTOM_LISTS[list_key]['values'] = dialog.result
            else:
                self.CUSTOM_LISTS[list_key] = dialog.result
            
            self.refresh_all_tabs_custom_lists()
            
            self.save_config()
            
            current_tab = self.get_current_tab()
            if current_tab:
                current_tab.log(f"✓ Список '{title}' обновлен")
    
    def edit_placeholders(self):
        """Открыть диалог редактирования плейсхолдеров"""
        dialog = PlaceholderEditorDialog(self.root, self.PLACEHOLDERS, self.CUSTOM_LISTS)
        self.root.wait_window(dialog.top)
        
        if dialog.result is not None:
            self.PLACEHOLDERS = dialog.result
            if dialog.lists_result is not None:
                self.CUSTOM_LISTS = dialog.lists_result
                self.refresh_all_tabs_custom_lists()
            self.save_config()
            
            # Логируем в текущую вкладку
            current_tab = self.get_current_tab()
            if current_tab:
                current_tab.log("✓ Настройки обновлены")
    
    def open_performance_settings(self):
        """Открыть окно настроек производительности"""
        PerformanceSettingsDialog(self.root, self)
    
    def log(self, message):
        """Логирование в текущую активную вкладку"""
        current_tab = self.get_current_tab()
        if current_tab:
            current_tab.log(message)
    
    def on_closing(self):
        """Обработчик закрытия главного окна"""
        try:
            active_tasks = [tab for tab in self.tabs if tab.is_processing]
            if active_tasks:
                result = messagebox.askyesno(
                    "Подтверждение",
                    f"Выполняется {len(active_tasks)} задач(и). Закрыть приложение?"
                )
                if not result:
                    return
            
            self.save_config()
            
            word_preload_manager.stop()
            
            for widget in self.root.winfo_children():
                if isinstance(widget, tk.Toplevel):
                    try:
                        widget.destroy()
                    except:
                        pass
            
            if WIN32COM_AVAILABLE:
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                except:
                    pass
            
            self.root.quit()
            self.root.destroy()
            
        except Exception as e:
            try:
                self.root.destroy()
            except:
                pass
        finally:
            import sys
            sys.exit(0)
    
    def process_documents_for_tab(self, tab):
        """Обработка документов для конкретной вкладки"""
        # Гарантируем что concurrent импорты доступны для multiprocessing
        _ensure_concurrent_imports()
        
        # Используем кэшированный pandas если загружен
        if self._pandas_loaded and self._pandas:
            pd = self._pandas
        else:
            import pandas as pd
        
        try:
            tab.log("\n" + "═" * 60)
            tab.log("🚀 НАЧАЛО ОБРАБОТКИ ДОКУМЕНТОВ")
            tab.log("═" * 60)
            
            excel_file = tab.excel_path.get()
            word_template = tab.word_template_path.get()
            excel_template = tab.excel_template_path.get()
            
            # Определяем какие шаблоны использовать
            use_word = bool(word_template)
            use_excel = bool(excel_template)
            
            if use_word and use_excel:
                tab.log(f"\n📝 Будут созданы документы обоих типов: Word и Excel")
            elif use_word:
                tab.log(f"\n📝 Тип документа: Word")
            else:
                tab.log(f"\n📊 Тип документа: Excel")
            
            output_folder = tab.output_folder.get()
            num_workers = self.worker_processes.get()
            
            # Определяем абсолютный путь к папке сохранения
            if not os.path.isabs(output_folder):
                output_folder = os.path.join(os.getcwd(), output_folder)
            
            os.makedirs(output_folder, exist_ok=True)
            
            tab.log(f"\n📊 Чтение Excel файла:")
            tab.log(f"   {excel_file}")
            df = pd.read_excel(excel_file, engine='openpyxl')
            
            tab.log(f"   ✓ Прочитано строк: {len(df)}")
            
            # Определяем колонки с датами по заголовкам поддерживаем колонки с датами по заголовкам
            date_columns = [col for col in df.columns if self.is_date_column(col)]
            if date_columns:
                tab.log(f"\n📅 Колонки с датами: {', '.join(date_columns)}")
            
            tab.log(f"\n📝 Используемые шаблоны:")
            if use_word:
                tab.log(f"   Word: {word_template}")
            if use_excel:
                tab.log(f"   Excel: {excel_template}")
            
            tab.log(f"\n⚡ Режим производительности:")
            tab.log(f"   Рабочих процессов: {num_workers}")
            if num_workers > 1:
                tab.log(f"   Параллельная обработка включена!")
            
            # Определяем обязательные столбцы Excel
            required_excel_columns = [ph["source_value"] for ph in self.PLACEHOLDERS 
                                     if ph["source_type"] == "excel" and ph["required"] and ph.get("active", True)]
            
            # Проверка столбцов
            missing = [col for col in required_excel_columns if col not in df.columns]
            if missing:
                tab.log(f"\n❌ КРИТИЧЕСКАЯ ОШИБКА: Отсутствуют обязательные столбцы:")
                for col in missing:
                    tab.log(f"   • {col}")
                self.root.after(0, lambda: messagebox.showerror("Ошибка", f"Отсутствуют столбцы в Excel:\n{', '.join(missing)}"))
                tab.is_processing = False
                tab.start_btn.configure(state="normal", text="▶ Начать обработку")
                return
            
            if required_excel_columns:
                tab.log(f"\n✓ Проверка обязательных столбцов пройдена ({len(required_excel_columns)} шт.)")
            
            # === ПОДГОТОВКА ДАННЫХ ДЛЯ ПАРАЛЛЕЛЬНОЙ ОБРАБОТКИ ===
            tab.log(f"\n🔄 Подготовка данных для обработки...")
            
            # Подготавливаем данные для каждой строки
            tasks = []
            for i in range(len(df)):
                row = df.iloc[i]
                row_data = {}
                
                # Копируем данные из Excel
                for col in df.columns:
                    value = row[col]
                    # Форматируем даты только в колонках с датами
                    if col in date_columns and pd.notna(value):
                        row_data[col] = self.to_date(value)
                    else:
                        row_data[col] = value
                
                # Логируем заголовок строки
                tab.log(f"\n" + "─" * 60)
                tab.log(f"📄 Обработка строки {i + 1} из {len(df)}")
                tab.log("─" * 60)
                tab.log(f"🔄 Обработка плейсхолдеров:")
                
                for ph in self.PLACEHOLDERS:
                    if not ph.get("active", True):
                        continue
                    
                    value = ""
                    original_value = ""
                    
                    if ph["source_type"] == "excel":
                        value = row.get(ph["source_value"], "")
                        if pd.isna(value):
                            value = ""
                        else:
                            value = str(value).strip()
                        original_value = value
                    elif ph["source_type"] == "dropdown":
                        dropdown_key = ph["source_value"]
                        if dropdown_key in tab.custom_list_vars:
                            value = tab.custom_list_vars[dropdown_key].get()
                        original_value = value
                    elif ph["source_type"] == "date":
                        value = tab.selected_date.get_date().strftime('%d.%m.%Y')
                        original_value = value
                    elif ph["source_type"] == "static":
                        value = ph["source_value"]
                        original_value = value
                    
                    # Применяем падеж
                    ph_case = ph.get("case", "nomn")
                    case_name = RUSSIAN_CASES.get(ph_case, "Именительный").split(" ")[0]
                    
                    if ph_case != "nomn" and value:
                        transformed_value = self.apply_case(value, ph_case)
                        if transformed_value != original_value:
                            tab.log(f"   ✓ {ph['name']} ({case_name}): '{original_value}' → '{transformed_value}'")
                            value = transformed_value
                        else:
                            tab.log(f"   • {ph['name']} ({case_name}): '{value}'")
                    else:
                        tab.log(f"   • {ph['name']} ({case_name}): '{value}'")
                    
                    row_data[ph["name"]] = value
                
                # Создаём задачи для Word шаблона (если заполнен)
                if use_word:
                    task_word = (i, row_data, word_template, output_folder, 
                               tab.filename_pattern.get(), required_excel_columns,
                               self.PLACEHOLDERS, tab.filename_column.get(), "word")
                    tasks.append(task_word)
                
                # Создаём задачи для Excel шаблона (если заполнен)
                if use_excel:
                    # Изменяем расширение в паттерне на .xlsx
                    excel_pattern = tab.filename_pattern.get().replace('.docx', '.xlsx')
                    task_excel = (i, row_data, excel_template, output_folder, 
                                excel_pattern, required_excel_columns,
                                self.PLACEHOLDERS, tab.filename_column.get(), "excel")
                    tasks.append(task_excel)
            
            tab.log(f"\n   ✓ Подготовлено {len(tasks)} задач\n")
            
            # === ПАРАЛЛЕЛЬНАЯ ОБРАБОТКА ===
            processed = 0
            with_empty = 0
            errors = []
            
            if num_workers == 1:
                # Последовательная обработка
                tab.log("📄 Последовательная обработка...")
                for task in tasks:
                    # Проверяем флаг остановки
                    if tab.should_stop:
                        tab.log("\n⚠️ Остановка обработки...")
                        break
                    
                    # Определяем тип задачи по последнему элементу кортежа
                    task_type = task[-1]
                    if task_type == "word":
                        result = _process_single_document(task[:-1])  # Убираем маркер типа
                    else:
                        result = _process_single_excel_document(task[:-1])
                    
                    if result['success']:
                        processed += 1
                        if result['is_incomplete']:
                            with_empty += 1
                        
                        # Обновляем прогресс
                        tab.update_progress(processed, len(tasks), f"Обработка документов: {processed}/{len(tasks)}")
                        
                        if processed % 20 == 0:
                            tab.log(f"   ✓ Обработано {processed}/{len(tasks)} документов...")
                    else:
                        errors.append(f"Строка {result['index'] + 1}: {result['error']}")
            else:
                # Параллельная обработка
                tab.log(f"⚡ Параллельная обработка на {num_workers} процессах...")
                tab.log("")
                
                with ProcessPoolExecutor(max_workers=num_workers) as executor:
                    # Отправляем задачи на выполнение с учётом типа
                    futures = {}
                    for task in tasks:
                        task_type = task[-1]
                        if task_type == "word":
                            future = executor.submit(_process_single_document, task[:-1])
                        else:
                            future = executor.submit(_process_single_excel_document, task[:-1])
                        futures[future] = task
                    
                    for future in as_completed(futures):
                        # Проверяем флаг остановки
                        if tab.should_stop:
                            tab.log("\n⚠️ Остановка обработки...")
                            # Отменяем оставшиеся задачи
                            for f in futures:
                                f.cancel()
                            break
                        
                        try:
                            result = future.result(timeout=300)
                            
                            # Выводим логи из результата
                            for log_msg in result.get('logs', []):
                                tab.log(log_msg)
                            
                            if result['success']:
                                processed += 1
                                if result['is_incomplete']:
                                    with_empty += 1
                                
                                # Обновляем прогресс
                                tab.update_progress(processed, len(tasks), f"Обработка документов: {processed}/{len(tasks)}")
                                
                                if processed % 20 == 0:
                                    tab.log(f"✓ Обработано {processed}/{len(tasks)} документов...")
                                    gc.collect()
                            else:
                                errors.append(f"Строка {result['index'] + 1}: {result['error']}")
                        
                        except Exception as e:
                            task = futures[future]
                            errors.append(f"Строка {task[0] + 1}: Критическая ошибка - {str(e)}")
            
            # === ИТОГИ ===
            tab.log("\n" + "═" * 60)
            if tab.should_stop:
                tab.log("⏹ ОБРАБОТКА ОСТАНОВЛЕНА ПОЛЬЗОВАТЕЛЕМ")
            elif errors:
                tab.log("⚠ ОБРАБОТКА ЗАВЕРШЕНА С ОШИБКАМИ")
            else:
                tab.log("✅ ОБРАБОТКА ЗАВЕРШЕНА УСПЕШНО!")
            tab.log("═" * 60)
            tab.log(f"📊 Статистика:")
            tab.log(f"   Всего обработано:          {processed} файлов")
            if len(tasks) > processed:
                tab.log(f"   Не обработано:             {len(tasks) - processed} файлов")
            tab.log(f"   Из них с пометкой _пусто:  {with_empty} файлов")
            if errors:
                tab.log(f"   Ошибок:                    {len(errors)}")
            tab.log(f"\n📁 Папка сохранения:")
            tab.log(f"   {os.path.abspath(output_folder)}")
            
            if errors and len(errors) <= 10:
                tab.log(f"\n❌ Ошибки:")
                for error in errors:
                    tab.log(f"   • {error}")
            elif errors:
                tab.log(f"\n❌ Ошибки (первые 10 из {len(errors)}):")
                for error in errors[:10]:
                    tab.log(f"   • {error}")
            
            tab.log("═" * 60)
            
            # Освобождаем память
            gc.collect()
            
            if tab.should_stop:
                self.root.after(0, lambda: messagebox.showinfo(
                    "Остановлено", 
                    f"Обработка остановлена пользователем.\n\nОбработано файлов: {processed} из {len(tasks)}\n\nПапка: {output_folder}"
                ))
            elif errors:
                self.root.after(0, lambda: messagebox.showwarning(
                    "Завершено с ошибками", 
                    f"Обработано: {processed} файлов\nОшибок: {len(errors)}\n\nПодробности в логе."
                ))
            else:
                self.root.after(0, lambda: messagebox.showinfo(
                    "Успех", 
                    f"Обработка завершена!\n\nВсего файлов: {processed}\nС пометкой '_пусто': {with_empty}\n\nПапка: {output_folder}"
                ))
            
        except Exception as e:
            if not tab.should_stop:
                tab.log("\n" + "═" * 60)
                tab.log("❌ КРИТИЧЕСКАЯ ОШИБКА!")
                tab.log("═" * 60)
                tab.log(f"{e}")
                import traceback
                tab.log(traceback.format_exc())
                tab.log("═" * 60)
                self.root.after(0, lambda: messagebox.showerror("Ошибка", f"Произошла ошибка:\n{e}"))
        
        finally:
            tab.is_processing = False
            tab.should_stop = False
            tab.start_btn.configure(text="▶ Начать обработку")
            gc.collect()
    
    def decline_female_surname(self, surname, case="nomn"):
        """
        Склонение женской фамилии по правилам русского языка
        
        Параметры:
            surname (str): фамилия в именительном падеже
            case (str): код падежа (nomn, gent, datv, accs, ablt, loct)
        
        Возвращает:
            str: склоненная фамилия или None, если фамилия не склоняется
        """
        if not surname or case == "nomn":
            return surname
        
        surname_lower = surname.lower()
        
        # Фамилии на -ова, -ева, -ёва (склоняются)
        if surname_lower.endswith(('ова', 'ева', 'ёва')):
            stem = surname_lower[:-1]
            # Определяем какое окончание (-ов-, -ев-, -ёв-)
            if surname_lower.endswith('ова'):
                base = 'ов'
            elif surname_lower.endswith('ёва'):
                base = 'ёв'
            else:
                base = 'ев'
            
            endings = {
                'gent': 'ой',
                'datv': 'ой',
                'accs': 'у',
                'ablt': 'ой',
                'loct': 'ой'
            }
            return stem + endings.get(case, 'а')
        
        # Фамилии на -ина, -ына (склоняются)
        if surname_lower.endswith(('ина', 'ына')):
            stem = surname_lower[:-1]
            endings = {
                'gent': 'ой',
                'datv': 'ой',
                'accs': 'у',
                'ablt': 'ой',
                'loct': 'ой'
            }
            return stem + endings.get(case, 'а')
        
        # Фамилии на -ская, -цкая (склоняются как прилагательные)
        if surname_lower.endswith(('ская', 'цкая')):
            stem = surname_lower[:-2]
            endings = {
                'gent': 'ой',
                'datv': 'ой',
                'accs': 'ую',
                'ablt': 'ой',
                'loct': 'ой'
            }
            return stem + endings.get(case, 'ая')
        
        # Фамилии на -ая (прилагательные)
        if surname_lower.endswith('ая'):
            stem = surname_lower[:-2]
            endings = {
                'gent': 'ой',
                'datv': 'ой',
                'accs': 'ую',
                'ablt': 'ой',
                'loct': 'ой'
            }
            return stem + endings.get(case, 'ая')
        
        # Остальные женские фамилии обычно не склоняются
        # (на согласный, -ых, -их, -енко и т.д.)
        return None
    
    def decline_male_surname(self, surname, case="nomn"):
        """
        Склонение мужской фамилии по правилам русского языка
        
        Параметры:
            surname (str): фамилия в именительном падеже
            case (str): код падежа (nomn, gent, datv, accs, ablt, loct)
        
        Возвращает:
            str: склоненная фамилия или None, если фамилия не склоняется
        """
        if not surname or case == "nomn":
            return surname
        
        surname_lower = surname.lower()
        
        # Проверка на несклоняемые фамилии
        # 1. Фамилии на -ых, -их
        if surname_lower.endswith(('ых', 'их')):
            return None
        
        # 2. Фамилии на -ко (украинские)
        if surname_lower.endswith('ко'):
            return None
        
        # 3. Фамилии на гласные -о, -е, -и, -у, -ю, -ы
        if surname_lower.endswith(('о', 'е', 'и', 'у', 'ю', 'ы')):
            return None
        
        # 4. Фамилии на -а с гласной перед ней (иностранные: Галуа, Моруа, Гарсиа)
        if surname_lower.endswith('а') and len(surname_lower) >= 2:
            if surname_lower[-2] in 'аеёиоуыэюя':
                return None
        
        # 5. Известные несклоняемые иностранные фамилии на согласную + а
        # (наиболее частые французские, испанские и другие)
        unsклоняемые_на_а = {
            'дюма', 'тома', 'золя', 'петрарка', 'неруда', 'гойя', 
            'сервантеса', 'гарсия', 'лорка', 'кафка', 'рабле'
        }
        if surname_lower in unsклоняемые_на_а:
            return None
        
        # Склонение по типам окончаний
        # Фамилии-прилагательные на -ой/-ый/-ский/-цкий
        if surname_lower.endswith(('ой', 'ый')):
            stem = surname_lower[:-2]
            endings = {
                'gent': 'ого',
                'datv': 'ому',
                'accs': 'ого',
                'ablt': 'ым',
                'loct': 'ом'
            }
            return stem + endings.get(case, surname_lower)
        
        # Фамилии на -ский, -цкий (прилагательные)
        if surname_lower.endswith(('ский', 'цкий')):
            stem = surname_lower[:-2]
            endings = {
                'gent': 'ого',
                'datv': 'ому',
                'accs': 'ого',
                'ablt': 'им',
                'loct': 'ом'
            }
            return stem + endings.get(case, surname_lower)
        
        # Фамилии на -а (с согласной перед ней) - склоняются
        if surname_lower.endswith('а') and len(surname_lower) >= 2:
            if surname_lower[-2] not in 'аеёиоуыэюя':
                stem = surname_lower[:-1]
                # Применяем правило ы/и после шипящих и заднеязычных
                if surname_lower[-2] in 'жчшщкгх':
                    i_ending = 'и'
                else:
                    i_ending = 'ы'
                endings = {
                    'gent': i_ending,
                    'datv': 'е',
                    'accs': 'у',
                    'ablt': 'ой',
                    'loct': 'е'
                }
                return stem + endings.get(case, surname_lower)
        
        # Фамилии на -я
        if surname_lower.endswith('я'):
            stem = surname_lower[:-1]
            endings = {
                'gent': 'и',
                'datv': 'е',
                'accs': 'ю',
                'ablt': 'ей',
                'loct': 'е'
            }
            return stem + endings.get(case, surname_lower)
        
        # Фамилии на согласный (включая мягкий знак и й)
        if surname_lower[-1] in 'бвгджзйклмнпрстфхцчшщь':
            # Типичные русские фамилии на -ов, -ев, -ёв
            if surname_lower.endswith(('ов', 'ев', 'ёв')):
                stem = surname_lower[:-2]
                endings = {
                    'gent': 'ова',
                    'datv': 'ову',
                    'accs': 'ова',
                    'ablt': 'овым',
                    'loct': 'ове'
                }
                if surname_lower.endswith('ёв'):
                    endings = {k: v.replace('ов', 'ёв') for k, v in endings.items()}
                elif surname_lower.endswith('ев'):
                    endings = {k: v.replace('ов', 'ев') for k, v in endings.items()}
                return stem + endings.get(case, surname_lower)
            
            # Фамилии на -ин, -ын
            elif surname_lower.endswith(('ин', 'ын')):
                stem = surname_lower[:-2]
                endings = {
                    'gent': 'ина',
                    'datv': 'ину',
                    'accs': 'ина',
                    'ablt': 'иным',
                    'loct': 'ине'
                }
                if surname_lower.endswith('ын'):
                    endings = {k: v.replace('ин', 'ын') for k, v in endings.items()}
                return stem + endings.get(case, surname_lower)
            
            # Фамилии на мягкий знак
            elif surname_lower.endswith('ь'):
                stem = surname_lower[:-1]
                endings = {
                    'gent': 'я',
                    'datv': 'ю',
                    'accs': 'я',
                    'ablt': 'ем',
                    'loct': 'е'
                }
                return stem + endings.get(case, surname_lower)
            
            # Слова с беглой гласной на -ел (Павел → Павла)
            elif len(surname_lower) >= 3 and surname_lower.endswith('ел') and surname_lower[-3] in 'бвгджзклмнпрстфхцчшщ':
                stem = surname_lower[:-2] + 'л'  # убираем 'ел', оставляем 'л'
                endings = {
                    'gent': 'а',
                    'datv': 'у',
                    'accs': 'а',
                    'ablt': 'ом',
                    'loct': 'е'
                }
                return stem + endings.get(case, '')
            
            # Имена и слова на -й (НЕ прилагательные!)
            # Анатолий → Анатолия, Евгений → Евгения, Андрей → Андрея
            elif surname_lower.endswith('й') and not surname_lower.endswith(('ой', 'ый', 'ский', 'цкий')):
                stem = surname_lower[:-1]
                endings = {
                    'gent': 'я',
                    'datv': 'ю',
                    'accs': 'я',
                    'ablt': 'ем',
                    'loct': 'е'
                }
                return stem + endings.get(case, surname_lower)
            
            # Фамилии на -ец (с беглой гласной: Кузнец → Кузнеца)
            elif surname_lower.endswith('ец'):
                stem = surname_lower[:-2]
                endings = {
                    'gent': 'ца',
                    'datv': 'цу',
                    'accs': 'ца',
                    'ablt': 'цем',
                    'loct': 'це'
                }
                return stem + endings.get(case, '')
            
            # Фамилии на -ок, -ёк (с беглой гласной: Игорёк → Игорька)
            elif surname_lower.endswith(('ок', 'ёк')):
                stem = surname_lower[:-2]  # убираем 'ок' или 'ёк'
                if surname_lower.endswith('ок'):
                    endings = {
                        'gent': 'ка',
                        'datv': 'ку',
                        'accs': 'ка',
                        'ablt': 'ком',
                        'loct': 'ке'
                    }
                else:
                    endings = {
                        'gent': 'ька',
                        'datv': 'ьку',
                        'accs': 'ька',
                        'ablt': 'ьком',
                        'loct': 'ьке'
                    }
                return stem + endings.get(case, '')
            
            # Прочие фамилии на согласный
            else:
                if surname_lower[-1] in 'жчшщц':
                    endings = {
                        'gent': 'а',
                        'datv': 'у',
                        'accs': 'а',
                        'ablt': 'ем',
                        'loct': 'е'
                    }
                else:
                    endings = {
                        'gent': 'а',
                        'datv': 'у',
                        'accs': 'а',
                        'ablt': 'ом',
                        'loct': 'е'
                    }
                return surname_lower + endings.get(case, '')
        
        # Если не подошло ни одно правило, возвращаем None
        return None
    
    def is_date_column(self, col_name):
        """Проверка, является ли колонка датой по заголовку"""
        col_lower = col_name.lower()
        date_keywords = ["дата", "д.р.", "д/р", "д.р", "date", "рождения", "рождение"]
        for keyword in date_keywords:
            if keyword in col_lower:
                return True
        return False
    
    def apply_case(self, value, case="nomn"):
        """Применение падежа к тексту"""
        import pandas as pd
        
        if not value or pd.isna(value):
            return ""
        
        if case == "nomn":
            return str(value).strip()
        
        value = str(value).strip()
        
        words = value.split()
        birth_year_indices = set()  # Индексы слов "года" и "рождения"
        
        for i, word in enumerate(words):
            if word.lower() == "года" and i + 1 < len(words) and words[i + 1].lower() == "рождения":
                if case == "gent":
                    birth_year_indices.add(i)
                    birth_year_indices.add(i + 1)
                else:
                    birth_year_indices.add(i + 1)
        
        # Предварительный анализ падежей всех слов
        word_case_info = []  # Список списков с информацией о каждом слове (все возможные разборы)
        for word in words:
            clean_word = word.rstrip(',.;:!?')
            try:
                parses = self.morph.parse(clean_word.lower())
                if parses:
                    # Сохраняем ВСЕ возможные разборы слова
                    word_parses = []
                    for p in parses:
                        word_info = {
                            'current_case': p.tag.case,  # Текущий падеж слова
                            'pos': p.tag.POS,  # Часть речи (NOUN, ADJF и т.д.)
                            'gender': p.tag.gender,  # Род
                            'number': p.tag.number  # Число
                        }
                        word_parses.append(word_info)
                    word_case_info.append(word_parses)
                else:
                    word_case_info.append(None)
            except:
                word_case_info.append(None)
        
        result_words = []
        skip_indices = set()  # Индексы слов, которые нужно пропустить (уже обработаны)
        
        for idx, word in enumerate(words):
            # Проверка: уже обработано в составе словосочетания
            if idx in skip_indices:
                continue
            
            if idx in birth_year_indices:
                result_words.append(word)
                continue
            
            is_capitalized = word and word[0].isupper()
            is_all_upper = word.isupper()
            
            trailing_punct = ""
            clean_word = word
            while clean_word and clean_word[-1] in ',.;:!?':
                trailing_punct = clean_word[-1] + trailing_punct
                clean_word = clean_word[:-1]
            
            is_date = False
            if '.' in clean_word:
                parts = clean_word.split('.')
                if len(parts) == 3 and all(p.isdigit() for p in parts):
                    if len(parts[0]) <= 2 and len(parts[1]) <= 2 and len(parts[2]) == 4:
                        is_date = True
            
            if is_date:
                result_words.append(word)
                continue
            
            is_number = clean_word.replace('-', '').isdigit()
            
            if is_number:
                result_words.append(word)
                continue
            
            # Проверка на инициалы (И., И.В., и т.п.)
            # Инициалы - это одна или несколько заглавных букв с точками
            clean_no_dots = clean_word.replace('.', '')
            # Одиночный инициал (И.)
            if len(clean_no_dots) == 1 and clean_no_dots.isalpha():
                result_words.append(word)
                continue
            # Составные инициалы (И.В., А.С. и т.п.)
            # Проверяем: все символы без точек - заглавные буквы и их 2-3
            if clean_no_dots.isupper() and clean_no_dots.isalpha() and 2 <= len(clean_no_dots) <= 4:
                result_words.append(word)
                continue
            
            # НОВАЯ ЛОГИКА: проверка согласованных словосочетаний
            # Если текущее слово уже в целевом падеже, проверяем, не является ли оно частью согласованного словосочетания
            current_parses = word_case_info[idx]
            
            # Проверяем, не является ли текущее слово фамилией
            # Признаки фамилии:
            # 1. Перед словом стоит инициал (А.В. Дроздов)
            # 2. После слова стоит инициал (Дроздов А.В.)
            # 3. Типичные окончания фамилий
            is_likely_surname = False
            
            # Проверка 1: перед текущим словом инициал
            if idx > 0:
                prev_word = words[idx - 1]
                prev_clean = prev_word.rstrip(',.;:!?')
                prev_no_dots = prev_clean.replace('.', '')
                # Проверяем, был ли предыдущий элемент инициалом
                is_prev_initial = False
                if len(prev_no_dots) == 1 and prev_no_dots.isalpha() and prev_no_dots.isupper():
                    is_prev_initial = True
                elif prev_no_dots.isupper() and prev_no_dots.isalpha() and 2 <= len(prev_no_dots) <= 4:
                    is_prev_initial = True
                
                if is_prev_initial and is_capitalized:
                    is_likely_surname = True
            
            # Проверка 2: после текущего слова инициал или имя/отчество
            if idx < len(words) - 1 and is_capitalized:
                next_word = words[idx + 1]
                next_clean = next_word.rstrip(',.;:!?')
                next_no_dots = next_clean.replace('.', '')
                # Проверяем, является ли следующий элемент инициалом
                is_next_initial = False
                if len(next_no_dots) == 1 and next_no_dots.isalpha() and next_no_dots.isupper():
                    is_next_initial = True
                elif next_no_dots.isupper() and next_no_dots.isalpha() and 2 <= len(next_no_dots) <= 4:
                    is_next_initial = True
                
                if is_next_initial:
                    is_likely_surname = True
                else:
                    # Проверяем, является ли следующее слово именем (с заглавной буквы, не инициал)
                    # или отчеством (типичные окончания)
                    next_lower = next_clean.lower()
                    is_next_capitalized = next_clean and next_clean[0].isupper()
                    
                    # Проверка на отчество
                    is_patronymic = next_lower.endswith(('ович', 'евич', 'ьич', 'овна', 'евна', 'ична', 'ьевна'))
                    
                    # Если после текущего слова идёт слово с заглавной (имя) или отчество
                    if is_next_capitalized or is_patronymic:
                        is_likely_surname = True
            
            # Проверка 3: типичные окончания русских фамилий (мужских и женских)
            if is_capitalized:
                word_lower_check = clean_word.lower()
                surname_endings = (
                    'ов', 'ев', 'ёв', 'ин', 'ын',  # мужские
                    'ова', 'ева', 'ёва', 'ина', 'ына',  # женские
                    'ский', 'цкий', 'ская', 'цкая',  # прилагательные
                    'ой', 'ый', 'ая'  # прилагательные
                )
                if word_lower_check.endswith(surname_endings):
                    is_likely_surname = True
            
            # Ищем разбор текущего слова, который соответствует целевому падежу
            current_info = None
            # ВАЖНО: если слово похоже на фамилию, не проверяем его на "уже в нужном падеже"
            # так как фамилии должны склоняться через специализированные функции
            if current_parses and not is_likely_surname:
                for parse in current_parses:
                    if parse['current_case'] == case:
                        current_info = parse
                        break
            
            if current_info:
                # Слово уже в нужном падеже
                # Проверяем, есть ли следующие слова, которые тоже в этом падеже (согласованное словосочетание)
                phrase_words = [word]
                phrase_end_idx = idx
                last_case_word_info = current_info  # Последнее слово с падежом (для проверки согласования)
                
                # Смотрим на следующие слова
                for next_idx in range(idx + 1, len(words)):
                    next_parses = word_case_info[next_idx]
                    next_word = words[next_idx]
                    next_clean = next_word.rstrip(',.;:!?')
                    
                    # Проверка на число - включаем в фразу и продолжаем поиск
                    is_next_number = next_clean.replace('-', '').isdigit()
                    if is_next_number:
                        phrase_words.append(next_word)
                        phrase_end_idx = next_idx
                        continue
                    
                    # Проверка на дату - включаем в фразу и продолжаем поиск
                    is_next_date = False
                    if '.' in next_clean:
                        parts = next_clean.split('.')
                        if len(parts) == 3 and all(p.isdigit() for p in parts):
                            if len(parts[0]) <= 2 and len(parts[1]) <= 2 and len(parts[2]) == 4:
                                is_next_date = True
                    if is_next_date:
                        phrase_words.append(next_word)
                        phrase_end_idx = next_idx
                        continue
                    
                    # Проверка на инициалы - включаем в фразу и ЗАВЕРШАЕМ поиск
                    # Инициалы обычно идут перед фамилиями, которые нужно склонять отдельно
                    clean_no_dots_next = next_clean.replace('.', '')
                    is_next_initial = False
                    if len(clean_no_dots_next) == 1 and clean_no_dots_next.isalpha():
                        is_next_initial = True
                    elif clean_no_dots_next.isupper() and clean_no_dots_next.isalpha() and 2 <= len(clean_no_dots_next) <= 4:
                        is_next_initial = True
                    if is_next_initial:
                        phrase_words.append(next_word)
                        phrase_end_idx = next_idx
                        break  # ВАЖНО: завершаем поиск, чтобы следующие слова (фамилии) обрабатывались отдельно
                    
                    # Ищем среди разборов следующего слова подходящий (в нужном падеже и согласованный)
                    next_info = None
                    if next_parses:
                        # Сначала проверяем предлоги и союзы
                        for parse in next_parses:
                            if parse['pos'] in ('PREP', 'CONJ'):
                                phrase_words.append(next_word)
                                phrase_end_idx = next_idx
                                next_info = 'skip'  # Маркер что нашли предлог/союз
                                break
                        
                        if next_info == 'skip':
                            continue
                        
                        # Ищем разбор в нужном падеже, согласованный по роду и числу
                        for parse in next_parses:
                            if parse['current_case'] == case:
                                # Проверка согласования
                                if (parse['pos'] in ('NOUN', 'ADJF', 'PRTF') and 
                                    last_case_word_info['gender'] == parse['gender'] and 
                                    last_case_word_info['number'] == parse['number']):
                                    next_info = parse
                                    break
                    
                    # Если нашли согласованное слово в целевом падеже
                    if next_info and next_info != 'skip':
                        phrase_words.append(next_word)
                        phrase_end_idx = next_idx
                        last_case_word_info = next_info  # Обновляем последнее слово с падежом
                    else:
                        break
                
                # Если нашли согласованное словосочетание (больше одного слова), оставляем его как есть
                if phrase_end_idx > idx:
                    result_words.extend(phrase_words)
                    # Помечаем все слова фразы как обработанные
                    for i in range(idx + 1, phrase_end_idx + 1):
                        skip_indices.add(i)
                    continue
                else:
                    # Одиночное слово уже в нужном падеже - оставляем как есть
                    result_words.append(word)
                    continue
            
            word_lower = clean_word.lower()
            result_word = None
            
            is_male_patronymic = word_lower.endswith(('ович', 'евич', 'ьич'))
            is_female_patronymic = word_lower.endswith(('овна', 'евна', 'ична', 'ьевна'))
            
            if is_male_patronymic:
                endings_map = {
                    'gent': 'а',
                    'datv': 'у',
                    'accs': 'а',
                    'ablt': 'ем',
                    'loct': 'е'
                }
                result_word = word_lower + endings_map.get(case, '')
            
            elif is_female_patronymic:
                stem = word_lower[:-1]  # убираем последнюю 'а'
                endings_map = {
                    'gent': 'ы',
                    'datv': 'е',
                    'accs': 'у',
                    'ablt': 'ой',
                    'loct': 'е'
                }
                result_word = stem + endings_map.get(case, 'а')
            
            elif is_capitalized and not is_female_patronymic and not is_male_patronymic:
                if word_lower.endswith(('ова', 'ева', 'ёва', 'ина', 'ына', 'ская', 'цкая', 'ая')):
                    declined = self.decline_female_surname(clean_word, case)
                else:
                    declined = self.decline_male_surname(clean_word, case)
                
                if declined:
                    result_word = declined
            
            if not result_word:
                parses = self.morph.parse(word_lower)
                if parses:
                    p = parses[0]
                    inflected = p.inflect({case})
                    if inflected:
                        result_word = inflected.word
            
            if not result_word:
                result_word = word_lower
            
            if is_all_upper:
                result_word = result_word.upper()
            elif is_capitalized:
                result_word = result_word.capitalize()
            
            result_word = result_word + trailing_punct
            
            result_words.append(result_word)
        
        return ' '.join(result_words)
    
    @staticmethod
    def to_date(value):
        """Преобразование в формат дд.мм.гггг"""
        import pandas as pd
        
        if pd.isna(value) or str(value).strip() == "":
            return ""
        
        value_str = str(value).strip()
        
        # Если значение уже в формате дд.мм.гггг, возвращаем как есть
        if len(value_str) == 10 and value_str[2] == '.' and value_str[5] == '.':
            try:
                datetime.strptime(value_str, '%d.%m.%Y')
                return value_str
            except:
                pass
        
        # Проверка на чистое время (без даты): 9:00, 9 ч 00 мин и т.п.
        # Если строка содержит двоеточие или "ч" (час), но не содержит точку, дефис или слэш (элементы даты)
        time_indicators = [':']
        date_indicators = ['.', '-', '/']
        has_time = any(ind in value_str for ind in time_indicators) or ' ч ' in value_str.lower() or value_str.lower().endswith(' ч')
        has_date = any(ind in value_str for ind in date_indicators)
        
        # Если есть признаки времени, но нет признаков даты - возвращаем как есть
        if has_time and not has_date:
            return value_str
        
        # Удаляем временную часть, если есть (например "2024-01-15 00:00:00" -> "2024-01-15")
        if ' ' in value_str:
            date_part = value_str.split(' ')[0]
            try:
                dt = pd.to_datetime(date_part, dayfirst=True, errors='coerce')
                if pd.notna(dt):
                    return dt.strftime('%d.%m.%Y')
            except:
                pass
        
        # Если значение числовое (serial date из Excel) - проверяем только если это явно datetime
        if isinstance(value, (int, float)):
            # Проверяем диапазон Excel serial dates (с 1900 года до ~2100)
            if 1 < value < 100000:
                try:
                    base_date = datetime(1899, 12, 30)
                    dt = base_date + timedelta(days=int(value))
                    # Проверяем, что результат - разумная дата (1900-2100 годы)
                    if 1900 <= dt.year <= 2100:
                        return dt.strftime('%d.%m.%Y')
                except:
                    pass
        
        # Пытаемся распарсить как дату (только если есть признаки даты в строке)
        if '/' in value_str or '-' in value_str or '.' in value_str:
            try:
                dt = pd.to_datetime(value_str, dayfirst=True, errors='coerce')
                if pd.notna(dt):
                    return dt.strftime('%d.%m.%Y')
            except:
                pass
        
        return value_str
    
    @staticmethod
    def replace_placeholders_in_paragraph(paragraph, replacements: dict):
        """Безопасная замена всех плейсхолдеров в параграфе с сохранением изображений и форматирования"""
        from docx.oxml.ns import qn
        
        for run in paragraph.runs:
            has_objects = False
            if hasattr(run._element, 'xpath'):
                drawings = run._element.xpath('.//w:drawing')
                pictures = run._element.xpath('.//w:pict')  # старый формат изображений
                has_objects = len(drawings) > 0 or len(pictures) > 0
            
            if has_objects:
                for text_elem in run._element.findall(qn('w:t')):
                    if text_elem.text:
                        modified_text = text_elem.text
                        for placeholder, replacement in replacements.items():
                            modified_text = modified_text.replace(placeholder, str(replacement))
                        text_elem.text = modified_text
            else:
                text = run.text
                for placeholder, replacement in replacements.items():
                    text = text.replace(placeholder, str(replacement))
                run.text = text
    
    def open_merge_window(self):
        """Открыть окно объединения документов"""
        MergeDocumentsWindow(self.root)
    
    def open_excel_constructor(self):
        """Открыть конструктор Excel"""
        ExcelConstructorWindow(self.root, self)

    @staticmethod
    def convert_numbering_to_text(doc):
        """
        Конвертирует автоматическую нумерацию Word в обычный текст.
        Каждый нумерованный пункт получит префикс вида "1.\u00A0" (номер + точка + неразрывный пробел).
        
        Args:
            doc: Document объект
        """
        from docx.oxml.ns import qn
        
        numbering_counters = {}
        
        for paragraph in doc.paragraphs:
            pPr = paragraph._element.find(qn('w:pPr'))
            if pPr is None:
                continue
                
            numPr = pPr.find(qn('w:numPr'))
            if numPr is None:
                continue
            
            numId_elem = numPr.find(qn('w:numId'))
            ilvl_elem = numPr.find(qn('w:ilvl'))
            
            if numId_elem is None:
                continue
            
            num_id = numId_elem.get(qn('w:val'))
            ilvl = int(ilvl_elem.get(qn('w:val'), '0'))
            
            counter_key = f"{num_id}_{ilvl}"
            
            if counter_key not in numbering_counters:
                numbering_counters[counter_key] = 1
            else:
                numbering_counters[counter_key] += 1
            
            current_number = numbering_counters[counter_key]
            
            number_text = f"{current_number}.\u00A0"  # \u00A0 - неразрывный пробел
            
            pPr.remove(numPr)
            
            if len(paragraph.runs) > 0:
                first_run = paragraph.runs[0]
                first_run.text = number_text + first_run.text
            else:
                paragraph.add_run(number_text)

    @staticmethod
    def merge_word_documents(file_paths, output_path, log_callback=None):
        """Объединение Word документов с сохранением форматирования и всех элементов"""
        from docx import Document
        from docx.enum.text import WD_BREAK
        
        if not file_paths:
            raise ValueError("Список файлов пуст")
        
        if log_callback:
            log_callback(f"Объединение {len(file_paths)} Word документов...")
        
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        merged_doc = Document(file_paths[0])
        
        # Конвертируем нумерацию в тексте и в первом документе
        GenerationDocApp.convert_numbering_to_text(merged_doc)
        
        for idx, file_path in enumerate(file_paths[1:], 1):
            paragraphs = merged_doc.paragraphs
            if paragraphs:
                last_para = paragraphs[-1]
                run = last_para.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                merged_doc.add_page_break()
            
            # Открываем документ для добавления
            doc = Document(file_path)
            
            # ВАЖНО: Конвертируем автонумерацию в обычный текст
            # Это решает проблему продолжения нумерации между документами
            GenerationDocApp.convert_numbering_to_text(doc)
            
            # ОПТИМИЗАЦИЯ: Используем напрямую элементы без deepcopy
            # deepcopy очень медленный для больших XML структур
            for element in doc.element.body:
                if element.tag.endswith('sectPr'):
                    continue
                
                # Перемещаем элемент напрямую вместо копирования
                merged_doc.element.body.append(element)
            
            # Очистка памяти после обработки каждого документа
            del doc
            if idx % 5 == 0:  # Каждые 5 документов
                gc.collect()
        
        merged_doc.save(output_path)
        
        # Финальная очистка
        del merged_doc
        gc.collect()
        
        if log_callback:
            log_callback("✓ Word документы успешно объединены")
    
    @staticmethod
    def pdf_has_text_layer(pdf_path, log_callback=None):
        """Проверяет, содержит ли PDF текстовый слой
        
        Args:
            pdf_path: путь к PDF файлу
            log_callback: функция для логирования (опционально)
            
        Returns:
            bool: True если есть извлекаемый текст, False если PDF - скан/изображение
        """
        if not PYMUPDF_AVAILABLE:
            # Fallback через pypdf
            try:
                from pypdf import PdfReader
                reader = PdfReader(pdf_path)
                total_text = ""
                page_count = len(reader.pages)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        total_text += text.strip()
                
                # Увеличенный порог: минимум 200 символов полезного текста на страницу
                # (исключаем мусорные символы, номера страниц и т.д.)
                useful_text = ''.join(c for c in total_text if c.isalnum() or c.isspace())
                avg_text_per_page = len(useful_text.strip()) / max(page_count, 1)
                
                if log_callback:
                    log_callback(f"    [DEBUG] Страниц: {page_count}, символов: {len(useful_text.strip())}, в среднем на страницу: {int(avg_text_per_page)}")
                
                # Если меньше 200 символов на страницу - считаем сканом
                return avg_text_per_page > 200
            except:
                # Если не можем проверить через pypdf, считаем сканом (нужен OCR)
                return False
        
        try:
            doc = fitz.open(pdf_path)
            total_text = ""
            has_images = False
            page_count = doc.page_count
            
            for page in doc:
                # Получаем текст
                total_text += page.get_text()
                
                # Проверяем наличие изображений
                images = page.get_images()
                if images:
                    has_images = True
            
            doc.close()
            
            # Фильтруем только полезный текст (буквы, цифры, пробелы)
            useful_text = ''.join(c for c in total_text if c.isalnum() or c.isspace())
            avg_text_per_page = len(useful_text.strip()) / max(page_count, 1)
            
            if log_callback:
                log_callback(f"    [DEBUG] Страниц: {page_count}, символов: {len(useful_text.strip())}, в среднем на страницу: {int(avg_text_per_page)}, есть изображения: {has_images}")
            
            # Если есть изображения и мало текста - это скан
            if has_images and avg_text_per_page < 200:
                return False
            
            # Если текста меньше 200 символов на страницу - считаем сканом
            return avg_text_per_page > 200
        except:
            # Если не можем проверить, считаем что нужен OCR
            return False
    
    @staticmethod
    def ocr_pdf(pdf_path, output_path=None, log_callback=None, ocr_resolution=1.5, enable_memory_optimization=True):
        """Выполняет OCR для PDF файла, создавая PDF с текстовым слоем
        
        Использует Windows OCR (встроен в Windows 10+) - никаких внешних моделей!
        
        Args:
            pdf_path: путь к исходному PDF
            output_path: путь для сохранения PDF с текстом (если None, перезаписывает исходный)
            log_callback: функция для логирования
            ocr_resolution: разрешение рендеринга для OCR (1.0-3.0). 
                           1.5 - оптимально, 2.0 - высокое качество, 3.0 - максимум
            enable_memory_optimization: включить оптимизацию памяти (очистка каждые 3 страницы)
            
        Returns:
            str: путь к PDF с текстовым слоем
        """
        import tempfile
        import asyncio
        
        if not PYMUPDF_AVAILABLE:
            raise ImportError("Для OCR требуется библиотека PyMuPDF: pip install pymupdf")
        if not WINDOWS_OCR_AVAILABLE:
            raise ImportError("Для OCR требуется библиотека winsdk: pip install winsdk\n(Требуется Windows 10 или новее)")
        if not REPORTLAB_AVAILABLE:
            raise ImportError("Для OCR требуется библиотека reportlab: pip install reportlab")
        if not PIL_AVAILABLE:
            raise ImportError("Для OCR требуется библиотека Pillow: pip install Pillow")
        
        if output_path is None:
            output_path = pdf_path
        
        if log_callback:
            log_callback(f"  OCR: обработка {os.path.basename(pdf_path)}...")
        
        # Инициализируем Windows OCR engine (кэшируем для повторного использования)
        if not hasattr(GenerationDocApp, '_ocr_engine'):
            try:
                # Пробуем создать OCR engine для русского языка
                try:
                    russian_lang = Language("ru")
                    GenerationDocApp._ocr_engine = OcrEngine.try_create(russian_lang)
                except Exception:
                    # Fallback на системные языки пользователя
                    GenerationDocApp._ocr_engine = OcrEngine.try_create_from_user_profile_languages()
                
                if GenerationDocApp._ocr_engine is None:
                    raise Exception("Не удалось инициализировать OCR engine")
                if log_callback:
                    log_callback(f"  OCR: движок инициализирован")
            except Exception as e:
                raise Exception(f"Ошибка инициализации Windows OCR: {str(e)}")
        
        ocr_engine = GenerationDocApp._ocr_engine
        
        # Открываем PDF с помощью PyMuPDF
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        
        if log_callback:
            log_callback(f"  OCR: {page_count} страниц для обработки")
        
        # Создаём новый PDF с OCR
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Регистрируем шрифт с поддержкой кириллицы
        # Пробуем найти системный шрифт Arial или DejaVu
        font_name = "Helvetica"  # fallback
        font_registered = False
        
        # Список путей к шрифтам с кириллицей (Windows)
        font_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/Arial.ttf",
            "C:/Windows/Fonts/times.ttf",
            "C:/Windows/Fonts/calibri.ttf",
            "C:/Windows/Fonts/segoeui.ttf",
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('CyrFont', font_path))
                    font_name = "CyrFont"
                    font_registered = True
                    if log_callback:
                        log_callback(f"  OCR: шрифт загружен: {os.path.basename(font_path)}")
                    break
                except Exception:
                    continue
        
        if not font_registered:
            if log_callback:
                log_callback(f"  OCR: предупреждение - используем Helvetica (кириллица может отображаться некорректно)")
        
        temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_pdf_path = temp_pdf.name
        temp_pdf.close()
        
        c = rl_canvas.Canvas(temp_pdf_path, pagesize=A4)
        
        # Список временных файлов для гарантированного удаления
        temp_files_cleanup = []
        
        try:
            for page_idx in range(page_count):
                if log_callback:
                    log_callback(f"  OCR: страница {page_idx + 1}/{page_count}...")
                
                page = doc[page_idx]
                
                # Получаем реальный размер страницы из оригинального PDF
                page_rect = page.rect
                page_width = page_rect.width
                page_height = page_rect.height
                
                # Устанавливаем размер страницы в canvas равным реальному размеру
                c.setPageSize((page_width, page_height))
                
                pix = None
                img = None
                temp_img_path = None
                
                try:
                    # Разрешение для OCR (настраиваемое)
                    mat = fitz.Matrix(ocr_resolution, ocr_resolution)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Конвертируем в PIL Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    img_width, img_height = img.size
                    
                    # Сохраняем во временный файл для reportlab
                    temp_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                    temp_img_path = temp_img.name
                    temp_img.close()
                    temp_files_cleanup.append(temp_img_path)
                    
                    img.save(temp_img_path, 'PNG')
                    
                    # Коэффициент масштабирования
                    scale_x = page_width / img_width
                    scale_y = page_height / img_height
                    
                    # Рисуем изображение на странице
                    c.drawImage(temp_img_path, 0, 0, width=page_width, height=page_height)
                    
                    try:
                        # Конвертируем PIL Image в формат для Windows OCR
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Создаём поток для изображения
                        img_bytes = io.BytesIO()
                        img.save(img_bytes, format='BMP')
                        img_data_bytes = img_bytes.getvalue()
                        img_bytes.close()
                        
                        # Асинхронная функция для OCR
                        async def perform_ocr_async():
                            stream = InMemoryRandomAccessStream()
                            writer = DataWriter(stream)
                            writer.write_bytes(img_data_bytes)
                            await writer.store_async()
                            stream.seek(0)
                            
                            decoder = await BitmapDecoder.create_async(stream)
                            software_bitmap = await decoder.get_software_bitmap_async(
                                BitmapPixelFormat.BGRA8,
                                BitmapAlphaMode.PREMULTIPLIED
                            )
                            
                            result = await ocr_engine.recognize_async(software_bitmap)
                            return result
                        
                        # Выполняем OCR
                        try:
                            loop = asyncio.new_event_loop()
                            asyncio.set_event_loop(loop)
                            ocr_result = loop.run_until_complete(perform_ocr_async())
                            loop.close()
                        except RuntimeError:
                            # Если уже есть running loop
                            ocr_result = asyncio.run(perform_ocr_async())
                        
                        # Обрабатываем результаты распознавания
                        if ocr_result:
                            # Сохраняем состояние и устанавливаем режим невидимого текста
                            # PDF оператор "3 Tr" = invisible text (только для поиска/копирования)
                            c.saveState()
                            c._code.append('3 Tr')  # Text render mode 3 = invisible
                            c.setFillColorRGB(0, 0, 0)
                            
                            for line in ocr_result.lines:
                                text = line.text
                                if text.strip():
                                    words = line.words
                                    if words:
                                        # Находим границы всех слов в строке
                                        min_x = min(w.bounding_rect.x for w in words)
                                        min_y = min(w.bounding_rect.y for w in words)
                                        max_x = max(w.bounding_rect.x + w.bounding_rect.width for w in words)
                                        max_y = max(w.bounding_rect.y + w.bounding_rect.height for w in words)
                                        
                                        # Масштабируем координаты
                                        x = min_x * scale_x
                                        y = page_height - max_y * scale_y
                                        height = (max_y - min_y) * scale_y
                                        
                                        # Размер шрифта по высоте
                                        font_size = max(height * 0.8, 8)
                                        
                                        c.setFont(font_name, font_size)
                                        c.drawString(x, y, text)
                            
                            # Восстанавливаем состояние
                            c.restoreState()
                    
                    except Exception as e:
                        if log_callback:
                            log_callback(f"  OCR: предупреждение на странице {page_idx + 1}: {str(e)}")
                    
                    c.showPage()
                
                finally:
                    # ОПТИМИЗАЦИЯ: Гарантированная очистка ресурсов после каждой страницы
                    if img:
                        try:
                            img.close()
                        except:
                            pass
                    if pix:
                        try:
                            pix = None  # PyMuPDF автоматически освобождает
                        except:
                            pass
                    
                    # Очистка памяти каждые 3 страницы (опционально)
                    if enable_memory_optimization and (page_idx + 1) % 3 == 0:
                        gc.collect()
        
        finally:
            # Гарантированное закрытие документа и сохранение
            doc.close()
            c.save()
            
            # Удаляем все временные файлы
            for temp_file in temp_files_cleanup:
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except:
                    pass
            
            # Финальная очистка памяти
            gc.collect()
        
        # Заменяем исходный файл
        import shutil
        shutil.move(temp_pdf_path, output_path)
        
        if log_callback:
            log_callback(f"  OCR: завершено")
        
        return output_path
    
    @staticmethod
    def merge_pdf_documents(file_paths, output_path, log_callback=None, use_ocr=True,
                            numbering_line1=None, numbering_line2=None, numbering_line3=None,
                            numbering_position='правый-нижний', numbering_border=True,
                            numbering_increment_mode='per_page'):
        """Объединение PDF документов с опциональным OCR для сканов
        
        Args:
            file_paths: список путей к PDF файлам
            output_path: путь к выходному файлу
            log_callback: функция для логирования
            use_ocr: применять ли OCR к PDF файлам без текстового слоя
            numbering_line1: первая строка штампа
            numbering_line2: вторая строка штампа (с автоинкрементом)
            numbering_line3: третья строка штампа
            numbering_position: позиция штампа
            numbering_border: рисовать рамку
            numbering_increment_mode: режим инкремента ('per_page' или 'per_document')
        
        Автоматически применяет OCR к PDF файлам без текстового слоя
        для обеспечения возможности копирования текста.
        
        Использует только Python библиотеки (без внешних программ).
        """
        import tempfile
        import shutil
        
        if not file_paths:
            raise ValueError("Список файлов пуст")
        
        if PdfMerger is None:
            raise ImportError("Требуется установить pypdf или PyPDF2: pip install pypdf")
        
        if log_callback:
            log_callback(f"Объединение {len(file_paths)} PDF документов...")
        
        # Проверяем доступность OCR (только Python библиотеки)
        ocr_status = get_ocr_status()
        ocr_ready = is_ocr_available() and use_ocr
        
        if use_ocr and not is_ocr_available() and log_callback:
            missing = []
            if not ocr_status['pymupdf']:
                missing.append("PyMuPDF (pip install pymupdf)")
            if not ocr_status['windows_ocr']:
                missing.append("winsdk (pip install winsdk) - требуется Windows 10+")
            if not ocr_status['reportlab']:
                missing.append("reportlab (pip install reportlab)")
            if not ocr_status['pillow']:
                missing.append("Pillow (pip install Pillow)")
            
            log_callback(f"  ⚠ OCR недоступен. Для установки:")
            for m in missing:
                log_callback(f"    {m}")
            log_callback(f"  ⚠ Сканированные PDF будут объединены без распознавания текста")
        elif not use_ocr and log_callback:
            log_callback(f"  ℹ Быстрый режим: OCR отключен")
        
        # Проверяем и применяем OCR к файлам без текстового слоя
        processed_files = []
        temp_files_to_cleanup = []
        
        # НУМЕРАЦИЯ: Если режим per_document или per_document_first_page, добавляем штампы ДО объединения
        current_line2 = numbering_line2
        add_numbering_before_merge = numbering_increment_mode in ['per_document', 'per_document_first_page'] and any([numbering_line1, numbering_line2, numbering_line3])
        
        for idx, pdf_file in enumerate(file_paths):
            if log_callback:
                log_callback(f"  Проверка файла {idx + 1}/{len(file_paths)}: {os.path.basename(pdf_file)}")
            
            # Проверяем наличие текстового слоя (только если OCR включен)
            if not use_ocr:
                # Быстрый режим - просто добавляем файл
                if log_callback:
                    log_callback(f"    ℹ Добавлен без OCR")
                processed_files.append(pdf_file)
                continue
            
            has_text = GenerationDocApp.pdf_has_text_layer(pdf_file, log_callback)
            
            if has_text:
                if log_callback:
                    log_callback(f"    ✓ Текстовый слой присутствует")
                processed_files.append(pdf_file)
            else:
                if not ocr_ready:
                    if log_callback:
                        log_callback(f"    ⚠ Текстовый слой отсутствует, OCR недоступен - используется оригинал")
                    processed_files.append(pdf_file)
                else:
                    if log_callback:
                        log_callback(f"    ⚠ Текстовый слой отсутствует, выполняется OCR...")
                    
                    # Создаём временный файл для OCR
                    temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                    temp_pdf_path = temp_pdf.name
                    temp_pdf.close()
                    temp_files_to_cleanup.append(temp_pdf_path)
                    
                    try:
                        GenerationDocApp.ocr_pdf(pdf_file, temp_pdf_path, log_callback)
                        processed_files.append(temp_pdf_path)
                        if log_callback:
                            log_callback(f"    ✓ OCR выполнен успешно")
                    except Exception as e:
                        if log_callback:
                            log_callback(f"    ⚠ Ошибка OCR: {str(e)}, используется оригинал")
                        processed_files.append(pdf_file)
        
        # НУМЕРАЦИЯ per_document/per_document_first_page: Добавляем штампы к каждому файлу ДО объединения
        files_to_merge = []
        if add_numbering_before_merge:
            for idx, pdf_file in enumerate(processed_files):
                # Создаем временный файл для PDF со штампами
                temp_numbered = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                temp_numbered_path = temp_numbered.name
                temp_numbered.close()
                temp_files_to_cleanup.append(temp_numbered_path)
                
                # Выбираем режим для добавления штампов внутри документа
                if numbering_increment_mode == 'per_document_first_page':
                    # Штамп только на первом листе
                    internal_mode = 'per_document_first_page'
                else:
                    # per_document: штамп на всех листах, но не инкрементируем
                    internal_mode = 'no_increment'
                
                # Добавляем штампы
                GenerationDocApp.add_numbering_to_existing_pdf(
                    pdf_file, temp_numbered_path,
                    numbering_line1, current_line2, numbering_line3,
                    numbering_position, numbering_border,
                    internal_mode,
                    log_callback
                )
                
                files_to_merge.append(temp_numbered_path)
                
                # Инкрементируем для следующего документа
                if current_line2:
                    current_line2 = GenerationDocApp.increment_line2(current_line2)
        else:
            files_to_merge = processed_files
        
        # Объединяем обработанные файлы
        merger = PdfMerger()
        
        try:
            for pdf_file in files_to_merge:
                merger.append(pdf_file)
            
            merger.write(output_path)
        
        finally:
            # ОПТИМИЗАЦИЯ: Гарантированное закрытие merger
            try:
                merger.close()
            except:
                pass
            
            # Удаляем временные файлы
            for temp_file in temp_files_to_cleanup:
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except:
                    pass
            
            # Очистка памяти после объединения
            gc.collect()
        
        if log_callback:
            log_callback("✓ PDF документы успешно объединены")
        
        # НУМЕРАЦИЯ per_page: Добавляем штампы ПОСЛЕ объединения
        if numbering_increment_mode == 'per_page' and any([numbering_line1, numbering_line2, numbering_line3]):
            GenerationDocApp.add_numbering_to_existing_pdf(
                output_path, output_path,
                numbering_line1, numbering_line2, numbering_line3,
                numbering_position, numbering_border,
                numbering_increment_mode, log_callback
            )
    
    @staticmethod
    def show_ocr_setup_dialog(parent):
        """Показывает диалог с инструкциями по установке OCR компонентов"""
        status = get_ocr_status()
        
        # Формируем сообщение о статусе
        status_lines = ["СТАТУС КОМПОНЕНТОВ OCR:\n"]
        
        components = [
            ("PyMuPDF (fitz)", status['pymupdf'], "pip install pymupdf"),
            ("Windows OCR (winsdk)", status['windows_ocr'], "pip install winsdk"),
            ("reportlab", status['reportlab'], "pip install reportlab"),
            ("Pillow", status['pillow'], "pip install Pillow"),
        ]
        
        for name, installed, install_cmd in components:
            mark = "✓" if installed else "✗"
            status_lines.append(f"  {mark} {name}: {'Установлено' if installed else 'НЕ УСТАНОВЛЕНО'}")
            if not installed:
                status_lines.append(f"      Установка: {install_cmd}")
        
        status_lines.append("")
        
        if is_ocr_available():
            status_lines.append("🎉 OCR ПОЛНОСТЬЮ ГОТОВ К РАБОТЕ!")
            status_lines.append("Сканированные PDF будут автоматически распознаваться.")
            status_lines.append("")
            status_lines.append("Используется Windows OCR (встроен в Windows 10+)")
            status_lines.append("Никаких внешних моделей не требуется!")
        else:
            status_lines.append("⚠️ OCR НЕ ДОСТУПЕН")
            status_lines.append("")
            status_lines.append("Для установки выполните:")
            status_lines.append("  pip install pymupdf winsdk reportlab Pillow")
            status_lines.append("")
            status_lines.append("Требования: Windows 10 или новее")
        
        messagebox.showinfo(
            "Статус OCR",
            "\n".join(status_lines),
            parent=parent
        )
        
        return is_ocr_available()
    
    @staticmethod
    def convert_word_to_pdf_direct(docx_file, pdf_file):
        """Прямая конвертация Word в PDF через COM-интерфейс Word
        
        Args:
            docx_file: путь к Word файлу
            pdf_file: путь к выходному PDF файлу
        """
        if not WIN32COM_AVAILABLE:
            raise ImportError(
                "Для конвертации требуется библиотека pywin32.\n"
                "Установите её командой: pip install pywin32"
            )
        
        word = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0  # Отключаем все диалоги
            
            doc = word.Documents.Open(docx_file)
            
            doc.SaveAs(pdf_file, FileFormat=17)
            
            doc.Close(False)
            
        except Exception as e:
            raise Exception(f"Ошибка при конвертации через Word COM: {str(e)}")
        
        finally:
            if word:
                try:
                    word.Quit()
                except:
                    pass
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                except:
                    pass
    
    @staticmethod
    def convert_word_to_pdf(file_paths, output_folder=None, log_callback=None,
                            numbering_line1=None, numbering_line2=None, numbering_line3=None,
                            numbering_position='правый-нижний', numbering_border=True,
                            numbering_increment_mode='per_document', progress_callback=None):
        """Конвертация Word документов в PDF с параллельной обработкой
        
        Args:
            file_paths: список путей к Word файлам
            output_folder: папка для сохранения PDF (если None, сохраняет рядом с исходным файлом)
            log_callback: функция для логирования
            numbering_line1: первая строка штампа
            numbering_line2: вторая строка штампа (с автоинкрементом)
            numbering_line3: третья строка штампа
            numbering_position: позиция штампа
            numbering_border: рисовать рамку
            numbering_increment_mode: режим инкремента ('per_document' или 'per_page')
            progress_callback: функция для обновления прогресса (current, total, message)
        
        Returns:
            список путей к созданным PDF файлам
        """
        _ensure_concurrent_imports()
        
        if not file_paths:
            raise ValueError("Список файлов пуст")
        
        if not DOCX2PDF_AVAILABLE and not WIN32COM_AVAILABLE:
            raise ImportError(
                "Для конвертации Word в PDF требуется одна из библиотек:\n"
                "1. pip install docx2pdf\n"
                "2. pip install pywin32\n\n"
                "Также требуется установленный Microsoft Word."
            )
        
        converted_files = []
        errors = []
        total = len(file_paths)
        
        if log_callback:
            log_callback(f"Начало конвертации {total} файлов...")
        
        try:
            max_workers = min(4, multiprocessing.cpu_count())
        except:
            max_workers = 2  # Безопасное значение по умолчанию
        
        tasks = [(docx_file, output_folder) for docx_file in file_paths]
        
        if len(tasks) == 1:
            result = _convert_single_pdf(tasks[0])
            if result['success']:
                converted_files.append(result['pdf_file'])
                if log_callback:
                    log_callback(f"  ✓ {os.path.basename(result['pdf_file'])}")
            else:
                errors.append(f"{os.path.basename(result['docx_file'])}: {result['error']}")
        else:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {executor.submit(_convert_single_pdf, task): task for task in tasks}
                
                completed = 0
                for future in as_completed(futures):
                    completed += 1
                    try:
                        result = future.result(timeout=300)  # 5 минут таймаут
                        
                        if result['success']:
                            converted_files.append(result['pdf_file'])
                            if log_callback:
                                log_callback(f"[{completed}/{total}] ✓ {os.path.basename(result['pdf_file'])}")
                        else:
                            errors.append(f"{os.path.basename(result['docx_file'])}: {result['error']}")
                            if log_callback:
                                log_callback(f"[{completed}/{total}] ✗ {os.path.basename(result['docx_file'])}: {result['error']}")
                        
                        # Обновляем прогресс
                        if progress_callback:
                            progress_callback(completed, total, f"Конвертация: {completed}/{total}")
                    
                    except Exception as e:
                        task = futures[future]
                        docx_file = task[0]
                        errors.append(f"{os.path.basename(docx_file)}: Критическая ошибка - {str(e)}")
                        if log_callback:
                            log_callback(f"[{completed}/{total}] ✗ {os.path.basename(docx_file)}: {str(e)}")
        
        if errors:
            error_msg = "Ошибки при конвертации:\n" + "\n".join(errors[:10])
            if len(errors) > 10:
                error_msg += f"\n... и ещё {len(errors) - 10} ошибок"
            
            if converted_files:
                if log_callback:
                    log_callback(f"\n⚠ Завершено с ошибками: {len(converted_files)} успешно, {len(errors)} ошибок")
                raise Warning(error_msg)
            else:
                raise Exception(error_msg)
        
        # Добавляем нумерацию к каждому PDF (если задана)
        if any([numbering_line1, numbering_line2, numbering_line3]):
            if log_callback:
                log_callback(f"\n[ШТАМП] Добавление нумерации к {len(converted_files)} файлам...")
            
            current_line2 = numbering_line2
            for idx, pdf_file in enumerate(converted_files, 1):
                if log_callback:
                    log_callback(f"  [{idx}/{len(converted_files)}] {os.path.basename(pdf_file)}...")
                
                # Выбираем режим для добавления штампов внутри документа
                if numbering_increment_mode == 'per_document_first_page':
                    internal_mode = 'per_document_first_page'
                else:
                    internal_mode = 'per_page'
                
                GenerationDocApp.add_numbering_to_existing_pdf(
                    pdf_file, pdf_file,
                    numbering_line1, current_line2, numbering_line3,
                    numbering_position, numbering_border,
                    internal_mode,
                    log_callback
                )
                
                # Инкрементируем line2 для следующего документа (если режим per_document или per_document_first_page)
                if numbering_increment_mode in ['per_document', 'per_document_first_page'] and current_line2:
                    current_line2 = GenerationDocApp.increment_line2(current_line2)
        
        return converted_files
    
    @staticmethod
    def convert_and_merge_word_to_pdf(file_paths, output_file, log_callback=None,
                                      numbering_line1=None, numbering_line2=None, numbering_line3=None,
                                      numbering_position='правый-нижний', numbering_border=True,
                                      numbering_increment_mode='per_document'):
        """Конвертация Word документов в PDF и объединение в один файл
        
        Args:
            file_paths: список путей к Word файлам
            output_file: путь к результирующему PDF файлу
            log_callback: функция для логирования
            numbering_line1: первая строка штампа
            numbering_line2: вторая строка штампа (с автоинкрементом)
            numbering_line3: третья строка штампа
            numbering_position: позиция штампа
            numbering_border: рисовать рамку
            numbering_increment_mode: режим инкремента ('per_document' или 'per_page')
        """
        if not file_paths:
            raise ValueError("Список файлов пуст")
        
        if not DOCX2PDF_AVAILABLE and not WIN32COM_AVAILABLE:
            raise ImportError(
                "Для конвертации Word в PDF требуется одна из библиотек:\n"
                "1. pip install docx2pdf\n"
                "2. pip install pywin32\n\n"
                "Также требуется установленный Microsoft Word."
            )
        
        if PdfMerger is None:
            raise ImportError("Требуется установить pypdf или PyPDF2: pip install pypdf")
        
        import tempfile
        temp_dir = tempfile.mkdtemp()
        temp_pdf_files = []
        errors = []
        
        try:
            total = len(file_paths)
            if log_callback:
                log_callback(f"Конвертация {total} документов...")
            
            for idx, docx_file in enumerate(file_paths, 1):
                try:
                    if log_callback:
                        log_callback(f"  [{idx}/{total}] {os.path.basename(docx_file)}...")
                    if not os.path.exists(docx_file):
                        raise FileNotFoundError(f"Файл не найден: {docx_file}")
                    
                    if not docx_file.lower().endswith('.docx'):
                        raise ValueError("Файл должен иметь расширение .docx")
                    
                    docx_file = os.path.abspath(docx_file)
                    
                    base_name = os.path.splitext(os.path.basename(docx_file))[0]
                    temp_pdf = os.path.join(temp_dir, base_name + ".pdf")
                    
                    if os.path.exists(temp_pdf):
                        try:
                            os.remove(temp_pdf)
                        except:
                            pass
                    
                    success = False
                    last_error = None
                    
                    if WIN32COM_AVAILABLE:
                        try:
                            GenerationDocApp.convert_word_to_pdf_direct(docx_file, temp_pdf)
                            success = True
                        except Exception as e:
                            last_error = f"win32com: {str(e)}"
                    
                    if not success and DOCX2PDF_AVAILABLE:
                        try:
                            convert(docx_file, temp_pdf)
                            success = True
                        except Exception as e:
                            if last_error:
                                last_error += f"; docx2pdf: {str(e)}"
                            else:
                                last_error = f"docx2pdf: {str(e)}"
                    
                    if success and os.path.exists(temp_pdf):
                        temp_pdf_files.append(temp_pdf)
                        if log_callback:
                            log_callback(f"    ✓ Успешно")
                    else:
                        if not last_error:
                            last_error = "PDF файл не был создан"
                        raise Exception(last_error)
                    
                except Exception as e:
                    error_text = str(e)
                    if "NoneType" in error_text or "COM" in error_text:
                        error_text += "\n💡 Попробуйте: 1) Закрыть все окна Word, 2) Запустить программу от администратора"
                    errors.append(f"{os.path.basename(docx_file)}: {error_text}")
            
            if not temp_pdf_files:
                raise Exception("Ошибки при конвертации:\n" + "\n".join(errors))
            
            if log_callback:
                log_callback(f"Объединение {len(temp_pdf_files)} PDF файлов...")
            
            merger = PdfMerger()
            
            try:
                for pdf_file in temp_pdf_files:
                    merger.append(pdf_file)
                
                merger.write(output_file)
            
            finally:
                # ОПТИМИЗАЦИЯ: Гарантированное закрытие merger
                try:
                    merger.close()
                except:
                    pass
                
                # Очистка памяти
                gc.collect()
            
            if log_callback:
                log_callback("✓ Объединение завершено")
            
            # Добавляем нумерацию (если задана)
            if any([numbering_line1, numbering_line2, numbering_line3]):
                GenerationDocApp.add_numbering_to_existing_pdf(
                    output_file, output_file,
                    numbering_line1, numbering_line2, numbering_line3,
                    numbering_position, numbering_border,
                    numbering_increment_mode, log_callback
                )
            
            if errors:
                raise Warning(f"Файл создан, но были ошибки при конвертации некоторых документов:\n" + "\n".join(errors))
            
        finally:
            # ОПТИМИЗАЦИЯ: Гарантированное удаление временной папки
            import shutil
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            
            # Финальная очистка памяти
            gc.collect()
    
    @staticmethod
    def increment_line2(line2_text):
        """Увеличивает числовое значение в строке 2
        
        Args:
            line2_text: строка с номером (например: "1", "АБВ/1319", "№ 1819-А")
            
        Returns:
            Строка с увеличенным номером
        """
        import re
        
        if not line2_text or not line2_text.strip():
            return line2_text
        
        # Ищем все числа в строке
        numbers = list(re.finditer(r'\d+', line2_text))
        
        if not numbers:
            return line2_text  # Если нет чисел, возвращаем как есть
        
        # Берем последнее число в строке
        last_number_match = numbers[-1]
        old_number = last_number_match.group()
        new_number = str(int(old_number) + 1)
        
        # Сохраняем ведущие нули
        if old_number.startswith('0') and len(old_number) > 1:
            new_number = new_number.zfill(len(old_number))
        
        # Заменяем последнее число
        result = line2_text[:last_number_match.start()] + new_number + line2_text[last_number_match.end():]
        return result
    
    @staticmethod
    def add_numbering_to_existing_pdf(input_pdf_path, output_pdf_path=None, 
                                      numbering_line1=None, numbering_line2=None, numbering_line3=None,
                                      numbering_position='правый-нижний', numbering_border=True,
                                      numbering_increment_mode='per_page', log_callback=None):
        """Добавляет штампы нумерации к существующему PDF файлу
        
        Args:
            input_pdf_path: путь к исходному PDF файлу
            output_pdf_path: путь к выходному файлу (если None, перезаписывает input_pdf_path)
            numbering_line1: первая строка штампа
            numbering_line2: вторая строка штампа (с автоинкрементом)
            numbering_line3: третья строка штампа
            numbering_position: позиция штампа
            numbering_border: рисовать рамку
            numbering_increment_mode: режим инкремента ('per_page' или 'per_document')
            log_callback: функция для логирования
        """
        # Проверяем существование входного файла
        if not os.path.exists(input_pdf_path):
            error_msg = f"[ШТАМП] ✗ Файл не найден: {input_pdf_path}"
            if log_callback:
                log_callback(error_msg)
            raise FileNotFoundError(error_msg)
        
        if not any([numbering_line1, numbering_line2, numbering_line3]):
            if log_callback:
                log_callback("[ШТАМП] Нумерация не задана, пропускаем")
            return input_pdf_path  # Если нумерация не задана, ничего не делаем
        
        if not REPORTLAB_AVAILABLE or not PIL_AVAILABLE:
            if log_callback:
                log_callback("[ШТАМП] ⚠ reportlab или Pillow не установлены, штампы не добавлены")
            return input_pdf_path
        
        try:
            from reportlab.pdfgen import canvas as rl_canvas
            from reportlab.lib.pagesizes import A4
            import tempfile
            
            if log_callback:
                log_callback(f"[ШТАМП] Добавление нумерации к PDF...")
                log_callback(f"  Строка 1: {numbering_line1}")
                log_callback(f"  Строка 2: {numbering_line2}")
                log_callback(f"  Строка 3: {numbering_line3}")
                log_callback(f"  Позиция: {numbering_position}")
                log_callback(f"  Рамка: {numbering_border}")
                log_callback(f"  Режим инкремента: {numbering_increment_mode}")
            
            # Импортируем PdfReader и PdfWriter
            from pypdf import PdfReader, PdfWriter
            
            reader = PdfReader(input_pdf_path)
            num_pages = len(reader.pages)
            
            if log_callback:
                log_callback(f"[ШТАМП] Обработка {num_pages} страниц...")
            
            # Создаем временный файл для overlay
            temp_overlay = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            temp_overlay_path = temp_overlay.name
            temp_overlay.close()
            
            # Создаем overlay PDF со штампами (начинаем с A4, но будем менять размер для каждой страницы)
            c = rl_canvas.Canvas(temp_overlay_path, pagesize=A4)
            
            current_line2 = numbering_line2
            
            for page_num in range(num_pages):
                # Получаем реальный размер текущей страницы
                current_page = reader.pages[page_num]
                
                # Получаем размер из mediabox
                page_box = current_page.mediabox
                page_width = float(page_box.width)
                page_height = float(page_box.height)
                
                # Получаем rotation страницы
                rotation = current_page.get('/Rotate', 0)
                
                if log_callback and page_num == 0:
                    log_callback(f"[ШТАМП] Размер страницы: {page_width:.1f}x{page_height:.1f}, поворот: {rotation}°")
                
                # Устанавливаем размер страницы overlay равным реальному размеру страницы
                c.setPageSize((page_width, page_height))
                
                # При режиме per_document_first_page штампуем только первую страницу
                if numbering_increment_mode == 'per_document_first_page' and page_num > 0:
                    # Пропускаем все страницы кроме первой, но все равно создаем пустую страницу в overlay
                    c.showPage()
                    continue
                
                # Добавляем штамп на страницу с правильными размерами и учётом rotation
                GenerationDocApp.add_numbering_stamp(
                    c, page_width, page_height,
                    numbering_line1, current_line2, numbering_line3,
                    numbering_position, numbering_border, rotation
                )
                
                if log_callback and page_num < 5:  # Логируем только первые 5 страниц
                    log_callback(f"  Страница {page_num + 1}: line2='{current_line2}'")
                
                # Инкрементируем для следующей страницы (в зависимости от режима)
                if numbering_increment_mode == 'no_increment':
                    # Не инкрементируем (для per_document внутри одного документа)
                    pass
                elif current_line2 and numbering_increment_mode in ['per_page', 'per_document']:
                    # Инкрементируем на каждой странице
                    current_line2 = GenerationDocApp.increment_line2(current_line2)
                
                c.showPage()
            
            c.save()
            
            # Объединяем overlay с оригиналом
            if output_pdf_path is None:
                output_pdf_path = input_pdf_path
            
            overlay_reader = PdfReader(temp_overlay_path)
            writer = PdfWriter()
            
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                overlay_page = overlay_reader.pages[page_num]
                page.merge_page(overlay_page)
                writer.add_page(page)
            
            # Создаем временный файл для результата
            temp_output = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            temp_output_path = temp_output.name
            temp_output.close()
            
            with open(temp_output_path, 'wb') as f:
                writer.write(f)
            
            # Заменяем оригинальный файл
            import shutil
            shutil.move(temp_output_path, output_pdf_path)
            
            # Удаляем временные файлы
            try:
                os.unlink(temp_overlay_path)
            except:
                pass
            
            if log_callback:
                log_callback(f"[ШТАМП] ✓ Нумерация успешно добавлена")
            
            return output_pdf_path
            
        except Exception as e:
            if log_callback:
                log_callback(f"[ШТАМП] ✗ Ошибка добавления штампов: {str(e)}")
            import traceback
            traceback.print_exc()
            return input_pdf_path
    
    @staticmethod
    def add_numbering_stamp(canvas_obj, page_width, page_height, line1=None, line2=None, line3=None, 
                          position='правый-нижний', draw_border=True, rotation=0):
        """Добавляет штамп нумерации на страницу PDF
        
        Args:
            canvas_obj: объект Canvas из reportlab
            page_width: ширина страницы
            page_height: высота страницы
            line1: первая строка текста
            line2: вторая строка текста (с автоинкрементом)
            line3: третья строка текста
            position: позиция штампа (правый-нижний, центр-нижний, левый-нижний, левый-верхний, центр-верхний, правый-верхний)
            draw_border: рисовать рамку вокруг текста (True/False)
            rotation: угол поворота страницы (0, 90, 180, 270)
        """
        if not any([line1, line2, line3]):
            return  # Если все параметры пустые, ничего не рисуем
        
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Регистрируем шрифт Times New Roman
        font_name = "Times-Roman"  # Стандартный шрифт Times
        try:
            # Пробуем загрузить системный Times New Roman для поддержки кириллицы
            times_paths = [
                "C:/Windows/Fonts/times.ttf",
                "C:/Windows/Fonts/Times.ttf",
                "C:/Windows/Fonts/timesnewroman.ttf",
            ]
            for font_path in times_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('TimesNewRoman', font_path))
                    font_name = 'TimesNewRoman'
                    break
        except:
            pass  # Используем стандартный Times-Roman
        
        # Параметры штампа
        font_size = 8
        line_height = 10  # Высота строки
        padding = 8  # Отступы внутри прямоугольника
        
        # Собираем строки для отображения
        lines = []
        if line1:
            lines.append(str(line1))
        if line2:
            lines.append(str(line2))
        if line3:
            lines.append(str(line3))
        
        if not lines:
            return
        
        # Вычисляем размеры прямоугольника
        max_text_width = 0
        canvas_obj.setFont(font_name, font_size)
        for line in lines:
            text_width = canvas_obj.stringWidth(line, font_name, font_size)
            max_text_width = max(max_text_width, text_width)
        
        rect_width = max_text_width + 2 * padding
        rect_height = len(lines) * line_height + 2 * padding
        
        # Определяем позицию в зависимости от параметра position
        margin = 20
        
        # ВАЖНО: При наличии rotation в PDF, нам нужно рисовать с учётом этого поворота
        # Используем canvas transformations для правильной ориентации
        
        # Вычисляем визуальные размеры (как видит пользователь)
        if rotation in [90, 270]:
            visual_width = page_height
            visual_height = page_width
        else:
            visual_width = page_width
            visual_height = page_height
        
        # Вычисляем позицию в визуальных координатах
        if 'правый' in position:
            visual_x = visual_width - rect_width - margin
        elif 'центр' in position:
            visual_x = (visual_width - rect_width) / 2
        else:  # левый
            visual_x = margin
        
        if 'нижний' in position:
            visual_y = margin
        else:  # верхний
            visual_y = visual_height - rect_height - margin
        
        # Применяем canvas transformations для rotation
        canvas_obj.saveState()
        
        if rotation == 90:
            # Поворот canvas на 90° по часовой
            canvas_obj.translate(page_width, 0)
            canvas_obj.rotate(90)
            rect_x = visual_x
            rect_y = visual_y
        elif rotation == 180:
            # Поворот на 180°
            canvas_obj.translate(page_width, page_height)
            canvas_obj.rotate(180)
            rect_x = visual_x
            rect_y = visual_y
        elif rotation == 270:
            # Поворот на 270° по часовой
            canvas_obj.translate(0, page_height)
            canvas_obj.rotate(270)
            rect_x = visual_x
            rect_y = visual_y
        else:
            # Без rotation
            rect_x = visual_x
            rect_y = visual_y
        
        # Рисуем прямоугольник (если нужна рамка)
        if draw_border:
            canvas_obj.setStrokeColorRGB(0, 0, 0)  # Черная рамка
            # Рисуем только рамку без заливки (прозрачный фон)
            canvas_obj.rect(rect_x, rect_y, rect_width, rect_height, fill=0, stroke=1)
        
        # Рисуем текст (центрированный)
        canvas_obj.setFillColorRGB(0, 0, 0)  # Черный текст
        canvas_obj.setFont(font_name, font_size)
        
        # Начальная позиция текста (сверху вниз)
        text_y = rect_y + rect_height - padding - font_size
        
        for line in lines:
            text_width = canvas_obj.stringWidth(line, font_name, font_size)
            # Центрируем текст
            text_x = rect_x + (rect_width - text_width) / 2
            canvas_obj.drawString(text_x, text_y, line)
            text_y -= line_height
        
        # Восстанавливаем состояние canvas
        canvas_obj.restoreState()
    
    @staticmethod
    def _image_to_pdf_with_reportlab(image_path, output_pdf_path, log_callback=None, max_image_size=None, fit_mode='центр',
                                     numbering_line1=None, numbering_line2=None, numbering_line3=None,
                                     numbering_position='правый-нижний', numbering_border=True, use_ocr=False):
        """Внутренняя функция для создания PDF из изображения с использованием reportlab
        
        Используется когда нужна нумерация или OCR с reportlab
        """
        import tempfile
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        
        img = None
        temp_img_path = None
        
        try:
            # Открываем изображение
            img = Image.open(image_path)
            
            # Ограничение размера изображения (опционально)
            if max_image_size and max(img.size) > max_image_size:
                ratio = max_image_size / max(img.size)
                new_size = tuple(int(dim * ratio) for dim in img.size)
                img = img.resize(new_size, Image.Resampling.LANCZOS)
                if log_callback:
                    log_callback(f"    ℹ Изображение уменьшено до {new_size} для экономии памяти")
            
            # Конвертируем в RGB если нужно
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            img_width, img_height = img.size
            page_width, page_height = A4
            
            # Применяем режим размещения на странице
            if fit_mode == 'заполнить':
                scaled_width = page_width
                scaled_height = page_height
                x_offset = 0
                y_offset = 0
                if log_callback:
                    log_callback(f"    ℹ Режим: растянуто на весь лист")
            elif fit_mode == 'вписать':
                scale = min(page_width / img_width, page_height / img_height)
                scaled_width = img_width * scale
                scaled_height = img_height * scale
                x_offset = (page_width - scaled_width) / 2
                y_offset = (page_height - scaled_height) / 2
                if log_callback:
                    log_callback(f"    ℹ Режим: вписано в лист")
            else:  # 'центр'
                scale = min(page_width / img_width, page_height / img_height)
                scaled_width = img_width * scale
                scaled_height = img_height * scale
                x_offset = (page_width - scaled_width) / 2
                y_offset = (page_height - scaled_height) / 2
            
            # Сохраняем во временный файл
            temp_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            temp_img_path = temp_img.name
            temp_img.close()
            img.save(temp_img_path, 'PNG')
            
            # Создаём PDF
            c = rl_canvas.Canvas(output_pdf_path, pagesize=A4)
            
            # Рисуем изображение
            c.drawImage(temp_img_path, x_offset, y_offset, width=scaled_width, height=scaled_height)
            
            # Добавляем штамп нумерации (если указан)
            if log_callback:
                log_callback(f"[DEBUG] Проверка нумерации: line1={numbering_line1}, line2={numbering_line2}, line3={numbering_line3}")
            
            if any([numbering_line1, numbering_line2, numbering_line3]):
                if log_callback:
                    log_callback(f"[ШТАМП] Добавление штампа на страницу...")
                    log_callback(f"  Строка 1: {numbering_line1}")
                    log_callback(f"  Строка 2: {numbering_line2}")
                    log_callback(f"  Строка 3: {numbering_line3}")
                    log_callback(f"  Позиция: {numbering_position}")
                    log_callback(f"  Рамка: {numbering_border}")
                
                GenerationDocApp.add_numbering_stamp(c, page_width, page_height, 
                                                    numbering_line1, numbering_line2, numbering_line3,
                                                    numbering_position, numbering_border)
                if log_callback:
                    log_callback(f"    ✓ Добавлен штамп нумерации")
            else:
                if log_callback:
                    log_callback(f"[DEBUG] Нумерация не применяется (все строки пустые)")
            
            c.showPage()
            c.save()
            
            return output_pdf_path
        
        finally:
            # Гарантированное закрытие ресурсов
            if img:
                try:
                    img.close()
                except:
                    pass
            
            # Удаляем временный файл
            if temp_img_path:
                try:
                    if os.path.exists(temp_img_path):
                        os.unlink(temp_img_path)
                except:
                    pass
    
    @staticmethod
    def image_to_pdf_simple(image_path, output_pdf_path, log_callback=None, max_image_size=None, fit_mode='центр',
                           numbering_line1=None, numbering_line2=None, numbering_line3=None,
                           numbering_position='правый-нижний', numbering_border=True):
        """Простая конвертация изображения в PDF без OCR (быстрый режим)
        
        Args:
            image_path: путь к файлу изображения
            output_pdf_path: путь для сохранения PDF
            log_callback: функция для логирования
            max_image_size: максимальный размер изображения (None = без ограничений)
            fit_mode: режим размещения изображения:
                     'center' - по центру с сохранением пропорций
                     'fill' - на весь лист (может обрезать)
                     'fit' - на весь лист с сохранением пропорций
            numbering_line1: строка 1 для штампа (опционально)
            numbering_line2: строка 2 для штампа с автоинкрементом (опционально)
            numbering_line3: строка 3 для штампа (опционально)
            numbering_position: позиция штампа (bottom_right, bottom_center, bottom_left, top_left, top_center, top_right)
            numbering_border: рисовать рамку вокруг штампа
            
        Returns:
            str: путь к созданному PDF файлу
        """
        if not PIL_AVAILABLE:
            raise ImportError("Требуется библиотека Pillow: pip install Pillow")
        
        img = None
        try:
            # Открываем изображение
            img = Image.open(image_path)
            
            # Ограничение размера изображения (опционально)
            if max_image_size and max(img.size) > max_image_size:
                ratio = max_image_size / max(img.size)
                new_size = tuple(int(dim * ratio) for dim in img.size)
                img = img.resize(new_size, Image.Resampling.LANCZOS)
                if log_callback:
                    log_callback(f"    ℹ Изображение уменьшено до {new_size} для экономии памяти")
            
            # Конвертируем в RGB если нужно
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Применяем режим размещения
            if fit_mode == 'заполнить':
                # На весь лист A4 (может обрезать)
                from reportlab.lib.pagesizes import A4
                a4_width, a4_height = A4
                # Изменяем размер до A4
                img = img.resize((int(a4_width), int(a4_height)), Image.Resampling.LANCZOS)
                if log_callback:
                    log_callback(f"    ℹ Режим: растянуто на весь лист")
            elif fit_mode == 'вписать':
                # На весь лист с сохранением пропорций
                from reportlab.lib.pagesizes import A4
                a4_width, a4_height = A4
                img_ratio = img.size[0] / img.size[1]
                a4_ratio = a4_width / a4_height
                
                if img_ratio > a4_ratio:
                    # Изображение шире -> подгоняем по ширине
                    new_width = int(a4_width)
                    new_height = int(a4_width / img_ratio)
                else:
                    # Изображение выше -> подгоняем по высоте
                    new_height = int(a4_height)
                    new_width = int(a4_height * img_ratio)
                
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                if log_callback:
                    log_callback(f"    ℹ Режим: вписано в лист с пропорциями")
            # else: 'center' - по умолчанию, ничего не меняем
            
            # Сохраняем как PDF
            img.save(output_pdf_path, 'PDF', resolution=100.0)
            
            return output_pdf_path
        
        finally:
            # ОПТИМИЗАЦИЯ: Гарантированное закрытие изображения
            if img:
                try:
                    img.close()
                except:
                    pass
    
    @staticmethod
    def image_to_pdf_with_ocr(image_path, output_pdf_path, log_callback=None, max_image_size=None, fit_mode='центр',
                              numbering_line1=None, numbering_line2=None, numbering_line3=None,
                              numbering_position='правый-нижний', numbering_border=True):
        """Конвертация изображения в PDF с OCR для создания текстового слоя
        
        Args:
            image_path: путь к файлу изображения
            output_pdf_path: путь для сохранения PDF
            log_callback: функция для логирования
            max_image_size: максимальный размер изображения (None = без ограничений)
            fit_mode: режим размещения (центр/заполнить/вписать)
            numbering_line1: строка 1 для штампа (опционально)
            numbering_line2: строка 2 для штампа с автоинкрементом (опционально)
            numbering_line3: строка 3 для штампа (опционально)
            numbering_position: позиция штампа
            numbering_border: рисовать рамку вокруг штампа
            
        Returns:
            str: путь к созданному PDF файлу
        """
        import tempfile
        
        if not PIL_AVAILABLE:
            raise ImportError("Требуется библиотека Pillow: pip install Pillow")
        
        if not REPORTLAB_AVAILABLE:
            # Fallback - просто сохраняем изображение как PDF без OCR
            return GenerationDocApp.image_to_pdf_simple(image_path, output_pdf_path, log_callback, max_image_size, fit_mode,
                                                        numbering_line1, numbering_line2, numbering_line3,
                                                        numbering_position, numbering_border)
        
        return GenerationDocApp._image_to_pdf_with_reportlab(
            image_path, output_pdf_path, log_callback, max_image_size, fit_mode,
            numbering_line1, numbering_line2, numbering_line3, 
            numbering_position, numbering_border, use_ocr=True
        )
    
    @staticmethod
    def convert_images_to_pdf(file_paths, output_folder=None, log_callback=None, use_ocr=True, 
                              max_image_size=4000, fit_mode='центр',
                              numbering_line1=None, numbering_line2=None, numbering_line3=None,
                              numbering_position='правый-нижний', numbering_border=True,
                              numbering_increment_mode='per_document', progress_callback=None):
        """Конвертация изображений в PDF с опциональным OCR и параллельной обработкой
        
        Args:
            file_paths: список путей к файлам изображений
            output_folder: папка для сохранения PDF (если None, сохраняет рядом с исходным файлом)
            log_callback: функция для логирования
            use_ocr: применять ли OCR к изображениям
            max_image_size: максимальный размер изображения (None = без ограничений, 4000 по умолчанию)
            fit_mode: режим размещения (центр/заполнить/вписать)
            numbering_line1: строка 1 для штампа
            numbering_line2: строка 2 для штампа с автоинкрементом
            numbering_line3: строка 3 для штампа
            numbering_position: позиция штампа
            numbering_border: рисовать рамку
            numbering_increment_mode: режим инкремента ('per_document' или 'per_page')
            progress_callback: функция для обновления прогресса (current, total, message)
        
        Returns:
            список путей к созданным PDF файлам
        """
        _ensure_concurrent_imports()
        
        if not file_paths:
            raise ValueError("Список файлов пуст")
        
        if not PIL_AVAILABLE:
            raise ImportError(
                "Для конвертации изображений в PDF требуется библиотека Pillow:\n"
                "pip install Pillow"
            )
        
        converted_files = []
        errors = []
        total = len(file_paths)
        
        if log_callback:
            log_callback(f"Начало конвертации {total} изображений...")
        
        # ОПТИМИЗАЦИЯ: Для изображений используем меньше потоков (OCR интенсивен)
        try:
            max_workers = min(2, multiprocessing.cpu_count())  # Максимум 2 потока для изображений
        except:
            max_workers = 1
        
        # Создаем задачи с автоинкрементом line2 (если нужно)
        tasks = []
        current_line2 = numbering_line2
        
        if log_callback and numbering_line2:
            log_callback(f"[ИНКРЕМЕНТ] Режим: {numbering_increment_mode}, базовое значение line2: '{numbering_line2}'")
        
        for idx, image_file in enumerate(file_paths):
            tasks.append((image_file, output_folder, use_ocr, max_image_size, fit_mode, 
                         numbering_line1, current_line2, numbering_line3, 
                         numbering_position, numbering_border))
            
            if log_callback and current_line2:
                log_callback(f"[ЗАДАЧА {idx+1}] {os.path.basename(image_file)} -> line2='{current_line2}'")
            
            # Инкрементируем line2 для следующего документа (если режим per_document)
            if numbering_increment_mode == 'per_document' and current_line2:
                old_line2 = current_line2
                current_line2 = GenerationDocApp.increment_line2(current_line2)
                if log_callback:
                    log_callback(f"[ИНКРЕМЕНТ] '{old_line2}' -> '{current_line2}'")
        
        # Обработка одного файла последовательно (быстрее для одного изображения)
        if len(tasks) == 1:
            result = _convert_single_image(tasks[0])
            if result['success']:
                converted_files.append(result['pdf_file'])
                if log_callback:
                    ocr_status = "с OCR" if use_ocr else "без OCR"
                    log_callback(f"  ✓ Создан ({ocr_status}): {os.path.basename(result['pdf_file'])}")
            else:
                errors.append(f"{os.path.basename(result['image_file'])}: {result['error']}")
                if log_callback:
                    log_callback(f"  ✗ Ошибка: {os.path.basename(result['image_file'])}: {result['error']}")
        else:
            # Параллельная обработка для нескольких файлов
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {executor.submit(_convert_single_image, task): task for task in tasks}
                
                completed = 0
                for future in as_completed(futures):
                    completed += 1
                    
                    # Обновляем прогресс
                    if progress_callback:
                        progress_callback(completed, total, f"Конвертация изображений → PDF: {completed}/{total}")
                    
                    try:
                        result = future.result(timeout=600)  # 10 минут таймаут для OCR
                        
                        if result['success']:
                            converted_files.append(result['pdf_file'])
                            if log_callback:
                                ocr_status = "с OCR" if use_ocr else "без OCR"
                                log_callback(f"[{completed}/{total}] ✓ Создан ({ocr_status}): {os.path.basename(result['pdf_file'])}")
                        else:
                            errors.append(f"{os.path.basename(result['image_file'])}: {result['error']}")
                            if log_callback:
                                log_callback(f"[{completed}/{total}] ✗ {os.path.basename(result['image_file'])}: {result['error']}")
                    
                    except Exception as e:
                        task = futures[future]
                        image_file = task[0]
                        errors.append(f"{os.path.basename(image_file)}: Критическая ошибка - {str(e)}")
                        if log_callback:
                            log_callback(f"[{completed}/{total}] ✗ {os.path.basename(image_file)}: {str(e)}")
                    
                    # ОПТИМИЗАЦИЯ: Очистка памяти после каждых 3 файлов
                    if completed % 3 == 0:
                        gc.collect()
        
        # Финальная очистка памяти
        gc.collect()
        
        if not converted_files:
            error_msg = "Ошибки при конвертации всех файлов:\n" + "\n".join(errors)
            if errors:
                raise Exception(error_msg)
            else:
                raise Exception("Не удалось сконвертировать файлы")
        
        if errors:
            error_msg = f"Успешно конвертировано: {len(converted_files)}/{total}\n\nОшибки:\n" + "\n".join(errors[:10])
            if len(errors) > 10:
                error_msg += f"\n... и ещё {len(errors) - 10} ошибок"
            
            if len(errors) == total:
                raise Exception(error_msg)
        
        return converted_files
    
    @staticmethod
    def convert_and_merge_images_to_pdf(file_paths, output_file, log_callback=None, use_ocr=True,
                                        max_image_size=4000, fit_mode='центр',
                                        numbering_line1=None, numbering_line2=None, numbering_line3=None,
                                        numbering_position='правый-нижний', numbering_border=True,
                                        numbering_increment_mode='per_document'):
        """Конвертация изображений в PDF с опциональным OCR и объединение в один файл
        
        Args:
            file_paths: список путей к файлам изображений
            output_file: путь к результирующему PDF файлу
            log_callback: функция для логирования
            use_ocr: применять ли OCR к изображениям
            max_image_size: максимальный размер изображения (None = без ограничений)
            fit_mode: режим размещения (центр/заполнить/вписать)
            numbering_line1: строка 1 для штампа
            numbering_line2: строка 2 для штампа с автоинкрементом
            numbering_line3: строка 3 для штампа
            numbering_position: позиция штампа
            numbering_border: рисовать рамку
            numbering_increment_mode: режим инкремента ('per_document' или 'per_page')
        """
        if not file_paths:
            raise ValueError("Список файлов пуст")
        
        if not PIL_AVAILABLE:
            raise ImportError(
                "Для конвертации изображений в PDF требуется библиотека Pillow:\n"
                "pip install Pillow"
            )
        
        if PdfMerger is None:
            raise ImportError("Требуется установить pypdf или PyPDF2: pip install pypdf")
        
        import tempfile
        import shutil
        temp_dir = tempfile.mkdtemp()
        temp_pdf_files = []
        errors = []
        
        try:
            total = len(file_paths)
            if log_callback:
                ocr_mode = "с OCR" if use_ocr else "без OCR (быстрый режим)"
                log_callback(f"Конвертация {total} изображений {ocr_mode}...")
                if numbering_line2:
                    log_callback(f"[ИНКРЕМЕНТ] Режим: {numbering_increment_mode}, базовое значение line2: '{numbering_line2}'")
            
            for idx, image_file in enumerate(file_paths, 1):
                try:
                    if log_callback:
                        log_callback(f"  [{idx}/{total}] {os.path.basename(image_file)}...")
                    
                    if not os.path.exists(image_file):
                        raise FileNotFoundError(f"Файл не найден: {image_file}")
                    
                    image_file = os.path.abspath(image_file)
                    
                    base_name = os.path.splitext(os.path.basename(image_file))[0]
                    temp_pdf = os.path.join(temp_dir, base_name + ".pdf")
                    
                    # Вычисляем текущий line2 с учетом режима инкремента
                    current_line2 = numbering_line2
                    if numbering_increment_mode == 'per_document' and numbering_line2 and idx > 1:
                        # Инкрементируем line2 для каждого документа
                        for i in range(idx - 1):
                            old_line2 = current_line2
                            current_line2 = GenerationDocApp.increment_line2(current_line2)
                            if log_callback and i == idx - 2:  # Логируем только последний инкремент
                                log_callback(f"[ИНКРЕМЕНТ] Документ {idx}: '{numbering_line2}' -> '{current_line2}'")
                    
                    if log_callback and current_line2:
                        log_callback(f"[ЗАДАЧА {idx}] Нумерация line2='{current_line2}'")
                    
                    # Используем конвертацию с OCR или без
                    if use_ocr:
                        GenerationDocApp.image_to_pdf_with_ocr(image_file, temp_pdf, log_callback, max_image_size, fit_mode,
                                                              numbering_line1, current_line2, numbering_line3,
                                                              numbering_position, numbering_border)
                    else:
                        GenerationDocApp.image_to_pdf_simple(image_file, temp_pdf, log_callback, max_image_size, fit_mode,
                                                            numbering_line1, current_line2, numbering_line3,
                                                            numbering_position, numbering_border)
                    
                    if os.path.exists(temp_pdf):
                        temp_pdf_files.append(temp_pdf)
                        if log_callback:
                            log_callback(f"    ✓ Готово")
                    else:
                        raise Exception("PDF файл не был создан")
                    
                except Exception as e:
                    error_text = str(e)
                    errors.append(f"{os.path.basename(image_file)}: {error_text}")
                    if log_callback:
                        log_callback(f"    ✗ Ошибка: {error_text}")
            
            if not temp_pdf_files:
                raise Exception("Ошибки при конвертации:\n" + "\n".join(errors))
            
            if log_callback:
                log_callback(f"Объединение {len(temp_pdf_files)} PDF файлов...")
            
            merger = PdfMerger()
            
            try:
                for pdf_file in temp_pdf_files:
                    merger.append(pdf_file)
                
                merger.write(output_file)
            
            finally:
                # ОПТИМИЗАЦИЯ: Гарантированное закрытие merger
                try:
                    merger.close()
                except:
                    pass
                
                # Очистка памяти
                gc.collect()
            
            if log_callback:
                log_callback("✓ Объединение завершено")
            
            if errors:
                raise Warning(f"Файл создан, но были ошибки при конвертации некоторых изображений:\n" + "\n".join(errors))
            
        finally:
            # ОПТИМИЗАЦИЯ: Гарантированное удаление временной папки
            import shutil
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            
            # Финальная очистка памяти
            gc.collect()
            # Финальная очистка памяти
            gc.collect()
    
    @staticmethod
    def split_pdf_file(input_pdf, output_folder, ranges_text, log_callback=None):
        """
        Разделение PDF файла на части по указанным диапазонам
        
        Args:
            input_pdf: путь к исходному PDF
            output_folder: папка для сохранения результатов
            ranges_text: диапазоны страниц в формате "1-2; 3-5; 6; 7-13"
            log_callback: функция для логирования
        
        Returns:
            list: список созданных файлов
        """
        from PyPDF2 import PdfReader, PdfWriter
        
        if log_callback:
            log_callback(f"Разделение PDF: {os.path.basename(input_pdf)}")
        
        pdf_reader = PdfReader(input_pdf)
        total_pages = len(pdf_reader.pages)
        base_name = os.path.splitext(os.path.basename(input_pdf))[0]
        created_files = []
        
        # Режим: разделение по отдельным страницам
        if ranges_text == "all_pages":
            if log_callback:
                log_callback(f"Режим: разделение по отдельным страницам (всего {total_pages} стр.)")
            
            ranges = [(i, i) for i in range(total_pages)]
        else:
            # Режим: разделение по диапазонам
            if not ranges_text or ranges_text.strip() == "":
                raise ValueError("Укажите диапазоны страниц!")
            
            if log_callback:
                log_callback(f"Режим: разделение по диапазонам ({ranges_text})")
            
            # Парсим диапазоны (разделитель ; или ,)
            ranges = []
            for part in ranges_text.replace(';', ',').split(','):
                part = part.strip()
                if not part:
                    continue
                
                if '-' in part:
                    # Диапазон вида 1-5
                    start, end = part.split('-', 1)
                    start = int(start.strip()) - 1  # Convert to 0-based
                    end = int(end.strip()) - 1
                    ranges.append((start, end))
                else:
                    # Одна страница
                    page = int(part.strip()) - 1
                    ranges.append((page, page))
            
            if not ranges:
                raise ValueError("Не удалось распознать диапазоны страниц!")
        
        # Создаём файлы для каждого диапазона
        for part_num, (start_page, end_page) in enumerate(ranges, 1):
            # Проверяем границы
            if start_page < 0 or end_page >= total_pages:
                if log_callback:
                    log_callback(f"  ⚠ Пропущен диапазон {start_page+1}-{end_page+1} (выходит за границы PDF)")
                continue
            
            output_filename = os.path.join(output_folder, f"{base_name}_part{part_num:03d}.pdf")
            pdf_writer = PdfWriter()
            
            for page_num in range(start_page, end_page + 1):
                pdf_writer.add_page(pdf_reader.pages[page_num])
            
            with open(output_filename, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            created_files.append(output_filename)
            if log_callback:
                if start_page == end_page:
                    log_callback(f"  ✓ Часть {part_num}: страница {start_page+1}")
                else:
                    log_callback(f"  ✓ Часть {part_num}: страницы {start_page+1}-{end_page+1}")
        
        return created_files
    
    @staticmethod
    def rotate_pdf_pages(input_pdf, output_pdf, angle=90, pages_range="", log_callback=None):
        """
        Поворот страниц PDF
        
        Args:
            input_pdf: путь к исходному PDF
            output_pdf: путь к выходному PDF
            angle: угол поворота (90, 180, 270)
            pages_range: диапазон страниц (пусто=все, "1,3,5" или "1-10" или "1-5,10,15-20")
            log_callback: функция для логирования
        """
        from PyPDF2 import PdfReader, PdfWriter
        
        if log_callback:
            log_callback(f"Поворот страниц PDF: {os.path.basename(input_pdf)}")
            log_callback(f"Угол: {angle}°")
        
        pdf_reader = PdfReader(input_pdf)
        pdf_writer = PdfWriter()
        total_pages = len(pdf_reader.pages)
        
        # Парсим диапазон страниц
        pages_to_rotate = set()
        if not pages_range or pages_range.strip() == "":
            pages_to_rotate = set(range(total_pages))
        else:
            for part in pages_range.split(','):
                part = part.strip()
                if '-' in part:
                    start, end = part.split('-')
                    start = int(start.strip()) - 1  # Convert to 0-based
                    end = int(end.strip()) - 1
                    pages_to_rotate.update(range(start, end + 1))
                else:
                    pages_to_rotate.add(int(part.strip()) - 1)
        
        # Фильтруем недопустимые номера
        pages_to_rotate = {p for p in pages_to_rotate if 0 <= p < total_pages}
        
        if log_callback:
            if len(pages_to_rotate) == total_pages:
                log_callback(f"Поворачиваются все страницы ({total_pages})")
            else:
                log_callback(f"Поворачиваются страницы: {len(pages_to_rotate)} из {total_pages}")
        
        # Применяем поворот
        for page_num in range(total_pages):
            page = pdf_reader.pages[page_num]
            if page_num in pages_to_rotate:
                page.rotate(int(angle))
            pdf_writer.add_page(page)
        
        with open(output_pdf, 'wb') as output_file:
            pdf_writer.write(output_file)
        
        if log_callback:
            log_callback(f"✓ PDF сохранен: {os.path.basename(output_pdf)}")
    
    @staticmethod
    def extract_pdf_pages(input_pdf, output_pdf, pages_range, log_callback=None):
        """
        Извлечение конкретных страниц из PDF
        
        Args:
            input_pdf: путь к исходному PDF
            output_pdf: путь к выходному PDF
            pages_range: диапазон страниц ("1-5", "1,3,5", "1-5,10,15-20")
            log_callback: функция для логирования
        """
        from PyPDF2 import PdfReader, PdfWriter
        
        if log_callback:
            log_callback(f"Извлечение страниц из PDF: {os.path.basename(input_pdf)}")
            log_callback(f"Диапазон: {pages_range}")
        
        if not pages_range or pages_range.strip() == "":
            raise ValueError("Укажите страницы для извлечения!")
        
        pdf_reader = PdfReader(input_pdf)
        pdf_writer = PdfWriter()
        total_pages = len(pdf_reader.pages)
        
        # Парсим диапазон страниц
        pages_to_extract = []
        for part in pages_range.split(','):
            part = part.strip()
            if '-' in part:
                start, end = part.split('-')
                start = int(start.strip()) - 1  # Convert to 0-based
                end = int(end.strip()) - 1
                pages_to_extract.extend(range(start, end + 1))
            else:
                pages_to_extract.append(int(part.strip()) - 1)
        
        # Фильтруем недопустимые номера и убираем дубликаты
        pages_to_extract = sorted(set(p for p in pages_to_extract if 0 <= p < total_pages))
        
        if not pages_to_extract:
            raise ValueError(f"Не найдено допустимых страниц! Всего страниц в PDF: {total_pages}")
        
        if log_callback:
            # Показываем первые и последние страницы
            if len(pages_to_extract) <= 10:
                pages_str = ', '.join(str(p+1) for p in pages_to_extract)
            else:
                first_five = ', '.join(str(p+1) for p in pages_to_extract[:5])
                last_five = ', '.join(str(p+1) for p in pages_to_extract[-5:])
                pages_str = f"{first_five} ... {last_five}"
            log_callback(f"Извлекается {len(pages_to_extract)} страниц: {pages_str}")
        
        # Извлекаем страницы
        for page_num in pages_to_extract:
            pdf_writer.add_page(pdf_reader.pages[page_num])
        
        with open(output_pdf, 'wb') as output_file:
            pdf_writer.write(output_file)
        
        if log_callback:
            log_callback(f"✓ PDF сохранен: {os.path.basename(output_pdf)} ({len(pages_to_extract)} стр.)")
    
    @staticmethod
    def extract_data_to_excel(file_paths, output_excel, method="full_text", 
                              regex_pattern=None, log_callback=None, progress_callback=None):
        """
        Извлекает данные из документов (Word/PDF) и сохраняет в Excel
        
        Args:
            file_paths: список путей к файлам для обработки
            output_excel: путь к выходному Excel файлу
            method: метод извлечения ('full_text', 'tables_only', 'regex')
            regex_pattern: регулярное выражение для извлечения (если method='regex')
            log_callback: функция для логирования
            progress_callback: функция для обновления прогресса
        """
        from openpyxl import Workbook
        from docx import Document
        import re
        
        if not file_paths:
            raise ValueError("Список файлов пуст")
        
        if log_callback:
            log_callback(f"Начало извлечения данных из {len(file_paths)} файлов...")
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Extracted Data"
        
        # Заголовки
        headers = ["Файл", "Путь", "Тип", "Извлеченный текст"]
        if method == "tables_only":
            headers.append("Таблица #")
        ws.append(headers)
        
        total = len(file_paths)
        extracted_count = 0
        errors = []
        
        for idx, file_path in enumerate(file_paths, 1):
            if progress_callback:
                progress_callback(idx, total, f"Обработка: {os.path.basename(file_path)}")
            
            try:
                filename = os.path.basename(file_path)
                file_ext = os.path.splitext(file_path)[1].lower()
                
                extracted_text = ""
                
                if file_ext in ['.docx', '.doc']:
                    # Извлечение из Word
                    doc = Document(file_path)
                    
                    if method == "full_text":
                        # Полный текст
                        extracted_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                        ws.append([filename, file_path, "Word", extracted_text])
                        
                    elif method == "tables_only":
                        # Только таблицы
                        for table_idx, table in enumerate(doc.tables, 1):
                            table_text = []
                            for row in table.rows:
                                row_text = [cell.text.strip() for cell in row.cells]
                                table_text.append(" | ".join(row_text))
                            extracted_text = "\n".join(table_text)
                            ws.append([filename, file_path, "Word", extracted_text, table_idx])
                        
                        if not doc.tables:
                            ws.append([filename, file_path, "Word", "(нет таблиц)", 0])
                    
                    elif method == "regex":
                        # По регулярному выражению
                        full_text = "\n".join([para.text for para in doc.paragraphs])
                        
                        if regex_pattern:
                            matches = re.findall(regex_pattern, full_text, re.MULTILINE | re.IGNORECASE)
                            if matches:
                                extracted_text = "\n".join([str(m) for m in matches])
                            else:
                                extracted_text = "(совпадений не найдено)"
                        else:
                            extracted_text = "(regex не указан)"
                        
                        ws.append([filename, file_path, "Word", extracted_text])
                
                elif file_ext == '.pdf':
                    # Извлечение из PDF
                    from PyPDF2 import PdfReader
                    
                    pdf = PdfReader(file_path)
                    
                    if method == "full_text":
                        # Полный текст
                        all_text = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                all_text.append(page_text)
                        extracted_text = "\n".join(all_text)
                        ws.append([filename, file_path, "PDF", extracted_text])
                    
                    elif method == "tables_only":
                        # PDF таблицы сложнее извлекать, используем текст
                        ws.append([filename, file_path, "PDF", "(извлечение таблиц из PDF требует библиотеку tabula-py)", 0])
                    
                    elif method == "regex":
                        # По регулярному выражению
                        all_text = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                all_text.append(page_text)
                        full_text = "\n".join(all_text)
                        
                        if regex_pattern:
                            matches = re.findall(regex_pattern, full_text, re.MULTILINE | re.IGNORECASE)
                            if matches:
                                extracted_text = "\n".join([str(m) for m in matches])
                            else:
                                extracted_text = "(совпадений не найдено)"
                        else:
                            extracted_text = "(regex не указан)"
                        
                        ws.append([filename, file_path, "PDF", extracted_text])
                
                else:
                    # Неподдерживаемый формат
                    ws.append([filename, file_path, "Unknown", "(неподдерживаемый формат)"])
                    if log_callback:
                        log_callback(f"  ⚠ Пропущен (неподдерживаемый формат): {filename}")
                    continue
                
                extracted_count += 1
                if log_callback:
                    log_callback(f"[{idx}/{total}] ✓ {filename}")
            
            except Exception as e:
                error_msg = f"{os.path.basename(file_path)}: {str(e)}"
                errors.append(error_msg)
                ws.append([os.path.basename(file_path), file_path, "Error", f"ОШИБКА: {str(e)}"])
                if log_callback:
                    log_callback(f"[{idx}/{total}] ✗ {error_msg}")
        
        # Автоматически подгоняем ширину столбцов
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = min(max_length + 2, 100)  # Ограничиваем максимум
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Сохраняем Excel
        wb.save(output_excel)
        
        if log_callback:
            log_callback(f"\n✅ Извлечение завершено!")
            log_callback(f"  Успешно обработано: {extracted_count} файлов")
            if errors:
                log_callback(f"  Ошибок: {len(errors)}")
                for error in errors[:5]:
                    log_callback(f"    - {error}")
                if len(errors) > 5:
                    log_callback(f"    ... и ещё {len(errors) - 5} ошибок")

class MergeDocumentsWindow:
    """Окно объединения документов с системой вкладок"""
    def __init__(self, parent):
        self.window = tk.Toplevel(parent)
        self.window.withdraw()
        self.window.title("Работа с документами")
        self.window.geometry("750x900")
        self.window.transient(parent)
        
        self.window.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (750 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (900 // 2)
        self.window.geometry(f"+{x}+{y}")
        
        self.window.deiconify()
        
        # Система вкладок
        self.tabs = []  # Список объектов MergeTabTask
        self.tab_counter = 0  # Счетчик для уникальных ID вкладок
        self.max_tabs = 5  # Максимальное количество вкладок
        
        self.create_widgets()
        
        self.add_tab()
    
    def create_widgets(self):
        """Создание интерфейса с системой вкладок"""
        title_frame = tk.Frame(self.window, bg=COLORS["primary"], height=50)
        title_frame.pack(fill=tk.X)
        
        title_label = tk.Label(
            title_frame,
            text="Работа с документами",
            font=FONTS["title"],
            bg=COLORS["primary"],
            fg="white"
        )
        title_label.pack(pady=12)
        
        tabs_container = tk.Frame(self.window, bg=COLORS["bg_secondary"])
        tabs_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=(5, 10))
        
        tabs_control = tk.Frame(tabs_container, bg=COLORS["bg_secondary"], height=35)
        tabs_control.pack(fill=tk.X, pady=(0, 5))
        
        self.add_tab_btn = tk.Button(
            tabs_control,
            text="➕",
            command=self.add_tab,
            bg=COLORS["success"],
            fg="white",
            font=("Segoe UI", 10),
            relief=tk.FLAT,
            cursor="hand2",
            width=2,
            height=1,
            activebackground=COLORS["success_hover"]
        )
        self.add_tab_btn.pack(side=tk.LEFT, padx=5)
        ToolTip(self.add_tab_btn, f"Добавить новую вкладку (макс. {self.max_tabs})")
        
        self.tab_count_label = tk.Label(
            tabs_control,
            text=f"Вкладок: 0/{self.max_tabs}",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_secondary"]
        )
        self.tab_count_label.pack(side=tk.LEFT, padx=10)
        
        # ═══════════════════════════════════════════════════════════
        # CUSTOM BUTTON TABS (КАК В PLACEHOLDEREDITOR)
        # ═══════════════════════════════════════════════════════════
        
        self.tabs_buttons_frame = tk.Frame(
            tabs_container,
            bg=COLORS["bg_secondary"],
            height=38
        )
        self.tabs_buttons_frame.pack(fill=tk.X, pady=(0, SPACING["sm"]))
        self.tabs_buttons_frame.pack_propagate(False)
        
        self.tabs_content_frame = tk.Frame(
            tabs_container,
            bg=COLORS["bg_secondary"]
        )
        self.tabs_content_frame.pack(fill=tk.BOTH, expand=True)
        
        self.tab_buttons = []
        self.tab_frames = []
        self.current_tab_index = 0  # Индекс активной вкладки
    
    def add_tab(self):
        """Добавление новой вкладки"""
        if len(self.tabs) >= self.max_tabs:
            messagebox.showinfo(
                "Ограничение", 
                f"Достигнуто максимальное количество вкладок ({self.max_tabs})",
                parent=self.window
            )
            return
        
        tab_frame = tk.Frame(self.tabs_content_frame, bg=COLORS["bg_secondary"])
        self.tab_frames.append(tab_frame)
        
        self.tab_counter += 1
        tab = MergeTabTask(tab_frame, self, self.tab_counter)
        self.tabs.append(tab)
        
        tab.tab_name = f"Задача {self.tab_counter}"
        
        self.create_tab_button(tab.tab_name, len(self.tabs) - 1)
        
        # Переключаемся на новую вкладку
        self.switch_tab(len(self.tabs) - 1)
        
        self.update_tab_controls()
    
    def close_tab(self, tab_index=None):
        """Закрытие вкладки"""
        if tab_index is None:
            tab_index = self.current_tab_index
        
        if len(self.tabs) <= 1:
            messagebox.showinfo("Информация", "Нельзя закрыть последнюю вкладку", parent=self.window)
            return
        
        tab = self.tabs[tab_index]
        if tab.is_processing:
            result = messagebox.askyesno(
                "Подтверждение",
                "На этой вкладке выполняется задача. Закрыть вкладку?",
                parent=self.window
            )
            if not result:
                return
        
        self.tab_buttons[tab_index][0].destroy()
        del self.tab_buttons[tab_index]
        
        self.tab_frames[tab_index].destroy()
        del self.tab_frames[tab_index]
        
        del self.tabs[tab_index]
        
        if tab_index >= len(self.tabs):
            tab_index = len(self.tabs) - 1
        self.switch_tab(tab_index)
        
        self.update_tab_controls()
    
    def create_tab_button(self, tab_name, tab_index):
        """Создание кнопки вкладки"""
        btn_container = tk.Frame(
            self.tabs_buttons_frame,
            bg=COLORS["bg_tertiary"],
            highlightthickness=1,
            highlightbackground=COLORS["border"]
        )
        btn_container.pack(side=tk.LEFT, padx=2)
        
        tab_btn = tk.Button(
            btn_container,
            text=tab_name,
            font=FONTS["button"],
            bg=COLORS["bg_tertiary"],
            activebackground=COLORS["primary_light"],
            fg=COLORS["text_primary"],
            activeforeground="white",
            relief=tk.FLAT,
            cursor="hand2",
            padx=10,
            pady=5,
            command=lambda: self.switch_tab(tab_index)
        )
        tab_btn.pack(side=tk.LEFT)
        
        close_btn = tk.Button(
            btn_container,
            text="✕",
            font=("Segoe UI", 8, "bold"),
            bg=COLORS["bg_tertiary"],
            activebackground=COLORS["danger"],
            fg=COLORS["text_secondary"],
            activeforeground="white",
            relief=tk.FLAT,
            cursor="hand2",
            padx=5,
            pady=5,
            command=lambda: self.close_tab(tab_index)
        )
        close_btn.pack(side=tk.LEFT)
        ToolTip(close_btn, "Закрыть вкладку")
        
        def show_context_menu(event):
            # Динамически находим текущий индекс этой кнопки
            current_index = None
            for i, (cont, btn, cls) in enumerate(self.tab_buttons):
                if btn == tab_btn:
                    current_index = i
                    break
            
            if current_index is None:
                return
            
            menu = ModernContextMenu(self.window)
            menu.add_command(
                label="Переименовать вкладку",
                command=lambda: self.rename_tab(current_index)
            )
            menu.add_separator()
            menu.add_command(
                label="Закрыть вкладку",
                command=lambda: self.close_tab(current_index)
            )
            menu.post(event.x_root, event.y_root)
        
        tab_btn.bind("<Button-3>", show_context_menu)
        close_btn.bind("<Button-2>", lambda e: self.close_tab(tab_index))
        
        self.tab_buttons.append((btn_container, tab_btn, close_btn))
    
    def switch_tab(self, tab_index):
        """Переключение между вкладками"""
        if tab_index < 0 or tab_index >= len(self.tabs):
            return
        
        self.current_tab_index = tab_index
        
        for i, (container, tab_btn, close_btn) in enumerate(self.tab_buttons):
            if i == tab_index:
                container.configure(
                    bg=COLORS["primary"],
                    highlightbackground=COLORS["primary"]
                )
                tab_btn.configure(
                    bg=COLORS["primary"],
                    fg="white",
                    font=("Segoe UI", 10, "bold")
                )
                close_btn.configure(
                    bg=COLORS["primary"],
                    fg="white"
                )
            else:
                container.configure(
                    bg=COLORS["bg_tertiary"],
                    highlightbackground=COLORS["border"]
                )
                tab_btn.configure(
                    bg=COLORS["bg_tertiary"],
                    fg=COLORS["text_primary"],
                    font=FONTS["button"]
                )
                close_btn.configure(
                    bg=COLORS["bg_tertiary"],
                    fg=COLORS["text_secondary"]
                )
        
        for i, frame in enumerate(self.tab_frames):
            if i == tab_index:
                frame.pack(fill=tk.BOTH, expand=True)
            else:
                frame.pack_forget()
    
    def rename_tab(self, tab_index):
        """Переименование вкладки"""
        if tab_index < 0 or tab_index >= len(self.tabs):
            return
        
        tab = self.tabs[tab_index]
        current_name = tab.tab_name
        
        # Диалог ввода нового имени
        dialog = SimpleInputDialog(
            self.window,
            "Переименовать вкладку",
            "Введите новое название вкладки:",
            default_value=current_name
        )
        self.window.wait_window(dialog.top)
        
        if dialog.result and dialog.result.strip():
            new_name = dialog.result.strip()
            tab.tab_name = new_name
            container, tab_btn, close_btn = self.tab_buttons[tab_index]
            tab_btn.configure(text=new_name)
            tab.tab_name = new_name
            self.notebook.tab(tab_index, text=f"{new_name} [x]")
    
    def update_tab_controls(self):
        """Обновление счетчика и состояния кнопки добавления"""
        count = len(self.tabs)
        self.tab_count_label.config(text=f"Вкладок: {count}/{self.max_tabs}")
        
        if count >= self.max_tabs:
            self.add_tab_btn.config(state="disabled")
        else:
            self.add_tab_btn.config(state="normal")

class ExcelConstructorWindow:
    """Конструктор для преобразования Excel файлов"""
    def __init__(self, parent, main_app):
        self.main_app = main_app
        self.window = tk.Toplevel(parent)
        self.window.withdraw()
        self.window.title("Конструктор Excel")
        self.window.geometry("1000x750")
        self.window.transient(parent)
        
        self.window.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (1000 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (750 // 2)
        self.window.geometry(f"+{x}+{y}")
        
        self.window.deiconify()
        self.window.grab_set()
        
        self.source_df = None
        self.source_file = None
        self.source_file_path = tk.StringVar(value="")
        
        self.create_widgets()
    
    def create_widgets(self):
        top_frame = tk.Frame(self.window, bg=COLORS["primary"], height=70)
        top_frame.pack(fill=tk.X, padx=0, pady=0)
        top_frame.pack_propagate(False)
        
        tk.Label(top_frame, text="Исходный файл:", bg=COLORS["primary"], fg="white", font=FONTS["heading"]).pack(side=tk.LEFT, padx=(15, 8), pady=18)
        
        # Поле с поддержкой drag-and-drop
        self.file_entry = ctk.CTkEntry(
            top_frame,
            textvariable=self.source_file_path,
            font=FONTS["body"],
            state="readonly",
            fg_color="white",
            text_color=COLORS["text_primary"],
            border_color=COLORS["border"],
            height=40,
            placeholder_text="Файл не выбран"
        )
        self.file_entry.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 8), pady=18)
        ToolTip(self.file_entry, "Путь к загруженному Excel файлу")
        
        # Настройка drag-and-drop для поля
        def on_excel_drop(file_path):
            self.load_excel_file(file_path)
        
        setup_file_drop(
            self.file_entry,
            self.source_file_path,
            file_types=['.xlsx', '.xls'],
            on_drop_callback=on_excel_drop,
            parent_window=self.window
        )
        
        self.rows_label = tk.Label(top_frame, text="", bg=COLORS["primary"], font=FONTS["small"], fg="white")
        self.rows_label.pack(side=tk.LEFT, padx=8)
        
        load_btn = create_modern_button(
            top_frame,
            text="Загрузить Excel",
            command=self.load_excel,
            style="primary",
            width=110,
            height=40,
            tooltip="Загрузить Excel файл для преобразования"
        )
        load_btn.pack(side=tk.RIGHT, padx=(5, 15), pady=18)
        
        self.preview_btn = create_modern_button(
            top_frame,
            text="👁 Просмотр",
            command=self.preview_source_file,
            style="success",
            width=95,
            height=40,
            tooltip="Предварительный просмотр загруженного файла"
        )
        self.preview_btn.configure(state=tk.DISABLED)
        self.preview_btn.pack(side=tk.RIGHT, padx=5, pady=18)
        
        # Основная рабочая область с фиксированными пропорциями
        main_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)
        
        right_frame = tk.LabelFrame(
            main_frame, 
            text=" Инструкция ", 
            bg=COLORS["bg_secondary"], 
            font=FONTS["heading"], 
            fg=COLORS["text_primary"], 
            width=360,
            relief=tk.SOLID,
            borderwidth=1
        )
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=False, padx=(8, 0))
        right_frame.pack_propagate(False)
        
        # Инструкция
        instruction_text = """Инструкция:

1. Загрузите Excel файл с исходными данными

2. Проверьте данные в таблице слева

3. Нажмите "Сформировать файл"

4. В новом окне выберите:
   • Плейсхолдер (название новой колонки)
   • Исходные колонки для объединения

5. Повторите для всех нужных плейсхолдеров

6. Нажмите "Создать файл"

Даты автоматически форматируются в дд.мм.гггг
Лишние пробелы удаляются"""
        
        tk.Label(
            right_frame,
            text=instruction_text,
            bg=COLORS["bg_tertiary"],
            font=FONTS["small"],
            fg=COLORS["text_primary"],
            justify=tk.LEFT,
            padx=12,
            pady=12,
            relief=tk.FLAT
        ).pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
        
        self.build_button = create_modern_button(
            right_frame,
            text="▶ Сформировать файл",
            command=self.open_file_builder,
            style="success",
            height=48,
            tooltip="Открыть окно построения файла с выбором колонок"
        )
        self.build_button.configure(state=tk.DISABLED)
        self.build_button.pack(fill=tk.X, padx=8, pady=8)
        
        left_frame = tk.LabelFrame(
            main_frame, 
            text=" Исходные данные ", 
            bg=COLORS["bg_secondary"], 
            font=FONTS["heading"], 
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 8))
        left_frame.pack_propagate(True)
        
        # Фрейм для Treeview с современным стилем
        tree_frame = tk.Frame(left_frame, bg=COLORS["bg_secondary"])
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
        
        tree_container, self.data_tree = create_modern_treeview(
            tree_frame,
            columns=[],  # Колонки будут настроены в display_preview
            height=18,
            selectable=True
        )
        tree_container.pack(fill=tk.BOTH, expand=True)
    
    def load_excel(self):
        """Загрузка Excel файла через диалог"""
        # Пытаемся использовать кэш из главного приложения
        if hasattr(self.main_app, '_pandas_loaded') and self.main_app._pandas_loaded:
            pd = self.main_app._pandas
        else:
            import pandas as pd
        
        file_path = filedialog.askopenfilename(
            title="Выберите Excel файл",
            filetypes=[("Excel файлы", "*.xlsx *.xls"), ("Все файлы", "*.*")]
        )
        
        if not file_path:
            return
        
        self.load_excel_file(file_path, pd)
    
    def load_excel_file(self, file_path, pd=None):
        """Загрузка Excel файла по пути (используется и для drag-and-drop)"""
        # Если pandas не передан, загружаем
        if pd is None:
            if hasattr(self.main_app, '_pandas_loaded') and self.main_app._pandas_loaded:
                pd = self.main_app._pandas
            else:
                import pandas as pd
        
        try:
            self.source_df = pd.read_excel(file_path)
            
            # Форматируем все значения (даты преобразуем в формат дд.мм.гггг)
            for col in self.source_df.columns:
                self.source_df[col] = self.source_df[col].apply(
                    lambda x: self.format_date_value(x) if pd.notna(x) else ""
                )
            
            # Заменяем NaN на пустые строки (на всякий случай)
            self.source_df = self.source_df.fillna("")
            
            self.source_file = file_path
            self.source_file_path.set(file_path)
            
            self.display_preview()
            
            # Активируем кнопки
            self.build_button.configure(state=tk.NORMAL)
            self.preview_btn.configure(state=tk.NORMAL)
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при загрузке файла:\n{str(e)}", parent=self.window)
    
    def preview_source_file(self):
        """Предварительный просмотр исходного Excel файла"""
        if not self.source_file or not os.path.exists(self.source_file):
            messagebox.showwarning("Предупреждение", "Сначала загрузите Excel файл!")
            return
        
        try:
            PreviewWindow(self.window, self.source_file, f"Просмотр: {os.path.basename(self.source_file)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{str(e)}")
    
    def display_preview(self):
        """Отображает все строки из файла"""
        if self.source_df is None:
            return
        
        self.data_tree.delete(*self.data_tree.get_children())
        
        columns = list(self.source_df.columns)
        self.data_tree["columns"] = columns
        self.data_tree["show"] = "headings"
        
        # Заголовки с современным стилем
        for col in columns:
            # Вычисляем оптимальную ширину на основе заголовка
            header_width = max(len(col) * 9, 120)
            self.data_tree.heading(col, text=col)
            # stretch=tk.NO для фиксации размера при изменении
            self.data_tree.column(col, width=header_width, minwidth=80, anchor=tk.W, stretch=tk.NO)
        
        # ВСЕ строки с форматированием дат
        for idx, row in self.source_df.iterrows():
            values = format_dataframe_row(row)
            insert_treeview_row(self.data_tree, values)
        
        self.rows_label.config(text=f"Строк: {len(self.source_df)}")
    
    def format_date_value(self, value):
        """Универсальное форматирование значения с обработкой дат и времени"""
        import pandas as pd
        
        if pd.isna(value) or value == "":
            return ""
        
        if isinstance(value, pd.Timestamp) or hasattr(value, 'strftime'):
            return value.strftime('%d.%m.%Y')
        
        value_str = str(value).strip()
        
        if not value_str:
            return ""
        
        if len(value_str) == 10 and value_str[2] == '.' and value_str[5] == '.':
            try:
                datetime.strptime(value_str, '%d.%m.%Y')
                return value_str
            except:
                pass
        
        # Проверка на чистое время (без даты): 9:00, 9 ч 00 мин и т.п.
        time_indicators = [':']
        date_indicators = ['.', '-', '/']
        has_time = any(ind in value_str for ind in time_indicators) or ' ч ' in value_str.lower() or value_str.lower().endswith(' ч')
        has_date = any(ind in value_str for ind in date_indicators)
        
        # Если есть признаки времени, но нет признаков даты - возвращаем как есть
        if has_time and not has_date:
            return value_str
        
        if ' ' in value_str:
            try:
                dt = pd.to_datetime(value_str, dayfirst=True, errors='coerce')
                if pd.notna(dt):
                    return dt.strftime('%d.%m.%Y')
            except:
                pass
        
        try:
            dt = pd.to_datetime(value_str, dayfirst=True, errors='coerce')
            if pd.notna(dt):
                return dt.strftime('%d.%m.%Y')
        except:
            pass
        
        return value_str
    
    def open_file_builder(self):
        """Открыть окно построения файла"""
        if self.source_df is None:
            messagebox.showwarning("Предупреждение", "Сначала загрузите Excel файл")
            return
        
        self.window.withdraw()
        
        FileBuilderWindow(self.main_app.root, self.source_df, self.main_app, self.source_file, self)

class FileBuilderWindow:
    """Окно построения файла с выбором колонок для плейсхолдеров"""
    def __init__(self, parent, source_df, main_app, source_file, constructor_window=None):
        self.source_df = source_df
        self.main_app = main_app
        self.source_file = source_file
        self.constructor_window = constructor_window  # Ссылка на окно конструктора
        self.column_mappings = []  # [{placeholder: "...", source_columns: [...]}, ...]
        
        self.window = tk.Toplevel(parent)
        self.window.title("Построение файла")
        self.window.geometry("900x700")
        self.window.transient(parent)
        self.window.grab_set()
        
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        self.create_widgets()
        
        self.window.update_idletasks()
        parent.update_idletasks()
        
        x = parent.winfo_rootx() + (parent.winfo_width() // 2) - (900 // 2)
        y = parent.winfo_rooty() + (parent.winfo_height() // 2) - (700 // 2)
        
        self.window.geometry(f"+{x}+{y}")
    
    def create_widgets(self):
        header_frame = tk.Frame(self.window, bg=COLORS["success"])
        header_frame.pack(fill=tk.X)
        
        if self.constructor_window:
            back_btn = tk.Button(
                header_frame,
                text="← Назад",
                command=self.back_to_constructor,
                bg=COLORS["success"],
                fg="white",
                font=FONTS["button"],
                relief=tk.FLAT,
                cursor="hand2",
                padx=15,
                pady=15,
                bd=0
            )
            back_btn.pack(side=tk.LEFT)
            
            ToolTip(back_btn, "Вернуться к загрузке Excel файла")
            
            def on_enter(e):
                back_btn.config(bg=COLORS["success_hover"])
            def on_leave(e):
                back_btn.config(bg=COLORS["success"])
            back_btn.bind("<Enter>", on_enter)
            back_btn.bind("<Leave>", on_leave)
        
        header_label = tk.Label(
            header_frame,
            text="Настройка колонок для нового файла",
            bg=COLORS["success"],
            fg="white",
            font=FONTS["title"],
            pady=15
        )
        header_label.pack(side=tk.LEFT, expand=True)
        
        info_frame = tk.Frame(self.window, bg=COLORS["bg_tertiary"])
        info_frame.pack(fill=tk.X, padx=12, pady=12)
        
        tk.Label(
            info_frame,
            text=f"Исходный файл: {os.path.basename(self.source_file)} | Строк: {len(self.source_df)} | Колонок: {len(self.source_df.columns)}",
            bg=COLORS["bg_tertiary"],
            font=FONTS["body"],
            fg=COLORS["text_primary"]
        ).pack(pady=8)
        
        main_frame = tk.PanedWindow(self.window, orient=tk.HORIZONTAL, bg=COLORS["bg_secondary"], sashwidth=5)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 12))
        
        left_panel = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        main_frame.add(left_panel, minsize=400)
        
        add_btn = create_modern_button(
            left_panel,
            text="+ Добавить колонку",
            command=self.add_column_mapping,
            style="primary",
            tooltip="Добавить новую колонку в результирующий файл"
        )
        add_btn.pack(fill=tk.X, pady=(8, 12), padx=8)
        
        canvas_frame = tk.Frame(left_panel, bg=COLORS["bg_secondary"])
        canvas_frame.pack(fill=tk.BOTH, expand=True, padx=8)
        
        canvas = tk.Canvas(canvas_frame, bg=COLORS["bg_secondary"], highlightthickness=0)
        scrollbar = tk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
        
        self.mappings_frame = tk.Frame(canvas, bg=COLORS["bg_secondary"])
        
        self.mappings_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=self.mappings_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.canvas = canvas
        
        def _on_mousewheel(event):
            # Получаем виджет под курсором
            x, y = event.x_root, event.y_root
            widget_under_cursor = self.window.winfo_containing(x, y)
            
            # Не перехватываем скролл для Listbox и Text (включая ScrolledText)
            if widget_under_cursor and (isinstance(widget_under_cursor, tk.Listbox) or isinstance(widget_under_cursor, tk.Text)):
                return  # Виджет сам обработает скролл
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def _on_mouse_button(event):
            if event.num == 4:
                canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                canvas.yview_scroll(1, "units")
        
        # Сохраняем обработчики для привязки к динамически создаваемым виджетам
        self._on_mousewheel = _on_mousewheel
        self._on_mouse_button = _on_mouse_button
        
        canvas.bind("<MouseWheel>", _on_mousewheel)
        canvas.bind("<Button-4>", _on_mouse_button)  # Linux - прокрутка вверх
        canvas.bind("<Button-5>", _on_mouse_button)  # Linux - прокрутка вниз
        canvas_frame.bind("<MouseWheel>", _on_mousewheel)
        canvas_frame.bind("<Button-4>", _on_mouse_button)
        canvas_frame.bind("<Button-5>", _on_mouse_button)
        self.mappings_frame.bind("<MouseWheel>", _on_mousewheel)
        self.mappings_frame.bind("<Button-4>", _on_mouse_button)
        self.mappings_frame.bind("<Button-5>", _on_mouse_button)
        
        def _bind_to_mousewheel(event):
            canvas.bind("<MouseWheel>", _on_mousewheel)
            canvas.bind("<Button-4>", _on_mouse_button)
            canvas.bind("<Button-5>", _on_mouse_button)
        
        def _unbind_from_mousewheel(event):
            canvas.unbind("<MouseWheel>")
            canvas.unbind("<Button-4>")
            canvas.unbind("<Button-5>")
        
        canvas_frame.bind('<Enter>', _bind_to_mousewheel)
        canvas_frame.bind('<Leave>', _unbind_from_mousewheel)
        
        right_panel = tk.Frame(main_frame, bg=COLORS["bg_secondary"])
        main_frame.add(right_panel, minsize=400)
        
        log_frame = tk.LabelFrame(
            right_panel,
            text=" Лог выполнения ",
            font=FONTS["heading"],
            padx=12,
            pady=12,
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1
        )
        log_frame.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
        
        self.log_text = ScrolledText(
            log_frame,
            height=8,
            wrap=tk.WORD,
            bg=COLORS["bg_secondary"],
            font=FONTS["body"],
            relief=tk.FLAT,
            borderwidth=0
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        def show_context_menu(event):
            menu = ModernContextMenu(self.log_text)
            menu.add_command(label="Копировать", command=self.copy_log_text)
            menu.add_command(label="Выделить всё", command=self.select_all_log)
            menu.post(event.x_root, event.y_root)
        
        self.log_text.bind("<Button-3>", show_context_menu)
        
        self.log_text.tag_config("info", foreground="#2C5F8D")
        self.log_text.tag_config("success", foreground="#2D7A3E", font=(FONTS["body"][0], FONTS["body"][1], "bold"))
        self.log_text.tag_config("warning", foreground="#C67C15")
        self.log_text.tag_config("error", foreground="#B53737", font=(FONTS["body"][0], FONTS["body"][1], "bold"))
        self.log_text.tag_config("header", foreground="#1A3A5C", font=(FONTS["body"][0], FONTS["body"][1], "bold"))
        self.log_text.tag_config("detail", foreground="#5D6D7E", font=(FONTS["small"][0], FONTS["small"][1]))
        
        buttons_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"])
        buttons_frame.pack(fill=tk.X, padx=12, pady=12)
        
        presets_frame = tk.Frame(buttons_frame, bg=COLORS["bg_secondary"])
        presets_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        save_preset_btn = create_icon_button(
            presets_frame,
            icon="💾",
            command=self.save_preset,
            tooltip="Сохранить пресет настроек",
            style="success",
            width=32,
            height=28
        )
        save_preset_btn.pack(side=tk.LEFT, padx=(0, 2))
        
        load_preset_btn = create_icon_button(
            presets_frame,
            icon="📂",
            command=self.load_preset,
            tooltip="Загрузить пресет",
            style="primary",
            width=32,
            height=28
        )
        load_preset_btn.pack(side=tk.LEFT)
        ToolTip(load_preset_btn, "Загрузить пресет")
        
        create_btn = create_modern_button(
            buttons_frame,
            text="✓ Создать файл",
            command=self.create_file,
            style="success",
            height=48,
            tooltip="Создать новый Excel файл с выбранными колонками"
        )
        create_btn.pack(side=tk.RIGHT)
    
    def bind_scroll_to_widget(self, widget):
        """Рекурсивно привязать события прокрутки к виджету и всем его дочерним элементам"""
        # Не привязываем к Listbox и Text виджетам
        if isinstance(widget, (tk.Listbox, tk.Text, ScrolledText)):
            return
        
        try:
            widget.bind("<MouseWheel>", self._on_mousewheel, add="+")
            widget.bind("<Button-4>", self._on_mouse_button, add="+")
            widget.bind("<Button-5>", self._on_mouse_button, add="+")
        except:
            pass
        
        # Рекурсивно для всех дочерних виджетов
        for child in widget.winfo_children():
            self.bind_scroll_to_widget(child)
    
    def add_log(self, message, tag="info"):
        """Добавить запись в лог"""
        self.log_text.insert(tk.END, message, tag)
        self.log_text.see(tk.END)
    
    def copy_log_text(self):
        """Копирование выделенного текста"""
        try:
            selected_text = self.log_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.window.clipboard_clear()
            self.window.clipboard_append(selected_text)
        except tk.TclError:
            pass
    
    def select_all_log(self):
        """Выделить весь текст в логе"""
        self.log_text.tag_add(tk.SEL, "1.0", tk.END)
        self.log_text.mark_set(tk.INSERT, "1.0")
        self.log_text.see(tk.INSERT)
    
    def back_to_constructor(self):
        """Вернуться к окну конструктора"""
        if self.constructor_window:
            self.window.destroy()
            self.constructor_window.window.deiconify()
            self.constructor_window.window.grab_set()
    
    def on_closing(self):
        """Обработчик закрытия окна через крестик - закрывает всё"""
        # Закрываем текущее окно
        self.window.destroy()
        
        # Закрываем окно конструктора, если оно есть
        if self.constructor_window:
            self.constructor_window.window.destroy()
    
    def save_preset(self):
        """Сохранить текущую конфигурацию как пресет"""
        if not self.column_mappings:
            messagebox.showwarning("Предупреждение", "Нет колонок для сохранения в пресет")
            return
        
        dialog = SimpleInputDialog(
            self.window,
            "Сохранить пресет",
            "Введите имя пресета:"
        )
        self.window.wait_window(dialog.top)
        preset_name = dialog.result
        
        if not preset_name:
            return
        
        config = {}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            except:
                pass
        
        if "excel_presets" not in config:
            config["excel_presets"] = {}
        
        config["excel_presets"][preset_name] = {
            "column_mappings": self.column_mappings
        }
        
        try:
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
            
            self.add_log(f"💾 Пресет '{preset_name}' успешно сохранен ({len(self.column_mappings)} колонок)\n", "success")
            messagebox.showinfo("Успешно", f"Пресет '{preset_name}' сохранен")
        except Exception as e:
            self.add_log(f"❌ Ошибка сохранения пресета: {str(e)}\n", "error")
            messagebox.showerror("Ошибка", f"Не удалось сохранить пресет:\n{str(e)}")
    
    def load_preset(self):
        """Загрузить сохраненный пресет"""
        config = {}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            except:
                pass
        
        excel_presets = config.get("excel_presets", {})
        
        if not excel_presets:
            messagebox.showinfo("Информация", "Нет сохраненных пресетов")
            return
        
        dialog = tk.Toplevel(self.window)
        dialog.withdraw()
        dialog.title("Загрузить пресет")
        dialog.geometry("500x400")
        dialog.minsize(400, 300)
        dialog.resizable(True, True)
        dialog.transient(self.window)
        
        dialog.update_idletasks()
        x = self.window.winfo_x() + (self.window.winfo_width() // 2) - (500 // 2)
        y = self.window.winfo_y() + (self.window.winfo_height() // 2) - (400 // 2)
        dialog.geometry(f"+{x}+{y}")
        
        dialog.deiconify()
        dialog.grab_set()
        
        tk.Label(
            dialog,
            text="Выберите пресет",
            font=FONTS["heading"],
            pady=15
        ).pack()
        
        # Список пресетов
        listbox_frame = tk.Frame(dialog)
        listbox_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        scrollbar = tk.Scrollbar(listbox_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        listbox = tk.Listbox(
            listbox_frame,
            font=FONTS["body"],
            yscrollcommand=scrollbar.set
        )
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=listbox.yview)
        
        # Заполняем список (только имена и количество колонок)
        preset_names = list(excel_presets.keys())
        for name in preset_names:
            preset = excel_presets[name]
            col_count = len(preset.get("column_mappings", []))
            listbox.insert(tk.END, f"{name} ({col_count} колонок)")
        
        # Tooltip для отображения деталей пресета при наведении
        hover_tooltip = None
        
        def show_preset_details(event):
            nonlocal hover_tooltip
            
            # Закрываем предыдущий tooltip
            if hover_tooltip:
                try:
                    hover_tooltip.destroy()
                except:
                    pass
                hover_tooltip = None
            
            # Получаем элемент под курсором
            index = listbox.nearest(event.y)
            if index < 0 or index >= len(preset_names):
                return
            
            preset_name = preset_names[index]
            preset = excel_presets[preset_name]
            mappings = preset.get("column_mappings", [])
            
            if not mappings:
                return
            
            # Создаем tooltip окно
            hover_tooltip = tk.Toplevel(dialog)
            hover_tooltip.wm_overrideredirect(True)
            hover_tooltip.wm_geometry(f"+{event.x_root + 10}+{event.y_root + 10}")
            
            tooltip_frame = tk.Frame(
                hover_tooltip,
                bg=COLORS["bg_tertiary"],
                relief=tk.SOLID,
                borderwidth=1
            )
            tooltip_frame.pack(fill=tk.BOTH, expand=True)
            
            # Заголовок
            tk.Label(
                tooltip_frame,
                text=f"Колонки в пресете '{preset_name}':",
                font=FONTS["button"],
                bg=COLORS["primary_dark"],
                fg="white",
                padx=10,
                pady=5
            ).pack(fill=tk.X)
            
            # Список колонок с прокруткой (максимум 10 строк)
            details_frame = tk.Frame(tooltip_frame, bg=COLORS["bg_tertiary"])
            details_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # Ограничиваем количество отображаемых колонок
            max_display = 15
            for i, mapping in enumerate(mappings[:max_display]):
                col_name = mapping.get('column_name', mapping.get('placeholder', ''))
                source_cols = ', '.join(mapping.get('source_columns', []))
                # Ограничиваем длину строки
                if len(source_cols) > 50:
                    source_cols = source_cols[:47] + "..."
                
                tk.Label(
                    details_frame,
                    text=f"• {col_name}: {source_cols}",
                    font=FONTS["small"],
                    bg=COLORS["bg_tertiary"],
                    fg=COLORS["text_primary"],
                    anchor="w",
                    justify=tk.LEFT
                ).pack(anchor="w", pady=1)
            
            # Если колонок больше, показываем "..."
            if len(mappings) > max_display:
                tk.Label(
                    details_frame,
                    text=f"... и ещё {len(mappings) - max_display} колонок",
                    font=FONTS["small"],
                    bg=COLORS["bg_tertiary"],
                    fg=COLORS["text_secondary"],
                    anchor="w",
                    justify=tk.LEFT
                ).pack(anchor="w", pady=1)
        
        def hide_preset_details(event):
            nonlocal hover_tooltip
            if hover_tooltip:
                try:
                    hover_tooltip.destroy()
                except:
                    pass
                hover_tooltip = None
        
        listbox.bind('<Motion>', show_preset_details)
        listbox.bind('<Leave>', hide_preset_details)
        
        selected_preset = [None]
        
        def on_load():
            if listbox.curselection():
                idx = listbox.curselection()[0]
                selected_preset[0] = preset_names[idx]
                dialog.destroy()
        
        def on_delete():
            if listbox.curselection():
                idx = listbox.curselection()[0]
                preset_name = preset_names[idx]
                
                if messagebox.askyesno("Подтверждение", f"Удалить пресет '{preset_name}'?", parent=dialog):
                    del config["excel_presets"][preset_name]
                    
                    try:
                        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                            json.dump(config, f, ensure_ascii=False, indent=2)
                        
                        listbox.delete(idx)
                        preset_names.pop(idx)
                        messagebox.showinfo("Успешно", f"Пресет '{preset_name}' удален", parent=dialog)
                    except Exception as e:
                        messagebox.showerror("Ошибка", f"Не удалось удалить пресет:\n{str(e)}", parent=dialog)
        
        buttons_frame = tk.Frame(dialog)
        buttons_frame.pack(pady=15)
        
        load_btn = create_modern_button(
            buttons_frame,
            text="Загрузить",
            command=on_load,
            style="success"
        )
        load_btn.pack(side=tk.LEFT, padx=5)
        
        delete_btn = create_modern_button(
            buttons_frame,
            text="Удалить",
            command=on_delete,
            style="danger"
        )
        delete_btn.pack(side=tk.LEFT, padx=5)
        
        cancel_btn = create_modern_button(
            buttons_frame,
            text="Отмена",
            command=dialog.destroy,
            style="primary"
        )
        cancel_btn.pack(side=tk.LEFT, padx=5)
        
        dialog.wait_window()
        
        if selected_preset[0]:
            self._apply_preset(excel_presets[selected_preset[0]])
    
    def _apply_preset(self, preset):
        """Применить пресет к текущему файлу"""
        mappings = preset.get("column_mappings", [])
        
        if not mappings:
            messagebox.showwarning("Предупреждение", "Пресет не содержит колонок")
            return
        
        available_columns = list(self.source_df.columns)
        missing_columns = []
        
        for mapping in mappings:
            source_cols = mapping.get('source_columns', [])
            for col in source_cols:
                if col not in available_columns and col not in missing_columns:
                    missing_columns.append(col)
        
        self.column_mappings = mappings.copy()
        self.refresh_mappings_display()
        
        self.add_log(f"📂 Пресет загружен ({len(mappings)} колонок)\n", "success")
        
        if missing_columns:
            self.add_log(f"⚠ Внимание! Следующие колонки отсутствуют в файле:\n", "warning")
            for col in missing_columns:
                self.add_log(f"   • {col}\n", "warning")
            self.add_log(f"   Необходимо доформировать конфигурацию вручную\n\n", "warning")
            
            messagebox.showwarning(
                "Внимание",
                f"Пресет загружен, но следующие колонки отсутствуют в текущем файле:\n\n" +
                "\n".join(f"• {col}" for col in missing_columns) +
                "\n\nНеобходимо проверить и доформировать конфигурацию вручную"
            )
        else:
            self.add_log(f"✅ Все колонки из пресета найдены в файле\n\n", "success")
            messagebox.showinfo("Успешно", f"Пресет успешно загружен!\n\nВсе колонки найдены в файле")
    
    def add_column_mapping(self):
        """Добавить новую колонку"""
        dialog = ColumnMappingDialog(self.window, self.source_df.columns, self.main_app.PLACEHOLDERS)
        if dialog.result:
            self.column_mappings.append(dialog.result)
            column_name = dialog.result.get('column_name', dialog.result.get('placeholder', ''))
            self.add_log(f"➕ Добавлена колонка: {column_name}\n", "info")
            self.add_log(f"   Исходные колонки: {', '.join(dialog.result['source_columns'])}\n", "detail")
            self.refresh_mappings_display()
    
    def edit_column_mapping(self, index):
        """Редактировать колонку"""
        current = self.column_mappings[index]
        old_name = current.get('column_name', current.get('placeholder', ''))
        dialog = ColumnMappingDialog(
            self.window,
            self.source_df.columns,
            self.main_app.PLACEHOLDERS,
            initial_data=current
        )
        if dialog.result:
            self.column_mappings[index] = dialog.result
            new_name = dialog.result.get('column_name', dialog.result.get('placeholder', ''))
            self.add_log(f"✏️ Отредактирована колонка: {old_name} → {new_name}\n", "warning")
            self.refresh_mappings_display()
    
    def delete_column_mapping(self, index):
        """Удалить колонку"""
        mapping = self.column_mappings[index]
        column_name = mapping.get('column_name', mapping.get('placeholder', ''))
        self.column_mappings.pop(index)
        self.add_log(f"❌ Удалена колонка: {column_name}\n", "warning")
        self.refresh_mappings_display()
    
    def refresh_mappings_display(self):
        """Обновить отображение списка колонок"""
        for widget in self.mappings_frame.winfo_children():
            widget.destroy()
        
        if not self.column_mappings:
            empty_label = tk.Label(
                self.mappings_frame,
                text="Нажмите '+ Добавить колонку' для начала работы",
                bg=COLORS["bg_secondary"],
                font=FONTS["body"],
                fg=COLORS["text_secondary"],
                pady=50
            )
            empty_label.pack()
            # Привязываем прокрутку и к пустому лейблу
            self.bind_scroll_to_widget(empty_label)
            return
        
        for idx, mapping in enumerate(self.column_mappings):
            frame = tk.Frame(self.mappings_frame, bg=COLORS["bg_tertiary"], relief=tk.SOLID, bd=1)
            frame.pack(fill=tk.X, pady=6, padx=0)
            
            header_frame = tk.Frame(frame, bg=COLORS["primary_dark"])
            header_frame.pack(fill=tk.X)
            
            column_name = mapping.get('column_name', mapping.get('placeholder', ''))
            tk.Label(
                header_frame,
                text=f"{idx + 1}. Колонка: {column_name}",
                bg=COLORS["primary_dark"],
                fg="white",
                font=FONTS["button"],
                anchor=tk.W,
                padx=10,
                pady=5
            ).pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            btn_frame = tk.Frame(header_frame, bg=COLORS["primary_dark"])
            btn_frame.pack(side=tk.RIGHT, padx=5)
            
            edit_btn = create_modern_button(
                btn_frame,
                text="◆",
                command=lambda i=idx: self.edit_column_mapping(i),
                style="warning",
                width=32,
                height=28,
                tooltip="Редактировать колонку"
            )
            edit_btn.pack(side=tk.LEFT, padx=2)
            
            del_btn = create_modern_button(
                btn_frame,
                text="✖",
                command=lambda i=idx: self.delete_column_mapping(i),
                style="danger",
                width=32,
                height=28,
                tooltip="Удалить колонку"
            )
            del_btn.pack(side=tk.LEFT, padx=2)
            
            content_frame = tk.Frame(frame, bg=COLORS["bg_tertiary"], height=60)
            content_frame.pack(fill=tk.X, padx=10, pady=10)
            content_frame.pack_propagate(False)
            
            # Контейнер с прокруткой для длинных списков колонок
            inner_scroll_frame = tk.Frame(content_frame, bg=COLORS["bg_tertiary"])
            inner_scroll_frame.pack(fill=tk.BOTH, expand=True)
            
            columns_text = " + ".join(mapping['source_columns'])
            # Ограничиваем отображение очень длинных списков
            if len(columns_text) > 200:
                columns_text = columns_text[:197] + "..."
            
            tk.Label(
                inner_scroll_frame,
                text=f"Исходные колонки: {columns_text}",
                bg=COLORS["bg_tertiary"],
                font=FONTS["body"],
                fg=COLORS["text_primary"],
                justify=tk.LEFT,
                anchor="nw"
            ).pack(anchor=tk.W, fill=tk.BOTH, expand=True)
            
            function_names = {
                "default": "Дефолтная функция",
                "first": "Только первое значение",
                "comma": "Объединение через запятую",
                "comma_space": "Объединение через ', '",
                "upper": "Верхний регистр",
                "lower": "Нижний регистр",
                "newline": "С новой строки",
                "custom": "Своя функция"
            }
            func_type = mapping.get('function_type', 'default')
            func_name = function_names.get(func_type, func_type)
            
            tk.Label(
                inner_scroll_frame,
                text=f"Функция: {func_name}",
                bg=COLORS["bg_tertiary"],
                font=FONTS["small"],
                fg=COLORS["text_secondary"]
            ).pack(anchor=tk.W, pady=(5, 0))
            
            # Привязываем прокрутку колесом мыши ко всем элементам карточки
            self.bind_scroll_to_widget(frame)
    
    def create_file(self):
        """Создать результирующий файл"""
        import pandas as pd
        
        if not self.column_mappings:
            messagebox.showwarning("Предупреждение", "Добавьте хотя бы одну колонку")
            return
        
        try:
            self.add_log("\n" + "="*50 + "\n", "header")
            self.add_log("🚀 НАЧАЛО ПОСТРОЕНИЯ ФАЙЛА\n", "header")
            self.add_log("="*50 + "\n\n", "header")
            
            result_df = pd.DataFrame()
            
            self.add_log(f"📊 Обработка {len(self.column_mappings)} колонок...\n\n", "info")
            
            for col_idx, mapping in enumerate(self.column_mappings, 1):
                column_name = mapping.get('column_name', mapping.get('placeholder', ''))
                source_cols = mapping['source_columns']
                function_type = mapping.get('function_type', 'default')
                custom_function = mapping.get('custom_function', '')
                
                self.add_log(f"📝 Колонка {col_idx}/{len(self.column_mappings)}: '{column_name}'\n", "success")
                self.add_log(f"   Источники данных: {', '.join(source_cols)}\n", "detail")
                
                # Определение типа функции для вывода в лог
                function_names = {
                    "default": "Дефолтная функция (автоформатирование)",
                    "first": "Только первое значение",
                    "comma": "Объединение через запятую",
                    "comma_space": "Объединение через ', '",
                    "upper": "Верхний регистр",
                    "lower": "Нижний регистр",
                    "newline": "С новой строки",
                    "custom": "Своя функция"
                }
                func_name = function_names.get(function_type, function_type)
                self.add_log(f"   Функция обработки: {func_name}\n", "detail")
                
                # Собираем значения
                values_list = []
                total_rows = len(self.source_df)
                
                for row_idx, row in self.source_df.iterrows():
                    # Логируем каждую строку
                    self.add_log(f"   Строка {row_idx + 1}:\n", "detail")
                    
                    # Собираем непустые значения из выбранных колонок для других функций
                    values = []
                    for col in source_cols:
                        value = str(row.get(col, "")).strip()
                        if value:
                            values.append(value)
                            self.add_log(f"      • Из '{col}': '{value}'\n", "detail")
                    
                    # Применяем выбранную функцию обработки
                    if function_type == "default":
                        # Дефолтная функция - с автоформатированием
                        processed_parts = []
                        for col in source_cols:
                            # Берем значение напрямую из строки данных
                            value = str(row.get(col, "")).strip()
                            
                            # Пропускаем пустые значения
                            if not value:
                                continue
                            
                            col_lower = col.lower()
                            
                            # Воинская часть
                            if self.is_military_unit_column(col_lower):
                                value_lower = str(value).lower()
                                if "войсковая часть" in value_lower:
                                    value = value_lower  # Приводим к нижнему регистру
                                elif value.isdigit() and len(value) in (4, 5):
                                    value = f"войсковая часть {value}"
                                    self.add_log(f"      ✓ Преобразовано в: '{value}'\n", "detail")
                                else:
                                    continue
                            # Дата рождения
                            elif self.is_birth_date_column(col_lower):
                                if self.is_date_like(value):
                                    value = self.format_date(value)
                                value = f"{value} года рождения"
                                self.add_log(f"      ✓ Преобразовано в: '{value}'\n", "detail")
                            # Обычные даты
                            elif self.is_date_column(col) and self.is_date_like(value):
                                old_value = value
                                value = self.format_date(value)
                                if old_value != value:
                                    self.add_log(f"      ✓ Дата отформатирована: '{old_value}' → '{value}'\n", "detail")
                            
                            processed_parts.append(value)
                        
                        combined = " ".join(processed_parts)
                        combined = " ".join(combined.split())  # Удаляем множественные пробелы
                        
                    elif function_type == "first":
                        # Только первое значение
                        combined = values[0] if values else ''
                        
                    elif function_type == "comma":
                        # Объединение через запятую
                        combined = ",".join(values)
                        
                    elif function_type == "comma_space":
                        # Объединение через запятую с пробелом
                        combined = ", ".join(values)
                        
                    elif function_type == "upper":
                        # Все в верхний регистр
                        combined = " ".join(values).upper()
                        
                    elif function_type == "lower":
                        # Все в нижний регистр
                        combined = " ".join(values).lower()
                        
                    elif function_type == "newline":
                        # Каждая колонка с новой строки
                        combined = "\n".join(values)
                        
                    elif function_type == "custom":
                        # Пользовательская функция
                        try:
                            local_vars = {'values': values}
                            
                            # Если код содержит return, оборачиваем в функцию
                            if 'return' in custom_function:
                                func_code = f"def user_func(values):\n"
                                for line in custom_function.split('\n'):
                                    func_code += f"    {line}\n"
                                exec(func_code, {}, local_vars)
                                combined = str(local_vars['user_func'](values))
                            else:
                                # Иначе выполняем как выражение
                                exec(custom_function, {}, local_vars)
                                # Ищем result или берём последнюю переменную
                                if 'result' in local_vars:
                                    combined = str(local_vars['result'])
                                else:
                                    # Берём последнее значение (кроме values)
                                    for key in reversed(list(local_vars.keys())):
                                        if key != 'values':
                                            combined = str(local_vars[key])
                                            break
                                    else:
                                        combined = ''
                        except Exception as e:
                            combined = f"ОШИБКА: {str(e)}"
                            self.add_log(f"      ❌ Ошибка в пользовательской функции: {str(e)}\n", "error")
                    
                    else:
                        # Неизвестная функция - объединяем через пробел
                        combined = " ".join(values)
                    
                    # Логируем результат для каждой строки
                    self.add_log(f"      ➡ Результат: '{combined}'\n", "success")
                    
                    values_list.append(combined)
                
                result_df[column_name] = values_list
                self.add_log(f"   ✅ Колонка успешно создана ({len(values_list)} строк)\n\n", "success")
            
            # Диалог сохранения
            self.add_log("="*50 + "\n", "header")
            self.add_log("💾 СОХРАНЕНИЕ ФАЙЛА\n", "header")
            self.add_log("="*50 + "\n\n", "header")
            
            output_file = filedialog.asksaveasfilename(
                title="Сохранить файл",
                defaultextension=".xlsx",
                filetypes=[("Excel файлы", "*.xlsx"), ("Все файлы", "*.*")]
            )
            
            if not output_file:
                self.add_log("❌ Сохранение отменено пользователем\n", "warning")
                return
            
            self.add_log(f"📁 Путь сохранения: {output_file}\n", "info")
            
            result_df.to_excel(output_file, index=False)
            
            file_size = os.path.getsize(output_file)
            file_size_kb = file_size / 1024
            self.add_log(f"✅ Файл успешно сохранен (размер: {file_size_kb:.2f} КБ)\n", "success")
            self.add_log(f"   Колонок: {len(result_df.columns)}\n", "detail")
            self.add_log(f"   Строк: {len(result_df)}\n\n", "detail")
            
            current_tab = self.main_app.get_current_tab()
            if current_tab:
                current_tab.excel_path.set(output_file)
                current_tab.last_excel_dir = os.path.dirname(output_file)
                
                try:
                    # Используем pandas для чтения заголовков
                    if hasattr(self.main_app, '_pandas_loaded') and self.main_app._pandas_loaded and self.main_app._pandas:
                        pd_module = self.main_app._pandas
                    else:
                        pd_module = pd
                    
                    df_headers = pd_module.read_excel(output_file, engine='openpyxl', nrows=0)
                    current_tab.excel_columns = list(df_headers.columns)
                    
                    column_values = [""] + current_tab.excel_columns
                    if hasattr(current_tab, 'filename_column_combo'):
                        current_tab.filename_column_combo.configure(values=column_values)
                        current_tab.filename_column_combo.set("")
                    
                    # Логируем и в построитель файлов и в активную вкладку
                    self.add_log(f"📥 Файл загружен в активную вкладку (колонок: {len(current_tab.excel_columns)})\n", "success")
                    current_tab.log(f"📥 Файл загружен из конструктора Excel: {os.path.basename(output_file)}")
                    current_tab.log(f"Найдено столбцов: {len(current_tab.excel_columns)}")
                except Exception as e:
                    self.add_log(f"⚠ Файл сохранен, но не удалось загрузить колонки: {str(e)}\n", "warning")
            else:
                self.add_log(f"⚠ Не удалось найти активную вкладку, файл сохранен но не загружен\n", "warning")
            
            self.add_log("="*50 + "\n", "header")
            self.add_log("🎉 ПОСТРОЕНИЕ ФАЙЛА ЗАВЕРШЕНО!\n", "success")
            self.add_log("="*50 + "\n", "header")
            
            # Предлагаем предварительный просмотр
            result = messagebox.askyesno(
                "Файл создан",
                f"Файл успешно создан!\n\n{os.path.basename(output_file)}\n\nКолонок: {len(result_df.columns)}\nСтрок: {len(result_df)}\n\nОткрыть предварительный просмотр?",
                parent=self.window
            )
            
            if result:
                try:
                    PreviewWindow(self.window, output_file, f"Просмотр: {os.path.basename(output_file)}")
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Не удалось открыть предварительный просмотр:\n{str(e)}", parent=self.window)
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при создании файла:\n{str(e)}\n\n{traceback.format_exc()}"
            self.add_log("\n" + "="*50 + "\n", "error")
            self.add_log("❌ ОШИБКА\n", "error")
            self.add_log("="*50 + "\n", "error")
            self.add_log(f"{error_msg}\n", "error")
            messagebox.showerror("Ошибка", error_msg)
    
    def is_military_unit_column(self, col_name):
        """Проверка, является ли колонка воинской частью"""
        keywords = ["в/ч", "воинская часть", "войсковая часть", "в/часть", "воинская", "войскавая", "войсковая"]
        col_name = col_name.lower().replace(" ", "")
        for keyword in keywords:
            if keyword.replace(" ", "") in col_name:
                return True
        return False
    
    def is_birth_date_column(self, col_name):
        """Проверка, является ли колонка датой рождения"""
        keywords = ["д.р.", "д/р", "дата рождения", "дата.р", "д.рождения", "датарождения", "дата_рождения"]
        col_name = col_name.lower().replace(" ", "").replace("_", "")
        for keyword in keywords:
            if keyword.replace(" ", "").replace("_", "").replace(".", "") in col_name.replace(".", ""):
                return True
        return False
    
    def is_date_column(self, col_name):
        """Проверка, является ли колонка датой по заголовку"""
        col_lower = col_name.lower()
        date_keywords = ["дата", "д.р.", "д/р", "д.р", "date"]
        for keyword in date_keywords:
            if keyword in col_lower:
                return True
        return False
    
    def is_date_like(self, value):
        """Проверка, похоже ли значение на дату"""
        # Числовая дата (serial date)
        if value.replace(".", "").isdigit() and len(value) >= 5:
            return True
        # Уже отформатированная дата
        if "/" in value or "-" in value:
            return True
        return False
    
    def format_date(self, value):
        """Форматирование даты в дд.мм.гггг"""
        import pandas as pd
        
        if not value or pd.isna(value) or str(value).strip() == "":
            return value
        
        value_str = str(value).strip()
        
        if len(value_str) == 10 and value_str[2] == '.' and value_str[5] == '.':
            try:
                datetime.strptime(value_str, '%d.%m.%Y')
                return value_str
            except:
                pass
        
        # Проверка на чистое время (без даты): 9:00, 9 ч 00 мин и т.п.
        time_indicators = [':']
        date_indicators = ['.', '-', '/']
        has_time = any(ind in value_str for ind in time_indicators) or ' ч ' in value_str.lower() or value_str.lower().endswith(' ч')
        has_date = any(ind in value_str for ind in date_indicators)
        
        # Если есть признаки времени, но нет признаков даты - возвращаем как есть
        if has_time and not has_date:
            return value_str
        
        if ' ' in value_str:
            value_str = value_str.split(' ')[0]
        
        # Проверяем, есть ли признаки даты (точки, дефисы, слэши)
        has_date_format = '/' in value_str or '-' in value_str or '.' in value_str
        
        try:
            # Если это число без признаков даты - возвращаем как есть
            if value_str.replace(".", "").isdigit() and not has_date_format:
                return value
            
            # Если есть признаки даты, пытаемся распарсить
            if has_date_format:
                date_val = pd.to_datetime(value_str, dayfirst=True, errors='coerce')
                if pd.notna(date_val):
                    return date_val.strftime('%d.%m.%Y')
        except:
            pass
        
        return value
    
    def format_date_value(self, value):
        """Универсальное форматирование значения с обработкой дат и времени"""
        import pandas as pd
        
        # Если значение пустое или NaN
        if pd.isna(value) or value == "":
            return ""
        
        # Если это Timestamp или datetime объект из pandas - форматируем сразу
        if isinstance(value, pd.Timestamp) or hasattr(value, 'strftime'):
            return value.strftime('%d.%m.%Y')
        
        # Преобразуем в строку
        value_str = str(value).strip()
        
        # Если пустая строка
        if not value_str:
            return ""
        
        # Если значение уже в правильном формате дд.мм.гггг - возвращаем как есть
        if len(value_str) == 10 and value_str[2] == '.' and value_str[5] == '.':
            try:
                datetime.strptime(value_str, '%d.%m.%Y')
                return value_str
            except:
                pass
        
        # Проверка на чистое время (без даты): 9:00, 9 ч 00 мин и т.п.
        time_indicators = [':']
        date_indicators = ['.', '-', '/']
        has_time = any(ind in value_str for ind in time_indicators) or ' ч ' in value_str.lower() or value_str.lower().endswith(' ч')
        has_date = any(ind in value_str for ind in date_indicators)
        
        # Если есть признаки времени, но нет признаков даты - возвращаем как есть
        if has_time and not has_date:
            return value_str
        
        # Если в строке есть время (пробел + время), убираем его
        if ' ' in value_str:
            try:
                dt = pd.to_datetime(value_str, dayfirst=True, errors='coerce')
                if pd.notna(dt):
                    return dt.strftime('%d.%m.%Y')
            except:
                pass
        
        # Проверяем, есть ли признаки даты (точки, дефисы, слэши)
        if '/' in value_str or '-' in value_str or '.' in value_str:
            try:
                dt = pd.to_datetime(value_str, dayfirst=True, errors='coerce')
                if pd.notna(dt):
                    return dt.strftime('%d.%m.%Y')
            except:
                pass
        
        # Если не получилось - возвращаем как есть
        return value_str

# ── КЛАССЫ РЕДАКТОРОВ ФАЙЛОВ ──────────────────────────────────────────
# Открытие документов через системные приложения

def open_word_document(file_path):
    """Открывает Word документ в Microsoft Word"""
    try:
        if WIN32COM_AVAILABLE:
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = True
                word.Documents.Open(os.path.abspath(file_path))
                return
            except Exception as e:
                pass
        
        # Fallback: открываем через ассоциированное приложение
        os.startfile(file_path)
        
    except Exception as e:
        messagebox.showerror(
            "Ошибка",
            f"Не удалось открыть документ:\n{str(e)}"
        )

class ExcelEditor:
    """Редактор Excel файлов"""
    def __init__(self, parent, file_path):
        self.file_path = file_path
        self.df = None
        self.modified = False
        
        self.window = tk.Toplevel(parent)
        self.window.withdraw()
        self.window.title(f"Редактор: {os.path.basename(file_path)}")
        self.window.geometry("1100x750")
        self.window.transient(parent)
        
        self.create_widgets()
        self.load_excel()
        
        self.window.deiconify()
    
    def create_widgets(self):
        """Создание виджетов редактора"""
        top_frame = tk.Frame(self.window, bg=COLORS["primary"], height=60)
        top_frame.pack(fill=tk.X)
        top_frame.pack_propagate(False)
        
        tk.Label(
            top_frame,
            text="📊 Редактор Excel файла",
            bg=COLORS["primary"],
            fg="white",
            font=FONTS["title"]
        ).pack(side=tk.LEFT, padx=15, pady=15)
        
        btn_frame = tk.Frame(top_frame, bg=COLORS["primary"])
        btn_frame.pack(side=tk.RIGHT, padx=15, pady=10)
        
        add_btn = create_modern_button(
            btn_frame,
            text="➕ Добавить строку",
            command=self.add_row,
            style="success",
            width=130,
            height=36,
            tooltip="Добавить новую строку в таблицу"
        )
        add_btn.pack(side=tk.LEFT, padx=3)
        
        delete_btn = create_modern_button(
            btn_frame,
            text="➖ Удалить строку",
            command=self.delete_row,
            style="warning",
            width=130,
            height=36,
            tooltip="Удалить выбранную строку"
        )
        delete_btn.pack(side=tk.LEFT, padx=3)
        
        save_btn = create_modern_button(
            btn_frame,
            text="💾 Сохранить",
            command=self.save_excel,
            style="success",
            width=110,
            height=36,
            tooltip="Сохранить изменения в файл"
        )
        save_btn.pack(side=tk.LEFT, padx=3)
        
        table_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"])
        table_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tree_container, self.tree = create_modern_treeview(
            table_frame,
            columns=[],  # Колонки будут настроены в load_excel
            height=20,
            selectable=True
        )
        tree_container.pack(fill=tk.BOTH, expand=True)
        
        self.tree.bind('<Double-1>', self.on_double_click)
        
        status_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"], height=30)
        status_frame.pack(fill=tk.X)
        status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(
            status_frame,
            text="Загрузка файла...",
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_secondary"],
            font=FONTS["small"],
            anchor=tk.W
        )
        self.status_label.pack(fill=tk.X, padx=10, pady=5)
    
    def load_excel(self):
        """Загрузка Excel файла"""
        try:
            import pandas as pd
            
            self.df = pd.read_excel(self.file_path, engine='openpyxl')
            
            columns = list(self.df.columns)
            self.tree['columns'] = columns
            self.tree['show'] = 'headings'
            
            for col in columns:
                self.tree.heading(col, text=str(col), command=lambda c=col: self.edit_header(c))
                max_width = max(
                    len(str(col)) * 8,
                    self.df[col].astype(str).str.len().max() * 8 if not self.df[col].empty else 50
                )
                # stretch=tk.NO позволяет растягивать столбцы вручную без авто-перераспределения
                self.tree.column(col, width=min(max_width, 300), anchor=tk.W, minwidth=50, stretch=tk.NO)
            
            self.refresh_tree()
            
            self.status_label.config(text=f"Загружено: {len(self.df)} строк, {len(columns)} столбцов")
            self.modified = False
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при загрузке Excel:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
            self.window.destroy()
    
    def refresh_tree(self):
        """Обновление отображения таблицы"""
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        for idx, row in self.df.iterrows():
            values = format_dataframe_row(row)
            insert_treeview_row(self.tree, values)
    
    def on_double_click(self, event):
        """Обработка двойного клика для редактирования ячейки"""
        import pandas as pd
        
        # Определяем выбранную ячейку
        item = self.tree.identify('item', event.x, event.y)
        column = self.tree.identify_column(event.x)
        
        if not item or not column:
            return
        
        row_idx = int(item)
        col_idx = int(column.replace('#', '')) - 1
        col_name = self.df.columns[col_idx]
        
        current_value = self.df.at[row_idx, col_name]
        if pd.isna(current_value):
            current_value = ""
        else:
            current_value = str(current_value)
        
        dialog = SimpleInputDialog(
            self.window,
            "Редактирование ячейки",
            f"Строка {row_idx + 1}, Колонка '{col_name}':\n\nВведите новое значение:",
            current_value
        )
        self.window.wait_window(dialog.top)
        new_value = dialog.result
        
        if new_value is not None:  # Пользователь не нажал Cancel
            try:
                # Пытаемся определить тип данных
                if new_value == "":
                    self.df.at[row_idx, col_name] = None
                else:
                    # Пробуем преобразовать в число
                    try:
                        if '.' in new_value or ',' in new_value:
                            new_value = new_value.replace(',', '.')
                            self.df.at[row_idx, col_name] = float(new_value)
                        else:
                            self.df.at[row_idx, col_name] = int(new_value)
                    except ValueError:
                        # Оставляем как строку
                        self.df.at[row_idx, col_name] = new_value
                
                self.modified = True
                self.refresh_tree()
                self.status_label.config(text=f"✎ Изменено: строка {row_idx + 1}, колонка '{col_name}'")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось изменить значение:\n{str(e)}", parent=self.window)
    
    def edit_header(self, col_name):
        """Редактирование заголовка колонки"""
        dialog = SimpleInputDialog(
            self.window,
            "Переименовать колонку",
            f"Текущее название: '{col_name}'\n\nВведите новое название:",
            col_name
        )
        self.window.wait_window(dialog.top)
        new_name = dialog.result
        
        if new_name and new_name != col_name:
            self.df.rename(columns={col_name: new_name}, inplace=True)
            
            columns = list(self.df.columns)
            self.tree['columns'] = columns
            
            for col in columns:
                self.tree.heading(col, text=str(col), command=lambda c=col: self.edit_header(c))
            
            self.refresh_tree()
            self.modified = True
            self.status_label.config(text=f"✎ Колонка переименована: '{col_name}' → '{new_name}'")
    
    def add_row(self):
        """Добавление новой строки"""
        import pandas as pd
        new_row = pd.DataFrame({col: [None] for col in self.df.columns})
        self.df = pd.concat([self.df, new_row], ignore_index=True)
        
        self.refresh_tree()
        self.modified = True
        self.status_label.config(text=f"➕ Добавлена строка {len(self.df)}")
    
    def delete_row(self):
        """Удаление выбранной строки"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите строку для удаления", parent=self.window)
            return
        
        row_idx = int(selected[0])
        
        result = messagebox.askyesno(
            "Подтверждение",
            f"Удалить строку {row_idx + 1}?",
            parent=self.window
        )
        
        if result:
            self.df = self.df.drop(row_idx).reset_index(drop=True)
            self.refresh_tree()
            self.modified = True
            self.status_label.config(text=f"➖ Удалена строка {row_idx + 1}")
    
    def save_excel(self):
        """Сохранение Excel файла"""
        try:
            dir_name = os.path.dirname(self.file_path)
            base_name = os.path.basename(self.file_path)
            name_without_ext, ext = os.path.splitext(base_name)
            
            # Убираем старую пометку (ред) если есть
            if name_without_ext.endswith('(ред)'):
                name_without_ext = name_without_ext[:-5].strip()
            
            new_name = f"{name_without_ext}(ред){ext}"
            new_path = os.path.join(dir_name, new_name)
            
            self.df.to_excel(new_path, index=False, engine='openpyxl')
            self.file_path = new_path
            self.modified = False
            self.status_label.config(text=f"✓ Файл сохранён: {len(self.df)} строк")
            
            messagebox.showinfo(
                "Успешно",
                f"Excel файл успешно сохранён!\n\nФайл: {new_name}",
                parent=self.window
            )
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при сохранении:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
    
    def close_editor(self):
        """Закрытие редактора"""
        if self.modified:
            result = messagebox.askyesnocancel(
                "Несохранённые изменения",
                "Сохранить изменения перед закрытием?",
                parent=self.window
            )
            if result is None:
                return
            elif result:
                self.save_excel()
        
        self.window.destroy()

class ImageEditor:
    """Редактор изображений с фильтрами, обрезкой и ластиком"""
    def __init__(self, parent, file_path):
        self.file_path = file_path
        self.original_image = None
        self.current_image = None
        self.display_image = None
        self.modified = False
        self.tool_mode = 'select'  # select, crop, eraser, eraser_area
        self.crop_start = None
        self.crop_rect = None
        self.eraser_size = 20
        self.eraser_cursor = None  # Визуальный курсор ластика
        self.last_eraser_pos = None
        self.history = []  # История изменений для отката
        self.history_position = -1
        self.image_offset = (0, 0)  # Смещение изображения на холсте
        self.zoom_scale = 1.0  # Масштаб для Ctrl+колёсико
        
        if not PIL_AVAILABLE:
            messagebox.showerror(
                "Модуль недоступен",
                "Для редактирования изображений необходим модуль Pillow.\n\n"
                "Установите командой:\npip install Pillow",
                parent=parent
            )
            return
        
        self.window = tk.Toplevel(parent)
        self.window.withdraw()
        self.window.title(f"Редактор: {os.path.basename(file_path)}")
        self.window.geometry("1200x800")
        self.window.transient(parent)
        
        self.create_widgets()
        self.load_image()
        
        self.window.deiconify()
    
    def create_widgets(self):
        """Создание виджетов редактора"""
        top_frame = tk.Frame(self.window, bg=COLORS["primary"], height=60)
        top_frame.pack(fill=tk.X)
        top_frame.pack_propagate(False)
        
        tk.Label(
            top_frame,
            text="🖼️ Редактор изображений",
            bg=COLORS["primary"],
            fg="white",
            font=("Segoe UI", 11, "bold")
        ).pack(side=tk.LEFT, padx=15, pady=15)
        
        btn_frame = tk.Frame(top_frame, bg=COLORS["primary"])
        btn_frame.pack(side=tk.RIGHT, padx=15, pady=10)
        
        tk.Button(
            btn_frame,
            text="💾 Сохранить",
            command=self.save_image,
            bg=COLORS["success"],
            fg="white",
            font=FONTS["button"],
            relief=tk.FLAT,
            cursor="hand2",
            padx=15,
            pady=5
        ).pack(side=tk.LEFT, padx=2)
        
        tk.Button(
            btn_frame,
            text="↶ Отменить",
            command=self.reset_image,
            bg=COLORS["warning"],
            fg="white",
            font=FONTS["button"],
            relief=tk.FLAT,
            cursor="hand2",
            padx=15,
            pady=5
        ).pack(side=tk.LEFT, padx=2)
        
        main_container = tk.Frame(self.window, bg=COLORS["bg_secondary"])
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tools_frame = tk.LabelFrame(
            main_container,
            text=" Инструменты ",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            padx=10,
            pady=10
        )
        tools_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        
        tk.Label(
            tools_frame,
            text="Режим работы:",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=(0, 5))
        
        self.tool_var = tk.StringVar(value='select')
        
        tk.Radiobutton(
            tools_frame,
            text="👆 Выбор",
            variable=self.tool_var,
            value='select',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        tk.Radiobutton(
            tools_frame,
            text="✂️ Обрезка",
            variable=self.tool_var,
            value='crop',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        tk.Radiobutton(
            tools_frame,
            text="🧹 Ластик (точка)",
            variable=self.tool_var,
            value='eraser',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        tk.Radiobutton(
            tools_frame,
            text="🧽 Ластик (область)",
            variable=self.tool_var,
            value='eraser_area',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        self.eraser_separator = ttk.Separator(tools_frame, orient=tk.HORIZONTAL)
        
        self.eraser_label = tk.Label(
            tools_frame,
            text="Размер ластика:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        )
        
        self.eraser_scale = tk.Scale(
            tools_frame,
            from_=5,
            to=100,
            orient=tk.HORIZONTAL,
            command=self.update_eraser_size,
            bg=COLORS["bg_secondary"]
        )
        self.eraser_scale.set(20)
        
        ttk.Separator(tools_frame, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=10)
        
        tk.Label(
            tools_frame,
            text="Фильтры:",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=(0, 5))
        
        filters = [
            ("Чёрно-белое", self.apply_grayscale),
            ("Негатив", self.apply_invert),
            ("Размытие", self.apply_blur),
            ("Резкость", self.apply_sharpen),
            ("Контраст +", self.apply_enhance_contrast),
            ("Яркость +", self.apply_enhance_brightness),
        ]
        
        for filter_name, filter_func in filters:
            tk.Button(
                tools_frame,
                text=filter_name,
                command=filter_func,
                font=FONTS["small"],
                bg="white",
                relief=tk.SOLID,
                borderwidth=1,
                cursor="hand2"
            ).pack(fill=tk.X, pady=2)
        
        canvas_frame = tk.Frame(main_container, bg="white")
        canvas_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.canvas = tk.Canvas(canvas_frame, bg="white", cursor="crosshair")
        v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL, command=self.canvas.yview)
        h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL, command=self.canvas.xview)
        
        self.canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.canvas.bind('<ButtonPress-1>', self.on_mouse_down)
        self.canvas.bind('<B1-Motion>', self.on_mouse_drag)
        self.canvas.bind('<ButtonRelease-1>', self.on_mouse_up)
        self.canvas.bind('<Motion>', self.on_mouse_move)
        self.canvas.bind('<Control-MouseWheel>', self.on_zoom)
        
        status_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"], height=30)
        status_frame.pack(fill=tk.X)
        status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(
            status_frame,
            text="Загрузка изображения...",
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_secondary"],
            font=FONTS["small"],
            anchor=tk.W
        )
        self.status_label.pack(fill=tk.X, padx=10, pady=5)
    
    def load_image(self):
        """Загрузка изображения"""
        try:
            from PIL import Image
            
            self.original_image = Image.open(self.file_path).convert("RGBA")
            self.current_image = self.original_image.copy()
            self.save_to_history()
            self.update_canvas()
            
            size = self.original_image.size
            self.status_label.config(text=f"Изображение загружено: {size[0]}x{size[1]} px")
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка загрузки изображения:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
            self.window.destroy()
    
    def update_canvas(self):
        """Обновление отображения на холсте"""
        from PIL import ImageTk
        
        canvas_width = self.canvas.winfo_width()
        canvas_height = self.canvas.winfo_height()
        
        if canvas_width <= 1:
            canvas_width = 800
        if canvas_height <= 1:
            canvas_height = 600
        
        img_width, img_height = self.current_image.size
        
        scale_w = (canvas_width - 40) / img_width
        scale_h = (canvas_height - 40) / img_height
        scale = min(scale_w, scale_h, 1.0) * self.zoom_scale
        
        if scale != 1.0:
            new_width = int(img_width * scale)
            new_height = int(img_height * scale)
            try:
                display_img = self.current_image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            except AttributeError:
                display_img = self.current_image.resize((new_width, new_height), Image.ANTIALIAS)
        else:
            display_img = self.current_image
        
        self.display_scale = display_img.size[0] / img_width
        self.photo = ImageTk.PhotoImage(display_img)
        
        img_canvas_x = canvas_width // 2
        img_canvas_y = canvas_height // 2
        self.image_offset = (
            img_canvas_x - (display_img.size[0] // 2),
            img_canvas_y - (display_img.size[1] // 2)
        )
        
        self.canvas.delete("image")
        self.canvas.create_image(
            img_canvas_x,
            img_canvas_y,
            image=self.photo,
            anchor=tk.CENTER,
            tags="image"
        )
        
        if self.canvas.find_withtag("eraser_cursor"):
            self.canvas.tag_raise("eraser_cursor")
        
        bbox = self.canvas.bbox(tk.ALL)
        if bbox:
            self.canvas.config(scrollregion=bbox)
    
    def change_tool(self):
        """Смена инструмента"""
        self.tool_mode = self.tool_var.get()
        
        if self.tool_mode in ('eraser', 'eraser_area'):
            self.canvas.config(cursor="none")
            self.eraser_separator.pack(fill=tk.X, pady=10)
            self.eraser_label.pack(anchor=tk.W, pady=(0, 5))
            self.eraser_scale.pack(fill=tk.X, pady=2)
        elif self.tool_mode == 'crop':
            self.canvas.config(cursor="crosshair")
            # Удаляем курсор ластика
            self.canvas.delete("eraser_cursor")
            self.eraser_separator.pack_forget()
            self.eraser_label.pack_forget()
            self.eraser_scale.pack_forget()
        else:
            self.canvas.config(cursor="hand2")  # Ладошка для перетаскивания
            # Удаляем курсор ластика
            self.canvas.delete("eraser_cursor")
            self.eraser_separator.pack_forget()
            self.eraser_label.pack_forget()
            self.eraser_scale.pack_forget()
        
        if self.tool_mode not in ('crop', 'eraser_area') and self.crop_rect:
            self.canvas.delete("crop_rect")
            self.crop_rect = None
            self.crop_start = None
    
    def update_eraser_size(self, value):
        """Обновление размера ластика"""
        self.eraser_size = int(float(value))
        self.canvas.delete("eraser_cursor")
        self.eraser_cursor = None
    
    def update_eraser_cursor(self, x=None, y=None):
        """Обновление визуального курсора ластика"""
        # Всегда сначала удаляем ВСЕ старые курсоры по тегу
        self.canvas.delete("eraser_cursor")
        self.eraser_cursor = None
        # Принудительно обновляем canvas чтобы удаление применилось
        self.canvas.update_idletasks()
        
        if self.tool_mode in ('eraser', 'eraser_area') and x is not None and y is not None:
            radius = self.eraser_size / 2
            self.eraser_cursor = self.canvas.create_oval(
                x - radius, y - radius,
                x + radius, y + radius,
                outline="red",
                width=1,
                tags="eraser_cursor"
            )
    
    def on_mouse_move(self, event):
        """Движение мыши (без нажатия кнопки)"""
        if self.tool_mode in ('eraser', 'eraser_area'):
            x = self.canvas.canvasx(event.x)
            y = self.canvas.canvasy(event.y)
            self.update_eraser_cursor(x, y)
    
    def on_mouse_down(self, event):
        """Нажатие мыши"""
        # Преобразуем экранные координаты в координаты холста с учетом прокрутки
        x = self.canvas.canvasx(event.x)
        y = self.canvas.canvasy(event.y)
        
        if self.tool_mode == 'select':
            # Режим перетаскивания
            self.canvas.scan_mark(event.x, event.y)
        elif self.tool_mode in ('crop', 'eraser_area'):
            self.crop_start = (x, y)
            if self.crop_rect:
                self.canvas.delete("crop_rect")
        elif self.tool_mode == 'eraser':
            self.use_eraser(x, y)
            self.last_eraser_pos = (x, y)
    
    def on_mouse_drag(self, event):
        """Движение мыши"""
        x = self.canvas.canvasx(event.x)
        y = self.canvas.canvasy(event.y)
        
        if self.tool_mode == 'select':
            self.canvas.scan_dragto(event.x, event.y, gain=1)
        elif self.tool_mode in ('crop', 'eraser_area') and self.crop_start:
            if self.crop_rect:
                self.canvas.delete("crop_rect")
            color = "red" if self.tool_mode == 'crop' else "blue"
            self.crop_rect = self.canvas.create_rectangle(
                self.crop_start[0], self.crop_start[1],
                x, y,
                outline=color, width=2, dash=(5, 5), tags="crop_rect"
            )
            if self.tool_mode == 'eraser_area':
                self.update_eraser_cursor(x, y)
        elif self.tool_mode == 'eraser':
            self.use_eraser(x, y)
            self.last_eraser_pos = (x, y)
            self.update_eraser_cursor(x, y)
    
    def on_mouse_up(self, event):
        """Отпускание мыши"""
        # Преобразуем экранные координаты в координаты холста с учетом прокрутки
        x = self.canvas.canvasx(event.x)
        y = self.canvas.canvasy(event.y)
        
        if self.tool_mode == 'crop' and self.crop_start:
            # Применяем обрезку
            self.apply_crop(self.crop_start[0], self.crop_start[1], x, y)
            self.crop_start = None
            if self.crop_rect:
                self.canvas.delete("crop_rect")
                self.crop_rect = None
        elif self.tool_mode == 'eraser_area' and self.crop_start:
            # Применяем стирание области
            self.erase_area(self.crop_start[0], self.crop_start[1], x, y)
            self.crop_start = None
            if self.crop_rect:
                self.canvas.delete("crop_rect")
                self.crop_rect = None
        elif self.tool_mode == 'eraser':
            self.last_eraser_pos = None
    
    def on_zoom(self, event):
        """Масштабирование через Ctrl+колёсико мыши"""
        if event.delta > 0:
            self.zoom_scale = min(self.zoom_scale * 1.1, 5.0)
        else:
            self.zoom_scale = max(self.zoom_scale / 1.1, 0.1)
        
        self.update_canvas()
        self.status_label.config(text=f"🔍 Масштаб: {int(self.zoom_scale * 100)}%")
    
    def save_to_history(self):
        """Сохранение текущего состояния в историю"""
        # Удаляем все последующие состояния если мы не в конце истории
        if self.history_position < len(self.history) - 1:
            self.history = self.history[:self.history_position + 1]
        
        self.history.append(self.current_image.copy())
        self.history_position += 1
        
        # Ограничиваем размер истории (максимум 20 состояний)
        if len(self.history) > 20:
            self.history.pop(0)
            self.history_position -= 1
    
    def use_eraser(self, x, y):
        """Использование ластика"""
        from PIL import ImageDraw
        
        bbox = self.canvas.bbox("image")
        if not bbox:
            return
        
        img_x = int((x - bbox[0]) / self.display_scale)
        img_y = int((y - bbox[1]) / self.display_scale)
        
        draw = ImageDraw.Draw(self.current_image)
        radius = int(self.eraser_size / (2 * self.display_scale))
        
        bbox = [
            img_x - radius, img_y - radius,
            img_x + radius, img_y + radius
        ]
        
        alpha = self.current_image.split()[3]
        alpha_draw = ImageDraw.Draw(alpha)
        alpha_draw.ellipse(bbox, fill=0)
        
        self.current_image.putalpha(alpha)
        
        self.modified = True
        self.update_canvas()
        self.status_label.config(text="🧹 Ластик применён")
    
    def erase_area(self, x1, y1, x2, y2):
        """Стирание прямоугольной области"""
        from PIL import ImageDraw
        
        bbox = self.canvas.bbox("image")
        if not bbox:
            return
        
        # Нормализуем координаты
        if x1 > x2:
            x1, x2 = x2, x1
        if y1 > y2:
            y1, y2 = y2, y1
        
        # Конвертируем в координаты изображения
        img_x1 = int((x1 - bbox[0]) / self.display_scale)
        img_y1 = int((y1 - bbox[1]) / self.display_scale)
        img_x2 = int((x2 - bbox[0]) / self.display_scale)
        img_y2 = int((y2 - bbox[1]) / self.display_scale)
        
        # Ограничиваем координаты
        img_x1 = max(0, img_x1)
        img_y1 = max(0, img_y1)
        img_x2 = min(self.current_image.width, img_x2)
        img_y2 = min(self.current_image.height, img_y2)
        
        if img_x2 - img_x1 < 2 or img_y2 - img_y1 < 2:
            return
        
        # Стираем область (делаем прозрачной)
        alpha = self.current_image.split()[3]
        alpha_draw = ImageDraw.Draw(alpha)
        alpha_draw.rectangle([img_x1, img_y1, img_x2, img_y2], fill=0)
        self.current_image.putalpha(alpha)
        
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text=f"🧽 Стёрта область: {img_x2-img_x1}x{img_y2-img_y1} px")
    
    def apply_crop(self, x1, y1, x2, y2):
        """Применение обрезки"""
        bbox = self.canvas.bbox("image")
        if not bbox:
            return
        
        if x1 > x2:
            x1, x2 = x2, x1
        if y1 > y2:
            y1, y2 = y2, y1
        
        img_x1 = int((x1 - bbox[0]) / self.display_scale)
        img_y1 = int((y1 - bbox[1]) / self.display_scale)
        img_x2 = int((x2 - bbox[0]) / self.display_scale)
        img_y2 = int((y2 - bbox[1]) / self.display_scale)
        
        img_x1 = max(0, img_x1)
        img_y1 = max(0, img_y1)
        img_x2 = min(self.current_image.width, img_x2)
        img_y2 = min(self.current_image.height, img_y2)
        
        if img_x2 - img_x1 < 10 or img_y2 - img_y1 < 10:
            messagebox.showwarning("Предупреждение", "Выберите область побольше", parent=self.window)
            return
        
        self.current_image = self.current_image.crop((img_x1, img_y1, img_x2, img_y2))
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text=f"✂️ Обрезано: {img_x2-img_x1}x{img_y2-img_y1} px")
    
    def apply_grayscale(self):
        """Применение чёрно-белого фильтра"""
        from PIL import ImageOps
        rgb = self.current_image.convert('RGB')
        gray = ImageOps.grayscale(rgb)
        self.current_image = gray.convert('RGBA')
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text="Применён фильтр: Чёрно-белое")
    
    def apply_invert(self):
        """Применение негатива"""
        from PIL import ImageOps
        rgb = self.current_image.convert('RGB')
        inverted = ImageOps.invert(rgb)
        self.current_image = inverted.convert('RGBA')
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text="Применён фильтр: Негатив")
    
    def apply_blur(self):
        """Применение размытия"""
        from PIL import ImageFilter
        self.current_image = self.current_image.filter(ImageFilter.BLUR)
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text="Применён фильтр: Размытие")
    
    def apply_sharpen(self):
        """Применение резкости"""
        from PIL import ImageFilter
        self.current_image = self.current_image.filter(ImageFilter.SHARPEN)
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text="Применён фильтр: Резкость")
    
    def apply_enhance_contrast(self):
        """Увеличение контраста"""
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(self.current_image)
        self.current_image = enhancer.enhance(1.5)
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text="Применён фильтр: Контраст +")
    
    def apply_enhance_brightness(self):
        """Увеличение яркости"""
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Brightness(self.current_image)
        self.current_image = enhancer.enhance(1.3)
        self.save_to_history()
        self.modified = True
        self.update_canvas()
        self.status_label.config(text="Применён фильтр: Яркость +")
    
    def reset_image(self):
        """Сброс к предыдущему состоянию или оригиналу"""
        if self.history_position > 0:
            # Откат к предыдущему состоянию
            self.history_position -= 1
            self.current_image = self.history[self.history_position].copy()
            self.modified = True
            self.update_canvas()
            self.status_label.config(text=f"↶ Отменено (позиция {self.history_position + 1}/{len(self.history)})")
        else:
            # Если история пуста, возвращаемся к оригиналу
            result = messagebox.askyesno(
                "Подтверждение",
                "Отменить все изменения и вернуться к оригиналу?",
                parent=self.window
            )
            if result:
                self.current_image = self.original_image.copy()
                self.history = [self.current_image.copy()]
                self.history_position = 0
                self.modified = False
                self.update_canvas()
                self.status_label.config(text="↶ Изображение восстановлено")
    
    def save_image(self):
        """Сохранение изображения"""
        try:
            dir_name = os.path.dirname(self.file_path)
            base_name = os.path.basename(self.file_path)
            name_without_ext, ext = os.path.splitext(base_name)
            
            if name_without_ext.endswith('(ред)'):
                name_without_ext = name_without_ext[:-5].strip()
            
            new_name = f"{name_without_ext}(ред){ext}"
            new_path = os.path.join(dir_name, new_name)
            
            if new_path.lower().endswith('.jpg') or new_path.lower().endswith('.jpeg'):
                save_img = self.current_image.convert('RGB')
            else:
                save_img = self.current_image
            
            save_img.save(new_path)
            self.file_path = new_path
            self.modified = False
            self.status_label.config(text="✓ Изображение сохранено")
            
            messagebox.showinfo(
                "Успешно",
                f"Изображение успешно сохранено!\n\nФайл: {new_name}",
                parent=self.window
            )
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при сохранении:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
    
    def close_editor(self):
        """Закрытие редактора"""
        if self.modified:
            result = messagebox.askyesnocancel(
                "Несохранённые изменения",
                "Сохранить изменения перед закрытием?",
                parent=self.window
            )
            if result is None:
                return
            elif result:
                self.save_image()
        
        self.window.destroy()

class PDFEditor:
    """Редактор PDF файлов с обрезкой и ластиком"""
    def __init__(self, parent, file_path):
        self.file_path = file_path
        self.pdf_doc = None
        self.current_page = 0
        self.total_pages = 0
        self.page_images = []  # Список PIL изображений страниц
        self.modified_pages = set()  # Набор индексов изменённых страниц
        self.tool_mode = 'select'  # select, crop, eraser, eraser_area
        self.crop_start = None
        self.crop_rect = None
        self.eraser_size = 20
        self.eraser_cursor = None  # Визуальный курсор ластика
        self.image_offset = (0, 0)  # Смещение изображения на холсте
        
        self.page_history = {}  # {page_num: [list of states]}
        self.history_positions = {}
        self.zoom_scale = 1.0  # Масштаб для Ctrl+колёсико
        
        if not PYMUPDF_AVAILABLE:
            messagebox.showerror(
                "Модуль недоступен",
                "Для редактирования PDF необходим модуль PyMuPDF.\n\n"
                "Установите командой:\npip install PyMuPDF",
                parent=parent
            )
            return
        
        if not PIL_AVAILABLE:
            messagebox.showerror(
                "Модуль недоступен",
                "Для редактирования PDF необходим модуль Pillow.\n\n"
                "Установите командой:\npip install Pillow",
                parent=parent
            )
            return
        
        self.window = tk.Toplevel(parent)
        self.window.withdraw()
        self.window.title(f"Редактор: {os.path.basename(file_path)}")
        self.window.geometry("1200x800")
        self.window.transient(parent)
        
        self.create_widgets()
        self.load_pdf()
        
        self.window.deiconify()
    
    def create_widgets(self):
        """Создание виджетов редактора"""
        top_frame = tk.Frame(self.window, bg=COLORS["primary"], height=60)
        top_frame.pack(fill=tk.X)
        top_frame.pack_propagate(False)
        
        tk.Label(
            top_frame,
            text="📕 Редактор PDF",
            bg=COLORS["primary"],
            fg="white",
            font=("Segoe UI", 11, "bold")
        ).pack(side=tk.LEFT, padx=15, pady=15)
        
        btn_frame = tk.Frame(top_frame, bg=COLORS["primary"])
        btn_frame.pack(side=tk.RIGHT, padx=15, pady=10)
        
        tk.Button(
            btn_frame,
            text="↶ Отменить",
            command=self.undo_page_change,
            bg=COLORS["warning"],
            fg="white",
            font=FONTS["button"],
            relief=tk.FLAT,
            cursor="hand2",
            padx=10,
            pady=5
        ).pack(side=tk.LEFT, padx=2)
        
        tk.Button(
            btn_frame,
            text="💾 Сохранить",
            command=self.save_pdf,
            bg=COLORS["success"],
            fg="white",
            font=FONTS["button"],
            relief=tk.FLAT,
            cursor="hand2",
            padx=15,
            pady=5
        ).pack(side=tk.LEFT, padx=2)
        
        main_container = tk.Frame(self.window, bg=COLORS["bg_secondary"])
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Панель инструментов слева
        tools_frame = tk.LabelFrame(
            main_container,
            text=" Инструменты ",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"],
            padx=10,
            pady=10
        )
        tools_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        
        tk.Label(
            tools_frame,
            text="Режим работы:",
            font=FONTS["heading"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=(0, 5))
        
        self.tool_var = tk.StringVar(value='select')
        
        tk.Radiobutton(
            tools_frame,
            text="👆 Выбор",
            variable=self.tool_var,
            value='select',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        tk.Radiobutton(
            tools_frame,
            text="✂️ Обрезка",
            variable=self.tool_var,
            value='crop',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        tk.Radiobutton(
            tools_frame,
            text="🧹 Ластик (точка)",
            variable=self.tool_var,
            value='eraser',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        tk.Radiobutton(
            tools_frame,
            text="🧽 Ластик (область)",
            variable=self.tool_var,
            value='eraser_area',
            command=self.change_tool,
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        ).pack(anchor=tk.W, pady=2)
        
        # Размер ластика
        self.eraser_separator = ttk.Separator(tools_frame, orient=tk.HORIZONTAL)
        
        self.eraser_label = tk.Label(
            tools_frame,
            text="Размер ластика:",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        )
        
        self.eraser_scale = tk.Scale(
            tools_frame,
            from_=5,
            to=100,
            orient=tk.HORIZONTAL,
            command=self.update_eraser_size,
            bg=COLORS["bg_secondary"]
        )
        self.eraser_scale.set(20)
        
        # Изначально НЕ показываем слайдер (т.к. по умолчанию tool_mode='select')
        # Он будет показан только при выборе инструментов ластика
        
        # Область холста
        canvas_frame = tk.Frame(main_container, bg="white")
        canvas_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.canvas = tk.Canvas(canvas_frame, bg="white", cursor="arrow")
        v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL, command=self.canvas.yview)
        h_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL, command=self.canvas.xview)
        
        self.canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # События мыши
        self.canvas.bind('<ButtonPress-1>', self.on_mouse_down)
        self.canvas.bind('<B1-Motion>', self.on_mouse_drag)
        self.canvas.bind('<ButtonRelease-1>', self.on_mouse_up)
        self.canvas.bind('<Motion>', self.on_mouse_move)
        self.canvas.bind('<Control-MouseWheel>', self.on_zoom)
        
        # Навигация по страницам
        nav_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"], height=50)
        nav_frame.pack(fill=tk.X, pady=(5, 0))
        nav_frame.pack_propagate(False)
        
        self.prev_btn = tk.Button(
            nav_frame,
            text="◀ Предыдущая",
            command=self.prev_page,
            font=FONTS["button"],
            bg=COLORS["primary"],
            fg="white",
            relief=tk.FLAT,
            cursor="hand2",
            state=tk.DISABLED
        )
        self.prev_btn.pack(side=tk.LEFT, padx=5, pady=10)
        
        self.page_label = tk.Label(
            nav_frame,
            text="Страница: 0 / 0",
            font=FONTS["body"],
            bg=COLORS["bg_secondary"]
        )
        self.page_label.pack(side=tk.LEFT, expand=True)
        
        self.next_btn = tk.Button(
            nav_frame,
            text="Следующая ▶",
            command=self.next_page,
            font=FONTS["button"],
            bg=COLORS["primary"],
            fg="white",
            relief=tk.FLAT,
            cursor="hand2",
            state=tk.DISABLED
        )
        self.next_btn.pack(side=tk.RIGHT, padx=5, pady=10)
        
        # Статус-бар
        status_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"], height=30)
        status_frame.pack(fill=tk.X)
        status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(
            status_frame,
            text="Загрузка PDF...",
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_secondary"],
            font=FONTS["small"],
            anchor=tk.W
        )
        self.status_label.pack(fill=tk.X, padx=10, pady=5)
    
    def load_pdf(self):
        """Загрузка PDF файла"""
        try:
            import fitz
            from PIL import Image
            
            self.pdf_doc = fitz.open(self.file_path)
            self.total_pages = len(self.pdf_doc)
            
            for page_num in range(self.total_pages):
                page = self.pdf_doc[page_num]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom для качества
                
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img = img.convert("RGBA")  # Для поддержки прозрачности при стирании
                self.page_images.append(img)
                
                self.page_history[page_num] = [img.copy()]
                self.history_positions[page_num] = 0
            
            self.current_page = 0
            self.update_page_display()
            
            if self.total_pages > 1:
                self.next_btn.config(state=tk.NORMAL)
            
            self.status_label.config(text=f"PDF загружен: {self.total_pages} страниц")
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка загрузки PDF:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
            self.window.destroy()
    
    def update_page_display(self):
        """Обновление отображения текущей страницы"""
        from PIL import ImageTk
        
        if not self.page_images:
            return
        
        current_img = self.page_images[self.current_page]
        
        # Масштабируем для отображения
        canvas_width = self.canvas.winfo_width()
        canvas_height = self.canvas.winfo_height()
        
        if canvas_width <= 1:
            canvas_width = 800
        if canvas_height <= 1:
            canvas_height = 600
        
        img_width, img_height = current_img.size
        
        scale_w = (canvas_width - 40) / img_width
        scale_h = (canvas_height - 40) / img_height
        scale = min(scale_w, scale_h, 1.0) * self.zoom_scale
        
        if scale != 1.0:
            new_width = int(img_width * scale)
            new_height = int(img_height * scale)
            try:
                display_img = current_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            except AttributeError:
                display_img = current_img.resize((new_width, new_height), Image.ANTIALIAS)
        else:
            display_img = current_img
        
        self.display_scale = display_img.size[0] / img_width
        self.photo = ImageTk.PhotoImage(display_img)
        
        # Вычисляем смещение изображения на холсте
        img_canvas_x = canvas_width // 2
        img_canvas_y = canvas_height // 2
        self.image_offset = (
            img_canvas_x - (display_img.size[0] // 2),
            img_canvas_y - (display_img.size[1] // 2)
        )
        
        # Удаляем только изображение страницы, а не все объекты canvas (чтобы сохранить курсор ластика)
        self.canvas.delete("page_image")
        self.canvas.create_image(
            img_canvas_x,
            img_canvas_y,
            image=self.photo,
            anchor=tk.CENTER,
            tags="page_image"
        )
        
        # Поднимаем курсор ластика на передний план
        if self.canvas.find_withtag("eraser_cursor"):
            self.canvas.tag_raise("eraser_cursor")
        
        bbox = self.canvas.bbox(tk.ALL)
        if bbox:
            self.canvas.config(scrollregion=bbox)
        
        self.page_label.config(text=f"Страница: {self.current_page + 1} / {self.total_pages}")
        
        self.prev_btn.config(state=tk.NORMAL if self.current_page > 0 else tk.DISABLED)
        self.next_btn.config(state=tk.NORMAL if self.current_page < self.total_pages - 1 else tk.DISABLED)
    
    def prev_page(self):
        """Предыдущая страница"""
        if self.current_page > 0:
            self.current_page -= 1
            self.update_page_display()
    
    def next_page(self):
        """Следующая страница"""
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.update_page_display()
    
    def change_tool(self):
        """Смена инструмента"""
        self.tool_mode = self.tool_var.get()
        
        if self.tool_mode in ('eraser', 'eraser_area'):
            self.canvas.config(cursor="none")
            self.eraser_separator.pack(fill=tk.X, pady=10)
            self.eraser_label.pack(anchor=tk.W, pady=(0, 5))
            self.eraser_scale.pack(fill=tk.X, pady=2)
        else:
            self.canvas.delete("eraser_cursor")
            self.eraser_separator.pack_forget()
            self.eraser_label.pack_forget()
            self.eraser_scale.pack_forget()
        
        if self.tool_mode == 'crop':
            self.canvas.config(cursor="crosshair")
        elif self.tool_mode == 'select':
            self.canvas.config(cursor="hand2")
        
        if self.tool_mode not in ('crop', 'eraser_area') and self.crop_rect:
            self.canvas.delete("crop_rect")
            self.crop_rect = None
            self.crop_start = None
    
    def update_eraser_size(self, value):
        """Обновление размера ластика"""
        self.eraser_size = int(float(value))
        # Удаляем текущий курсор, чтобы он пересоздался с новым размером
        self.canvas.delete("eraser_cursor")
        self.eraser_cursor = None
    
    def update_eraser_cursor(self, x=None, y=None):
        """Обновление визуального курсора ластика"""
        self.canvas.delete("eraser_cursor")
        self.eraser_cursor = None
        self.canvas.update_idletasks()
        
        if self.tool_mode in ('eraser', 'eraser_area') and x is not None and y is not None:
            radius = self.eraser_size / 2
            self.eraser_cursor = self.canvas.create_oval(
                x - radius, y - radius,
                x + radius, y + radius,
                outline="red",
                width=1,
                tags="eraser_cursor"
            )
    
    def on_mouse_move(self, event):
        """Движение мыши (без нажатия кнопки)"""
        if self.tool_mode in ('eraser', 'eraser_area'):
            x = self.canvas.canvasx(event.x)
            y = self.canvas.canvasy(event.y)
            self.update_eraser_cursor(x, y)
    
    def on_mouse_down(self, event):
        """Нажатие мыши"""
        x = self.canvas.canvasx(event.x)
        y = self.canvas.canvasy(event.y)
        
        if self.tool_mode == 'select':
            self.canvas.scan_mark(event.x, event.y)
        elif self.tool_mode in ('crop', 'eraser_area'):
            self.crop_start = (x, y)
            if self.crop_rect:
                self.canvas.delete("crop_rect")
        elif self.tool_mode == 'eraser':
            self.use_eraser(x, y)
    
    def on_mouse_drag(self, event):
        """Движение мыши"""
        # Преобразуем экранные координаты в координаты холста с учетом прокрутки
        x = self.canvas.canvasx(event.x)
        y = self.canvas.canvasy(event.y)
        
        if self.tool_mode == 'select':
            self.canvas.scan_dragto(event.x, event.y, gain=1)
        elif self.tool_mode in ('crop', 'eraser_area') and self.crop_start:
            if self.crop_rect:
                self.canvas.delete("crop_rect")
            color = "red" if self.tool_mode == 'crop' else "blue"
            self.crop_rect = self.canvas.create_rectangle(
                self.crop_start[0], self.crop_start[1],
                x, y,
                outline=color, width=2, dash=(5, 5), tags="crop_rect"
            )
            if self.tool_mode == 'eraser_area':
                self.update_eraser_cursor(x, y)
        elif self.tool_mode == 'eraser':
            self.use_eraser(x, y)
            self.update_eraser_cursor(x, y)
    
    def on_mouse_up(self, event):
        """Отпускание мыши"""
        x = self.canvas.canvasx(event.x)
        y = self.canvas.canvasy(event.y)
        
        if self.tool_mode == 'crop' and self.crop_start:
            self.apply_crop(self.crop_start[0], self.crop_start[1], x, y)
            self.crop_start = None
            if self.crop_rect:
                self.canvas.delete("crop_rect")
                self.crop_rect = None
        elif self.tool_mode == 'eraser_area' and self.crop_start:
            self.erase_area(self.crop_start[0], self.crop_start[1], x, y)
            self.crop_start = None
            if self.crop_rect:
                self.canvas.delete("crop_rect")
                self.crop_rect = None
    
    def on_zoom(self, event):
        """Масштабирование через Ctrl+колёсико мыши"""
        # event.delta > 0 - крутим вверх (приближение)
        # event.delta < 0 - крутим вниз (отдаление)
        if event.delta > 0:
            self.zoom_scale = min(self.zoom_scale * 1.1, 5.0)
        else:
            self.zoom_scale = max(self.zoom_scale / 1.1, 0.1)
        
        self.update_page_display()
        self.status_label.config(text=f"🔍 Масштаб: {int(self.zoom_scale * 100)}%")
    
    def use_eraser(self, x, y):
        """Использование ластика"""
        from PIL import ImageDraw
        
        if self.current_page not in self.modified_pages:
            self.save_page_to_history()
        
        bbox = self.canvas.bbox("page_image")
        if not bbox:
            return
        
        current_img = self.page_images[self.current_page]
        
        img_x = int((x - bbox[0]) / self.display_scale)
        img_y = int((y - bbox[1]) / self.display_scale)
        
        draw = ImageDraw.Draw(current_img)
        radius = int(self.eraser_size / (2 * self.display_scale))
        
        bbox = [
            img_x - radius, img_y - radius,
            img_x + radius, img_y + radius
        ]
        
        draw.ellipse(bbox, fill=(255, 255, 255, 255))
        
        self.modified_pages.add(self.current_page)
        self.update_page_display()
        self.status_label.config(text=f"🧹 Ластик применён на странице {self.current_page + 1}")
    
    def erase_area(self, x1, y1, x2, y2):
        """Стирание прямоугольной области"""
        from PIL import ImageDraw
        
        self.save_page_to_history()
        
        bbox = self.canvas.bbox("page_image")
        if not bbox:
            return
        
        # Нормализуем координаты
        if x1 > x2:
            x1, x2 = x2, x1
        if y1 > y2:
            y1, y2 = y2, y1
        
        current_img = self.page_images[self.current_page]
        
        # Конвертируем в координаты изображения
        img_x1 = int((x1 - bbox[0]) / self.display_scale)
        img_y1 = int((y1 - bbox[1]) / self.display_scale)
        img_x2 = int((x2 - bbox[0]) / self.display_scale)
        img_y2 = int((y2 - bbox[1]) / self.display_scale)
        
        # Ограничиваем координаты
        img_x1 = max(0, img_x1)
        img_y1 = max(0, img_y1)
        img_x2 = min(current_img.width, img_x2)
        img_y2 = min(current_img.height, img_y2)
        
        if img_x2 - img_x1 < 2 or img_y2 - img_y1 < 2:
            return
        
        # Стираем область (закрашиваем белым)
        draw = ImageDraw.Draw(current_img)
        draw.rectangle([img_x1, img_y1, img_x2, img_y2], fill=(255, 255, 255, 255))
        
        self.modified_pages.add(self.current_page)
        self.update_page_display()
        self.status_label.config(text=f"🧽 Стёрта область на странице {self.current_page + 1}: {img_x2-img_x1}x{img_y2-img_y1} px")
    
    def apply_crop(self, x1, y1, x2, y2):
        """Применение обрезки"""
        self.save_page_to_history()
        
        bbox = self.canvas.bbox("page_image")
        if not bbox:
            return
        
        if x1 > x2:
            x1, x2 = x2, x1
        if y1 > y2:
            y1, y2 = y2, y1
        
        current_img = self.page_images[self.current_page]
        
        img_x1 = int((x1 - bbox[0]) / self.display_scale)
        img_y1 = int((y1 - bbox[1]) / self.display_scale)
        img_x2 = int((x2 - bbox[0]) / self.display_scale)
        img_y2 = int((y2 - bbox[1]) / self.display_scale)
        img_x1 = max(0, img_x1)
        img_y1 = max(0, img_y1)
        img_x2 = min(current_img.width, img_x2)
        img_y2 = min(current_img.height, img_y2)
        
        if img_x2 - img_x1 < 10 or img_y2 - img_y1 < 10:
            messagebox.showwarning("Предупреждение", "Выберите область побольше", parent=self.window)
            return
        
        self.page_images[self.current_page] = current_img.crop((img_x1, img_y1, img_x2, img_y2))
        self.modified_pages.add(self.current_page)
        self.update_page_display()
        self.status_label.config(text=f"✂️ Обрезана страница {self.current_page + 1}")
    
    def save_pdf(self):
        """Сохранение PDF файла"""
        try:
            import fitz
            from PIL import Image
            import io
            
            dir_name = os.path.dirname(self.file_path)
            base_name = os.path.basename(self.file_path)
            name_without_ext, ext = os.path.splitext(base_name)
            
            # Убираем старую пометку (ред) если есть
            if name_without_ext.endswith('(ред)'):
                name_without_ext = name_without_ext[:-5].strip()
            
            new_name = f"{name_without_ext}(ред){ext}"
            new_path = os.path.join(dir_name, new_name)
            
            new_doc = fitz.open()
            
            for page_num, img in enumerate(self.page_images):
                # Конвертируем PIL изображение в байты
                img_rgb = img.convert("RGB")
                img_bytes = io.BytesIO()
                img_rgb.save(img_bytes, format='JPEG', quality=95)
                img_bytes.seek(0)
                
                img_doc = fitz.open(stream=img_bytes, filetype="jpeg")
                pdf_bytes = img_doc.convert_to_pdf()
                img_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
                new_doc.insert_pdf(img_pdf)
            
            new_doc.save(new_path, garbage=4, deflate=True)
            new_doc.close()
            
            self.file_path = new_path
            self.modified_pages.clear()
            self.status_label.config(text="✓ PDF сохранён")
            
            messagebox.showinfo(
                "Успешно",
                f"PDF файл успешно сохранён!\n\nФайл: {new_name}",
                parent=self.window
            )
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при сохранении PDF:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
    
    def save_page_to_history(self):
        """Сохранение текущего состояния страницы в историю"""
        page_num = self.current_page
        current_img = self.page_images[page_num].copy()
        
        if page_num not in self.page_history:
            self.page_history[page_num] = []
            self.history_positions[page_num] = -1
        
        if self.history_positions[page_num] < len(self.page_history[page_num]) - 1:
            self.page_history[page_num] = self.page_history[page_num][:self.history_positions[page_num] + 1]
        
        self.page_history[page_num].append(current_img)
        self.history_positions[page_num] += 1
        
        if len(self.page_history[page_num]) > 20:
            self.page_history[page_num].pop(0)
            self.history_positions[page_num] -= 1
    
    def undo_page_change(self):
        """Отмена последнего изменения для текущей страницы"""
        page_num = self.current_page
        
        if page_num not in self.page_history or not self.page_history[page_num]:
            self.status_label.config(text="⚠️ Нет истории изменений для этой страницы")
            return
        
        if self.history_positions[page_num] <= 0:
            self.status_label.config(text="⚠️ Достигнуто начало истории изменений")
            return
        
        # Откатываемся к предыдущему состоянию
        self.history_positions[page_num] -= 1
        self.page_images[page_num] = self.page_history[page_num][self.history_positions[page_num]].copy()
        self.modified_pages.add(page_num)
        self.update_page_display()
        self.status_label.config(text=f"↶ Отменено (позиция {self.history_positions[page_num] + 1}/{len(self.page_history[page_num])})")
    
    def close_editor(self):
        """Закрытие редактора"""
        if self.modified_pages:
            result = messagebox.askyesnocancel(
                "Несохранённые изменения",
                "Сохранить изменения перед закрытием?",
                parent=self.window
            )
            if result is None:
                return
            elif result:
                self.save_pdf()
        
        if self.pdf_doc:
            self.pdf_doc.close()
        
        self.window.destroy()

# ── КЛАСС ДЛЯ ПРЕДВАРИТЕЛЬНОГО ПРОСМОТРА ФАЙЛОВ ──────────────────────
class PreviewWindow:
    """Окно предварительного просмотра файлов различных типов"""
    def __init__(self, parent, file_path, title="Предварительный просмотр", data_manager=None):
        self.file_path = file_path
        self.temp_pdf_path = None  # Для хранения пути к временному PDF файлу
        self.temp_docx_path = None  # Для хранения пути к временному DOCX файлу
        self.pdf_doc = None  # Для хранения PDF документа
        self.parent = parent
        self.data_manager = data_manager
        self.zoom_level = 0.7  # Начальный масштаб для PDF (0.7 = 35%)
        
        self.window = tk.Toplevel(parent)
        self.window.withdraw()
        self.window.title(title)
        self.window.geometry("900x700")
        self.window.transient(parent)
        
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        self.window.update_idletasks()
        parent.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (900 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (700 // 2)
        self.window.geometry(f"+{x}+{y}")
        
        self.create_widgets()
        self.load_preview()
        
        self.window.deiconify()
    
    def on_closing(self):
        """Обработчик закрытия окна с очисткой временных файлов"""
        try:
            if hasattr(self, 'pdf_doc') and self.pdf_doc:
                self.pdf_doc.close()
            
            if hasattr(self, 'temp_pdf_path') and self.temp_pdf_path and os.path.exists(self.temp_pdf_path):
                try:
                    os.unlink(self.temp_pdf_path)
                except:
                    pass
            
            if hasattr(self, 'temp_docx_path') and self.temp_docx_path and os.path.exists(self.temp_docx_path):
                try:
                    os.unlink(self.temp_docx_path)
                except:
                    pass
        except:
            pass
        
        self.window.destroy()
    
    def create_widgets(self):
        """Создание виджетов окна"""
        top_frame = tk.Frame(self.window, bg=COLORS["primary"], height=60)
        top_frame.pack(fill=tk.X)
        top_frame.pack_propagate(False)
        
        file_name = os.path.basename(self.file_path)
        file_ext = os.path.splitext(file_name)[1].lower()
        
        icon = "📄"
        if file_ext in ['.docx', '.doc']:
            icon = "📝"
        elif file_ext in ['.xlsx', '.xls']:
            icon = "📊"
        elif file_ext == '.pdf':
            icon = "📕"
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            icon = "🖼"
        
        tk.Label(
            top_frame,
            text=f"{icon}  {file_name}",
            bg=COLORS["primary"],
            fg="white",
            font=("Segoe UI", 11, "bold")
        ).pack(side=tk.LEFT, padx=15, pady=15)
        
        edit_btn = tk.Button(
            top_frame,
            text="✏️ Редактировать",
            command=self.open_editor,
            bg=COLORS["success"],
            fg="white",
            font=FONTS["button"],
            relief=tk.FLAT,
            cursor="hand2",
            padx=15,
            pady=5
        )
        edit_btn.pack(side=tk.RIGHT, padx=(0, 5), pady=15)
        
        main_frame = tk.Frame(self.window, bg=COLORS["bg_secondary"])
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        file_ext = os.path.splitext(self.file_path)[1].lower()
        
        if file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.pdf', '.docx', '.doc']:
            # Для изображений, PDF и Word создаем Canvas с прокруткой
            self.canvas_frame = tk.Frame(main_frame)
            self.canvas_frame.pack(fill=tk.BOTH, expand=True)
            
            self.canvas = tk.Canvas(self.canvas_frame, bg="white")
            v_scrollbar = tk.Scrollbar(self.canvas_frame, orient=tk.VERTICAL, command=self.canvas.yview)
            h_scrollbar = tk.Scrollbar(self.canvas_frame, orient=tk.HORIZONTAL, command=self.canvas.xview)
            
            self.canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
            
            v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
            self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            # Привязываем обработчик Ctrl+колесико мыши для масштабирования
            self.canvas.bind("<Control-MouseWheel>", self.on_mouse_wheel)
            
            # Привязываем обработчики для перемещения изображения мышью (drag-to-pan)
            self.canvas.bind("<ButtonPress-1>", self.on_canvas_click)
            self.canvas.bind("<B1-Motion>", self.on_canvas_drag)
            self.canvas.bind("<ButtonRelease-1>", self.on_canvas_release)
            
            self.preview_widget = self.canvas
            
            # Для PDF и Word добавляем навигацию между страницами
            if file_ext in ['.pdf', '.docx', '.doc']:
                self.pdf_current_page = 0
                self.pdf_total_pages = 0
                self.pdf_images = []  # Кэш изображений страниц
                
                nav_frame = tk.Frame(main_frame, bg=COLORS["bg_secondary"], height=50)
                nav_frame.pack(fill=tk.X, pady=(5, 0))
                nav_frame.pack_propagate(False)
                
                self.prev_page_btn = tk.Button(
                    nav_frame,
                    text="◀ Предыдущая",
                    command=self.prev_pdf_page,
                    font=FONTS["button"],
                    bg=COLORS["primary"],
                    fg="white",
                    relief=tk.FLAT,
                    cursor="hand2",
                    state=tk.DISABLED
                )
                self.prev_page_btn.pack(side=tk.LEFT, padx=5, pady=10)
                
                self.page_label = tk.Label(
                    nav_frame,
                    text="Страница: 0 / 0",
                    font=FONTS["body"],
                    bg=COLORS["bg_secondary"]
                )
                self.page_label.pack(side=tk.LEFT, expand=True)
                
                self.next_page_btn = tk.Button(
                    nav_frame,
                    text="Следующая ▶",
                    command=self.next_pdf_page,
                    font=FONTS["button"],
                    bg=COLORS["primary"],
                    fg="white",
                    relief=tk.FLAT,
                    cursor="hand2",
                    state=tk.DISABLED
                )
                self.next_page_btn.pack(side=tk.RIGHT, padx=5, pady=10)
        elif file_ext in ['.xlsx', '.xls']:
            # Для Excel создаём современную таблицу Treeview
            tree_container, self.tree = create_modern_treeview(
                main_frame,
                columns=[],
                height=20
            )
            tree_container.pack(fill=tk.BOTH, expand=True)
            
            self.preview_widget = self.tree
        else:
            # Для текстовых файлов используем ScrolledText
            self.preview_text = ScrolledText(
                main_frame,
                wrap=tk.WORD,
                font=("Consolas", 10),
                bg="white",
                fg=COLORS["text_primary"],
                relief=tk.SOLID,
                borderwidth=1
            )
            self.preview_text.pack(fill=tk.BOTH, expand=True)
            enable_field_shortcuts(self.preview_text, readonly=True)
            add_context_menu(self.preview_text, readonly=True)
            enable_text_selection_in_disabled(self.preview_text)
            
            self.preview_widget = self.preview_text
    
    def load_preview(self):
        """Загрузка и отображение предварительного просмотра"""
        file_ext = os.path.splitext(self.file_path)[1].lower()
        
        try:
            if file_ext in ['.docx', '.doc']:
                self.preview_word()
            elif file_ext == '.pdf':
                self.preview_pdf()
            elif file_ext in ['.xlsx', '.xls']:
                self.preview_excel()
            elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                self.preview_image()
            else:
                if hasattr(self, 'preview_text'):
                    self.preview_text.insert(tk.END, "Предварительный просмотр для этого типа файла не поддерживается.")
                    self.preview_text.config(state=tk.DISABLED)
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при загрузке файла:\n{str(e)}\n\n{traceback.format_exc()}"
            if hasattr(self, 'preview_text'):
                self.preview_text.insert(tk.END, error_msg)
                self.preview_text.config(state=tk.DISABLED)
            else:
                messagebox.showerror("Ошибка", error_msg, parent=self.window)
                self.on_closing()
    
    def preview_word(self):
        """Предварительный просмотр Word документа с выделением активных плейсхолдеров"""
        if not PYMUPDF_AVAILABLE:
            messagebox.showerror(
                "Библиотека недоступна",
                "Для визуального просмотра Word документов необходима библиотека PyMuPDF.\n\n"
                "Установите командой:\n"
                "pip install PyMuPDF\n\n"
                "После установки перезапустите приложение.",
                parent=self.window
            )
            self.on_closing()
            return
        
        try:
            import fitz
            import tempfile
            from docx import Document
            from docx.shared import RGBColor
            
            active_placeholders = []
            try:
                # Пытаемся получить data_manager (либо переданный напрямую, либо через parent)
                data_manager = self.data_manager
                if not data_manager:
                    current = self.parent
                    while current and not data_manager:
                        if hasattr(current, 'data_manager'):
                            data_manager = current.data_manager
                            break
                        if hasattr(current, 'master'):
                            current = current.master
                        else:
                            break
                
                if data_manager:
                    if hasattr(data_manager, 'PLACEHOLDERS'):
                        all_placeholders = data_manager.PLACEHOLDERS
                        # Берём только активные плейсхолдеры
                        active_placeholders = [
                            ph['name'] for ph in all_placeholders 
                            if ph.get('active', True)
                        ]
            except Exception as e:
                pass
            
            temp_docx_fd, temp_docx_path = tempfile.mkstemp(suffix='.docx', prefix='word_highlighted_')
            os.close(temp_docx_fd)
            
            doc = Document(self.file_path)
            
            # Выделяем плейсхолдеры в документе желтым цветом
            if active_placeholders:
                self.page_label.config(text="Выделение плейсхолдеров...")
                self.window.update()
                
                # Функция для разбиения run и выделения только плейсхолдеров
                def split_and_highlight_run(para, run, placeholders):
                    """
                    Разбивает run на части, выделяя только сами плейсхолдеры (слова),
                    а не весь текст run
                    """
                    text = run.text
                    if not text:
                        return 0
                    
                    # Находим все вхождения плейсхолдеров в этом run
                    matches = []
                    for placeholder in placeholders:
                        pattern = create_placeholder_pattern(placeholder)
                        for match in re.finditer(pattern, text, re.IGNORECASE):
                            matches.append((match.start(), match.end(), placeholder))
                    
                    if not matches:
                        return 0
                    
                    # Сортируем по позиции
                    matches.sort(key=lambda x: x[0])
                    
                    style = run.style
                    bold = run.bold
                    italic = run.italic
                    underline = run.underline
                    font_name = run.font.name
                    font_size = run.font.size
                    
                    # Разбиваем текст на части
                    parts = []
                    last_end = 0
                    for start, end, ph in matches:
                        # Текст до плейсхолдера
                        if start > last_end:
                            parts.append((text[last_end:start], False, None))
                        # Сам плейсхолдер
                        parts.append((text[start:end], True, ph))
                        last_end = end
                    # Остаток после последнего плейсхолдера
                    if last_end < len(text):
                        parts.append((text[last_end:], False, None))
                    
                    # Находим позицию run в параграфе (по XML элементу)
                    run_element = run._element
                    parent_element = run_element.getparent()
                    run_index_in_xml = list(parent_element).index(run_element)
                    
                    # Удаляем оригинальный run
                    parent_element.remove(run_element)
                    
                    highlighted_count = 0
                    for part_text, is_placeholder, ph_name in parts:
                        new_run = para.add_run(part_text)
                        
                        # Копируем форматирование (без цвета текста)
                        if style:
                            new_run.style = style
                        # Копируем bold/italic/underline только для НЕ-плейсхолдеров
                        if not is_placeholder:
                            if bold is not None:
                                new_run.bold = bold
                            if italic is not None:
                                new_run.italic = italic
                            if underline is not None:
                                new_run.underline = underline
                        if font_name:
                            new_run.font.name = font_name
                        if font_size:
                            new_run.font.size = font_size
                        
                        # Выделяем плейсхолдеры ТОЛЬКО жирным + подчеркиванием (БЕЗ цвета)
                        if is_placeholder:
                            new_run.bold = True
                            new_run.underline = True
                            highlighted_count += 1
                        
                        # Вставляем в правильную позицию
                        parent_element.insert(run_index_in_xml, new_run._element)
                        run_index_in_xml += 1
                    
                    return highlighted_count
                
                # Проходим по всем параграфам
                highlighted_words = 0
                for para_idx, para in enumerate(doc.paragraphs):
                    para_text = para.text
                    if para_text and any(re.search(create_placeholder_pattern(ph), para_text, re.IGNORECASE) for ph in active_placeholders):
                        for run_idx in range(len(para.runs) - 1, -1, -1):
                            run = para.runs[run_idx]
                            highlighted_words += split_and_highlight_run(para, run, active_placeholders)
                
                # Проходим по таблицам
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            for para in cell.paragraphs:
                                para_text = para.text
                                if para_text and any(re.search(create_placeholder_pattern(ph), para_text, re.IGNORECASE) for ph in active_placeholders):
                                    for run_idx in range(len(para.runs) - 1, -1, -1):
                                        run = para.runs[run_idx]
                                        highlighted_words += split_and_highlight_run(para, run, active_placeholders)
            
            doc.save(temp_docx_path)
            
            # Теперь конвертируем в PDF
            self.page_label.config(text="Конвертация в PDF...")
            self.window.update()
            
            temp_pdf_path = None
            
            try:
                # Пробуем конвертировать через win32com (Windows)
                if WIN32COM_AVAILABLE:
                    import win32com.client
                    import pythoncom
                    
                    try:
                        pythoncom.CoInitialize()
                        
                        # Используем EnsureDispatch для надёжной работы с COM
                        try:
                            word = win32com.client.gencache.EnsureDispatch("Word.Application")
                        except:
                            # Fallback на обычный Dispatch
                            word = win32com.client.Dispatch("Word.Application")
                        
                        word.Visible = False
                        
                        temp_pdf_fd, temp_pdf_path = tempfile.mkstemp(suffix='.pdf', prefix='word_preview_')
                        os.close(temp_pdf_fd)
                        
                        # Открываем модифицированный Word документ
                        word_doc = word.Documents.Open(os.path.abspath(temp_docx_path))
                        
                        word_doc.SaveAs(os.path.abspath(temp_pdf_path), FileFormat=17)
                        
                        word_doc.Close()
                        word.Quit()
                        
                        pythoncom.CoUninitialize()
                    
                    except Exception as word_error:
                        # Закрываем Word даже если была ошибка
                        try:
                            if 'word' in locals():
                                word.Quit()
                            pythoncom.CoUninitialize()
                        except:
                            pass
                        
                        raise Exception(f"Не удалось конвертировать документ через Word: {word_error}")
                    
                elif DOCX2PDF_AVAILABLE:
                    # Используем docx2pdf
                    from docx2pdf import convert
                    
                    temp_pdf_fd, temp_pdf_path = tempfile.mkstemp(suffix='.pdf', prefix='word_preview_')
                    os.close(temp_pdf_fd)
                    
                    convert(temp_docx_path, temp_pdf_path)
                else:
                    # Если нет ни одного конвертера
                    messagebox.showerror(
                        "Конвертер недоступен",
                        "Для визуального просмотра Word необходим конвертер.\n\n"
                        "Windows: Установите Microsoft Word\n"
                        "Или установите: pip install docx2pdf\n\n"
                        "После установки перезапустите приложение.",
                        parent=self.window
                    )
                    # Удаляем временный docx
                    try:
                        os.unlink(temp_docx_path)
                    except:
                        pass
                    self.on_closing()
                    return
                
                self.temp_pdf_path = temp_pdf_path
                self.temp_docx_path = temp_docx_path
                    
            except Exception as e:
                if temp_pdf_path and os.path.exists(temp_pdf_path):
                    try:
                        os.unlink(temp_pdf_path)
                    except:
                        pass
                if temp_docx_path and os.path.exists(temp_docx_path):
                    try:
                        os.unlink(temp_docx_path)
                    except:
                        pass
                raise e
            
            # Открываем PDF для отображения
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                pdf_doc = fitz.open(temp_pdf_path)
                self.pdf_total_pages = len(pdf_doc)
                self.pdf_doc = pdf_doc
                
                self.page_label.config(text=f"Страница: 1 / {self.pdf_total_pages}")
                
                # Активируем кнопки навигации
                if self.pdf_total_pages > 1:
                    self.next_page_btn.config(state=tk.NORMAL)
                
                self.show_pdf_page(0)
            else:
                raise Exception("Не удалось создать временный PDF файл")
                
        except Exception as e:
            error_msg = f"Не удалось открыть предпросмотр Word документа.\n\n"
            error_msg += f"Ошибка: {str(e)}\n\n"
            error_msg += "Возможные причины:\n"
            error_msg += "• Microsoft Word не установлен\n"
            error_msg += "• Word уже открыт (закройте все окна Word)\n"
            error_msg += "• Недостаточно прав доступа\n"
            
            messagebox.showerror("Ошибка предпросмотра", error_msg, parent=self.window)
            self.on_closing()
    
    def preview_pdf(self):
        """Предварительный просмотр PDF документа с визуальным отображением страниц"""
        if not PYMUPDF_AVAILABLE:
            messagebox.showerror(
                "Библиотека недоступна",
                "Для визуального просмотра PDF необходима библиотека PyMuPDF.\n\n"
                "Установите командой:\n"
                "pip install PyMuPDF\n\n"
                "После установки перезапустите приложение.",
                parent=self.window
            )
            self.window.destroy()
            return
        
        try:
            import fitz
            
            doc = fitz.open(self.file_path)
            self.pdf_total_pages = len(doc)
            self.pdf_doc = doc
            
            zoom_percent = int(self.zoom_level * 50)  # 2.0 = 100%
            self.page_label.config(text=f"Страница: 1 / {self.pdf_total_pages}  (Масштаб: {zoom_percent}%)")
            
            if self.pdf_total_pages > 1:
                self.next_page_btn.config(state=tk.NORMAL)
            
            self.show_pdf_page(0)
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при открытии PDF:\n\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
            self.window.destroy()
    
    def show_pdf_page(self, page_num):
        """Отображение указанной страницы PDF"""
        try:
            from PIL import Image, ImageTk
            
            if page_num < len(self.pdf_images) and self.pdf_images[page_num] is not None:
                img = self.pdf_images[page_num]
            else:
                # Конвертируем страницу в изображение
                page = self.pdf_doc[page_num]
                
                # Увеличиваем разрешение в соответствии с текущим масштабом
                mat = fitz.Matrix(self.zoom_level, self.zoom_level)
                # Убираем явное указание colorspace - пусть PyMuPDF использует нативный из PDF
                pix = page.get_pixmap(matrix=mat)
                
                # Конвертируем в PIL Image
                img_data = pix.tobytes("ppm")
                img = Image.frombytes("RGB", [pix.width, pix.height], img_data)
                
                while len(self.pdf_images) <= page_num:
                    self.pdf_images.append(None)
                self.pdf_images[page_num] = img
            
            self.canvas.update_idletasks()
            canvas_width = self.canvas.winfo_width()
            canvas_height = self.canvas.winfo_height()
            
            if canvas_width <= 1:
                canvas_width = 860
            if canvas_height <= 1:
                canvas_height = 550
            
            # Используем изображение как есть (с учётом zoom_level из fitz.Matrix)
            # Не масштабируем под размер canvas, позволяем прокрутку
            resized_img = img
            
            # Конвертируем и отображаем
            self.current_pdf_photo = ImageTk.PhotoImage(resized_img)
            self.canvas.delete("all")
            
            canvas_center_x = canvas_width // 2
            canvas_center_y = canvas_height // 2
            
            self.canvas.create_image(canvas_center_x, canvas_center_y, image=self.current_pdf_photo, anchor=tk.CENTER)
            
            bbox = self.canvas.bbox(tk.ALL)
            if bbox:
                self.canvas.config(scrollregion=bbox)
            
            self.pdf_current_page = page_num
            zoom_percent = int(self.zoom_level * 50)  # 2.0 = 100%, 1.0 = 50%, 4.0 = 200%
            self.page_label.config(text=f"Страница: {page_num + 1} / {self.pdf_total_pages}  (Масштаб: {zoom_percent}%)")
            
            # Управление кнопками навигации
            self.prev_page_btn.config(state=tk.NORMAL if page_num > 0 else tk.DISABLED)
            self.next_page_btn.config(state=tk.NORMAL if page_num < self.pdf_total_pages - 1 else tk.DISABLED)
            
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при отображении страницы:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
    
    def prev_pdf_page(self):
        """Предыдущая страница PDF"""
        if self.pdf_current_page > 0:
            self.show_pdf_page(self.pdf_current_page - 1)
    
    def next_pdf_page(self):
        """Следующая страница PDF"""
        if self.pdf_current_page < self.pdf_total_pages - 1:
            self.show_pdf_page(self.pdf_current_page + 1)
    
    def on_mouse_wheel(self, event):
        """Обработчик Ctrl+колесико мыши для масштабирования"""
        if not self.pdf_doc:
            return
        
        if event.delta > 0:
            self.change_zoom(0.2)  # Увеличиваем на 20%
        elif event.delta < 0:
            self.change_zoom(-0.2)  # Уменьшаем на 20%
    
    def change_zoom(self, delta):
        """Изменение масштаба документа"""
        # Ограничиваем масштаб от 0.5x до 5.0x
        new_zoom = self.zoom_level + delta
        new_zoom = max(0.5, min(5.0, new_zoom))
        
        if new_zoom != self.zoom_level:
            self.zoom_level = new_zoom
            self.pdf_images = []
            # Перерисовываем текущую страницу
            self.show_pdf_page(self.pdf_current_page)
            zoom_percent = int(self.zoom_level * 50)  # 2.0 = 100%, 1.0 = 50%, 4.0 = 200%
            self.page_label.config(
                text=f"Страница: {self.pdf_current_page + 1} / {self.pdf_total_pages}  (Масштаб: {zoom_percent}%)"
            )
    
    def on_canvas_click(self, event):
        """Обработчик нажатия левой кнопки мыши - начало перетаскивания"""
        self.canvas.scan_mark(event.x, event.y)
        self.canvas.config(cursor="fleur")
    
    def on_canvas_drag(self, event):
        """Обработчик перемещения мыши с зажатой кнопкой - перетаскивание изображения"""
        self.canvas.scan_dragto(event.x, event.y, gain=1)
    
    def on_canvas_release(self, event):
        """Обработчик отпускания кнопки мыши - конец перетаскивания"""
        self.canvas.config(cursor="")
    
    def preview_excel(self):
        """Предварительный просмотр Excel файла"""
        try:
            import pandas as pd
        except ImportError:
            # Fallback для текстового виджета если нет pandas
            if hasattr(self, 'preview_text'):
                self.preview_text.insert(tk.END, "Модуль pandas не установлен.\n")
                self.preview_text.insert(tk.END, "Для просмотра Excel файлов установите: pip install pandas openpyxl")
                self.preview_text.config(state=tk.DISABLED)
            return
        
        try:
            df = pd.read_excel(self.file_path, nrows=100, engine='openpyxl')
            
            self.tree["columns"] = list(df.columns)
            self.tree["show"] = "headings"
            
            # Заголовки столбцов с современным стилем
            for col in df.columns:
                self.tree.heading(col, text=str(col))
                # Автоматическая ширина столбца
                max_width = max(
                    len(str(col)) * 9,
                    df[col].astype(str).str.len().max() * 9 if not df[col].empty else 60
                )
                # stretch=tk.NO позволяет растягивать столбцы вручную без авто-перераспределения
                self.tree.column(col, width=min(max_width, 350), anchor=tk.W, minwidth=50, stretch=tk.NO)
            
            for idx, row in df.iterrows():
                values = format_dataframe_row(row)
                insert_treeview_row(self.tree, values)
            
            total_rows = len(df)
            if total_rows >= 100:
                messagebox.showinfo(
                    "Информация",
                    f"Показаны первые 100 строк из файла.\nВсего строк может быть больше.",
                    parent=self.window
                )
        except Exception as e:
            # Если что-то пошло не так, показываем ошибку
            messagebox.showerror("Ошибка", f"Не удалось загрузить Excel файл:\n{str(e)}", parent=self.window)
            self.window.destroy()
    
    def open_editor(self):
        """Открывает редактор в зависимости от типа файла"""
        file_ext = os.path.splitext(self.file_path)[1].lower()
        
        try:
            if file_ext in ['.docx', '.doc']:
                open_word_document(self.file_path)
            elif file_ext in ['.xlsx', '.xls']:
                ExcelEditor(self.window, self.file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                ImageEditor(self.window, self.file_path)
            elif file_ext == '.pdf':
                PDFEditor(self.window, self.file_path)
            else:
                messagebox.showwarning(
                    "Недоступно",
                    "Редактирование этого типа файлов не поддерживается.",
                    parent=self.window
                )
        except Exception as e:
            import traceback
            error_msg = f"Ошибка при открытии редактора:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
    
    def preview_image(self):
        """Предварительный просмотр изображения"""
        if not PIL_AVAILABLE:
            messagebox.showerror(
                "Ошибка",
                "Модуль Pillow не установлен.\nДля просмотра изображений установите: pip install Pillow",
                parent=self.window
            )
            self.window.destroy()
            return
        
        try:
            from PIL import Image, ImageTk
            
            img = Image.open(self.file_path)
            original_size = img.size
            
            # Функция для отображения изображения после отрисовки canvas
            def display_image():
                try:
                    # Принудительно обновляем окно чтобы canvas получил реальные размеры
                    self.canvas.update_idletasks()
                    
                    canvas_width = self.canvas.winfo_width()
                    canvas_height = self.canvas.winfo_height()
                    
                    # Если canvas еще не отрисован, используем размеры из geometry
                    if canvas_width <= 1:
                        canvas_width = 860  # 900 - 40 (отступы)
                    if canvas_height <= 1:
                        canvas_height = 600  # 700 - 100 (верх и низ)
                    
                    # Масштабируем изображение, сохраняя пропорции
                    img_width, img_height = img.size
                    
                    # Вычисляем коэффициент масштабирования
                    scale_w = (canvas_width - 40) / img_width
                    scale_h = (canvas_height - 40) / img_height
                    scale = min(scale_w, scale_h, 1.0)  # Не увеличиваем, только уменьшаем
                    
                    if scale < 1.0:
                        new_width = int(img_width * scale)
                        new_height = int(img_height * scale)
                        # Используем LANCZOS или ANTIALIAS для совместимости
                        try:
                            resized_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        except AttributeError:
                            try:
                                resized_img = img.resize((new_width, new_height), Image.ANTIALIAS)
                            except AttributeError:
                                # Для старых версий PIL
                                resized_img = img.resize((new_width, new_height))
                    else:
                        resized_img = img
                    
                    # Конвертируем в формат Tkinter
                    self.photo = ImageTk.PhotoImage(resized_img)
                    
                    self.canvas.delete("all")
                    
                    # Отображаем на Canvas по центру
                    img_x = max(canvas_width // 2, resized_img.size[0] // 2)
                    img_y = max(canvas_height // 2, resized_img.size[1] // 2)
                    
                    self.canvas.create_image(
                        img_x,
                        img_y,
                        image=self.photo,
                        anchor=tk.CENTER
                    )
                    
                    bbox = self.canvas.bbox(tk.ALL)
                    if bbox:
                        self.canvas.config(scrollregion=bbox)
                    
                    info_text = f"{os.path.basename(self.file_path)} - {original_size[0]}x{original_size[1]} px"
                    self.window.title(f"Просмотр: {info_text}")
                    
                except Exception as e:
                    import traceback
                    error_msg = f"Не удалось отобразить изображение:\n{str(e)}\n\n{traceback.format_exc()}"
                    messagebox.showerror("Ошибка", error_msg, parent=self.window)
            
            # Принудительно обновляем окно перед загрузкой изображения
            self.window.update_idletasks()
            # Откладываем отображение на 150ms, чтобы canvas успел полностью отрисоваться
            self.window.after(150, display_image)
            
        except Exception as e:
            import traceback
            error_msg = f"Не удалось загрузить изображение:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Ошибка", error_msg, parent=self.window)
            self.window.destroy()

class ColumnMappingDialog:
    """Диалог выбора колонок для плейсхолдера"""
    def __init__(self, parent, source_columns, placeholders, initial_data=None):
        self.result = None
        self.source_columns = list(source_columns)
        self.placeholders = placeholders
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Настройка колонки")
        self.dialog.geometry("750x850")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self.create_widgets(initial_data)
        
        self.dialog.update_idletasks()
        parent.update_idletasks()
        
        window_width = 750
        window_height = 850
        
        x = parent.winfo_rootx() + (parent.winfo_width() // 2) - (window_width // 2)
        y = parent.winfo_rooty() + (parent.winfo_height() // 2) - (window_height // 2)
        
        screen_height = self.dialog.winfo_screenheight()
        if y < 0:
            y = 20  # Отступ от верхнего края экрана
        elif y + window_height > screen_height:
            y = screen_height - window_height - 20
        
        self.dialog.geometry(f"+{x}+{y}")
        
        self.dialog.protocol("WM_DELETE_WINDOW", self.cancel)
        
        self.dialog.wait_window()
    
    def create_widgets(self, initial_data):
        # ═══════════════════════════════════════════════════════════
        # ЗАГОЛОВОК ОКНА
        # ═══════════════════════════════════════════════════════════
        title_frame = tk.Frame(self.dialog, bg=COLORS["primary"], height=60)
        title_frame.pack(fill=tk.X, side=tk.TOP)
        title_frame.pack_propagate(False)
        
        title_label = tk.Label(
            title_frame,
            text="⚙ Настройка колонки",
            font=FONTS["title"],
            bg=COLORS["primary"],
            fg="white"
        )
        title_label.pack(pady=15)
        
        btn_frame = tk.Frame(self.dialog, bg=COLORS["bg_tertiary"])
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=20, pady=15)
        
        cancel_btn = create_modern_button(
            btn_frame,
            text="✗ Отмена",
            command=self.cancel,
            style="secondary",
            width=110,
            height=40,
            tooltip="Отменить создание колонки"
        )
        cancel_btn.pack(side=tk.LEFT)
        
        save_btn = create_modern_button(
            btn_frame,
            text="✓ Сохранить",
            command=self.save,
            style="success",
            width=110,
            height=40,
            tooltip="Сохранить настройки колонки"
        )
        save_btn.pack(side=tk.LEFT, padx=(10, 0))
        
        canvas = tk.Canvas(self.dialog, bg=COLORS["bg_secondary"], highlightthickness=0)
        scrollbar = tk.Scrollbar(self.dialog, orient="vertical", command=canvas.yview)
        main_frame = tk.Frame(canvas, bg=COLORS["bg_secondary"])
        
        main_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=main_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=20, pady=(20, 10))
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=(20, 10))
        
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        def bind_to_mousewheel(widget):
            """Рекурсивно привязываем прокрутку ко всем виджетам"""
            widget.bind("<MouseWheel>", on_mousewheel)
            for child in widget.winfo_children():
                bind_to_mousewheel(child)
        
        # Привязываем ко всем виджетам
        bind_to_mousewheel(self.dialog)
        
        # Отвязываем события при закрытии окна
        def cleanup():
            try:
                self.dialog.unbind_all("<MouseWheel>")
            except:
                pass
        
        self.cleanup = cleanup
        
        # ═══════════════════════════════════════════════════════════
        # 1. ВЫБОР НАЗВАНИЯ КОЛОНКИ
        # ═══════════════════════════════════════════════════════════
        name_section = tk.LabelFrame(
            main_frame,
            text=" 1. Название новой колонки ",
            bg=COLORS["bg_secondary"],
            font=FONTS["heading"],
            fg=COLORS["text_primary"],
            padx=15,
            pady=10,
            relief=tk.SOLID,
            borderwidth=1
        )
        name_section.pack(fill=tk.X, pady=(0, 15))
        
        # Радиокнопки выбора типа названия
        self.name_type = tk.StringVar(value="placeholder")
        
        tk.Radiobutton(
            name_section,
            text="Использовать плейсхолдер",
            variable=self.name_type,
            value="placeholder",
            bg=COLORS["bg_secondary"],
            font=FONTS["body"],
            command=self.update_name_widgets
        ).pack(anchor=tk.W, pady=(0, 5))
        
        # Комбобокс для плейсхолдеров
        excel_placeholders = [p['name'] for p in self.placeholders if p.get('source_type') == 'excel']
        self.placeholder_var = tk.StringVar()
        self.placeholder_combo = ctk.CTkComboBox(
            name_section,
            variable=self.placeholder_var,
            values=excel_placeholders if excel_placeholders else ["Нет доступных плейсхолдеров"],
            font=FONTS["body"],
            state="readonly",
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            button_color=COLORS["primary"],
            button_hover_color=COLORS["primary_hover"],
            dropdown_fg_color=COLORS["bg_primary"],
            height=32
        )
        self.placeholder_combo.pack(fill=tk.X, padx=(25, 0), pady=(0, 10))
        if excel_placeholders:
            self.placeholder_combo.set(excel_placeholders[0])
        set_combobox_cursor(self.placeholder_combo)
        ToolTip(self.placeholder_combo, "Выберите плейсхолдер из списка")
        
        tk.Radiobutton(
            name_section,
            text="Своё название колонки",
            variable=self.name_type,
            value="custom",
            bg=COLORS["bg_secondary"],
            font=FONTS["body"],
            command=self.update_name_widgets
        ).pack(anchor=tk.W, pady=(0, 5))
        
        self.custom_name_var = tk.StringVar()
        self.custom_name_entry = ctk.CTkEntry(
            name_section,
            textvariable=self.custom_name_var,
            font=FONTS["body"],
            fg_color=COLORS["bg_primary"],
            border_color=COLORS["border"],
            height=32,
            state="disabled"
        )
        self.custom_name_entry.pack(fill=tk.X, padx=(25, 0))
        enable_field_shortcuts(self.custom_name_entry)
        add_context_menu(self.custom_name_entry)
        
        # ═══════════════════════════════════════════════════════════
        # 2. ВЫБОР ФУНКЦИИ ОБРАБОТКИ
        # ═══════════════════════════════════════════════════════════
        function_section = tk.LabelFrame(
            main_frame,
            text=" 2. Функция обработки данных ",
            bg=COLORS["bg_secondary"],
            font=FONTS["heading"],
            fg=COLORS["text_primary"],
            padx=15,
            pady=10,
            relief=tk.SOLID,
            borderwidth=1
        )
        function_section.pack(fill=tk.X, pady=(0, 15))
        
        self.function_type = tk.StringVar(value="default")
        
        functions = [
            ("default", "Дефолтная функция", "Объединение через пробел с автоформатированием дат и в/ч"),
            ("first", "Только первое значение", "Берётся значение только из первой выбранной колонки"),
            ("comma", "Объединение через запятую", "Значения объединяются через запятую: значение1,значение2"),
            ("comma_space", "Объединение через запятую с пробелом", "Значения объединяются: значение1, значение2"),
            ("upper", "Все в верхний регистр", "ВСЕ БУКВЫ ЗАГЛАВНЫЕ"),
            ("lower", "Все в нижний регистр", "все буквы строчные"),
            ("newline", "Каждая колонка с новой строки", "Значения разделяются переносом строки"),
            ("custom", "Своя функция (Python код)", "Введите Python код для обработки значений")
        ]
        
        for value, label, desc in functions:
            tk.Radiobutton(
                function_section,
                text=f"{label}",
                variable=self.function_type,
                value=value,
                bg=COLORS["bg_secondary"],
                font=FONTS["body"],
                command=self.update_function_widgets
            ).pack(anchor=tk.W, pady=(5, 0))
            
            tk.Label(
                function_section,
                text=f"     {desc}",
                bg=COLORS["bg_secondary"],
                font=FONTS["small"],
                fg=COLORS["text_secondary"],
                wraplength=600,
                justify=tk.LEFT
            ).pack(anchor=tk.W, pady=(0, 5))
        
        # Поле для кастомной функции
        custom_func_frame = tk.Frame(function_section, bg=COLORS["bg_secondary"])
        custom_func_frame.pack(fill=tk.BOTH, expand=True, padx=(25, 0), pady=(5, 0))
        
        tk.Label(
            custom_func_frame,
            text="Python функция (values - список значений из выбранных колонок):",
            bg=COLORS["bg_secondary"],
            font=FONTS["small"],
            fg=COLORS["text_primary"]
        ).pack(anchor=tk.W, pady=(0, 5))
        
        self.custom_function_text = ScrolledText(
            custom_func_frame,
            height=6,
            width=60,
            font=FONTS["mono"],
            state="disabled",
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"],
            relief=tk.FLAT,
            borderwidth=1,
            insertbackground=COLORS["primary"]
        )
        self.custom_function_text.pack(fill=tk.BOTH, expand=True)
        enable_field_shortcuts(self.custom_function_text)
        add_context_menu(self.custom_function_text)
        ToolTip(self.custom_function_text, "Введите Python код для обработки значений колонок")
        
        help_btn = create_modern_button(
            custom_func_frame,
            text="Показать примеры функций",
            command=self.show_function_help,
            style="primary",
            tooltip="Посмотреть примеры пользовательских функций"
        )
        help_btn.pack(anchor=tk.W, pady=(5, 0))
        
        # ═══════════════════════════════════════════════════════════
        # 3. ВЫБОР ИСХОДНЫХ КОЛОНОК
        # ═══════════════════════════════════════════════════════════
        columns_section = tk.LabelFrame(
            main_frame,
            text=" 3. Выбор исходных колонок ",
            bg=COLORS["bg_secondary"],
            font=FONTS["heading"],
            fg=COLORS["text_primary"],
            padx=15,
            pady=10,
            relief=tk.SOLID,
            borderwidth=1
        )
        columns_section.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        tk.Label(
            columns_section,
            text="Выберите колонки для объединения (можно несколько):",
            bg=COLORS["bg_secondary"],
            font=FONTS["body"],
            fg=COLORS["text_primary"]
        ).pack(anchor=tk.W, pady=(0, 10))
        
        # Скроллируемый список чекбоксов
        checkbox_canvas = tk.Canvas(
            columns_section, 
            bg=COLORS["card_bg"], 
            highlightthickness=1,
            highlightbackground=COLORS["border"],
            height=150
        )
        checkbox_scrollbar = tk.Scrollbar(
            columns_section, 
            orient="vertical", 
            command=checkbox_canvas.yview,
            width=12
        )
        checkbox_frame = tk.Frame(checkbox_canvas, bg=COLORS["card_bg"])
        
        checkbox_frame.bind(
            "<Configure>",
            lambda e: checkbox_canvas.configure(scrollregion=checkbox_canvas.bbox("all"))
        )
        
        checkbox_canvas.create_window((0, 0), window=checkbox_frame, anchor="nw")
        checkbox_canvas.configure(yscrollcommand=checkbox_scrollbar.set)
        
        checkbox_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        checkbox_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Прокрутка мышью для чекбоксов
        def on_checkbox_mousewheel(event):
            checkbox_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            return "break"  # Останавливаем всплытие события
        
        # Привязываем к checkbox canvas и всем чекбоксам
        def bind_checkbox_mousewheel(widget):
            widget.bind("<MouseWheel>", on_checkbox_mousewheel)
            for child in widget.winfo_children():
                bind_checkbox_mousewheel(child)
        
        bind_checkbox_mousewheel(checkbox_frame)
        checkbox_canvas.bind("<MouseWheel>", on_checkbox_mousewheel)
        
        # Чекбоксы для каждой колонки
        self.column_vars = {}
        for col in self.source_columns:
            var = tk.BooleanVar(value=False)
            if initial_data and col in initial_data.get('source_columns', []):
                var.set(True)
            
            cb = tk.Checkbutton(
                checkbox_frame,
                text=col,
                variable=var,
                bg=COLORS["card_bg"],
                fg=COLORS["text_primary"],
                activebackground=COLORS["card_bg"],
                activeforeground=COLORS["primary"],
                selectcolor=COLORS["card_bg"],
                font=FONTS["body"],
                anchor=tk.W,
                cursor="hand2",
                command=self.update_order_list
            )
            cb.pack(fill=tk.X, padx=5, pady=2)
            cb.bind("<MouseWheel>", on_checkbox_mousewheel)
            ToolTip(cb, f"Использовать колонку '{col}' в объединении")
            self.column_vars[col] = var
        
        # ═══════════════════════════════════════════════════════════
        # 4. ПОРЯДОК ОБЪЕДИНЕНИЯ
        # ═══════════════════════════════════════════════════════════
        order_section = tk.LabelFrame(
            main_frame,
            text=" 4. Порядок объединения колонок ",
            bg=COLORS["bg_secondary"],
            font=FONTS["heading"],
            fg=COLORS["text_primary"],
            padx=15,
            pady=10,
            relief=tk.SOLID,
            borderwidth=1
        )
        order_section.pack(fill=tk.X, pady=(0, 0))
        
        order_container = tk.Frame(order_section, bg=COLORS["bg_secondary"])
        order_container.pack(fill=tk.BOTH, pady=(5, 0))
        
        # Список с порядком колонок
        self.order_listbox = tk.Listbox(
            order_container,
            font=FONTS["body"],
            height=5,
            selectmode=tk.SINGLE,
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"],
            selectbackground=COLORS["primary"],
            selectforeground="white",
            relief=tk.SOLID,
            borderwidth=1,
            highlightthickness=0
        )
        self.order_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        ToolTip(self.order_listbox, "Текущий порядок объединения колонок")
        
        order_buttons_frame = tk.Frame(order_container, bg=COLORS["bg_secondary"])
        order_buttons_frame.pack(side=tk.LEFT, fill=tk.Y)
        
        up_btn = create_modern_button(
            order_buttons_frame,
            text="↑ Вверх",
            command=self.move_up,
            style="primary",
            width=80,
            tooltip="Переместить выбранную колонку вверх"
        )
        up_btn.pack(pady=(0, 5))
        
        down_btn = create_modern_button(
            order_buttons_frame,
            text="↓ Вниз",
            command=self.move_down,
            style="primary",
            width=80,
            tooltip="Переместить выбранную колонку вниз"
        )
        down_btn.pack()
        
        # ═══════════════════════════════════════════════════════════
        # ИНИЦИАЛИЗАЦИЯ ДАННЫХ
        # ═══════════════════════════════════════════════════════════
        if initial_data:
            if initial_data.get('name_type') == 'custom':
                self.name_type.set('custom')
                self.custom_name_var.set(initial_data.get('column_name', ''))
            else:
                self.name_type.set('placeholder')
                self.placeholder_var.set(initial_data.get('column_name', ''))
            
            self.function_type.set(initial_data.get('function_type', 'default'))
            if initial_data.get('custom_function'):
                self.custom_function_text.delete(1.0, tk.END)
                self.custom_function_text.insert(1.0, initial_data['custom_function'])
            
            for col in initial_data.get('source_columns', []):
                self.order_listbox.insert(tk.END, col)
        
        self.update_name_widgets()
        self.update_function_widgets()
    
    def update_name_widgets(self):
        """Включить/выключить виджеты названия колонки"""
        if self.name_type.get() == "placeholder":
            self.placeholder_combo.configure(state="readonly")
            self.custom_name_entry.configure(state="disabled")
        else:
            self.placeholder_combo.configure(state="disabled")
            self.custom_name_entry.configure(state="normal")
    
    def update_function_widgets(self):
        """Включить/выключить виджеты функции"""
        if self.function_type.get() == "custom":
            self.custom_function_text.config(state="normal", bg=COLORS["card_bg"])
        else:
            self.custom_function_text.config(state="disabled", bg=COLORS["bg_tertiary"])
    
    def show_function_help(self):
        """Показать примеры функций"""
        help_text = """📚 ПРИМЕРЫ ПОЛЬЗОВАТЕЛЬСКИХ ФУНКЦИЙ

Доступные переменные:
• values - список значений из выбранных колонок

Способы возврата результата:
• return результат  (если используете return)
• result = результат  (если НЕ используете return)

═══════════════════════════════════════════════════

Пример 1: Объединение через точку с запятой
result = '; '.join(values)

Пример 2: Взять только первые 3 значения
result = ' '.join(values[:3])

Пример 3: Убрать пустые значения и объединить
filtered = [v for v in values if v.strip()]
result = ', '.join(filtered)

Пример 4: Добавить префикс к каждому значению
result = ' | '.join([f"№{v}" for v in values])

Пример 5: Склеить без пробелов
result = ''.join(values)

Пример 6: Взять из второй колонки (если есть)
result = values[1] if len(values) > 1 else ''

Пример 7: Форматирование с условием (с return)
if len(values) > 1:
    return f"{values[0]} ({values[1]})"
else:
    return values[0] if values else ''

Пример 8: Каждое слово с заглавной буквы
result = ' '.join(values).title()

Пример 9: Добавить номер по порядку
parts = [f"{i+1}. {v}" for i, v in enumerate(values)]
result = '\\n'.join(parts)

Пример 10: Обработка с проверкой на пустоту
filtered = []
for v in values:
    if v and len(v) > 3:
        filtered.append(v.upper())
result = ' / '.join(filtered) if filtered else 'Нет данных'

Пример 11: Расчёты (если значения - числа)
numbers = [float(v) for v in values if v.replace('.', '').isdigit()]
result = str(sum(numbers)) if numbers else '0'

Пример 12: Замена текста
result = ' '.join(values).replace('старое', 'новое')
"""
        
        help_window = tk.Toplevel(self.dialog)
        help_window.title("Примеры функций")
        help_window.geometry("750x650")
        help_window.transient(self.dialog)
        help_window.configure(bg=COLORS["bg_primary"])
        
        # Современный заголовок
        header_frame = tk.Frame(help_window, bg=COLORS["bg_secondary"], height=60)
        header_frame.pack(fill=tk.X, pady=(0, 15))
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(
            header_frame,
            text="📚 Примеры функций",
            font=(FONTS["heading"][0], 16, "bold"),
            bg=COLORS["bg_secondary"],
            fg=COLORS["text_primary"]
        )
        title_label.pack(expand=True)
        
        # Контейнер для текста
        text_container = tk.Frame(help_window, bg=COLORS["bg_primary"])
        text_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 15))
        
        text_widget = ScrolledText(
            text_container,
            font=("Consolas", 9),
            wrap=tk.WORD,
            padx=15,
            pady=15,
            bg=COLORS["card_bg"],
            fg=COLORS["text_primary"],
            relief=tk.SOLID,
            borderwidth=1,
            insertbackground=COLORS["primary"]
        )
        text_widget.pack(fill=tk.BOTH, expand=True)
        enable_field_shortcuts(text_widget, readonly=True)
        add_context_menu(text_widget, readonly=True)
        text_widget.insert(1.0, help_text)
        text_widget.config(state="disabled")
        
        # Контейнер для кнопки
        button_frame = tk.Frame(help_window, bg=COLORS["bg_primary"])
        button_frame.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        close_btn = create_modern_button(
            button_frame,
            text="Закрыть",
            command=help_window.destroy,
            style="primary",
            width=120
        )
        close_btn.pack()
    
    def update_order_list(self):
        """Обновить список порядка при изменении выбора"""
        selected = [col for col, var in self.column_vars.items() if var.get()]
        
        current_order = list(self.order_listbox.get(0, tk.END))
        
        new_order = [col for col in current_order if col in selected]
        
        for col in selected:
            if col not in new_order:
                new_order.append(col)
        
        self.order_listbox.delete(0, tk.END)
        for col in new_order:
            self.order_listbox.insert(tk.END, col)
    
    def move_up(self):
        """Переместить выбранную колонку вверх"""
        selection = self.order_listbox.curselection()
        if not selection or selection[0] == 0:
            return
        
        index = selection[0]
        item = self.order_listbox.get(index)
        self.order_listbox.delete(index)
        self.order_listbox.insert(index - 1, item)
        self.order_listbox.selection_set(index - 1)
    
    def move_down(self):
        """Переместить выбранную колонку вниз"""
        selection = self.order_listbox.curselection()
        if not selection or selection[0] == self.order_listbox.size() - 1:
            return
        
        index = selection[0]
        item = self.order_listbox.get(index)
        self.order_listbox.delete(index)
        self.order_listbox.insert(index + 1, item)
        self.order_listbox.selection_set(index + 1)
    
    def cancel(self):
        """Отмена с очисткой событий"""
        self.cleanup()
        self.dialog.destroy()
    
    def save(self):
        if self.name_type.get() == "placeholder":
            column_name = self.placeholder_var.get()
            if not column_name:
                messagebox.showwarning("Предупреждение", "Выберите плейсхолдер")
                return
        else:
            column_name = self.custom_name_var.get().strip()
            if not column_name:
                messagebox.showwarning("Предупреждение", "Введите название колонки")
                return
        
        # Собираем выбранные колонки в порядке из listbox
        selected_columns = list(self.order_listbox.get(0, tk.END))
        
        if not selected_columns:
            messagebox.showwarning("Предупреждение", "Выберите хотя бы одну исходную колонку")
            return
        
        custom_function = None
        if self.function_type.get() == "custom":
            custom_function = self.custom_function_text.get(1.0, tk.END).strip()
            if not custom_function:
                messagebox.showwarning("Предупреждение", "Введите код функции")
                return
            
            # Простая проверка синтаксиса
            try:
                compile(custom_function, '<string>', 'exec')
            except SyntaxError as e:
                messagebox.showerror("Ошибка синтаксиса", f"Ошибка в коде функции:\n{e}")
                return
        
        self.result = {
            "column_name": column_name,
            "name_type": self.name_type.get(),
            "source_columns": selected_columns,
            "function_type": self.function_type.get(),
            "custom_function": custom_function
        }
        
        self.cleanup()
        self.dialog.destroy()

def main():
    # Защита для multiprocessing в Windows
    _ensure_concurrent_imports()
    multiprocessing.freeze_support()
    
    if TKDND_AVAILABLE:
        try:
            root = TkinterDnD.Tk()
        except:
            root = tk.Tk()
    else:
        root = tk.Tk()
    
    setup_global_entry_shortcuts(root)
    
    app = GenerationDocApp(root)
    root.mainloop()

def setup_global_entry_shortcuts(root):
    """Настройка глобальных горячих клавиш для всех Entry виджетов"""
    
    def get_actual_widget(widget):
        """Получить реальный tk.Entry или tk.Text из CustomTkinter виджета"""
        if hasattr(widget, '_entry'):
            return widget._entry
        elif hasattr(widget, '_textbox'):
            return widget._textbox
        elif hasattr(widget, 'winfo_class'):
            wclass = widget.winfo_class()
            if wclass in ('Entry', 'Text'):
                return widget
        return widget
    
    def find_focused_widget():
        """Найти виджет, который сейчас в фокусе"""
        try:
            focused = root.focus_get()
            if focused:
                actual = get_actual_widget(focused)
                return actual
        except:
            pass
        return None
    
    def universal_key_handler(event):
        """Универсальный обработчик клавиш с поддержкой Ctrl+буква"""
        # State: 4 = Control на Linux/Mac, 12 = Control на Windows (8 + 4)
        # 8 = NumLock, 4 = Control
        is_ctrl = (event.state & 0x4) != 0
        
        if not is_ctrl:
            return None  # Пропускаем обычные клавиши
        
        # Используем keycode для независимости от раскладки
        # Windows keycodes: A=65, C=67, V=86, X=88
        keycode = event.keycode
        keysym_lower = event.keysym.lower()
        char_lower = event.char.lower() if event.char else ""
        
        # Также проверяем keysym и char для совместимости
        
        # Ctrl+C (keycode 67 на Windows, 54 на Linux)
        if keycode == 67 or keycode == 54 or keysym_lower in ('c', 'с') or char_lower in ('c', 'с'):
            return handle_copy(event)
        
        # Ctrl+X (keycode 88 на Windows, 53 на Linux)
        elif keycode == 88 or keycode == 53 or keysym_lower in ('x', 'х') or char_lower in ('x', 'х'):
            return handle_cut(event)
        
        # Ctrl+V (keycode 86 на Windows, 55 на Linux)
        elif keycode == 86 or keycode == 55 or keysym_lower in ('v', 'м') or char_lower in ('v', 'м'):
            return handle_paste(event)
        
        # Ctrl+A (keycode 65 на Windows, 38 на Linux)
        elif keycode == 65 or keycode == 38 or keysym_lower in ('a', 'ф') or char_lower in ('a', 'ф'):
            return handle_select_all(event)
        
        return None  # Пропускаем другие комбинации
    
    def handle_copy(event):
        widget = find_focused_widget()
        if widget is None:
            return None
        
        try:
            wclass = widget.winfo_class()
            
            if wclass == 'Text' or isinstance(widget, (tk.Text, ScrolledText)):
                if widget.tag_ranges(tk.SEL):
                    text = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
                    root.clipboard_clear()
                    root.clipboard_append(text)
                    root.update()
            elif wclass == 'Entry' or isinstance(widget, tk.Entry):
                if widget.selection_present():
                    text = widget.selection_get()
                    root.clipboard_clear()
                    root.clipboard_append(text)
                    root.update()
            return "break"
        except:
            pass
        return None
    
    def handle_cut(event):
        widget = find_focused_widget()
        if widget is None:
            return None
        
        try:
            wclass = widget.winfo_class()
            
            if wclass == 'Text' or isinstance(widget, (tk.Text, ScrolledText)):
                if widget.tag_ranges(tk.SEL):
                    text = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
                    root.clipboard_clear()
                    root.clipboard_append(text)
                    root.update()
                    widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
            elif wclass == 'Entry' or isinstance(widget, tk.Entry):
                if widget.selection_present():
                    text = widget.selection_get()
                    root.clipboard_clear()
                    root.clipboard_append(text)
                    root.update()
                    widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
            return "break"
        except:
            pass
        return None
    
    def handle_paste(event):
        widget = find_focused_widget()
        if widget is None:
            return None
        
        try:
            clipboard_text = root.clipboard_get()
            wclass = widget.winfo_class()
            
            if wclass == 'Text' or isinstance(widget, (tk.Text, ScrolledText)):
                if widget.tag_ranges(tk.SEL):
                    widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
                widget.insert(tk.INSERT, clipboard_text)
            elif wclass == 'Entry' or isinstance(widget, tk.Entry):
                if widget.selection_present():
                    widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
                widget.insert(widget.index(tk.INSERT), clipboard_text)
            return "break"
        except:
            pass
        return None
    
    def handle_select_all(event):
        widget = find_focused_widget()
        if widget is None:
            return None
        
        try:
            wclass = widget.winfo_class()
            
            if wclass == 'Text' or isinstance(widget, (tk.Text, ScrolledText)):
                widget.tag_remove(tk.SEL, "1.0", tk.END)
                widget.tag_add(tk.SEL, "1.0", tk.END)
                widget.mark_set(tk.INSERT, "1.0")
                widget.see(tk.INSERT)
            elif wclass == 'Entry' or isinstance(widget, tk.Entry):
                widget.select_range(0, tk.END)
                widget.icursor(tk.END)
            return "break"
        except:
            pass
        return None
    
    # Привязываем универсальный обработчик ко всем событиям клавиатуры
    root.bind_all("<KeyPress>", universal_key_handler, add=True)

# ══════════════════════════════════════════════════════════════════════
# ОКНО ВСТАВКИ ПЛЕЙСХОЛДЕРОВ (ВИЗУАЛЬНЫЙ РЕДАКТОР)
# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    main()
